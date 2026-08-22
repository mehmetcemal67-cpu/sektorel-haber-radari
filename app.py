import streamlit as st
import pandas as pd
import requests
import concurrent.futures
import xml.etree.ElementTree as ET
import re, html, json
import sqlite3, hashlib
from pathlib import Path
from datetime import datetime, timedelta, timezone, date
from urllib.parse import urlparse
from email.utils import parsedate_to_datetime
from io import BytesIO
from bs4 import BeautifulSoup
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn


# ============================================================
# V104 — MEVCUT BÖLÜMLERİ OLGUNLAŞTIRMA
# 1) Aynı olayın daha güçlü tekilleştirilmesi
# 2) Durum bilgisinin URL'ye değil olay kimliğine de dayanması
# 3) "Dünden Beri Ne Değişti?" yalnız gerçek/maddi değişiklikler
# 4) Vardiya Başlangıç Özeti: seçici, 5–8 tekil gelişme
# ============================================================

_V104_ENTITY_HINTS = {
    'tüik','tcmb','epdk','tübitak','kosgeb','tse','türkpatent','ssb','aselsan','tusaş','tusas',
    'roketsan','havelsan','baykar','togg','mke','kaan','kızılelma','kizilelma','hisar','siper',
    'teknofest','tcg','anadolu','turksat','türksat','thk','thy','turkcell','türk telekom'
}
_V104_GENERIC_EVENT_WORDS = {
    'türkiye','türk','sanayi','teknoloji','haber','son','yeni','ilk','bugün','dün','açıklama',
    'açıkladı','belirtti','duyurdu','başladı','gerçekleşti','oldu','edildi','yapıldı','kapsamında',
    'milyon','milyar','bin','yüzde','oran','veri','verileri','program','proje'
}

def _v104_event_tokens(title, summary=''):
    """
    Olay eşleştirmesinde yalnız başlık kelimelerine bağlı kalmaz.
    Kurum/ürün/yer/özgül sayı ve eylem çekirdeğini kısa bir imzaya dönüştürür.
    """
    title_n=norm(title)
    summary_n=norm(summary)
    title_tokens=set(_title_tokens(title))
    # Özellikle başlıkta geçen ayırt edici kurum/ürün kelimelerini koru.
    entities={x for x in re.findall(r'[a-z0-9çğıöşü]+',title_n)
              if x in _V104_ENTITY_HINTS or (len(x)>=5 and x not in _V104_GENERIC_EVENT_WORDS)}
    nums=set(re.findall(r'\b\d+(?:[.,]\d+)?\b',title_n))
    # Başlık çok kısaysa özetten sınırlı destek al.
    if len(title_tokens)<4:
        extra=[x for x in _title_tokens(summary_n) if x not in _V104_GENERIC_EVENT_WORDS]
        title_tokens.update(extra[:5])
    return set(title_tokens)|entities|nums

def _v104_event_similarity(a_title,a_summary,b_title,b_summary):
    a=_v104_event_tokens(a_title,a_summary)
    b=_v104_event_tokens(b_title,b_summary)
    if not a or not b:
        return 0.0
    jac=len(a&b)/max(1,len(a|b))
    # Aynı kurum/ürün çekirdeği varsa eşleşmeyi destekle; tek başına yeterli sayma.
    ae={x for x in a if x in _V104_ENTITY_HINTS}
    be={x for x in b if x in _V104_ENTITY_HINTS}
    entity_bonus=0.12 if (ae&be) else 0.0
    # Aynı özgül sayı/tarih varsa küçük destek.
    an={x for x in a if re.fullmatch(r'\d+(?:[.,]\d+)?',x)}
    bn={x for x in b if re.fullmatch(r'\d+(?:[.,]\d+)?',x)}
    number_bonus=0.06 if (an&bn) else 0.0
    return min(1.0,jac+entity_bonus+number_bonus)

def _v104_event_representatives(df):
    """Mevcut Olay_ID'leri ikinci kez birleştirerek yanlış çoğalmayı azaltır."""
    if df is None or df.empty:
        return df
    x=df.copy()
    if 'Tarih_dt' in x.columns:
        x['Tarih_dt']=pd.to_datetime(x['Tarih_dt'],utc=True,errors='coerce')
    # Önce mevcut kümelerden temsilci çıkar.
    if 'Olay_ID' in x.columns:
        reps=[]
        for oid,g in x.groupby('Olay_ID',dropna=False):
            g=g.sort_values('Tarih_dt',ascending=False,na_position='last')
            r=g.iloc[0].copy()
            r['_v104_members']=list(g.index)
            r['_v104_summary']=' '.join(g.get('İçerik_Özeti',pd.Series(dtype=str)).fillna('').astype(str).head(5))
            reps.append(r)
        reps=pd.DataFrame(reps)
    else:
        reps=x.copy()
        reps['_v104_members']=[[i] for i in reps.index]
        reps['_v104_summary']=reps.get('İçerik_Özeti','')

    # Ters indeks: performansı korur.
    token_index={}
    clusters=[]
    for _,r in reps.sort_values('Tarih_dt',ascending=False,na_position='last').iterrows():
        toks=_v104_event_tokens(r.get('Başlık',''),r.get('_v104_summary',''))
        candidates=set()
        for t in toks:
            candidates.update(token_index.get(t,set()))
        best=None; best_sim=0.0
        for ci in candidates:
            cr=clusters[ci]['rep']
            sim=_v104_event_similarity(
                r.get('Başlık',''),r.get('_v104_summary',''),
                cr.get('Başlık',''),cr.get('_v104_summary','')
            )
            if sim>best_sim:
                best_sim=sim; best=ci
        threshold=0.52 if len(toks)>=6 else 0.60
        if best is None or best_sim<threshold:
            best=len(clusters)
            clusters.append({'rep':r,'members':list(r['_v104_members'])})
        else:
            clusters[best]['members'].extend(r['_v104_members'])
            # En yeni temsilciyi koru.
            try:
                if pd.to_datetime(r.get('Tarih_dt'),utc=True,errors='coerce') > pd.to_datetime(clusters[best]['rep'].get('Tarih_dt'),utc=True,errors='coerce'):
                    clusters[best]['rep']=r
            except Exception:
                pass
        for t in toks:
            token_index.setdefault(t,set()).add(best)

    rows=[]
    for ci,c in enumerate(clusters,1):
        g=x.loc[list(dict.fromkeys(c['members']))].copy()
        g=g.sort_values('Tarih_dt',ascending=False,na_position='last')
        rep=g.iloc[0].copy()
        rep['Olay_ID']=f'V104-{ci:04d}'
        rep['Olay_Haber_Sayisi']=len(g)
        domains={str(v) for v in g.get('Domain',pd.Series(dtype=str)).tolist() if str(v).strip()}
        rep['Olay_Kaynak_Sayisi']=max(len(domains),int(pd.to_numeric(g.get('Olay_Kaynak_Sayisi',0),errors='coerce').fillna(0).max() or 0))
        # Aynı olayın en yüksek risk/teyit bilgisini kaybetme.
        rep['Risk_Skoru']=int(pd.to_numeric(g.get('Risk_Skoru',0),errors='coerce').fillna(0).max() or 0)
        rows.append(rep)
    return pd.DataFrame(rows).drop(columns=['_v104_members','_v104_summary'],errors='ignore')

def _v104_event_status_sets():
    """
    Durumu hem URL hem başlık hem de olay-token imzasıyla indeksler.
    Aynı olay farklı kaynaktan görünse bile kullanıcı işlemi kaybolmaz.
    """
    cached=st.session_state.get('_v104_status_cache')
    if cached is not None:
        return cached
    result={'imp':set(),'akt':set(),'notes':set(),'pres':set()}
    if not _init_history_db():
        return result
    try:
        with _history_connect() as conn:
            for table,key in [
                ('important_basket','imp'),('osint_report_basket','akt'),
                ('note_history','notes'),('presentation_basket','pres')
            ]:
                for title,url in conn.execute(f"SELECT title,url FROM {table}").fetchall():
                    title=str(title or ''); url=str(url or '').strip()
                    if url: result[key].add('U:'+url)
                    tk=title_key(title)
                    if tk: result[key].add('T:'+tk)
                    sig=' '.join(sorted(_v104_event_tokens(title,'')))
                    if sig: result[key].add('E:'+sig)
    except Exception:
        pass
    st.session_state['_v104_status_cache']=result
    return result

def _v104_row_status_keys(r):
    url=str(r.get('URL',r.get('url','')) or '').strip()
    title=str(r.get('Başlık',r.get('title','')) or '')
    summary=str(r.get('İçerik_Özeti',r.get('summary','')) or '')
    keys=set()
    if url: keys.add('U:'+url)
    tk=title_key(title)
    if tk: keys.add('T:'+tk)
    sig=' '.join(sorted(_v104_event_tokens(title,summary)))
    if sig: keys.add('E:'+sig)
    return keys

def _v63_add_status_badges(df):
    """V104 — tüm tablolarda olay bazlı güvenilir Durum sütunu."""
    if df is None or df.empty:
        return df
    out=df.copy()
    sets=_v104_event_status_sets()
    def badge(r):
        keys=_v104_row_status_keys(r)
        b=[]
        if keys & sets['pres']:  b.append('🖥️ Sunum Sepetinde')
        if keys & sets['imp']:   b.append('📌 Önemli Gelişmelerde')
        if keys & sets['notes']: b.append('📝 Bilgi Notu Yapıldı')
        if keys & sets['akt']:   b.append('📁 AKT Sepetinde')
        return ' • '.join(b) if b else '—'
    out['Durum']=out.apply(badge,axis=1)
    return out

def _v73_invalidate_status_cache():
    st.session_state.pop('_v73_status_sets_cache',None)
    st.session_state.pop('_v104_status_cache',None)

def _v104_material_change(prev,cur):
    """Kaynak sayısındaki sıradan artışı tek başına 'yeni bilgi' saymaz."""
    prev_risk=int(prev.get('risk_score') or 0)
    cur_risk=int(cur.get('risk_score') or 0)
    risk_up=(cur_risk>=prev_risk+15) or (_risk_rank(cur.get('risk_status',''))>_risk_rank(prev.get('risk_status','')))
    verify_up=_verification_rank(cur.get('verification',''))>_verification_rank(prev.get('verification',''))

    prev_text=(prev.get('title') or '')+' '+(prev.get('summary') or '')
    cur_text=(cur.get('title') or '')+' '+(cur.get('summary') or '')
    prev_tokens=set(_history_tokens(prev_text))
    cur_tokens=set(_history_tokens(cur_text))
    new_tokens=cur_tokens-prev_tokens

    # Gerçek yeni bilgi için yalnız genel kelimeler değil, sayı/kurum/özgül içerik aranır.
    prev_nums=set(re.findall(r'\b\d+(?:[.,]\d+)?\b',prev_text))
    cur_nums=set(re.findall(r'\b\d+(?:[.,]\d+)?\b',cur_text))
    new_nums=cur_nums-prev_nums
    meaningful=[t for t in new_tokens if len(t)>=5 and t not in _V104_GENERIC_EVENT_WORDS]
    materially_updated=(len(meaningful)>=8) or (len(new_nums)>=1 and len(meaningful)>=3)

    return risk_up,verify_up,materially_updated,meaningful,new_nums

def _compare_since_previous(df,current_scan_id=None):
    """
    V104 — yalnız gerçek değişiklikler:
    yeni olay / maddi yeni bilgi / risk artışı / teyit güçlenmesi.
    Aynı olayın yeni bir sitede tekrar yayımlanması tek başına değişiklik değildir.
    """
    current=_v104_event_representatives(df)
    prev_id=_previous_scan_id(current_scan_id)
    previous=_load_scan_events(prev_id)
    if current is None or current.empty:
        return pd.DataFrame(),None,None
    if previous.empty:
        return pd.DataFrame(),prev_id,None

    prev_records=[]
    for _,p in previous.iterrows():
        rec=p.to_dict()
        rec['tokens']=_history_tokens((rec.get('title') or '')+' '+(rec.get('summary') or ''))
        prev_records.append(rec)

    changes=[]
    for _,r in current.iterrows():
        c={
            'title':str(r.get('Başlık','') or ''),'source':str(r.get('Kaynak','') or ''),
            'url':str(r.get('URL','') or ''),'category':str(r.get('Kategori','') or ''),
            'summary':str(r.get('İçerik_Özeti','') or ''),'risk_score':int(r.get('Risk_Skoru',0) or 0),
            'risk_status':str(r.get('Risk_Durumu','') or ''),'verification':str(r.get('Doğrulama','') or ''),
            'source_count':int(r.get('Olay_Kaynak_Sayisi',1) or 1)
        }
        best=None; best_sim=0.0
        for p in prev_records:
            sim=_v104_event_similarity(c['title'],c['summary'],p.get('title',''),p.get('summary',''))
            if c['url'] and c['url']==str(p.get('url','') or ''):
                sim=max(sim,0.98)
            if sim>best_sim:
                best_sim=sim; best=p

        if best is None or best_sim<0.50:
            changes.append({
                'Değişim':'🆕 YENİ OLAY','Başlık':c['title'],'Kaynak':c['source'],'Kategori':c['category'],
                'Risk':c['risk_score'],'Önceki Risk':'—','Kaynak Sayısı':c['source_count'],
                'Açıklama':'Önceki taramada aynı olaya ilişkin yeterli eşleşme bulunmamıştır.',
                'URL':c['url'],'_priority':100+c['risk_score']
            })
            continue

        risk_up,verify_up,material,new_words,new_nums=_v104_material_change(best,c)
        if risk_up:
            kind='⚠️ RİSK ARTTI'
            expl=f"Risk {int(best.get('risk_score') or 0)}/100 seviyesinden {c['risk_score']}/100 seviyesine yükselmiştir."
            priority=95+c['risk_score']
        elif verify_up:
            kind='✅ TEYİT GÜÇLENDİ'
            expl=f"Doğrulama seviyesi {best.get('verification','')} düzeyinden {c['verification']} düzeyine yükselmiştir."
            priority=90+c['risk_score']
        elif material:
            kind='🔄 YENİ BİLGİ'
            bits=[]
            if new_nums: bits.append('yeni sayısal veri: '+', '.join(sorted(new_nums)[:4]))
            if new_words: bits.append('yeni içerik: '+', '.join(sorted(new_words)[:6]))
            expl='; '.join(bits) if bits else 'Olay hakkında maddi yeni bilgi tespit edilmiştir.'
            priority=80+c['risk_score']
        else:
            continue

        changes.append({
            'Değişim':kind,'Başlık':c['title'],'Kaynak':c['source'],'Kategori':c['category'],
            'Risk':c['risk_score'],'Önceki Risk':int(best.get('risk_score') or 0),
            'Kaynak Sayısı':c['source_count'],'Açıklama':expl,'URL':c['url'],'_priority':priority
        })

    out=pd.DataFrame(changes)
    if not out.empty:
        # Aynı olayın değişiklik listesinde de yalnız bir kez görünmesi.
        out['_sig']=out.apply(lambda r:' '.join(sorted(_v104_event_tokens(r.get('Başlık',''),r.get('Açıklama','')))),axis=1)
        out=out.sort_values(['_priority','Risk'],ascending=[False,False]).drop_duplicates('_sig',keep='first')
        out=out.drop(columns=['_priority','_sig'],errors='ignore')
    prev_time=str(previous.iloc[0].get('scanned_at','')) if not previous.empty else None
    return out,prev_id,prev_time

def _v104_shift_priority(r):
    text=norm(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')} {r.get('Kategori','')}")
    official=_is_official_radar_row(r) or _verification_rank(r.get('Doğrulama',''))>=4
    data_terms=['tüik','tcmb','epdk','istatistik','veri','endeks','oran','kapasite kullanım',
                'sanayi üretimi','ihracat','ithalat','istihdam','milyar','milyon','yüzde','%']
    strategic_terms=['yatırım','üretim','tesis','fabrika','çip','yarı iletken','yapay zeka','yapay zekâ',
                     'siber','ar-ge','arge','patent','teşvik','kritik teknoloji','otomotiv','enerji']
    defence_terms=['savunma','aselsan','tusaş','tusas','roketsan','havelsan','baykar','kaan','kızılelma',
                   'füze','iha','siha','uzay','uydu','teknofest']
    has_data=any(x in text for x in data_terms) or bool(re.search(r'\d',text))
    strategic=any(x in text for x in strategic_terms)
    defence=any(x in text for x in defence_terms)
    critical=(int(r.get('Risk_Skoru',0) or 0)>=70 or r.get('Duygu')=='Negatif' or
              bool(critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti',''))))
    verified=_verification_rank(r.get('Doğrulama',''))>=3

    # Kullanıcının istediği açık öncelik sırası.
    if official and has_data: tier=5
    elif official: tier=5
    elif strategic: tier=4
    elif critical: tier=3
    elif defence: tier=2
    elif verified: tier=1
    else: tier=0

    score=tier*100
    score+=min(int(r.get('Risk_Skoru',0) or 0),100)
    score+=min(int(r.get('Olay_Kaynak_Sayisi',0) or 0)*5,20)
    if has_data: score+=15
    return score,tier

def _shift_start_summary(df,current_scan_id=None):
    """
    V104 — sabah ilk bakış: 5–8 adet gerçekten önemli, olay bazında tekil gelişme.
    Öncelik: resmî veri/açıklama > stratejik sanayi-teknoloji > kritik negatif >
    savunma/uzay/teknoloji programı > yüksek teyitli yeni gelişme.
    """
    if df is None or df.empty:
        return {},pd.DataFrame(),""

    baseline,baseline_label,baseline_scan_id=_shift_baseline(current_scan_id)
    x=df.copy()
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')
    since=x[(x['Tarih_dt'].isna()) | (x['Tarih_dt']>=baseline)].copy() if baseline is not None else x.copy()

    changes,_,_=_compare_since_previous(df,current_scan_id)
    new_events=int(changes['Tür'].astype(str).str.contains('YENİ OLAY').sum()) if not changes.empty else 0
    risk_up=int(changes['Tür'].astype(str).str.contains('RİSK ARTTI').sum()) if not changes.empty else 0
    verify_up=int(changes['Tür'].astype(str).str.contains('TEYİT').sum()) if not changes.empty else 0

    reps=_v104_event_representatives(since) if not since.empty else pd.DataFrame()
    if not reps.empty:
        scored=reps.apply(_v104_shift_priority,axis=1)
        reps['_V104_Puan']=[v[0] for v in scored]
        reps['_V104_Kademe']=[v[1] for v in scored]
        # Önemsiz/generic içerik sabah özetini doldurmasın.
        eligible=reps[reps['_V104_Kademe']>0].copy()
        eligible=eligible.sort_values(['_V104_Puan','Tarih_dt'],ascending=[False,False],na_position='last')
        # En az 5 uygun olay varsa 5; güçlü aday çoksa en fazla 8.
        top_n=min(8,max(5,min(len(eligible),8))) if len(eligible)>=5 else len(eligible)
        top=eligible.head(top_n).drop(columns=['_V104_Puan','_V104_Kademe'],errors='ignore')
    else:
        top=pd.DataFrame()

    high=0; osb=0
    if not reps.empty:
        high=int((reps.get('Risk_Durumu',pd.Series(dtype=str))=='Yüksek Risk').sum())
        osb=sum(bool(critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti',''))) for _,r in reps.iterrows())

    stats={
        'new_news':len(since),
        'new_important_events':new_events,
        'high_risk':high,
        'risk_up':risk_up,
        'verify_up':verify_up,
        'osb':osb,
        'baseline_label':baseline_label
    }
    return stats,top,baseline_label

# ============================================================
# /V104
# ============================================================

# V106 — V105 düzeltmesi:
# Yalnızca görünür 'Resmî Açıklama – Medya Karşılaştırması' paneli kaldırılmıştır.
# Sorgu/negatif tarama yardımcı fonksiyonları ve V104 çekirdeği korunmuştur.

# V107 — Sepete eklemede aynı olayın mevcut taramadaki kaynakları otomatik zenginleştirilir.
# Yerel/kısa sürüm seçilse bile resmî/ana akım/daha ayrıntılı sürüm ve kritik veriler birleştirilir.
# Ek ağ isteği yapılmaz; V106 kararlı çekirdeği korunur.

# V108 — V107 zenginleştirme TypeError düzeltmesi:
# _v107_unique_sentences içindeki hashlenemeyen set, frozenset olarak saklanmaktadır.
# V107 olay/kaynak zenginleştirme mantığı korunmuştur.


# ============================================================
# V109 — PANEL GELİŞTİRMELERİ
# ============================================================

def _v109_numbers(text):
    return set(re.findall(
        r'\b\d+(?:[.,]\d+)?(?:\s*(?:%|yüzde|milyon|milyar|bin|adet|tl|dolar|euro|avro|mw|gw|gwh|mwh|km))?',
        norm(text)
    ))

def _v109_sentences(text):
    out=[]
    for s in _sentence_chunks(_clean_note_text(text)):
        s=_clean_note_text(s).strip()
        if len(s)>=35:
            out.append(s)
    return out

def _v109_direct_difference(prev,cur,kind):
    prev_text=_clean_note_text((prev.get('title') or '')+' '+(prev.get('summary') or ''))
    cur_text=_clean_note_text((cur.get('title') or '')+' '+(cur.get('summary') or ''))

    if 'RİSK ARTTI' in kind:
        return (
            f"Önceki taramada risk düzeyi {int(prev.get('risk_score') or 0)}/100 iken, "
            f"yeni taramada {int(cur.get('risk_score') or 0)}/100 seviyesine yükselmiştir."
        )
    if 'TEYİT GÜÇLENDİ' in kind:
        return (
            f"Önceki taramada doğrulama düzeyi “{prev.get('verification','')}” iken, "
            f"yeni taramada “{cur.get('verification','')}” seviyesine yükselmiştir."
        )
    if 'YENİ OLAY' in kind:
        return 'Önceki karşılaştırma taramasında aynı olaya ilişkin yeterli eşleşme bulunmamaktadır; gelişme yeni olay olarak değerlendirilmiştir.'

    prev_nums=_v109_numbers(prev_text)
    cur_nums=_v109_numbers(cur_text)
    new_nums=[x for x in cur_nums if x not in prev_nums]
    prev_tok=set(_history_tokens(prev_text))

    candidates=[]
    for s in _v109_sentences(cur.get('summary') or cur_text):
        toks=set(_history_tokens(s))
        fresh=[t for t in toks-prev_tok if len(t)>=5 and t not in _V104_GENERIC_EVENT_WORDS]
        nums=_v109_numbers(s)-prev_nums
        score=len(fresh)*2+len(nums)*5
        if score>0:
            candidates.append((score,s))
    if candidates:
        best=max(candidates,key=lambda x:x[0])[1]
        best=_v66_formalize_sentence_endings(best).strip()
        if best and best[-1] not in '.!?': best+='.'
        return (
            'Önceki taramada yer almayan yeni sayısal/olgusal bilgi tespit edilmiştir: '+best
            if new_nums else
            'Önceki taramaya göre yeni ayrıntı eklenmiştir: '+best
        )

    if new_nums:
        return 'Önceki taramada bulunmayan yeni sayısal bilgiler açıklanmıştır: '+', '.join(sorted(new_nums)[:5])+'.'
    return 'Olayın içeriğinde önceki taramaya göre anlamlı yeni ayrıntılar tespit edilmiştir.'

def _compare_since_previous(df,current_scan_id=None):
    """V109 — yalnız gerçek değişiklikler ve doğrudan 'Ne Değişti?' açıklaması."""
    current=_v104_event_representatives(df)
    prev_id=_previous_scan_id(current_scan_id)
    previous=_load_scan_events(prev_id)
    if current is None or current.empty:
        return pd.DataFrame(),None,None
    if previous.empty:
        return pd.DataFrame(),prev_id,None

    prev_records=[p.to_dict() for _,p in previous.iterrows()]
    changes=[]
    for _,r in current.iterrows():
        c={
            'title':str(r.get('Başlık','') or ''),'source':str(r.get('Kaynak','') or ''),
            'url':str(r.get('URL','') or ''),'category':str(r.get('Kategori','') or ''),
            'summary':str(r.get('İçerik_Özeti','') or ''),'risk_score':int(r.get('Risk_Skoru',0) or 0),
            'risk_status':str(r.get('Risk_Durumu','') or ''),'verification':str(r.get('Doğrulama','') or ''),
            'source_count':int(r.get('Olay_Kaynak_Sayisi',1) or 1)
        }

        best=None; best_sim=0.0
        for pr in prev_records:
            sim=_v104_event_similarity(c['title'],c['summary'],pr.get('title',''),pr.get('summary',''))
            if c['url'] and c['url']==str(pr.get('url','') or ''):
                sim=max(sim,0.98)
            if sim>best_sim:
                best_sim=sim; best=pr

        if best is None or best_sim<0.50:
            kind='🆕 YENİ OLAY'; priority=100+c['risk_score']; prev_risk='—'
            diff=_v109_direct_difference({},c,kind)
        else:
            risk_up,verify_up,material,_,_=_v104_material_change(best,c)
            if risk_up:
                kind='⚠️ RİSK ARTTI'; priority=95+c['risk_score']
            elif verify_up:
                kind='✅ TEYİT GÜÇLENDİ'; priority=90+c['risk_score']
            elif material:
                kind='🔄 YENİ BİLGİ'; priority=80+c['risk_score']
            else:
                continue
            prev_risk=int(best.get('risk_score') or 0)
            diff=_v109_direct_difference(best,c,kind)

        changes.append({
            'Değişim':kind,'Ne Değişti?':diff,'Başlık':c['title'],'Kaynak':c['source'],
            'Kategori':c['category'],'Risk':c['risk_score'],'Önceki Risk':prev_risk,
            'Kaynak Sayısı':c['source_count'],'URL':c['url'],'_priority':priority
        })

    out=pd.DataFrame(changes)
    if not out.empty:
        out['_sig']=out.apply(
            lambda r:' '.join(sorted(_v104_event_tokens(r.get('Başlık',''),r.get('Ne Değişti?','')))),
            axis=1
        )
        out=out.sort_values(['_priority','Risk'],ascending=[False,False]).drop_duplicates('_sig',keep='first')
        out=out.drop(columns=['_priority','_sig'],errors='ignore')
    prev_time=str(previous.iloc[0].get('scanned_at','')) if not previous.empty else None
    return out,prev_id,prev_time

def _v109_chronology_events(df):
    if df is None or df.empty:
        return pd.DataFrame()
    x=df.copy()
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')
    rows=[]
    groups=x.groupby('Olay_ID',dropna=False) if 'Olay_ID' in x.columns else [(f'ROW-{i}',x.iloc[[i]]) for i in range(len(x))]
    for oid,g in groups:
        g=g.sort_values('Tarih_dt',ascending=False,na_position='last')
        recs=g.to_dict('records')
        rep=max(recs,key=_v107_source_quality) if recs else g.iloc[0].to_dict()
        latest=g.iloc[0]
        rep=dict(rep)
        rep['Tarih_dt']=latest.get('Tarih_dt')
        rep['Tarih']=latest.get('Tarih')
        domains={str(v) for v in g.get('Domain',pd.Series(dtype=str)).tolist() if str(v).strip()}
        sources=list(dict.fromkeys(str(v) for v in g.get('Kaynak',pd.Series(dtype=str)).tolist() if str(v).strip()))
        rep['Kaynak Sayısı']=max(len(domains),len(sources),int(rep.get('Olay_Kaynak_Sayisi',0) or 0))
        rep['Haber Sayısı']=len(g)
        rep['Kaynaklar']=' • '.join(sources[:8])
        rep['_Olay_ID']=str(oid)
        rows.append(rep)
    return pd.DataFrame(rows).sort_values('Tarih_dt',ascending=False,na_position='last').reset_index(drop=True)

def _v109_event_sources(df,event_id):
    if df is None or df.empty or not event_id or 'Olay_ID' not in df.columns:
        return pd.DataFrame()
    x=df[df['Olay_ID'].astype(str)==str(event_id)].copy()
    if x.empty: return x
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')
    return x.sort_values('Tarih_dt',ascending=False,na_position='last')

def _v109_official_source_type(r):
    text=norm(f"{r.get('Kaynak','')} {r.get('Domain','')} {r.get('Başlık','')} {r.get('URL','')}")
    if 'tuik' in text or 'tüik' in text or 'turkiye istatistik' in text: return 'TÜİK'
    if 'tubitak' in text or 'tübitak' in text: return 'TÜBİTAK'
    if 'kosgeb' in text: return 'KOSGEB'
    if 'turkpatent' in text or 'türkpatent' in text: return 'TÜRKPATENT'
    if re.search(r'\btse\b',text) or 'türk standartları' in text: return 'TSE'
    if 'ssb.gov' in text or 'savunma sanayii başkan' in text or 'savunma sanayii baskan' in text: return 'SSB'
    if 'sanayi.gov' in text or 'sanayi ve teknoloji bakan' in text: return 'Bakanlık'
    if 'resmigazete' in text or 'resmî gazete' in text or 'resmi gazete' in text: return 'Resmî Gazete'
    return 'Diğer Resmî'

def _official_radar_rows(df):
    if df is None or df.empty:
        return pd.DataFrame()
    x=df[df.apply(_is_official_radar_row,axis=1)].copy()
    if x.empty: return x
    x['Kurum Türü']=x.apply(_v109_official_source_type,axis=1)
    x=x.sort_values('Tarih_dt',ascending=False,na_position='last')
    return x.drop_duplicates(subset=['URL','Başlık'])

# ============================================================
# /V109
# ============================================================

# V109 — doğrudan fark, olay bazlı kronoloji ve kurum türü filtreli Resmî Kaynak Radarı.
# V108 kaynak zenginleştirme ve V106 kararlı çekirdek korunmuştur.

st.set_page_config(page_title='Sanayi & Teknoloji OSINT Radarı', page_icon='🛡️', layout='wide')

# ============================================================
# V55 — ŞİFRE KORUMASI
# V54 STABLE işlevlerine dokunmaz; yalnızca uygulama girişini korur.
# Streamlit Secrets:
# APP_PASSWORD = "guclu-sifreniz"
# ============================================================
def _v55_password_gate():
    try:
        expected = str(st.secrets["APP_PASSWORD"])
    except Exception:
        st.error(
            "🔐 Uygulama şifresi tanımlanmamış. "
            "Streamlit App Settings → Secrets bölümüne APP_PASSWORD ekleyin."
        )
        st.stop()

    if st.session_state.get("_v55_authenticated", False):
        return

    st.title("🔐 Sanayi ve Teknoloji OSINT Radar")
    st.caption("Devam etmek için uygulama şifresini girin.")

    with st.form("_v55_login_form", clear_on_submit=False):
        entered = st.text_input("Şifre", type="password")
        submitted = st.form_submit_button("Giriş Yap", use_container_width=True)

    if submitted:
        import hmac
        if hmac.compare_digest(str(entered), expected):
            st.session_state["_v55_authenticated"] = True
            st.rerun()
        else:
            st.error("Şifre hatalı.")

    st.stop()

_v55_password_gate()



HEADERS={'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/151 Safari/537.36'}

# -----------------------------
# KONU EVRENİ
# -----------------------------
TOPIC_TERMS = [
    # Sanayi / üretim
    'sanayi','sanayi üretimi','imalat','üretim','fabrika','tesis','organize sanayi','OSB','endüstri',
    'makine','makine sanayii','endüstriyel otomasyon','otomasyon','robotik','endüstri 4.0','mesleki üretim',
    'kapasite','kapasite kullanım','yatırım','yatırım teşvik','yerli üretim','yerlileştirme','millileştirme',
    'tedarik zinciri','tedarikçi','lojistik','depo','depolama','tersane','gemi inşa','denizcilik',
    # Teknoloji / dijital
    'teknoloji','teknolojik','Ar-Ge','Arge','araştırma geliştirme','inovasyon','patent','faydalı model',
    'dijital dönüşüm','endüstri 4.0','yapay zeka','yapay zekâ','makine öğrenmesi','derin öğrenme','yazılım',
    'siber güvenlik','siber saldırı','veri sızıntısı','veri merkezi','bulut','cloud','saas','yazılım şirketi',
    'çip','mikroçip','yarı iletken','semiconductor','işlemci','wafer','elektronik','pcb','sensör',
    'telekom','5G','6G','fiber','internet altyapısı','kuantum','blokzincir','blockchain','fintech',
    # İleri teknoloji / sağlık teknolojileri
    'biyoteknoloji','biyomedikal','nanoteknoloji','medikal cihaz','sağlık teknolojisi','gen tedavisi',
    'malzeme','ileri malzeme','kompozit','karbon fiber','3D yazıcı','eklemeli imalat','batarya teknolojisi',
    # Savunma / havacılık / uzay
    'savunma sanayii','savunma sanayi','savunma teknolojisi','ASELSAN','TUSAŞ','TUSAS','ROKETSAN','HAVELSAN',
    'Baykar','Bayraktar','İHA','SİHA','drone','insansız hava aracı','insansız deniz aracı','KAAN','Kızılelma',
    'HİSAR','SİPER','füze','roket','radar','elektronik harp','elektronik destek','komuta kontrol',
    'mühimmat','zırhlı araç','tank','denizaltı','fırkateyn','korvet','helikopter','havacılık','uçak',
    'havacılık sanayii','uzay','uydu','uydu teknolojisi','roket fırlatma','fırlatma sistemi','Türkiye Uzay Ajansı',
    # Otomotiv / mobilite
    'otomotiv','TOGG','elektrikli araç','hibrit araç','otonom araç','sürücüsüz araç','batarya','şarj',
    'şarj istasyonu','mobilite','raylı sistem','lokomotif','metro','demiryolu','lastik','yan sanayi',
    # Enerji / kimya / kaynak
    'enerji','enerji depolama','güneş enerjisi','solar','rüzgar enerjisi','hidrojen','yakıt hücresi',
    'nükleer enerji','nükleer santral','petrol','doğalgaz','LNG','elektrik üretimi','şebeke','kimya',
    'petrokimya','plastik','polimer','demir çelik','çelik','metal','alüminyum','bakır','madencilik','maden',
    # Diğer üretim sektörleri
    'tekstil','hazır giyim','gıda teknolojisi','gıda sanayii','tarım teknolojisi','akıllı tarım','seracılık',
    'su ürünleri','inşaat teknolojisi','çimento','cam','seramik','kağıt','ambalaj','mobilya',
    # Ekosistem / kamu / girişim
    'TÜBİTAK','KOSGEB','Sanayi ve Teknoloji Bakanlığı','TSE','TürkPatent','TEKNOFEST','teknopark',
    'girişim','girişimcilik','startup','start-up','venture capital','yatırım turu','teknoloji transferi',
    'teknoloji geliştirme bölgesi','Ar-Ge merkezi','tasarım merkezi','OSBÜK','ihracat','ithalat','yüksek teknoloji',
    'orta yüksek teknoloji','kritik teknoloji','stratejik ürün','stratejik yatırım'
]

NEGATIVE_TERMS = [
    'iflas','konkordato','zarar açıkladı','net zarar','üretim durdu','üretim durduruldu',
    'fabrika kapandı','fabrika kapanıyor','işten çıkarma','işçi çıkarma','toplu işten çıkarma',
    'grev','lokavt','soruşturma','dava açıldı','dava','ceza','para cezası','geri çağırma',
    'recall','arıza','kaza','patlama','yangın','siber saldırı','veri sızıntısı','hacklendi',
    'fidye yazılımı','ambargo','yaptırım','ihracat yasağı','ithalat yasağı','lisans reddi',
    'ruhsat iptali','sözleşme feshi','ihale iptal','ihale iptal edildi','askıya alındı',
    'ertelendi','gecikme','teslim edilemedi','testi geçemedi','kapasite kaybı','daralma',
    'sert düşüş','pazar kaybı','maliyet artışı','tedarik sorunu','tedarik krizi','çip krizi',
    'kıtlık','blokaj','güvenlik açığı','kritik zafiyet','zafiyet','usulsüzlük','yolsuzluk',
    'vurgun','casusluk','can kaybı','ölüm','yaralanma','ifşa edildi'
]
HIGH_RISK_TERMS = [
    'iflas','konkordato','üretim durdu','fabrika kapandı','toplu işten çıkarma','siber saldırı',
    'veri sızıntısı','fidye yazılımı','ambargo','yaptırım','ihracat yasağı','lisans reddi',
    'ruhsat iptali','sözleşme feshi','ihale iptal edildi','patlama','yangın','can kaybı','ölüm',
    'kritik zafiyet','yolsuzluk','usulsüzlük','casusluk','tedarik krizi','çip krizi'
]

CATEGORIES={
 'Savunma & Havacılık':['savunma','aselsan','tusaş','tusas','roketsan','havelsan','baykar','bayraktar','iha','siha','kaan','kızılelma','füze','roket','havacılık'],
 'Dijital & Yapay Zeka':['yapay zeka','yapay zekâ','siber','yazılım','5g','6g','veri merkezi','bulut','kuantum'],
 'Yarı İletken & Elektronik':['çip','mikroçip','yarı iletken','işlemci','elektronik','wafer','pcb'],
 'Otomotiv & Mobilite':['otomotiv','togg','elektrikli araç','batarya','şarj'],
 'Enerji':['enerji','hidrojen','güneş','rüzgar','nükleer','enerji depolama'],
 'Sanayi & Üretim':['sanayi','imalat','üretim','fabrika','osb','makine','robotik','otomasyon','demir çelik','kimya'],
 'Uzay & İleri Teknoloji':['uzay','uydu','tua','nanoteknoloji','biyoteknoloji'],
 'Kurumsal Ekosistem':['tübitak','kosgeb','sanayi ve teknoloji bakanlığı','türkpatent','teknopark','teknofest']
}

TR_MAIN=[
 'aa.com.tr','trthaber.com','ntv.com.tr','cnnturk.com','haberturk.com','hurriyet.com.tr','milliyet.com.tr',
 'sabah.com.tr','sozcu.com.tr','cumhuriyet.com.tr','karar.com','yenisafak.com','star.com.tr','aksam.com.tr',
 'turkiyegazetesi.com.tr','t24.com.tr','haber7.com','haberler.com','ensonhaber.com','gazeteduvar.com.tr',
 'odatv.com','medyascope.tv','tv100.com','tgrthaber.com.tr','mynet.com','dunya.com','ekonomim.com',
 'bloomberght.com','paraanaliz.com','bigpara.com','fortuneturkey.com','doviz.com','haberler.com'
]
TR_TECH=[
 'webrazzi.com','shiftdelete.net','donanimhaber.com','chip.com.tr','log.com.tr','technopat.net',
 'hardwareplus.com.tr','turk-internet.com','savunmasanayist.com','savunmatr.com','defenceturk.net',
 'defencehere.com','c4defence.com','savunmahaber.com','gdh.digital','stratejikortak.com','m5dergi.com','mavivatan.net'
]
TR_OFFICIAL=[
 'sanayi.gov.tr','tubitak.gov.tr','kosgeb.gov.tr','tse.org.tr','turkpatent.gov.tr','tua.gov.tr','ticaret.gov.tr',
 'uab.gov.tr','aselsan.com','tusas.com','roketsan.com.tr','havelsan.com.tr','baykartech.com','togg.com.tr','tei.com.tr','tai.com.tr'
]
GR=[
 'kathimerini.gr','protothema.gr','news247.gr','tovima.gr','enikos.gr','naftemporiki.gr','skai.gr','capital.gr',
 'defence-point.gr','defencereview.gr','militaire.gr','pronews.gr','newsbreak.gr','pentapostagma.gr','hellasjournal.com'
]
SOCIAL=['x.com','twitter.com','youtube.com','linkedin.com','facebook.com','instagram.com']

SOURCE_ALIASES={
 'aa':'aa.com.tr','anadolu ajansı':'aa.com.tr','anadolu agency':'aa.com.tr','trt haber':'trthaber.com','trt':'trthaber.com',
 'ntv':'ntv.com.tr','cnn türk':'cnnturk.com','cnn turk':'cnnturk.com','habertürk':'haberturk.com','hürriyet':'hurriyet.com.tr',
 'milliyet':'milliyet.com.tr','sabah':'sabah.com.tr','sözcü':'sozcu.com.tr','cumhuriyet':'cumhuriyet.com.tr','karar':'karar.com',
 'yeni şafak':'yenisafak.com','türkiye gazetesi':'turkiyegazetesi.com.tr','t24':'t24.com.tr','haberler':'haberler.com',
 'dünya':'dunya.com','ekonomim':'ekonomim.com','bloomberg ht':'bloomberght.com','webrazzi':'webrazzi.com',
 'shiftdelete':'shiftdelete.net','donanımhaber':'donanimhaber.com','technopat':'technopat.net','savunma sanayi st':'savunmasanayist.com',
 'savunma sanayi':'savunmasanayist.com','defence türk':'defenceturk.net','defence turk':'defenceturk.net','defencehere':'defencehere.com',
 'c4 defence':'c4defence.com','c4defence':'c4defence.com','sanayi ve teknoloji bakanlığı':'sanayi.gov.tr','tübitak':'tubitak.gov.tr',
 'kosgeb':'kosgeb.gov.tr','türkpatent':'turkpatent.gov.tr','türkiye uzay ajansı':'tua.gov.tr','aselsan':'aselsan.com',
 'tusaş':'tusas.com','tusas':'tusas.com','roketsan':'roketsan.com.tr','havelsan':'havelsan.com.tr','baykar':'baykartech.com','togg':'togg.com.tr'
}

def norm(s):
    return re.sub(r'\s+',' ',str(s or '').lower()).strip()

def title_key(s):
    return re.sub(r'[^\w\s]',' ',norm(s)).strip()[:180]

def domain(url):
    try: return urlparse(url).netloc.lower().replace('www.','')
    except: return ''

def parse_dt(v):
    if not v: return None
    s=str(v).strip()
    for x in (s.replace('Z','+00:00'),s):
        try:
            d=datetime.fromisoformat(x)
            if d.tzinfo is None: d=d.replace(tzinfo=timezone.utc)
            return d.astimezone(timezone.utc)
        except: pass
    try:
        d=parsedate_to_datetime(s)
        if d.tzinfo is None: d=d.replace(tzinfo=timezone.utc)
        return d.astimezone(timezone.utc)
    except: return None

def _to_utc_datetime(value):
    """datetime / pandas.Timestamp / string değerlerini güvenli biçimde UTC-aware datetime'a çevirir."""
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass

    try:
        if isinstance(value, pd.Timestamp):
            ts = value
            if ts.tzinfo is None:
                ts = ts.tz_localize('UTC')
            else:
                ts = ts.tz_convert('UTC')
            return ts.to_pydatetime()
    except Exception:
        pass

    if isinstance(value, datetime):
        if value.tzinfo is None:
            return value.replace(tzinfo=timezone.utc)
        return value.astimezone(timezone.utc)

    try:
        ts = pd.to_datetime(value, utc=True, errors='coerce')
        if pd.isna(ts):
            return None
        return ts.to_pydatetime()
    except Exception:
        return None


def fmt_dt(d):
    d = _to_utc_datetime(d)
    return d.astimezone().strftime('%d.%m.%Y %H:%M:%S') if d else 'Tarih/saat bilinmiyor'

def infer_source(source_name='',source_url='',article_url=''):
    d=domain(source_url)
    if d and d not in ('news.google.com','google.com'): return d
    n=norm(source_name)
    for a,d in SOURCE_ALIASES.items():
        if a in n: return d
    # domain adının yayıncı adına gömülü olması
    for d in TR_MAIN+TR_TECH+TR_OFFICIAL+GR:
        stem=d.split('.')[0]
        if stem and stem in re.sub(r'[^a-z0-9ğüşöçıİĞÜŞÖÇ]','',n): return d
    return domain(article_url)

def source_group(d):
    d=domain(d)
    if d in TR_OFFICIAL: return '🇹🇷 Resmi / Kurumsal'
    if d in TR_TECH: return '🇹🇷 Türk Teknoloji / Savunma'
    if d in TR_MAIN: return '🇹🇷 Türk Medyası / Ekonomi'
    if d in GR: return '🇬🇷 Yunan Medyası — Türk Savunma'
    if d in SOCIAL: return '📱 Açık Sosyal / İndeks'
    return '🌍 Diğer / Açık Kaynak'

def source_rank(d):
    d=domain(d)
    if d in TR_OFFICIAL: return 500
    if d in TR_TECH: return 450
    if d in TR_MAIN: return 400
    if d in GR: return 300
    if d in SOCIAL: return 250
    return 100

def relevant(text,user_query=''):
    t=norm(text)
    if any(x in t for x in TOPIC_TERMS): return True
    uq=re.split(r'\bOR\b|,|\n',user_query or '',flags=re.I)
    generic={'sanayi','teknoloji','üretim','yatırım','enerji','türkiye','türk','haber'}
    return any(len(x.strip())>2 and norm(x.strip()) not in generic and norm(x.strip()) in t for x in uq)

def greek_defense(text):
    t=norm(text)
    terms=['turkey','türkiye','turkish','türk','τουρκ','aselsan','tusaş','tusas','roketsan','havelsan','baykar','bayraktar','kaan','kızılelma','siper','hisar','iha','siha','drone','uav','missile','fighter','frigate','submarine','defense','defence','savunma','άμυνα']
    return any(x in t for x in terms)


OSB_FIRE_LOCATION_TERMS = [
    'osb','organize sanayi','organize sanayi bölgesi','organize sanayi bölgesinde',
    'organize sanayi bölgesindeki','organize sanayi sitesinde'
]
INDUSTRIAL_LOCATION_TERMS = [
    'osb','organize sanayi','organize sanayi bölgesi','organize sanayi sitesinde',
    'fabrika','fabrikada','fabrikasında','tesis','tesiste','üretim tesisi','sanayi tesisi',
    'imalathane','atölye','depo','üretim alanı','sanayi sitesi'
]
CRITICAL_INCIDENT_TERMS = [
    'yangın','yangını','yangin','alev','alevler','yanıyor','yaniyor','yandı','yandi',
    'patlama','patladı','patladi','infilak','infilak etti','parlama',
    'fabrika yangını','tesis yangını','fabrika patlaması','tesis patlaması'
]

def is_osb_fire(title, snippet=''):
    """Geriye dönük uyumluluk: OSB + yangın bağlamı."""
    t=norm(f'{title} {snippet}')
    return (
        any(term in t for term in OSB_FIRE_LOCATION_TERMS)
        and any(term in t for term in ['yangın','yangını','yangin','alev','alevler','yanıyor','yaniyor','yandı','yandi'])
    )

def critical_industrial_incident(title, snippet=''):
    """
    Özel kırmızı alarm için:
    - OSB içi yangın
    - OSB içi patlama
    - OSB dışı fabrika/tesis yangını
    - OSB dışı fabrika/tesis patlaması
    """
    t=norm(f'{title} {snippet}')
    has_location=any(term in t for term in INDUSTRIAL_LOCATION_TERMS)
    has_incident=any(_v89_has_term(t,term) for term in CRITICAL_INCIDENT_TERMS)
    if not (has_location and has_incident):
        return None

    osb=any(term in t for term in OSB_FIRE_LOCATION_TERMS)
    fire=any(term in t for term in ['yangın','yangını','yangin','alev','alevler','yanıyor','yaniyor','yandı','yandi'])
    explosion=any(term in t for term in ['patlama','patladı','patladi','infilak','infilak etti','parlama'])

    if osb and explosion:
        return '💥 OSB PATLAMA'
    if osb and fire:
        return '🔥 OSB YANGINI'
    if explosion:
        return '💥 FABRİKA/TESİS PATLAMASI'
    if fire:
        return '🔥 FABRİKA/TESİS YANGINI'
    return '🚨 KRİTİK SANAYİ OLAYI'


NEGATION_OR_RESOLUTION_PHRASES = [
    'olmadı','olmadığı','bulunmadı','bulunmadığı','yaşanmadı','gerçekleşmedi',
    'etkilenmedi','etkilenmediği','risk bulunmuyor','risk yok','tehdit yok',
    'iptal edilmedi','kapanmadı','durmadı','sona erdi','kaldırıldı',
    'giderildi','çözüldü','önlendi','engellendi','bertaraf edildi'
]

POSITIVE_SIGNAL_TERMS = [
    'arttı','artış','yükseldi','yükseliş','rekor','büyüdü','büyüme',
    'yatırım','yatırım kararı','yeni yatırım','ihracat arttı','ihracat artışı',
    'kapasite arttı','kapasite artışı','üretim arttı','üretim artışı',
    'devreye alındı','faaliyete geçti','başarıyla','başarılı',
    'anlaşma imzalandı','sözleşme imzalandı','teslim edildi',
    'teşvik','destek','hibe','istihdam artışı','yeni istihdam'
]

SEVERE_NEGATIVE_TERMS = {
    'iflas','konkordato','üretim durdu','fabrika kapandı','toplu işten çıkarma',
    'siber saldırı','veri sızıntısı','fidye yazılımı','ambargo','yaptırım',
    'ihracat yasağı','lisans reddi','ruhsat iptali','sözleşme feshi',
    'ihale iptal edildi','patlama','yangın','can kaybı','ölüm',
    'kritik zafiyet','yolsuzluk','usulsüzlük','casusluk','tedarik krizi','çip krizi'
}

def _term_regex(term):
    # Alt-string kaynaklı "ceza/cezasız", "dava/davalar" vb. yanlış eşleşmeleri azalt.
    escaped=re.escape(term)
    if ' ' in term:
        return re.compile(escaped,re.I)
    return re.compile(r'(?<!\w)'+escaped+r'(?!\w)',re.I)

def _physical_incident_is_real(term, context):
    if term not in {'yangın','patlama','ölüm'}:
        return True
    incident_markers=[
        'çıktı','çıkan','meydana geldi','meydana gelen','patladı','infilak',
        'alev','yaralandı','yaralı','hasar','müdahale','söndürüldü',
        'kontrol altına','tahliye','hayatını kaybetti','öldü'
    ]
    return any(x in context for x in incident_markers)

def _active_adverse_terms(terms, text):
    t=norm(text)
    active=[]
    for term in terms:
        rx=_term_regex(term)
        matches=list(rx.finditer(t))
        if not matches:
            continue

        term_active=False
        for m in matches:
            lo=max(0,m.start()-90); hi=min(len(t),m.end()+90)
            ctx=t[lo:hi]

            # Fiziksel olay kelimesi yalnız kavramsal/önleyici bir kullanımdaysa alarm verme.
            if not _physical_incident_is_real(term,ctx):
                continue

            # "yaptırım kaldırıldı", "ihlal yaşanmadı", "üretim durmadı" gibi bağlamları bastır.
            # Gerçekleşmiş yangın/patlama/can kaybı ise "kontrol altına alındı" gibi sonraki olumlu
            # gelişmeler olayın negatif niteliğini ortadan kaldırmaz.
            if term not in {'yangın','patlama','can kaybı','ölüm'}:
                if any(p in ctx for p in NEGATION_OR_RESOLUTION_PHRASES):
                    continue

            term_active=True
            break

        if term_active:
            active.append(term)
    return active

def _sentence_chunks(text):
    txt=re.sub(r'\s+',' ',str(text or '')).strip()
    if not txt:
        return []
    return [x.strip() for x in re.split(r'(?<=[.!?;:])\s+',txt) if x.strip()]

def _negated_in_context(term, sentence):
    s=norm(sentence)
    # Olumsuzluk/çözülme ifadeleri, ilgili risk kelimesinin yakın çevresindeyse baskılanır.
    negators=[
        'değil','değildir','olmadı','olmadığı','bulunmadı','bulunmadığı',
        'yaşanmadı','gerçekleşmedi','etkilenmedi','etkilenmediği',
        'risk yok','tehdit yok','iptal edilmedi','kapanmadı','durmadı',
        'giderildi','çözüldü','önlendi','engellendi','kaldırıldı','sona erdi'
    ]
    return any(n in s for n in negators)

def _positive_strength(text):
    t=norm(text)
    positive_terms=[
        'arttı','artış','yükseldi','yükseliş','rekor','büyüdü','büyüme',
        'yeni yatırım','yatırım kararı','yatırım yaptı','yatırım yapacak',
        'ihracat arttı','ihracat artışı','kapasite artışı','kapasite arttı',
        'üretim artışı','üretim arttı','devreye alındı','faaliyete geçti',
        'başarıyla','başarılı','anlaşma imzalandı','sözleşme imzalandı',
        'teslim edildi','teşvik','destek','hibe','istihdam artışı','yeni istihdam',
        'pazar payı arttı','gelir arttı','kâr arttı','kar arttı'
    ]
    return sum(1 for x in positive_terms if x in t)

# V48 — Ekonomik/operasyonel haber dilinde sık görülen, önceki sözlükte kolay
# kaçabilen negatif sinyaller. Bunlar tam sayfa indirmeden RSS içerik/özetinde aranır.
V48_NEGATIVE_PHRASES = [
    'düştü','düşüş','azaldı','azalış','geriledi','gerileme','daraldı','daralma',
    'zarar açıkladı','zarar etti','net zarar','faaliyet zararı','kayıp yaşadı',
    'satışlar düştü','satışlar azaldı','satışlarda düşüş','satışlarda azalma',
    'üretim düştü','üretim azaldı','üretimde düşüş','üretimde azalma',
    'ihracat düştü','ihracat azaldı','ihracatta düşüş','ihracatta gerileme',
    'siparişler düştü','siparişler azaldı','siparişlerde düşüş',
    'kapasite düştü','kapasite azaldı','kapasite kullanım oranı düştü',
    'istihdam azaldı','istihdam düştü','istihdam kaybı',
    'işten çıkarma','işçi çıkarma','personel azaltma','toplu işten çıkarma',
    'maliyet arttı','maliyet artışı','maliyet baskısı','girdi maliyetleri arttı',
    'fiyat baskısı','finansman maliyeti','nakit sıkıntısı','likidite sıkıntısı',
    'talep düştü','talep azaldı','talep daralması','talepte daralma',
    'pazar payı düştü','pazar payı kaybı','rekabet gücü kaybı',
    'beklentinin altında','beklentilerin altında','hedefin altında','hedefin gerisinde',
    'kriz','aksama','kesinti','arıza','gecikme','ertelendi','iptal edildi',
    'faaliyet durdu','üretim durdu','üretime ara verdi','üretime ara verildi',
    'fabrika kapandı','tesis kapandı','kapanma kararı',
    'iflas','konkordato','haciz','borç krizi',
    'soruşturma','inceleme başlatıldı','ceza verildi','para cezası',
    'yasaklandı','yasak','geri çağırma','ürün geri çağırma',
    'siber saldırı','veri sızıntısı','veri ihlali','kritik açık','güvenlik açığı',
    'tedarik sorunu','tedarik krizi','tedarik zinciri aksaması',
    'kaza','yangın','patlama','yaralandı','can kaybı'
]


# V49 — Yapısal / eleştirel negatiflik katmanı
# Haber sayfasına gitmez; yalnızca eldeki Başlık + RSS içerik/özet üzerinde çalışır.
V49_STRUCTURAL_NEGATIVE = [
    'tehlikeli gidiş','tehlikeli seyir','olumsuz gidiş','olumsuz seyir',
    'kötü gidiş','kötüye gidiş','kötüleşiyor','kötüleşme',
    'alarm veriyor','alarm zilleri','kan kaybediyor','kan kaybı',
    'ivme kaybediyor','ivme kaybı','güç kaybediyor','güç kaybı',
    'rekabet gücü zayıflıyor','rekabet gücü geriliyor','rekabet gücü kaybı',
    'zayıf seyir','zayıflama','zayıflıyor','yavaşlıyor','yavaşlama',
    'sıkıntılı süreç','sıkıntılı dönem','kritik süreç','kritik eşik',
    'sorun büyüyor','sorunlar büyüyor','sorun devam ediyor','sorun sürüyor',
    'risk artıyor','riskler artıyor','baskı artıyor','baskı altında',
    'istenilen seviyede değil','istenen seviyede değil',
    'beklenen seviyede değil','yeterli değil','yetersiz kaldı','yetersiz kalıyor',
    'teşvik yetmiyor','teşvikler yetmiyor','destek yetmiyor','destekler yetmiyor',
    'sadece teşvik vermekle olmuyor','sadece destek vermekle olmuyor',
    'çözüm olmuyor','çözüm değil','sürdürülebilir değil',
    'endişe yaratıyor','endişe veriyor','kaygı yaratıyor','kaygı veriyor',
    'uyarı geldi','uyarı yaptı','uyardı','dikkat çekti',
    'olumsuz tablo','karamsar tablo','zorlu görünüm','zayıf görünüm',
    'darboğaz','çıkmaz','kırılganlık','kırılgan hale geldi'
]


# V50 — Geniş Negatif Bölümü
# Kullanıcı açısından "Negatif" yalnızca gerçekleşmiş kötü olay değildir.
# Eleştirel, uyarıcı, yetersizlik bildiren, politika/sektör performansını sorgulayan
# ve yapısal sorun işaret eden haberler de aynı Negatif bölümüne girer.
V50_CRITICAL_NEGATIVE = [
    'eleştirdi','eleştiri','eleştirildi','tepki gösterdi','tepki çekti',
    'itiraz etti','itiraz','uyarıda bulundu','uyarı yaptı','uyardı',
    'dikkat çekti','dikkat çekiyor','endişesini dile getirdi','kaygısını dile getirdi',
    'yeterli değil','yetersiz','yetersiz kaldı','yetersiz kalıyor',
    'eksik kaldı','eksiklik','yetmiyor','yetmedi','karşılamıyor',
    'çözüm değil','çözüm olmadı','çözüm olmuyor','sonuç vermiyor','sonuç vermedi',
    'etkili değil','etkisiz','başarısız','başarısızlık',
    'hedefin gerisinde','hedeflerin gerisinde','beklentinin altında','beklentilerin altında',
    'istenen seviyede değil','istenilen seviyede değil','arzu edilen seviyede değil',
    'sorunlu','sorunlar','sorun devam ediyor','sorun sürüyor','sorun büyüyor',
    'risk taşıyor','risk oluşturuyor','risk yaratıyor','tehdit oluşturuyor',
    'sürdürülebilir değil','kırılgan','kırılganlık','darboğaz',
    'rekabet sorunu','rekabet gücü kaybı','rekabet gücü zayıflıyor',
    'verimlilik sorunu','finansmana erişim sorunu','nitelikli iş gücü sorunu',
    'maliyet baskısı','finansman baskısı','kur baskısı',
    'sanayici zorlanıyor','sektör zorlanıyor','firmalar zorlanıyor',
    'üretici zorlanıyor','ihracatçı zorlanıyor',
    'teşvikler yetersiz','destekler yetersiz','teşvik yetmiyor','destek yetmiyor',
    'politika yetersiz','politikalar yetersiz','düzenleme yetersiz',
    'önlem yetersiz','tedbir yetersiz','önlemler yetersiz','tedbirler yetersiz'
]

def _v50_critical_negative_signals(text):
    t = norm(text)
    found = set()
    for phrase in V50_CRITICAL_NEGATIVE:
        if phrase in t:
            # Açık biçimde reddedilen eleştirileri yanlış negatif yapma.
            idx = t.find(phrase)
            ctx = t[max(0, idx-65):idx+len(phrase)+65] if idx >= 0 else t
            if any(x in ctx for x in [
                'eleştiri yok','sorun yok','risk yok','yetersiz değil',
                'başarısız değil','kırılgan değil','zorlanmıyor'
            ]):
                continue
            found.add(phrase)
    return found

V49_PERSISTENCE_PATTERNS = [
    r'\b\d+\s*(?:çeyrektir|çeyrek boyunca|aydır|ay boyunca|yıldır|yıl boyunca|haftadır)\b',
    r'\buzun süredir\b',
    r'\bsüregelen\b',
    r'\bdevam eden\b',
    r'\bsürmekte olan\b',
    r'\bkronik\b'
]

def _v49_structural_negative_signals(text):
    t = norm(text)
    found = set()

    for phrase in V49_STRUCTURAL_NEGATIVE:
        if phrase in t:
            found.add(phrase)

    persistent = any(re.search(pat, t, re.I) for pat in V49_PERSISTENCE_PATTERNS)

    # Süre ifadesi tek başına negatif değildir. Ancak yapısal negatif bir ifade
    # veya başka bir negatif sinyal ile birlikteyse ağırlık kazanır.
    return found, persistent

V48_STRONG_NEGATIVE = [
    'üretim durdu','faaliyet durdu','fabrika kapandı','tesis kapandı',
    'toplu işten çıkarma','iflas','konkordato','siber saldırı','veri sızıntısı',
    'veri ihlali','yangın','patlama','can kaybı','hayatını kaybetti',
    'ihracat yasağı','yaptırım','ambargo'
]

V48_DIRECTION_PATTERNS = [
    r'(?:yüzde|%)\s*\d+(?:[.,]\d+)?\s*(?:düştü|azaldı|geriledi|daraldı)',
    r'\d+(?:[.,]\d+)?\s*(?:%|yüzde)\s*(?:düştü|azaldı|geriledi|daraldı)',
    r'(?:üretim|ihracat|satış|sipariş|istihdam|kapasite|talep|gelir|kâr|kar)\w*\s+.{0,55}\b(?:düştü|azaldı|geriledi|daraldı)\b',
    r'\b(?:geçen yıla|önceki yıla|geçen aya|önceki aya)\s+göre.{0,70}\b(?:düştü|azaldı|geriledi|daraldı)\b'
]


def _v89_has_term(text, term):
    """
    Negatif terimleri alt-dize ile değil kelime/ifade sınırıyla arar.
    Böylece 'kaza' -> 'kazandı/kazanç/kazanım' eşleşmesi oluşmaz.
    """
    t=norm(text)
    term=norm(term)
    if not term:
        return False
    # _term_regex mevcut negatif analiz motorunun güvenli eşleştiricisidir.
    try:
        return bool(_term_regex(term).search(t))
    except Exception:
        # Fallback: tek kelimede Unicode kelime sınırı, çok kelimede sınırlandırılmış ifade.
        return bool(re.search(r'(?<!\w)'+re.escape(term)+r'(?!\w)',t,re.I))

def _v48_extra_negative_signals(text):
    t=norm(text)
    found=set()
    for phrase in V48_NEGATIVE_PHRASES:
        if _v89_has_term(t,phrase):
            # "düşmedi / azalmadı / gerilemedi" gibi açık olumsuzlamaları alma.
            m=_term_regex(phrase).search(t)
            pos=m.start() if m else -1
            ctx=t[max(0,pos-60):pos+len(phrase)+60] if pos>=0 else t
            if any(x in ctx for x in [
                'düşmedi','azalmadı','gerilemedi','daralmadı','iptal edilmedi',
                'aksama olmadı','kesinti olmadı','etkilenmedi','risk yok'
            ]):
                continue
            found.add(phrase)

    directional=False
    for pat in V48_DIRECTION_PATTERNS:
        if re.search(pat,t,re.I):
            directional=True
            found.add('sayısal/yönsel düşüş')
            break
    return found,directional

def _negative_sentence_analysis(title, snippet):
    """
    V48 hızlı hassas analiz:
    - Başlık + RSS içerik/özet birlikte
    - mevcut negatif/risk sözlükleri
    - geniş ekonomik/operasyonel sözlük
    - sayısal/yönsel düşüş tespiti
    - olumsuzlama kontrolü
    """
    title_n=norm(title)
    full=f"{title}. {snippet}"
    sentences=_sentence_chunks(full)

    active_neg=set()
    active_risk=set()
    title_neg=set()
    title_risk=set()
    strong_event=False

    physical_terms={'yangın','patlama','can kaybı','ölüm'}
    physical_markers=[
        'çıktı','çıkan','meydana geldi','meydana gelen','patladı','infilak',
        'alev','yaralandı','yaralı','hasar','müdahale','söndürüldü',
        'kontrol altına','tahliye','hayatını kaybetti','öldü'
    ]

    for s in sentences:
        sn=norm(s)
        for term in NEGATIVE_TERMS:
            if not _term_regex(term).search(sn):
                continue
            if term in physical_terms:
                if not any(m in sn for m in physical_markers):
                    continue
            elif _negated_in_context(term,sn):
                continue
            active_neg.add(term)
            if _term_regex(term).search(title_n):
                title_neg.add(term)

        for term in HIGH_RISK_TERMS:
            if not _term_regex(term).search(sn):
                continue
            if term in physical_terms:
                if not any(m in sn for m in physical_markers):
                    continue
            elif _negated_in_context(term,sn):
                continue
            active_risk.add(term)
            if _term_regex(term).search(title_n):
                title_risk.add(term)

    extra,directional=_v48_extra_negative_signals(full)
    active_neg.update(extra)

    for phrase in extra:
        if phrase!='sayısal/yönsel düşüş' and phrase in title_n:
            title_neg.add(phrase)

    # V49: klasik düşüş/zarar kelimesi bulunmasa bile eleştirel ve yapısal
    # kötüleşme dili ayrıca yakalanır.
    structural,persistent=_v49_structural_negative_signals(full)
    active_neg.update(structural)
    for phrase in structural:
        if phrase in title_n:
            title_neg.add(phrase)

    # V50: eleştirel/uyarıcı/yetersizlik bildiren içerikler de doğrudan
    # mevcut Negatif havuzuna eklenir. Ayrı kategori oluşturulmaz.
    critical_negative=_v50_critical_negative_signals(full)
    active_neg.update(critical_negative)
    for phrase in critical_negative:
        if phrase in title_n:
            title_neg.add(phrase)

    if any(x in norm(full) for x in V48_STRONG_NEGATIVE):
        strong_event=True

    return active_neg,active_risk,title_neg,title_risk,strong_event,directional,structural,persistent,critical_negative

def classify(title,snippet,source_domain=''):
    full=f'{title} {snippet}'
    t=norm(full)

    neg_set,risk_set,title_neg,title_risk,strong_event,directional,structural,persistent,critical_negative=_negative_sentence_analysis(title,snippet)
    neg=sorted(neg_set)
    risk=sorted(risk_set)

    cat='Genel Sanayi / Teknoloji'
    for c,ks in CATEGORIES.items():
        if any(k in t for k in ks):
            cat=c
            break

    score=5
    reasons=[]

    if neg:
        score += min(30,6*len(neg))
        score += min(14,5*len(title_neg))
        reasons.append(f'{len(neg)} doğrulanmış negatif sinyal')

    if directional:
        score += 8
        reasons.append('ölçülebilir düşüş/gerileme')

    if structural:
        score += min(16, 7 + 3*len(structural))
        reasons.append('yapısal/eleştirel olumsuzluk')

    if critical_negative:
        # Eleştirel yaklaşım doğrudan Negatif bölümüne girecek kadar ağırlık alır,
        # fakat tek başına Yüksek Risk sayılmaz.
        score += min(15, 8 + 2*len(critical_negative))
        reasons.append('eleştirel/uyarıcı yaklaşım')

    if persistent and (structural or critical_negative or neg_set):
        score += 8
        reasons.append('olumsuzluğun sürekliliği')

    if risk:
        score += min(32,9*len(risk))
        score += min(14,5*len(title_risk))
        reasons.append(f'{len(risk)} yüksek risk sinyali')

    if strong_event:
        score += 14
        reasons.append('doğrudan ağır olumsuz olay')

    # Gerçek negatiflik varsa sektörel etki skoru eklenir.
    if neg or risk:
        if any(x in t for x in ['üretim','fabrika','tesis','istihdam','kapasite','ihracat','tedarik','satış','sipariş']):
            score += 6
            reasons.append('üretim/ekonomi etkisi')
        if any(x in t for x in ['savunma','kritik altyapı','enerji','siber','yarı iletken','çip']):
            score += 7
            reasons.append('stratejik/kritik sektör etkisi')

    positive_count=_positive_strength(full)
    severe_active=strong_event or any(x in norm(full) for x in V48_STRONG_NEGATIVE)

    # V48 farkı: olumlu sinyal gerçek negatifliği SİLMEZ.
    # Yalnızca ağır risk yoksa skoru sınırlı ölçüde dengeler.
    if positive_count and neg and not severe_active:
        score=max(0,score-min(8,2*positive_count))
        reasons.append('karma/olumlu unsurlar mevcut')

    score=max(0,min(100,score))

    # En kritik değişiklik: gerçek ve bağlamsal negatif sinyal bulunduysa,
    # yüksek risk olmasa dahi haber Negatif olabilir.
    sentiment='Negatif' if neg else 'Nötr'

    if severe_active and (risk or score>=55):
        status='Yüksek Risk'
    elif risk and score>=68:
        status='Yüksek Risk'
    elif (structural or critical_negative) and neg:
        status='Negatif'
    elif neg and score>=18:
        status='Negatif'
    else:
        status='Normal'

    if status=='Normal' and not neg:
        reasons=['olumsuz risk sinyali tespit edilmedi']

    # V87 — çok sınırlı yanlış-negatif koruması.
    # Açık başarı/madalya/ödül ve normal test ilerlemesi haberleri,
    # başlıkta gerçek bir olumsuzluk yoksa negatif değildir.
    _hn=norm(title)
    _positive_head=bool(re.search(
        r'(madalya\s+kazan|ödül\s+kazan|şampiyon|rekor\s+kır|başarıyla|'
        r'başarı\s+elde|testleri?\s+devam\s+ediyor|test\s+süreci\s+devam)',
        _hn,re.I
    ))
    _bad_head=bool(re.search(
        r'(başarısız|kaza|yangın|patlama|ölüm|yaralan|iptal|gecik|arıza|'
        r'iflas|saldırı|eleştir|yetersiz|kriz|sorun|tehlike|zarar|kayıp|'
        r'geriledi|azaldı|düştü|ceza|yaptırım)',
        _hn,re.I
    ))
    if _positive_head and not _bad_head:
        sentiment='Nötr'
        status='Normal'
        score=min(score,12)
        neg=[]
        risk=[]
        reasons=['açık başarı veya normal test/program ilerlemesi; negatif değildir']

    return sentiment,score,status,neg,risk,cat,reasons

def _v89_negative_selfcheck():
    """Basit regresyon kontrolleri; panelde gösterilmez."""
    cases=[
        ('Türk öğrenciler uluslararası yarışmada 15 madalya kazandı','Nötr'),
        ('Şirket yılın ilk yarısında güçlü kazanç açıkladı','Nötr'),
        ('Yeni teknoloji kazanımı ihracat kapasitesini artırdı','Nötr'),
    ]
    for h,expected in cases:
        try:
            sent,_,_,_,_,_,_=classify(h,'')
            if sent!=expected:
                return False
        except Exception:
            return False
    return True


def rss(query, timeout=7):
    try:
        r=requests.get('https://news.google.com/rss/search',params={'q':query,'hl':'tr','gl':'TR','ceid':'TR:tr'},headers=HEADERS,timeout=timeout)
        r.raise_for_status(); root=ET.fromstring(r.content); out=[]
        for it in root.findall('.//item'):
            src=it.find('source')
            out.append({
                'title':html.unescape(it.findtext('title') or ''),
                'url':it.findtext('link') or '',
                'date':it.findtext('pubDate') or '',
                'snippet':BeautifulSoup(it.findtext('description') or '','html.parser').get_text(' ',strip=True),
                'source':src.text if src is not None else '',
                'source_url':src.get('url','') if src is not None else ''
            })
        return out
    except Exception:
        return []

def ddgs_text(q):
    try:
        from ddgs import DDGS
    except Exception:
        try: from duckduckgo_search import DDGS
        except Exception: return []
    try:
        with DDGS() as d: return list(d.text(q,region='tr-tr',timelimit='d',max_results=40))
    except Exception: return []

def gdelt(q, timespan):
    try:
        r=requests.get('https://api.gdeltproject.org/api/v2/doc/doc',params={'query':q,'mode':'artlist','maxrecords':250,'format':'json','sort':'HybridRel','timespan':timespan},headers=HEADERS,timeout=8)
        r.raise_for_status(); return r.json().get('articles',[]) or []
    except Exception: return []

def period_window(hours):
    if hours<=3: return '6h'
    if hours<=24: return '1d'
    if hours<=48: return '2d'
    if hours<=168: return '7d'
    return '30d'

def _query_terms(user_query):
    parts=re.split(r'\bOR\b|,|;|\n',user_query or '',flags=re.I)
    out=[]; seen=set()
    for x in parts:
        x=x.strip().strip('"').strip("'")
        if len(x)>=3 and norm(x) not in seen:
            seen.add(norm(x)); out.append(x)
    return out

def build_turkish_queries(when, user_query=''):
    # Geniş arama evreni: tek dev sorgu yerine konu kümeleri paralel taranır.
    # Böylece kapsam genişlerken Google News sorguları aşırı ağırlaşmaz.
    groups=[
        '(sanayi OR imalat OR üretim OR fabrika OR tesis OR OSB OR "organize sanayi" OR endüstri)',
        '(makine OR otomasyon OR robotik OR "endüstri 4.0" OR kapasite OR "kapasite kullanım")',
        '(teknoloji OR inovasyon OR "Ar-Ge" OR Arge OR patent OR "dijital dönüşüm" OR teknopark)',
        '("yapay zeka" OR "yapay zekâ" OR "makine öğrenmesi" OR yazılım OR SaaS OR bulut)',
        '("siber güvenlik" OR "siber saldırı" OR "veri sızıntısı" OR kuantum OR blockchain OR fintech)',
        '(çip OR mikroçip OR "yarı iletken" OR semiconductor OR işlemci OR wafer OR elektronik OR PCB OR sensör)',
        '("savunma sanayii" OR "savunma sanayi" OR ASELSAN OR TUSAŞ OR ROKETSAN OR HAVELSAN OR Baykar OR Bayraktar)',
        '(İHA OR SİHA OR drone OR KAAN OR Kızılelma OR HİSAR OR SİPER OR füze OR roket OR radar OR "elektronik harp")',
        '(havacılık OR "havacılık sanayii" OR uçak OR helikopter OR uzay OR uydu OR "roket fırlatma" OR "Türkiye Uzay Ajansı")',
        '(otomotiv OR TOGG OR "elektrikli araç" OR "hibrit araç" OR "otonom araç" OR batarya OR şarj OR mobilite)',
        '(enerji OR "enerji depolama" OR "güneş enerjisi" OR "rüzgar enerjisi" OR hidrojen OR "yakıt hücresi" OR "nükleer enerji")',
        '(kimya OR petrokimya OR plastik OR polimer OR "demir çelik" OR çelik OR metal OR alüminyum OR bakır)',
        '(madencilik OR maden OR tekstil OR "gıda teknolojisi" OR "gıda sanayii" OR "tarım teknolojisi" OR seracılık)',
        '(lojistik OR "tedarik zinciri" OR tersane OR "gemi inşa" OR denizcilik OR demiryolu OR "raylı sistem")',
        '(biyoteknoloji OR biyomedikal OR nanoteknoloji OR "medikal cihaz" OR "sağlık teknolojisi" OR "ileri malzeme" OR kompozit)',
        '(TÜBİTAK OR KOSGEB OR "Sanayi ve Teknoloji Bakanlığı" OR TürkPatent OR TEKNOFEST OR "yatırım teşvik" OR "teknoloji transferi")',
        '(startup OR "start-up" OR girişimcilik OR "yatırım turu" OR "venture capital" OR "Ar-Ge merkezi" OR "tasarım merkezi")',
        '(ihracat OR ithalat OR "yüksek teknoloji" OR "orta yüksek teknoloji" OR "kritik teknoloji" OR "stratejik ürün" OR yerlileştirme)'
    ]
    qs=[f'Türkiye {g} when:{when}' for g in groups]
    # Kullanıcının kutuya eklediği ÖZEL terimler ayrıca taranır.
    # Performans: varsayılan geniş evrende zaten bulunan terimleri ikinci kez sorgulamayız.
    # Böylece normal kullanımda 38 civarı sorgu yerine yaklaşık 18 ana sorgu çalışır;
    # kullanıcı gerçekten yeni bir terim eklerse yalnızca o terim(ler) ek sorgu olur.
    built_in={norm(x) for x in TOPIC_TERMS}
    generic={'sanayi','teknoloji','üretim','imalat','fabrika','türkiye','türk'}
    custom=[
        x for x in _query_terms(user_query)
        if norm(x) not in generic and norm(x) not in built_in
    ]
    for term in custom[:8]:
        qs.append(f'Türkiye ("{term}") when:{when}')
    return qs


# -----------------------------
# V41 — RESMÎ KAYNAK / İSTATİSTİK RADARI
# -----------------------------
OFFICIAL_RADAR_DOMAINS = [
    'sanayi.gov.tr','tubitak.gov.tr','kosgeb.gov.tr','turkpatent.gov.tr','tse.org.tr',
    'ssb.gov.tr','tuik.gov.tr','tcmb.gov.tr','ticaret.gov.tr','epdk.gov.tr','teias.gov.tr',
    'tua.gov.tr'
]

PRIMARY_STATS_DOMAINS = [
    'tuik.gov.tr','tcmb.gov.tr','ticaret.gov.tr','sanayi.gov.tr','ssb.gov.tr',
    'epdk.gov.tr','teias.gov.tr','tim.org.tr','osd.org.tr','odmd.org.tr'
]

STATISTIC_TERMS = [
    'sanayi üretim','sanayi üretimi','üretim endeksi','imalat sanayi',
    'kapasite kullanım','kapasite kullanım oranı','kko',
    'ihracat','dış ticaret','dış ticaret istatistik',
    'otomotiv üretim','otomotiv ihracat','araç üretim',
    'savunma ihracat','savunma ve havacılık ihracat',
    'elektrik üretim','enerji üretim','kurulu güç','tüketim',
    'yatırım teşvik','teşvik belgesi','sabit yatırım',
    'ar-ge','arge','araştırma geliştirme','yenilik','patent başvuru',
    'teknoloji istatistik','bilişim','girişim','yüksek teknoloji'
]

def build_official_radar_queries(when):
    """Genel medya taramasından ayrı, birincil/resmî kaynak sorguları."""
    gov_sites='('+' OR '.join('site:'+d for d in OFFICIAL_RADAR_DOMAINS)+')'
    return [
        f'(sanayi OR teknoloji OR üretim OR yatırım OR ihracat OR savunma OR Ar-Ge OR patent) {gov_sites} when:{when}',
        f'("basın açıklaması" OR duyuru OR açıklandı OR yayımlandı OR rapor OR veri OR istatistik) {gov_sites} when:{when}'
    ]

def build_statistics_queries(when):
    """Günlük sayısal veri yayımlarını yakalamaya dönük dar ve hızlı ek sorgular."""
    sites='('+' OR '.join('site:'+d for d in PRIMARY_STATS_DOMAINS)+')'
    return [
        f'("sanayi üretimi" OR "kapasite kullanım" OR ihracat OR "dış ticaret" OR "otomotiv üretimi") {sites} when:{when}',
        f'("savunma ihracatı" OR "enerji üretimi" OR "kurulu güç" OR "yatırım teşvik" OR "Ar-Ge") {sites} when:{when}'
    ]

def _is_official_radar_row(r):
    d=domain(r.get('Domain','') or r.get('URL',''))
    srcn=norm(r.get('Kaynak',''))
    if d in OFFICIAL_RADAR_DOMAINS or d in PRIMARY_STATS_DOMAINS:
        return True
    names=['sanayi ve teknoloji bakanlığı','tübitak','tubitak','kosgeb','türkpatent','turkpatent',
           'tse','savunma sanayii başkanlığı','ssb','tüik','tuik','tcmb','ticaret bakanlığı',
           'epdk','teiaş','teias','türkiye uzay ajansı']
    return any(x in srcn for x in names)


# -----------------------------
# V52 — GÜNÜN EN DEĞERLİ 10 GELİŞMESİ
# -----------------------------
V52_STRATEGIC_TERMS=[
    'savunma','savunma sanayii','tusaş','aselsan','roketsan','havelsan','baykar',
    'kaan','kızılelma','füze','hava savunma','siber','kritik altyapı',
    'yapay zeka','yarı iletken','çip','nükleer','enerji','otomotiv',
    'yatırım','fabrika','üretim','ihracat','arge','ar-ge','teknoloji yatırımı',
    'kritik mineral','nadir toprak','tedarik zinciri'
]

def _v52_event_value_table(df,n=10):
    """
    Olay bazlı 0-100 Değer Skoru.
    Gerçek okunma/tıklanma verisi mevcut akışta bulunmadığından uydurulmaz.
    Bunun yerine erişilebilen güçlü vekiller kullanılır:
    önem/risk, kaynak yayılımı, resmî teyit, güncellik, stratejik önem,
    negatif/eleştirel etki ve aynı olayın haber yoğunluğu.
    """
    cols=['Sıra','Değer_Skoru','Tarih','Gelişme','Neden_Değerli',
          'Kaynak_Sayısı','Haber_Sayısı','Resmî_Teyit','Risk','URL']
    if df is None or df.empty:
        return pd.DataFrame(columns=cols)

    now=pd.Timestamp.now(tz='UTC')
    items=[]

    for oid,g in df.groupby('Olay_ID',dropna=False):
        g=g.sort_values('Tarih_dt',ascending=False).copy()
        rep=g.iloc[0]
        maxrisk=int(pd.to_numeric(g.get('Risk_Skoru',0),errors='coerce').fillna(0).max())
        domains={domain(x) for x in g.get('Domain',pd.Series(dtype=str)).astype(str) if x}
        source_count=max(1,len(domains))
        news_count=len(g)
        official=any(_is_official_radar_row(r) for _,r in g.iterrows())

        latest=pd.to_datetime(g['Tarih_dt'],utc=True,errors='coerce').max()
        age_h=max(0.0,(now-latest).total_seconds()/3600) if pd.notna(latest) else 24.0
        recency=max(0.0,1.0-min(age_h,24.0)/24.0)

        text=norm(' '.join(
            (g['Başlık'].fillna('').astype(str)+' '+g['İçerik_Özeti'].fillna('').astype(str)).head(6).tolist()
        ))
        strategic_hits=sum(1 for x in V52_STRATEGIC_TERMS if x in text)
        strategic=min(1.0,strategic_hits/3.0)

        negative=bool(
            (g.get('Duygu',pd.Series(index=g.index,dtype=str))=='Negatif').any()
            or (g.get('Risk_Durumu',pd.Series(index=g.index,dtype=str))=='Yüksek Risk').any()
        )

        # 0-100: kullanıcının istediği kıstaslara göre dengeli ağırlık.
        risk_part=min(25.0,maxrisk*0.25)
        spread_part=min(20.0,5.0*source_count + max(0,news_count-source_count)*1.5)
        official_part=15.0 if official else 0.0
        recency_part=10.0*recency
        strategic_part=15.0*strategic
        impact_part=10.0 if negative else (5.0 if maxrisk>=35 else 0.0)

        # Gerçek click/read metriği yoksa "çok sayıda bağımsız kaynakta yankı"
        # popülerlik vekili olarak en fazla 5 puan taşır.
        popularity_proxy=min(5.0,max(0,source_count-1)*1.5 + max(0,news_count-2)*0.5)

        score=int(round(min(100,risk_part+spread_part+official_part+recency_part+
                            strategic_part+impact_part+popularity_proxy)))

        why=[]
        if source_count>=4: why.append(f'{source_count} farklı kaynakta geniş yankı')
        elif source_count>=2: why.append(f'{source_count} farklı kaynakta yer aldı')
        if official: why.append('resmî/birincil kaynak teyidi')
        if maxrisk>=70: why.append('yüksek risk/önem')
        elif maxrisk>=35: why.append('dikkat gerektiren etki')
        if strategic>=0.67: why.append('stratejik sanayi-teknoloji konusu')
        elif strategic>0: why.append('sanayi-teknoloji açısından ilgili')
        if negative: why.append('negatif/eleştirel etki')
        if recency>=0.75: why.append('çok güncel')
        if not why: why.append('güncel olay yoğunluğu')

        items.append({
            'Değer_Skoru':score,
            'Tarih':rep.get('Tarih',''),
            'Gelişme':rep.get('Başlık',''),
            'Neden_Değerli':' • '.join(why[:5]),
            'Kaynak_Sayısı':source_count,
            'Haber_Sayısı':news_count,
            'Resmî_Teyit':'Evet' if official else 'Hayır',
            'Risk':maxrisk,
            'URL':rep.get('URL','')
        })

    out=pd.DataFrame(items)
    if out.empty: return pd.DataFrame(columns=cols)
    out=out.sort_values(['Değer_Skoru','Kaynak_Sayısı','Haber_Sayısı','Tarih'],
                        ascending=[False,False,False,False]).head(n).reset_index(drop=True)
    out.insert(0,'Sıra',range(1,len(out)+1))
    return out[cols]


def _v53_find_event_row(df, value_row):
    """Top-10 satırını ana dataframe'deki temsilci haberle eşleştirir."""
    if df is None or df.empty:
        return None
    url=str(value_row.get('URL','') or '')
    title=norm(value_row.get('Gelişme',''))
    if url and 'URL' in df.columns:
        m=df[df['URL'].astype(str)==url]
        if not m.empty:
            return m.iloc[0]
    if title:
        m=df[df['Başlık'].astype(str).map(norm)==title]
        if not m.empty:
            return m.iloc[0]
    return None

def _v54_content_sentences(text,title=''):
    """Tam haber metninden menü/tekrar/gürültüyü azaltarak bilgi taşıyan cümleleri seçer."""
    clean=_clean_note_text(text)
    if not clean:
        return []
    title_n=norm(title)
    raw=_sentence_chunks(clean)
    out=[]; seen=set()
    noise=[
        'çerez','cookie','reklam','abonelik','bildirimleri aç','tüm hakları saklıdır',
        'gizlilik politikası','kullanım koşulları','facebook','instagram','twitter',
        'whatsapp','telegram','son dakika haberleri için'
    ]
    for s in raw:
        s=_clean_note_text(s)
        sn=norm(s)
        if len(s)<35 or len(s)>650: continue
        if any(x in sn for x in noise): continue
        if title_n and sn==title_n: continue
        key=re.sub(r'\W+','',sn)[:260]
        if not key or key in seen: continue
        seen.add(key)
        out.append(s)
    return out

def _v54_article_summary(detail,fallback_row,max_sentences=4):
    """
    Haberin içeriğini 2-4 bilgi yoğun cümlede özetler.
    Değer skoru, kaynak sayısı, resmî teyit gibi sıralama metadatasını özete katmaz.
    """
    title=_clean_note_text((detail or {}).get('title','') or fallback_row.get('Başlık',''))
    text=_clean_note_text((detail or {}).get('text','') or fallback_row.get('İçerik_Özeti',''))
    sents=_v54_content_sentences(text,title)

    if not sents:
        fallback=_clean_note_text(fallback_row.get('İçerik_Özeti',''))
        return fallback[:900].strip() if fallback else title

    # İlk anlamlı cümle bağlamı korur. Sonraki cümleler bilgi yoğunluğuna göre seçilir.
    selected=[sents[0]]
    candidates=[]
    for idx,s in enumerate(sents[1:],1):
        sn=norm(s)
        score=0
        if re.search(r'\b\d+(?:[.,]\d+)?\b',s): score+=4
        if any(x in sn for x in [
            'açıkladı','duyurdu','belirtti','bildirdi','ifade etti','kaydetti',
            'üretim','ihracat','ithalat','yatırım','istihdam','kapasite','satış',
            'sözleşme','anlaşma','teslim','tedarik','teşvik','destek','proje',
            'yangın','patlama','hasar','yaralı','kayıp','siber','veri',
            'arttı','azaldı','düştü','yükseldi','geriledi','başladı','tamamlandı'
        ]): score+=3
        if any(x in sn for x in [
            'bakanlık','tüik','ssb','şirket','firma','kurum','başkanlığı',
            'genel müdür','bakan','başkanı'
        ]): score+=1
        # Çok erken cümlelere hafif öncelik.
        score+=max(0,3-min(idx,3))
        candidates.append((score,idx,s))

    for _,_,s in sorted(candidates,key=lambda x:(-x[0],x[1])):
        if s not in selected:
            selected.append(s)
        if len(selected)>=max_sentences:
            break

    # Haber akışını bozmayacak şekilde özgün sıraya döndür.
    order={s:i for i,s in enumerate(sents)}
    selected=sorted(selected,key=lambda s:order.get(s,999))
    result=_join_sentences_naturally(selected)

    # Tek olayın 45 satırlık toplam özeti şişirmesini önle.
    return result[:1500].strip()

def _v54_deep_top10_summary(df,value10,max_lines=45):
    """
    Yalnızca Top-10 olayın temsilci haber sayfalarını butona basılınca zenginleştirir.
    Normal tarama hızını etkilemez. Her olay içerik odaklı özetlenir.
    """
    if value10 is None or value10.empty:
        return "Bugünün en değerli gelişmeleri arasında özet oluşturulabilecek içerik bulunamadı."

    lines=[
        f"Sanayi ve teknoloji gündeminde günün en değerli {len(value10)} gelişmesine ilişkin durum özeti aşağıda sunulmuştur.",
        ""
    ]

    for _,v in value10.head(10).iterrows():
        if len(lines)>=max_lines-3:
            break
        rank=int(v.get('Sıra',0) or 0)
        row=_v53_find_event_row(df,v)
        title=_clean_note_text(v.get('Gelişme',''))

        if row is None:
            detail_text=title
        else:
            try:
                # Ağ/tam metin işlemi SADECE özet butonuna basıldığında bu 10 haber için çalışır.
                detail=article_detail(row.to_dict() if hasattr(row,'to_dict') else row)
            except Exception:
                detail=None
            detail_text=_v54_article_summary(detail,row,4)

        lines.append(f"{rank}. {title}")
        if detail_text:
            lines.append(detail_text)
        lines.append("")

    lines.append(
        "Söz konusu gelişmelerin yeni açıklamalar ve ilave açık kaynak verileri doğrultusunda takip edilmesi önem taşımaktadır."
    )
    return '\n'.join(lines[:max_lines])

def make_v54_top10_summary_docx(df,value10,text=None):
    text=text or _v54_deep_top10_summary(df,value10,45)
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    doc.styles['Normal'].font.name='Times New Roman'
    doc.styles['Normal'].font.size=Pt(11)

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run('BUGÜNÜN SANAYİ VE TEKNOLOJİ DURUM ÖZETİ')
    rr.bold=True; rr.font.size=Pt(14)

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().astimezone().strftime('%d.%m.%Y %H:%M'))

    for line in text.splitlines():
        if not line.strip(): continue
        bp=doc.add_paragraph()
        bp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after=Pt(5)
        r=bp.add_run(line)
        if re.match(r'^\d+\.\s',line):
            r.bold=True

    bio=BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.getvalue()



# -----------------------------
# V58 — ANALİTİK TAKİP ÜÇLÜSÜ
# 1) Olay Yaşam Döngüsü
# 2) Takip Edilecek Açık Hususlar
# 3) Teyit / Çelişki Matrisi
# Ek web isteği yapmaz; mevcut tarama sonuçlarını kullanır.
# -----------------------------

V58_RESOLUTION_TERMS=[
    'kontrol altına alındı','söndürüldü','sona erdi','tamamlandı','çözüldü',
    'giderildi','yeniden başladı','üretim yeniden başladı','faaliyet yeniden başladı',
    'normalleşti','normal seyrine döndü','tahliye sona erdi','arıza giderildi',
    'erişim sağlandı','sistem yeniden devreye alındı'
]

V58_ESCALATION_TERMS=[
    'arttı','büyüdü','genişledi','yayılıyor','devam ediyor','sürüyor',
    'üretim durdu','faaliyet durdu','tahliye','ikinci patlama','yeni patlama',
    'can kaybı','yaralı sayısı','hasar arttı','soruşturma başlatıldı',
    'acil durum','kriz','kesinti sürüyor'
]

def _v58_event_groups(df):
    if df is None or df.empty or 'Olay_ID' not in df.columns:
        return {}
    groups={}
    for oid,g in df.groupby('Olay_ID',dropna=False):
        groups[str(oid)]=g.sort_values('Tarih_dt',ascending=True).copy()
    return groups

def _v58_event_stage(g):
    """Olayın mevcut taramadaki izlerine göre yaşam döngüsü aşaması."""
    if g is None or g.empty:
        return 'İlk Sinyal'

    text=norm(' '.join(
        (g['Başlık'].fillna('').astype(str)+' '+g['İçerik_Özeti'].fillna('').astype(str)).tolist()
    ))
    source_count=max(1,g['Domain'].astype(str).replace('',pd.NA).dropna().nunique()) if 'Domain' in g.columns else 1
    news_count=len(g)
    official=any(_is_official_radar_row(r) for _,r in g.iterrows())

    if any(x in text for x in V58_RESOLUTION_TERMS):
        return '✅ Sonuçlandı'
    if official:
        return '🟢 Teyit Edildi'
    if source_count>=2 or news_count>=3 or any(x in text for x in V58_ESCALATION_TERMS):
        return '🟠 Gelişiyor'
    return '🔵 İlk Sinyal'

def _v58_stage_reason(g,stage):
    source_count=max(1,g['Domain'].astype(str).replace('',pd.NA).dropna().nunique()) if 'Domain' in g.columns else 1
    news_count=len(g)
    official=any(_is_official_radar_row(r) for _,r in g.iterrows())
    text=norm(' '.join(
        (g['Başlık'].fillna('').astype(str)+' '+g['İçerik_Özeti'].fillna('').astype(str)).tolist()
    ))
    reasons=[]
    if official: reasons.append('resmî/birincil açıklama mevcut')
    if source_count>=2: reasons.append(f'{source_count} farklı kaynak')
    if news_count>=3: reasons.append(f'{news_count} haber kaydı')
    if any(x in text for x in V58_RESOLUTION_TERMS): reasons.append('sonuç/normalleşme ifadesi')
    elif any(x in text for x in V58_ESCALATION_TERMS): reasons.append('devam/etki artışı sinyali')
    if not reasons: reasons.append('tek/erken kaynak sinyali')
    return ' • '.join(reasons)

def _v58_event_lifecycle_table(df,limit=25):
    cols=['Tarih','Aşama','Başlık','Kategori','Kaynak_Sayısı','Haber_Sayısı',
          'Doğrulama','Risk_Skoru','Aşama_Gerekçesi','URL']
    groups=_v58_event_groups(df)
    rows=[]
    for oid,g in groups.items():
        latest=g.sort_values('Tarih_dt',ascending=False).iloc[0]
        stage=_v58_event_stage(g)
        source_count=max(1,g['Domain'].astype(str).replace('',pd.NA).dropna().nunique()) if 'Domain' in g.columns else 1
        rows.append({
            'Tarih':latest.get('Tarih',''),
            'Aşama':stage,
            'Başlık':latest.get('Başlık',''),
            'Kategori':latest.get('Kategori',''),
            'Kaynak_Sayısı':source_count,
            'Haber_Sayısı':len(g),
            'Doğrulama':latest.get('Doğrulama',''),
            'Risk_Skoru':int(pd.to_numeric(g['Risk_Skoru'],errors='coerce').fillna(0).max()) if 'Risk_Skoru' in g.columns else 0,
            'Aşama_Gerekçesi':_v58_stage_reason(g,stage),
            'URL':latest.get('URL',''),
            '_stage_rank':{'🟠 Gelişiyor':4,'🟢 Teyit Edildi':3,'🔵 İlk Sinyal':2,'✅ Sonuçlandı':1}.get(stage,0),
            '_dt':pd.to_datetime(latest.get('Tarih_dt'),utc=True,errors='coerce')
        })
    if not rows:
        return pd.DataFrame(columns=cols)
    out=pd.DataFrame(rows).sort_values(['_stage_rank','Risk_Skoru','_dt'],ascending=[False,False,False])
    return out.head(limit).drop(columns=['_stage_rank','_dt'],errors='ignore')

def _v58_open_questions_for_group(g):
    """
    'Bilinmiyor' iddiası üretmez; mevcut içerikte ayrıca teyit/izleme gerektiren
    alanları analist kontrol listesi olarak önerir.
    """
    text=norm(' '.join(
        (g['Başlık'].fillna('').astype(str)+' '+g['İçerik_Özeti'].fillna('').astype(str)).tolist()
    ))
    qs=[]

    if any(_v89_has_term(text,x) for x in ['yangın','patlama','infilak','kaza']):
        qs += [
            'Olayın kesin nedeni ve teknik inceleme sonucu',
            'Can kaybı/yaralı ve maddi hasarın resmî bilançosu',
            'Üretim/faaliyet sürekliliğine etkisi ve normale dönüş takvimi'
        ]
    if any(x in text for x in ['siber','veri sızıntısı','veri ihlali','fidye','güvenlik açığı']):
        qs += [
            'Etkilenen sistem/veri kapsamının kesinleştirilmesi',
            'İhlalin kaynağı ve alınan düzeltici tedbirler',
            'Operasyonel hizmetlere etkisinin sürüp sürmediği'
        ]
    if any(x in text for x in ['yatırım','fabrika kurulacak','tesis kurulacak','teşvik']):
        qs += [
            'Yatırım tutarı, kapasitesi ve finansman yapısının teyidi',
            'Yatırım/üretime geçiş takvimi',
            'İstihdam ve yerli tedarik etkisinin netleşmesi'
        ]
    if any(x in text for x in ['ihracat','sözleşme','anlaşma','sipariş','teslimat','savunma']):
        qs += [
            'Sözleşme/anlaşmanın kapsamı ve parasal büyüklüğü',
            'Teslimat/uygulama takvimi',
            'Karşı taraf veya resmî makam teyidi'
        ]
    if any(x in text for x in ['üretim düştü','daralma','geriledi','azaldı','maliyet baskısı','rekabet gücü']):
        qs += [
            'Olumsuz eğilimin geçici mi yapısal mı olduğunun izlenmesi',
            'Bir sonraki resmî veri setinde eğilimin devam edip etmediği',
            'Sektör/şirket bazında üretim, ihracat ve istihdam etkisi'
        ]

    official=any(_is_official_radar_row(r) for _,r in g.iterrows())
    source_count=max(1,g['Domain'].astype(str).replace('',pd.NA).dropna().nunique()) if 'Domain' in g.columns else 1
    if not official:
        qs.append('Resmî/birincil kaynak açıklaması')
    if source_count<2:
        qs.append('İkinci bağımsız kaynaktan teyit')

    if not qs:
        qs=[
            'Gelişmenin kapsamının yeni açıklamalarla netleşmesi',
            'Resmî/birincil kaynak teyidi',
            'Sanayi/teknoloji alanındaki somut etkisinin izlenmesi'
        ]

    # Sıralı tekilleştirme, en fazla 4 açık husus.
    out=[]
    seen=set()
    for q in qs:
        k=norm(q)
        if k in seen: continue
        seen.add(k); out.append(q)
        if len(out)>=4: break
    return out

def _v58_open_issues_table(df,limit=20):
    cols=['Tarih','Başlık','Aşama','Takip_Edilecek_Açık_Hususlar','Risk_Skoru','Doğrulama','URL']
    groups=_v58_event_groups(df)
    rows=[]
    for oid,g in groups.items():
        latest=g.sort_values('Tarih_dt',ascending=False).iloc[0]
        stage=_v58_event_stage(g)
        # Sonuçlanan olaylar açık hususlar listesinin altında kalsın; aktif olaylar öne çıksın.
        qs=_v58_open_questions_for_group(g)
        risk=int(pd.to_numeric(g['Risk_Skoru'],errors='coerce').fillna(0).max()) if 'Risk_Skoru' in g.columns else 0
        rows.append({
            'Tarih':latest.get('Tarih',''),
            'Başlık':latest.get('Başlık',''),
            'Aşama':stage,
            'Takip_Edilecek_Açık_Hususlar':' • '.join(qs),
            'Risk_Skoru':risk,
            'Doğrulama':latest.get('Doğrulama',''),
            'URL':latest.get('URL',''),
            '_active':0 if stage=='✅ Sonuçlandı' else 1,
            '_dt':pd.to_datetime(latest.get('Tarih_dt'),utc=True,errors='coerce')
        })
    if not rows:
        return pd.DataFrame(columns=cols)
    out=pd.DataFrame(rows).sort_values(['_active','Risk_Skoru','_dt'],ascending=[False,False,False])
    return out.head(limit).drop(columns=['_active','_dt'],errors='ignore')

def _v58_numeric_claims(g):
    """Kaynak bazında başlık+özetten sayısal iddiaları çıkarır."""
    claims=[]
    pat=re.compile(r'(?:%\s*)?\b\d+(?:[.,]\d+)?\b(?:\s*%|\s*(?:milyon|milyar|bin|adet|kişi|yaralı|ölü|mw|gw|ton|tl|dolar|euro|avro))?',re.I)
    for _,r in g.iterrows():
        txt=f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')}"
        nums={x.strip() for x in pat.findall(txt) if x.strip()}
        claims.append((str(r.get('Kaynak','')),nums))
    return claims

def _v58_conflict_status(g):
    source_count=max(1,g['Domain'].astype(str).replace('',pd.NA).dropna().nunique()) if 'Domain' in g.columns else 1
    official=any(_is_official_radar_row(r) for _,r in g.iterrows())
    claims=_v58_numeric_claims(g)

    nonempty=[nums for _,nums in claims if nums]
    numeric_conflict=False
    if len(nonempty)>=2:
        # Birden fazla kaynağın sayısal kümeleri tamamen ayrışıyorsa uyar.
        for i in range(len(nonempty)):
            for j in range(i+1,len(nonempty)):
                if nonempty[i] and nonempty[j] and nonempty[i].isdisjoint(nonempty[j]):
                    numeric_conflict=True
                    break
            if numeric_conflict: break

    text=norm(' '.join(
        (g['Başlık'].fillna('').astype(str)+' '+g['İçerik_Özeti'].fillna('').astype(str)).tolist()
    ))
    verbal_conflict=(
        ('can kaybı yok' in text and ('can kaybı' in text.replace('can kaybı yok','') or 'hayatını kaybetti' in text))
        or ('yaralı yok' in text and 'yaralandı' in text)
        or ('üretim durdu' in text and ('üretim devam ediyor' in text or 'üretim sürüyor' in text))
    )

    if numeric_conflict or verbal_conflict:
        return '🔴 Çelişkili Bilgi','Kaynaklar arasında sayı/olgu farklılığı tespit edildi; manuel teyit önerilir.'
    if official:
        return '🟢 Resmî Teyitli','Resmî/birincil kaynak mevcut.'
    if source_count>=2:
        return '🟢 Çoklu Kaynak','En az iki farklı kaynak aynı olayı destekliyor.'
    return '🟡 Tek Kaynak','İkinci bağımsız veya resmî teyit henüz görünmüyor.'

def _v58_verification_matrix(df,limit=25):
    cols=['Tarih','Başlık','Teyit_Durumu','Teyit_Açıklaması','Kaynak_Sayısı',
          'Haber_Sayısı','Risk_Skoru','URL']
    groups=_v58_event_groups(df)
    rows=[]
    rank={'🔴 Çelişkili Bilgi':4,'🟡 Tek Kaynak':3,'🟢 Çoklu Kaynak':2,'🟢 Resmî Teyitli':1}
    for oid,g in groups.items():
        latest=g.sort_values('Tarih_dt',ascending=False).iloc[0]
        status,reason=_v58_conflict_status(g)
        source_count=max(1,g['Domain'].astype(str).replace('',pd.NA).dropna().nunique()) if 'Domain' in g.columns else 1
        risk=int(pd.to_numeric(g['Risk_Skoru'],errors='coerce').fillna(0).max()) if 'Risk_Skoru' in g.columns else 0
        rows.append({
            'Tarih':latest.get('Tarih',''),
            'Başlık':latest.get('Başlık',''),
            'Teyit_Durumu':status,
            'Teyit_Açıklaması':reason,
            'Kaynak_Sayısı':source_count,
            'Haber_Sayısı':len(g),
            'Risk_Skoru':risk,
            'URL':latest.get('URL',''),
            '_rank':rank.get(status,0),
            '_dt':pd.to_datetime(latest.get('Tarih_dt'),utc=True,errors='coerce')
        })
    if not rows:
        return pd.DataFrame(columns=cols)
    out=pd.DataFrame(rows).sort_values(['_rank','Risk_Skoru','_dt'],ascending=[False,False,False])
    return out.head(limit).drop(columns=['_rank','_dt'],errors='ignore')

# -----------------------------
# V51 — RESMÎ AÇIKLAMA / MEDYA KARŞILAŞTIRMASI
# -----------------------------
_COMPARE_STOP={
    've','ile','bir','bu','şu','için','da','de','mi','mı','mu','mü','olan','olarak',
    'son','yeni','göre','daha','çok','ise','ile','ancak','fakat','tarafından','dedi',
    'açıkladı','açıklama','haber','gelişme','türkiye','türk'
}

def _compare_tokens(text):
    t=norm(text)
    toks=re.findall(r'[a-z0-9çğıöşü]{3,}',t)
    return {x for x in toks if x not in _COMPARE_STOP}

def _event_similarity(a,b):
    """Başlık + kısa içerik üzerinden hızlı olay benzerliği; ağ isteği yapmaz."""
    at=_compare_tokens(f"{a.get('Başlık','')} {str(a.get('İçerik_Özeti',''))[:500]}")
    bt=_compare_tokens(f"{b.get('Başlık','')} {str(b.get('İçerik_Özeti',''))[:500]}")
    if not at or not bt: return 0.0
    inter=len(at & bt)
    union=max(1,len(at | bt))
    j=inter/union
    title_a=_compare_tokens(a.get('Başlık',''))
    title_b=_compare_tokens(b.get('Başlık',''))
    tj=len(title_a & title_b)/max(1,min(len(title_a),len(title_b))) if title_a and title_b else 0
    return 0.55*tj+0.45*j

def _short_claim(r,limit=220):
    txt=_clean_note_text(r.get('İçerik_Özeti',''))
    if not txt or norm(txt)==norm(r.get('Başlık','')):
        txt=_clean_note_text(r.get('Başlık',''))
    sents=_sentence_chunks(txt)
    if sents:
        txt=' '.join(sents[:2])
    return txt[:limit].strip()

def _comparison_difference(media,official):
    """İki kısa metindeki belirgin yön/iddia farklarını özetler; LLM/ağ çağrısı yok."""
    mt=norm(f"{media.get('Başlık','')} {media.get('İçerik_Özeti','')}")
    ot=norm(f"{official.get('Başlık','')} {official.get('İçerik_Özeti','')}")
    pairs=[
        (['tamamen durdu','üretim durdu','faaliyet durdu'],['kısmi','belirli bölüm','geçici','kısa süre','devam ediyor'],'Medya daha geniş bir durma/aksama bildirirken resmî açıklama etkinin kısmi veya geçici olduğunu belirtiyor.'),
        (['yangın','patlama','kaza'],['kontrol altına','söndürüldü','müdahale edildi'],'Resmî açıklama olayın kontrol/müdahale durumuna ilişkin ek bilgi içeriyor.'),
        (['can kaybı','öldü','hayatını kaybetti'],['can kaybı yok','can kaybı bulunmuyor'],'Can kaybına ilişkin medya ve resmî açıklama arasında farklı ifade bulunuyor.'),
        (['yaralı','yaralandı'],['yaralı yok','yaralanan yok'],'Yaralanma bilgisine ilişkin farklı ifade bulunuyor.'),
        (['veri sızıntısı','veri ihlali'],['etkilenmedi','sınırlı','belirli kullanıcı'],'Resmî açıklama olayın kapsamını medya anlatımına göre sınırlandırıyor/netleştiriyor.'),
        (['kriz','tehlike','alarm'],['normal','rutin','planlandığı','devam ediyor'],'Medya daha olumsuz/uyarıcı bir çerçeve kullanırken resmî açıklama daha sınırlı veya olağan bir durum tarif ediyor.')
    ]
    for mkeys,okeys,msg in pairs:
        if any(x in mt for x in mkeys) and any(x in ot for x in okeys):
            return msg

    mn=set(re.findall(r'(?:%\s*)?\d+(?:[.,]\d+)?',mt))
    on=set(re.findall(r'(?:%\s*)?\d+(?:[.,]\d+)?',ot))
    if mn and on and mn!=on:
        return 'Medya ve resmî açıklamada yer alan sayısal bilgiler farklılık gösteriyor; rakamların ayrıca kontrol edilmesi önerilir.'

    return 'Aynı olaya ilişkin resmî açıklama bulundu. Belirgin bir çelişki otomatik olarak tespit edilmedi; ayrıntılar birlikte kontrol edilebilir.'

def _official_media_comparison(df):
    """
    Sabit panel için medya haberlerini aynı taramadaki resmî/birincil içeriklerle eşleştirir.
    Ek web isteği yoktur; mevcut Resmî Kaynak Radarı verisini kullanır.
    """
    cols=['Tarih','Medya_Kaynağı','Medya_Haberi','Resmî_Kaynak','Resmî_Açıklama',
          'Karşılaştırma','Eşleşme','Medya_URL','Resmî_URL']
    if df is None or df.empty:
        return pd.DataFrame(columns=cols)

    officials=df[df.apply(_is_official_radar_row,axis=1)].copy()
    media=df[~df.apply(_is_official_radar_row,axis=1)].copy()
    if officials.empty or media.empty:
        return pd.DataFrame(columns=cols)

    rows=[]
    # Performans için resmî havuz zaten küçüktür; medya tarafında en yeni 250 içerik yeterli.
    media=media.sort_values('Tarih_dt',ascending=False).head(250)
    officials=officials.sort_values('Tarih_dt',ascending=False).head(80)

    for _,m in media.iterrows():
        best=None; best_score=0.0
        mdt=pd.to_datetime(m.get('Tarih_dt'),utc=True,errors='coerce')
        for _,o in officials.iterrows():
            odt=pd.to_datetime(o.get('Tarih_dt'),utc=True,errors='coerce')
            if pd.notna(mdt) and pd.notna(odt):
                if abs((mdt-odt).total_seconds()) > 72*3600:
                    continue
            score=_event_similarity(m,o)
            if score>best_score:
                best_score=score; best=o
        # Yanlış eşleşmeyi azaltmak için ölçülü eşik.
        if best is None or best_score<0.30:
            continue

        rows.append({
            'Tarih':m.get('Tarih',''),
            'Medya_Kaynağı':m.get('Kaynak',''),
            'Medya_Haberi':_short_claim(m),
            'Resmî_Kaynak':best.get('Kaynak',''),
            'Resmî_Açıklama':_short_claim(best),
            'Karşılaştırma':_comparison_difference(m,best),
            'Eşleşme':int(round(best_score*100)),
            'Medya_URL':m.get('URL',''),
            'Resmî_URL':best.get('URL','')
        })

    if not rows:
        return pd.DataFrame(columns=cols)
    out=pd.DataFrame(rows).drop_duplicates(subset=['Medya_URL','Resmî_URL'])
    return out.sort_values(['Eşleşme','Tarih'],ascending=[False,False])

def _contains_number_or_rate(text):
    t=str(text or '')
    return bool(re.search(
        r'(?<!\w)(?:%\\s*)?\\d+(?:[.,]\\d+)?(?:\\s*%|\\s*(?:milyon|milyar|trilyon|bin|adet|ton|mw|gw|gwh|twh|tl|₺|dolar|euro|avro))?',
        t,flags=re.I
    ))

def _critical_numbers(text, limit=4):
    t=re.sub(r'\\s+',' ',str(text or ''))
    pats=re.findall(
        r'(?:%\\s*\\d+(?:[.,]\\d+)?|\\d+(?:[.,]\\d+)?\\s*%|'
        r'\\d+(?:[.,]\\d+)?\\s*(?:milyon|milyar|trilyon|bin)\\s*(?:TL|₺|dolar|euro|avro)?|'
        r'\\d+(?:[.,]\\d+)?\\s*(?:MW|GW|GWh|TWh|ton|adet))',
        t,flags=re.I
    )
    out=[]
    for p in pats:
        p=p.strip()
        if p and p not in out:
            out.append(p)
        if len(out)>=limit:
            break
    return ', '.join(out)

def _important_statistics_rows(df):
    """Bugün yayımlanan, sanayi/teknoloji açısından sayısal veri taşıyan içerikleri seçer."""
    if df is None or df.empty:
        return pd.DataFrame()

    x=df.copy()
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')
    local_tz=datetime.now().astimezone().tzinfo
    today_local=datetime.now().astimezone().date()

    def is_today(v):
        try:
            return v is not None and pd.notna(v) and v.tz_convert(local_tz).date()==today_local
        except Exception:
            return False

    def stat_match(r):
        text=norm(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')} {r.get('Kategori','')}")
        term_hit=any(term in text for term in STATISTIC_TERMS)
        number_hit=_contains_number_or_rate(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')}")
        return term_hit and number_hit

    mask=x.apply(stat_match,axis=1)
    today_mask=x['Tarih_dt'].apply(is_today)
    result=x[mask & today_mask].copy()

    # Eğer yayın saati eksik gelmişse ama resmî/statistik kaynağı ve veri içeriği varsa dışarıda bırakma.
    missing_date=x['Tarih_dt'].isna()
    fallback=x[mask & missing_date & x.apply(_is_official_radar_row,axis=1)].copy()
    result=pd.concat([result,fallback],ignore_index=False).drop_duplicates(subset=['URL','Başlık'])

    if result.empty:
        return result

    result['Kritik_Sayı']=result.apply(
        lambda r:_critical_numbers(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')}"),
        axis=1
    )
    result['Birincil_Kaynak']=result.apply(lambda r:'✅' if _is_official_radar_row(r) else '—',axis=1)
    result=result.sort_values('Tarih_dt',ascending=False,na_position='last')
    return result

def _official_radar_rows(df):
    if df is None or df.empty:
        return pd.DataFrame()
    x=df[df.apply(_is_official_radar_row,axis=1)].copy()
    if x.empty:
        return x
    x=x.sort_values('Tarih_dt',ascending=False,na_position='last')
    return x.drop_duplicates(subset=['URL','Başlık'])

def _two_sentence_summary(text):
    sents=_detail_sentences(str(text or ''),'')
    if not sents:
        raw=_clean_note_text(text)
        return raw[:500]
    return ' '.join(sents[:2])

def _presentation_candidates(df,n=5):
    """Sunuma girmeye değer 5 başlık: stratejik önem + risk + resmîlik + sayısal veri + güncellik."""
    if df is None or df.empty:
        return pd.DataFrame()
    x=df.copy()
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')

    def score(r):
        text=norm(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')} {r.get('Kategori','')}")
        s=int(r.get('Risk_Skoru',0) or 0)//3
        if r.get('Risk_Durumu')=='Yüksek Risk': s+=18
        if _is_official_radar_row(r): s+=18
        if _contains_number_or_rate(text): s+=8
        if any(k in text for k in ['yatırım','ihracat','kapasite','savunma','yarı iletken','çip','yapay zeka',
                                   'enerji','otomotiv','uzay','ar-ge','arge','üretim','teşvik']): s+=14
        if critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti','')): s+=16
        try: s+=min(int(r.get('Olay_Kaynak_Sayisi',0) or 0)*3,12)
        except Exception: pass
        return s

    x['_Sunum_Puanı']=x.apply(score,axis=1)
    x=x.sort_values(['_Sunum_Puanı','Tarih_dt'],ascending=[False,False],na_position='last')
    if 'Olay_ID' in x.columns:
        x=x.drop_duplicates(subset=['Olay_ID'],keep='first')
    else:
        x=x.drop_duplicates(subset=['Başlık'],keep='first')
    x=x.head(n).copy()
    x['Sunum_Başlığı']=x['Başlık'].astype(str)
    x['2_Cümle_Özet']=x['İçerik_Özeti'].apply(_two_sentence_summary)
    x['Kritik_Sayı']=x.apply(
        lambda r:_critical_numbers(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')}") or '—',
        axis=1
    )
    return x.drop(columns=['_Sunum_Puanı'],errors='ignore')

def build_negative_queries(when):
    return [
        f'Türkiye (iflas OR konkordato OR "üretim durdu" OR "fabrika kapandı" OR "işten çıkarma" OR grev OR soruşturma OR dava OR ceza OR "geri çağırma" OR "siber saldırı" OR "veri sızıntısı" OR yaptırım OR ambargo OR "ihale iptal" OR ertelendi OR gecikme OR "tedarik krizi" OR daralma OR zafiyet OR usulsüzlük OR yolsuzluk) (sanayi OR teknoloji OR üretim OR fabrika OR savunma OR otomotiv OR enerji OR şirket OR tesis OR proje) when:{when}',
        f'Türkiye ((OSB OR "organize sanayi" OR fabrika OR tesis OR "sanayi sitesi") (yangın OR yangını OR alev OR patlama OR patladı OR infilak)) when:{when}'
    ]

def build_greek_queries(when):
    site='('+' OR '.join('site:'+x for x in GR)+')'
    return [
        f'(Turkey OR Türkiye OR Turkish OR Τουρκία OR τουρκική) (defense OR defence OR savunma OR άμυνα OR arms) {site} when:{when}',
        f'(Baykar OR Bayraktar OR ASELSAN OR TUSAŞ OR Roketsan OR HAVELSAN OR KAAN OR Kızılelma OR SİPER OR HİSAR) {site} when:{when}',
        f'(Turkey OR Turkish OR Τουρκία) (drone OR UAV OR missile OR fighter OR frigate OR submarine OR defense industry) {site} when:{when}'
    ]

def build_social_queries(when):
    site='('+' OR '.join('site:'+x for x in SOCIAL)+')'
    return [
        f'(Türkiye OR Türk) (sanayi OR teknoloji OR üretim OR savunma OR yapay zeka OR siber) {site} when:{when}',
        f'(ASELSAN OR TUSAŞ OR ROKETSAN OR HAVELSAN OR Baykar OR TOGG OR TÜBİTAK) {site} when:{when}',
        f'(iflas OR üretim durdu OR fabrika kapandı OR soruşturma OR siber saldırı OR yaptırım) (sanayi OR teknoloji OR savunma) {site} when:{when}'
    ]

def normalize_rows(raw, cutoff, mode, user_query):
    out=[]; reasons={'zaman':0,'konu':0,'kaynak':0,'yunan':0,'gecersiz':0}
    for r in raw:
        url=(r.get('url') or r.get('link') or '').strip(); title=html.unescape((r.get('title') or '').strip())
        if not url or not title: reasons['gecersiz']+=1; continue
        dt=parse_dt(r.get('date') or r.get('publishedAt') or r.get('seendate'))
        # Tüm tarihleri UTC-aware datetime olarak karşılaştır. Bazı RSS/arama
        # sağlayıcıları timezone bilgisi olmadan tarih döndürebildiği için
        # doğrudan datetime karşılaştırması TypeError üretebilir.
        if dt:
            try:
                if dt.tzinfo is None:
                    dt=dt.replace(tzinfo=timezone.utc)
                else:
                    dt=dt.astimezone(timezone.utc)
                cutoff_utc = cutoff if cutoff.tzinfo is not None else cutoff.replace(tzinfo=timezone.utc)
                cutoff_utc = cutoff_utc.astimezone(timezone.utc)
                if dt < cutoff_utc:
                    reasons['zaman']+=1
                    continue
            except (TypeError, ValueError, AttributeError):
                # Tarih karşılaştırılamıyorsa haberi düşürme; aşağıda
                # bilinmeyen tarih olarak sıralanmasına izin ver.
                dt=None
        if not dt and mode=='turkish':
            # tarih yoksa hızlı bakışta atmayalım; sadece sıralamada alta al.
            pass
        snippet=html.unescape((r.get('snippet') or r.get('body') or r.get('description') or '').strip())
        src=r.get('source') or ''
        d=infer_source(src,r.get('source_url',''),url)
        t=f'{title} {snippet}'
        if mode=='greek':
            if d not in GR or not greek_defense(t): reasons['yunan']+=1; continue
        elif mode=='social':
            if d not in SOCIAL: reasons['kaynak']+=1; continue
            if not relevant(t,user_query): reasons['konu']+=1; continue
        elif mode=='global':
            if not relevant(t,user_query): reasons['konu']+=1; continue
        else:
            # Türk batch'inde kaynak filtresi YOK. Arama zaten Türkiye odaklı.
            # Bu, Google News'in yayıncı URL'sini Google domaininde tuttuğu durumlarda
            # Türk haberlerinin 0'a düşmesini engeller. Türk kaynakları sıralamada öne çıkar.
            if not relevant(t,user_query): reasons['konu']+=1; continue
        sentiment,score,status,neg,risk,cat,risk_reasons=classify(title,snippet,d)
        out.append({
            'Tarih_dt':dt,'Tarih':fmt_dt(dt),'Başlık':title,'İçerik_Özeti':snippet or title,
            'URL':url,'RSS_URL':url,'Kaynak':(src if norm(src) not in {'google haberler','google news','google'} else (d or src or 'Açık Kaynak')),
            'Yayıncı_URL':(r.get('source_url') or '').strip(),'Yayıncı':src or d or 'Açık Kaynak',
            'Domain':d,'Kaynak_Grubu':source_group(d),
            'Kategori':cat,'Duygu':sentiment,'Skor':score,'Risk_Skoru':score,'Risk_Durumu':status,
            'Risk_Gerekçesi':'; '.join(risk_reasons),'Negatif_Sinyaller':neg,'Risk_Sinyalleri':risk,
            'Seç':False,'Görsel_URL':'','_mode':mode
        })
    return out,reasons


def source_reliability(domain_name, source_name=''):
    d=domain(domain_name); n=norm(source_name)
    if d in TR_OFFICIAL: return '🟢 A — Birincil / resmî'
    if d in TR_MAIN or d in TR_TECH: return '🟢 A — Güvenilir medya'
    if d in GR: return '🔵 B — Yunan medya'
    if d in SOCIAL: return '🟠 C — Sosyal / indeks'
    return '🟡 B — Açık kaynak'



def dedupe(rows):
    """URL ve başlık anahtarına göre hızlı tekilleştirme; kronolojik sıralamayı korur."""
    out=[]
    urls=set()
    titles=set()
    for r in rows:
        u=str(r.get('URL','') or '')
        k=title_key(str(r.get('Başlık','') or ''))
        if u and u in urls:
            continue
        if k and k in titles:
            continue
        if u:
            urls.add(u)
        if k:
            titles.add(k)
        out.append(r)

    out.sort(
        key=lambda x:(
            x.get('Tarih_dt') is not None,
            _to_utc_datetime(x.get('Tarih_dt')) or datetime.min.replace(tzinfo=timezone.utc),
            source_rank(x.get('Domain',''))
        ),
        reverse=True
    )
    return out


def _title_tokens(text):
    """Başlıktan olay eşleştirmesi için anlamlı token kümesi üretir."""
    txt=norm(text)
    txt=re.sub(r'[^\wçğıöşüÇĞİÖŞÜ]+',' ',txt)
    stop={
        've','ile','bir','bu','da','de','için','son','yeni','türkiye','türk','haberi','haber',
        'açıklama','dedi','oldu','olarak','olan','milyon','milyar','bin','yüzde','ile ilgili'
    }
    return {x for x in txt.split() if len(x)>=3 and x not in stop}


def _event_signature(title):
    """
    Aynı/çok benzer haber başlıklarını hızlı gruplamak için deterministik imza.
    İlk 6 ayırt edici token kullanılır. O(n²) SequenceMatcher taraması yerine
    ters indeks kullanacağımız için yüzlerce haberde çok daha hızlıdır.
    """
    toks=sorted(_title_tokens(title))
    return ' '.join(toks[:6])


def _jaccard(a,b):
    if not a or not b:
        return 0.0
    inter=len(a & b)
    union=len(a | b)
    return inter/union if union else 0.0


def source_reliability(source_domain,source_name=''):
    d=domain(source_domain); n=norm(source_name)
    if d in TR_OFFICIAL: return '🟢 A — Birincil / resmî'
    if d in TR_MAIN or d in TR_TECH: return '🟢 A — Güvenilir medya'
    if d in GR: return '🔵 B — Yunan medya'
    if d in SOCIAL: return '🟠 C — Sosyal / indeks'
    return '🟡 B — Açık kaynak'


def enrich_rows(rows):
    """
    HIZLI analitik katman.
    Önceki sürümde her haber diğer bütün haberlerle SequenceMatcher üzerinden
    karşılaştırılıyordu ve doğrulama için ikinci kez O(n²) tarama yapılıyordu.
    Bu sürüm ters token indeksi + olay grubu istatistikleri kullanır.
    """
    if not rows:
        return rows

    # 1) Tarih + temel risk sınıflaması: O(n)
    for r in rows:
        r['Tarih_dt']=_to_utc_datetime(r.get('Tarih_dt'))
        sentiment,score,status,neg,risk,cat,reasons=classify(
            r.get('Başlık',''),r.get('İçerik_Özeti',''),r.get('Domain','')
        )
        r['Duygu']=sentiment
        r['Risk_Skoru']=score
        r['Risk_Durumu']=status
        r['Negatif_Sinyaller']=neg
        r['Risk_Sinyalleri']=risk
        r['Risk_Gerekçesi']='; '.join(reasons)
        r['Kaynak_Güvenilirliği']=source_reliability(r.get('Domain',''),r.get('Kaynak',''))
        r['_tokens']=_title_tokens(r.get('Başlık',''))

    # 2) Olay kümelemesi: ters token indeksi.
    # Her haber yalnızca ortak token taşıyan sınırlı sayıdaki önceki adayla karşılaştırılır.
    token_index={}
    event_representative={}
    next_event=1

    for idx,r in enumerate(rows):
        toks=r['_tokens']
        candidate_events=set()
        for tok in toks:
            candidate_events.update(token_index.get(tok,set()))

        best_event=None
        best_score=0.0
        for eid in candidate_events:
            rep_tokens=event_representative[eid]
            score=_jaccard(toks,rep_tokens)
            if score>best_score:
                best_score=score
                best_event=eid

        # Aynı olay için Jaccard eşiği. Çok kısa başlıklarda biraz daha sıkı.
        threshold=0.48 if len(toks)>=6 else 0.58
        if best_event is None or best_score < threshold:
            best_event=f'OLAY-{next_event:03d}'
            next_event+=1
            event_representative[best_event]=set(toks)

        r['Olay_ID']=best_event
        for tok in toks:
            token_index.setdefault(tok,set()).add(best_event)

    # 3) Olay istatistikleri bir kez hesaplanır: O(n)
    groups={}
    for r in rows:
        groups.setdefault(r['Olay_ID'],[]).append(r)

    event_meta={}
    for eid,g in groups.items():
        domains={x.get('Domain') for x in g if x.get('Domain')}
        times=[x.get('Tarih_dt') for x in g if x.get('Tarih_dt') is not None]
        official=any(domain(x.get('Domain','')) in TR_OFFICIAL for x in g)
        social_only=all(domain(x.get('Domain','')) in SOCIAL for x in g if x.get('Domain')) if domains else False

        if official:
            verification='🟢 Resmî açıklama / birincil kaynak'
        elif len(domains)>=2 or len(g)>=3:
            verification='🟢 Çoklu kaynakla destekleniyor'
        elif social_only:
            verification='🟠 Sosyal medya / tek kaynak'
        elif any(domain(x.get('Domain','')) in TR_MAIN+TR_TECH+GR for x in g):
            verification='🟡 Tek medya kaynağı'
        else:
            verification='🟡 Tek/açık kaynak'

        event_meta[eid]={
            'sources':len(domains),
            'first':fmt_dt(min(times)) if times else '',
            'last':fmt_dt(max(times)) if times else '',
            'verification':verification
        }

    for r in rows:
        meta=event_meta[r['Olay_ID']]
        r['Olay_Kaynak_Sayisi']=meta['sources']
        r['Olay_İlk_Görülme']=meta['first'] or r.get('Tarih','')
        r['Olay_Son_Görülme']=meta['last'] or r.get('Tarih','')
        r['Doğrulama']=meta['verification']
        r.pop('_tokens',None)

    return rows

def build_event_summary(df):
    if df.empty: return pd.DataFrame()
    items=[]
    for oid,g in df.groupby('Olay_ID',dropna=False):
        g=g.sort_values('Tarih_dt',ascending=False)
        head=str(g.iloc[0].get('Başlık',''))
        risk=int(g['Risk_Skoru'].max())
        cat=str(g.iloc[0].get('Kategori',''))
        sources=', '.join(dict.fromkeys(str(x) for x in g['Kaynak'].tolist()))
        items.append({'Olay_ID':oid,'Öne Çıkan Başlık':head,'Kategori':cat,'Haber Sayısı':len(g),'Kaynak Sayısı':g['Domain'].nunique(),'Risk':risk,'İlk Görülme':g['Olay_İlk_Görülme'].min(),'Son Görülme':g['Olay_Son_Görülme'].max(),'Kaynaklar':sources})
    return pd.DataFrame(items).sort_values(['Risk','Son Görülme'],ascending=[False,False])

def trend_table(df):
    if df.empty: return pd.DataFrame()
    x=df.copy(); x['Saat']=x['Tarih_dt'].apply(lambda d: d.strftime('%Y-%m-%d %H:00') if d else 'Bilinmiyor')
    return x.groupby(['Kategori']).size().reset_index(name='Haber').sort_values('Haber',ascending=False)

def watchlist_hits(df, terms):
    terms=[norm(x) for x in re.split(r',|\n|;',terms or '') if len(norm(x))>=2]
    if df.empty or not terms: return pd.DataFrame()
    mask=df.apply(lambda r:any(t in norm(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')}") for t in terms),axis=1)
    return df[mask].copy()

def _repair_mojibake_utf8(text):
    """
    'TÃ¼rkiye', 'genÃ§', 'katÄ±lÄ±m', 'baÅarÄ±' gibi UTF-8'in yanlış
    Latin-1/Windows-1252 olarak çözülmesinden doğan bozulmaları düzeltir.
    Doğru Türkçe metne dokunmamaya çalışır.
    """
    s=str(text or '')
    if not s:
        return s

    suspicious=('Ã','Ä','Å','Â','â€','â€™','â€œ','â€','â€“','â€”','\x80','\x81','\x8d','\x8f','\x90','\x9d','\x9f')
    if not any(x in s for x in suspicious):
        return s

    # Önce Windows-1252 mojibake işaretlerini bayt değerlerine geri çevirebilmek
    # için özel karakter -> byte haritası oluştur.
    cp1252_rev={}
    for b in range(256):
        try:
            ch=bytes([b]).decode('cp1252')
            cp1252_rev[ch]=b
        except Exception:
            pass

    def char_to_byte(ch):
        o=ord(ch)
        if o <= 255:
            return o
        return cp1252_rev.get(ch)

    # UTF-8 olabilecek bayt dizilerini parça parça düzelt; doğru Unicode
    # karakterler (ör. gerçek “ ’ ğ ş) sınır olarak korunur.
    out=[]
    buf=[]
    def flush():
        nonlocal buf
        if not buf:
            return
        raw=bytes(buf)
        original=''.join(chr(b) for b in buf)
        try:
            decoded=raw.decode('utf-8')
            # Yalnız gerçekten mojibake işaretlerini azaltıyorsa kabul et.
            before=sum(original.count(x) for x in ('Ã','Ä','Å','Â'))
            after=sum(decoded.count(x) for x in ('Ã','Ä','Å','Â'))
            out.append(decoded if after < before else original)
        except Exception:
            out.append(original)
        buf=[]

    for ch in s:
        b=char_to_byte(ch)
        if b is None:
            flush()
            out.append(ch)
        else:
            buf.append(b)
    flush()
    fixed=''.join(out)

    # Çok katmanlı bozulma varsa en fazla iki tur daha dene.
    for _ in range(2):
        if not any(x in fixed for x in ('Ã','Ä','Å','Â')):
            break
        try:
            candidate=fixed.encode('latin1').decode('utf-8')
            if sum(candidate.count(x) for x in ('Ã','Ä','Å','Â')) < sum(fixed.count(x) for x in ('Ã','Ä','Å','Â')):
                fixed=candidate
            else:
                break
        except Exception:
            break
    return fixed

def _clean_note_text(value):
    """
    V78 Word-safe metin temizliği:
    - mojibake'i bayt düzeyinde onarır,
    - Türkçe karakterleri Unicode NFC biçiminde korur,
    - DOCX/XML açısından sorunlu kontrol/görünmez karakterleri temizler.
    """
    import html as _html
    import unicodedata as _unicodedata

    text=BeautifulSoup(str(value or ''),'html.parser').get_text(' ',strip=True)
    text=_html.unescape(text)
    text=_repair_mojibake_utf8(text)

    # Kalan yaygın tipografik bozulmalar.
    replacements={
        'â€™':'’','â€˜':'‘','â€œ':'“','â€':'”',
        'â€“':'–','â€”':'—','â€¦':'…','Â ':' ','Â':''
    }
    for bad,good in replacements.items():
        text=text.replace(bad,good)

    for bad in ('\u00ad','\u200b','\u200c','\u200d','\ufeff'):
        text=text.replace(bad,'')

    # XML 1.0 geçersiz kontrol karakterlerini at.
    text=''.join(
        ch for ch in text
        if ch in ('\t','\n','\r') or ord(ch)>=32
    )

    text=_unicodedata.normalize('NFC',text)
    text=re.sub(r'\s+',' ',text).strip()
    return text

def _sentence_split_tr(text):
    text=_clean_note_text(text)
    if not text:
        return []
    parts=re.split(r'(?<=[.!?])\s+(?=[A-ZÇĞİÖŞÜ0-9“"])',text)
    return [p.strip() for p in parts if len(p.strip())>20]

def _unique_sentences(sentences):
    out=[]; seen=set()
    for s in sentences:
        k=norm(s)
        if not k or k in seen:
            continue
        seen.add(k); out.append(s)
    return out

def _note_source_sentence(r):
    title=_clean_note_text(r.get('Başlık',''))
    source=_clean_note_text(r.get('Kaynak','Açık Kaynak'))
    when=_clean_note_text(r.get('Tarih',''))
    cat=_clean_note_text(r.get('Kategori',''))
    if when:
        return f"{when} tarihinde {source} tarafından yayımlanan “{title}” başlıklı içerik, {cat.lower() if cat else 'sanayi ve teknoloji'} alanındaki gelişmelere ilişkindir."
    return f"{source} tarafından yayımlanan “{title}” başlıklı içerik, {cat.lower() if cat else 'sanayi ve teknoloji'} alanındaki gelişmelere ilişkindir."

def _detail_sentences(text, title=''):
    """Haber gövdesinden bilgi taşıyan cümleleri temizler; ayrıntıyı korur."""
    text=_clean_note_text(text)
    if not text:
        return []
    raw=_sentence_split_tr(text)
    title_n=norm(title)
    boiler=[
        'çerez','cookie','abonelik','abone ol','reklam','tüm hakları saklıdır',
        'gizlilik politikası','kullanım koşulları','google news','bildirimleri aç',
        'uygulamamızı indirin','facebook','instagram','twitter','whatsapp',
        'son dakika haberleri için','haberlerimizi takip'
    ]
    out=[]; seen=set()
    for s in raw:
        sn=norm(s)
        if len(s)<28 or sn==title_n or any(b in sn for b in boiler):
            continue
        key=' '.join(sn.split()[:16])
        if key in seen:
            continue
        seen.add(key)
        out.append(s.strip())
    return out


def _sent_score(s):
    """Bilgi yoğun cümlelere öncelik verir."""
    n=norm(s)
    score=0
    if re.search(r'\b\d+(?:[.,]\d+)?\b', s): score+=3
    if any(x in n for x in ['tarih','saat','yıl','ay','gün','bugün','dün']): score+=2
    if any(x in n for x in ['bakan','başkan','valilik','belediye','şirket','kurum','bakanlık','müdür','yetkili','açıkladı','bildirdi','belirtti']): score+=3
    if any(x in n for x in ['nedeni','sebebi','sonucu','sonuç','etki','hasar','zarar','yaralı','hayatını kaybetti','tahliye','müdahale','kontrol altına']): score+=3
    if any(x in n for x in ['üretim','kapasite','yatırım','ihracat','ithalat','tesis','fabrika','osb','teknoloji','savunma','enerji']): score+=2
    return score


def _join_sentences_naturally(sentences):
    """Kaynak cümlelerini bilgi kaybı olmadan okunabilir paragraf akışına getirir."""
    if not sentences:
        return ''
    out=[]
    for s in sentences:
        s=s.strip()
        if not s:
            continue
        if s[-1] not in '.!?':
            s+='.'
        out.append(s)
    return ' '.join(out)


def _compose_single_article_note(row, detail):
    """
    Tek haberi 'haber özeti' gibi değil, ayrıntılı bilgi notu gibi ele alır:
    konu/olay -> gelişmeler -> açıklamalar/veriler -> mevcut durum/sonuç.
    Ara başlık kullanmaz.
    """
    title=_clean_note_text(detail.get('title') or row.get('Başlık',''))
    source=_clean_note_text(detail.get('source') or row.get('Kaynak','Açık Kaynak'))
    published=_clean_note_text(detail.get('published') or row.get('Tarih',''))
    fulltext=_clean_note_text(detail.get('text') or row.get('İçerik_Özeti','') or title)
    sentences=_detail_sentences(fulltext,title)

    # Haber sırasını esas al. İlk cümleler olayın başlangıcını çoğunlukla verir.
    # Çok uzun haberlerde bilgi yoğun cümleleri de mutlaka koru.
    if len(sentences)>45:
        first=sentences[:22]
        rest=sentences[22:]
        important=sorted(enumerate(rest), key=lambda z:_sent_score(z[1]), reverse=True)[:18]
        important=[s for _,s in sorted(important,key=lambda z:z[0])]
        sentences=first+important

    intro=(
        f"{published} tarihinde {source} tarafından yayımlanan “{title}” başlıklı haberde, "
        if published else
        f"{source} tarafından yayımlanan “{title}” başlıklı haberde, "
    )

    if not sentences:
        fallback=_clean_note_text(row.get('İçerik_Özeti','') or title)
        return intro + (fallback[0].lower()+fallback[1:] if len(fallback)>1 else fallback)

    # İlk 1-2 cümle olayın girişini oluşturur; geri kalanı kronolojik/haber sırasıyla devam eder.
    first=sentences[:2]
    remaining=sentences[2:]
    opening=_join_sentences_naturally(first)
    if opening:
        opening=opening[0].lower()+opening[1:]
    para1=intro+opening

    # Uzun haberlerde okunabilirlik için doğal paragraf bölmeleri.
    chunks=[]
    chunk_size=7
    for i in range(0,len(remaining),chunk_size):
        part=remaining[i:i+chunk_size]
        txt=_join_sentences_naturally(part)
        if txt:
            chunks.append(txt)

    parts=[para1]+chunks

    # Son cümlede yalnızca kaynakta aktarılan çerçeveye dayan.
    last_context=sentences[-3:] if len(sentences)>=3 else sentences
    conclusion=(
        "Bu çerçevede, haberde aktarılan son durum itibarıyla "
        + _join_sentences_naturally(last_context)
    )
    # Son üç cümleyi gövdede zaten kullandığımız için birebir tekrar çok fazlaysa genel, temkinli kapanış kullan.
    if len(norm(conclusion))>900:
        conclusion="Bu çerçevede gelişmenin, haberde aktarılan mevcut durum ve ilgili kurumların sonraki açıklamaları doğrultusunda izlenmesi önem taşımaktadır."
    parts.append(conclusion)

    return '\n\n'.join(parts)


def _compose_prose_note(df):
    """
    Seçilen gerçek haber sayfalarının tam metninden ayrıntılı bilgi notu oluşturur.
    'Giriş/Gelişme/Sonuç' başlıkları yazılmaz; anlatı doğal olarak bu sırada ilerler.
    """
    if df is None or df.empty:
        return '', []

    x=df.copy()
    if 'Tarih_dt' in x.columns:
        x['Tarih_dt']=pd.to_datetime(x['Tarih_dt'],utc=True,errors='coerce')
        x=x.sort_values('Tarih_dt',ascending=True,na_position='last')

    enriched=[]
    for _,r in x.iterrows():
        row=r.to_dict()
        detail=article_detail(row)
        enriched.append((row,detail))

    if len(enriched)==1:
        row,detail=enriched[0]
        note=_compose_single_article_note(row,detail)
        return note,enriched

    # Çoklu haberde kısa bir doğal giriş, ardından her haber kronolojik sırada ayrıntılı işlenir.
    source_names=[]
    for row,detail in enriched:
        s=_clean_note_text(detail.get('source') or row.get('Kaynak',''))
        if s and s not in source_names:
            source_names.append(s)

    opening=(
        f"Seçilen {len(enriched)} açık kaynak haberi birlikte değerlendirildiğinde, konuya ilişkin gelişmeler "
        f"{len(source_names)} farklı kaynağın aktardığı bilgiler çerçevesinde kronolojik bir seyir göstermektedir. "
        f"Aşağıdaki anlatımda haberlerde yer alan olaylar, açıklamalar, kişi ve kurumlar, teknik ve sayısal veriler, "
        f"neden-sonuç ilişkileri ile bildirilen etkiler mümkün olduğunca ayrıntılı biçimde korunmuştur."
    )

    blocks=[opening]
    for row,detail in enriched:
        blocks.append(_compose_single_article_note(row,detail))

    blocks.append(
        "Mevcut açık kaynak bilgileri birlikte değerlendirildiğinde, gelişmenin bundan sonraki seyri bakımından "
        "ilgili kurum ve kuruluşların yeni açıklamalarının, resmî duyuruların ve farklı açık kaynaklardan gelecek "
        "teyitlerin izlenmesi önem taşımaktadır. Bu bilgi notunda kaynak haberlerde yer almayan bir husus olgu olarak eklenmemiştir."
    )
    return '\n\n'.join(blocks), enriched

def make_analyst_docx(df, title='BİLGİ NOTU'):
    """
    V66: Başlıksız üç aşamalı bilgi notu yapısı:
    1) İlk paragraf kısa özet,
    2) devam eden paragraf(lar) ayrıntı/rakam/istatistik/gelişme,
    3) son paragraf sonuç ve kısa değerlendirme.
    Metin daima 'Arz olunur.' ile tamamlanır.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    styles=doc.styles
    styles['Normal'].font.name='Times New Roman'; styles['Normal'].font.size=Pt(12)
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(_clean_note_text(title)); r.bold=True; r.font.size=Pt(14)
    p=doc.add_paragraph(); p.add_run('Tarih: ').bold=True
    p.add_run(datetime.now().astimezone().strftime('%d.%m.%Y'))

    enriched=[]
    x=df.copy() if df is not None else pd.DataFrame()
    if 'Tarih_dt' in x.columns:
        x['Tarih_dt']=pd.to_datetime(x['Tarih_dt'],utc=True,errors='coerce')
        x=x.sort_values('Tarih_dt',ascending=True,na_position='last')
    for _,rr in x.iterrows():
        row=rr.to_dict()
        try:
            detail=article_detail(row)
        except Exception:
            detail={}
        enriched.append((row,detail))

    all_sent=[]
    for row,detail in enriched:
        title_text=_clean_note_text(detail.get('title') or row.get('Başlık',''))
        body=_clean_note_text(detail.get('text') or row.get('İçerik_Özeti') or title_text)
        all_sent.extend(_akt_clean_sentences(title_text,body))

    # Yakın tekrarları temizle.
    uniq=[]; seen=[]
    for sent in all_sent:
        sent=_clean_note_text(sent)
        key=norm(sent)
        toks=set(key.split())
        if not key: continue
        dup=False
        for old in seen[-35:]:
            union=len(toks|old)
            if union and len(toks&old)/union>=0.78:
                dup=True; break
        if not dup:
            uniq.append(sent.strip()); seen.append(toks)

    def add_body(text):
        bp=doc.add_paragraph()
        bp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.first_line_indent=Cm(1.25)
        bp.paragraph_format.line_spacing=1.15
        bp.paragraph_format.space_after=Pt(8)
        safe_text=_repair_mojibake_utf8(_clean_note_text(text))
        bp.add_run(_v66_formalize_sentence_endings(safe_text))

    if uniq:
        # İlk paragraf: haberin kısa özeti. Başlık yazılmaz.
        intro=_join_sentences_naturally(uniq[:2])
        add_body(intro)

        # Gelişme bölümü: başlık kullanılmadan, ayrıntı/rakam/istatistikler korunarak devam eder.
        detail_s=uniq[2:]
        if not detail_s:
            detail_s=uniq

        # Uzun haberlerde ayrıntıları iki paragraf halinde dağıtarak okunabilirliği koru.
        detail_s=detail_s[:18]
        if len(detail_s)<=9:
            add_body(_join_sentences_naturally(detail_s))
        else:
            add_body(_join_sentences_naturally(detail_s[:9]))
            add_body(_join_sentences_naturally(detail_s[9:18]))

        # Son paragraf: sonuç + kısa/temkinli değerlendirme; ayrı başlık yoktur.
        tail=_join_sentences_naturally(uniq[-3:])
        if tail:
            conclusion=(
                f"Mevcut bilgiler çerçevesinde, {tail[0].lower()+tail[1:]} "
                "Gelişmenin sanayi ve teknoloji alanındaki muhtemel etkilerinin, ilgili kurum ve kuruluşların "
                "yeni açıklamaları ile resmî veriler doğrultusunda takip edilmesinin uygun olacağı değerlendirilmektedir."
            )
        else:
            conclusion=(
                "Mevcut bilgiler çerçevesinde gelişmenin sanayi ve teknoloji alanındaki etkilerinin, ilgili kurum "
                "ve kuruluşların yeni açıklamaları ile resmî veriler doğrultusunda takip edilmesinin uygun olacağı değerlendirilmektedir."
            )
        add_body(conclusion)
    else:
        add_body('Seçilen habere ilişkin ayrıntılı içerik temin edilememiştir.')
        add_body(
            'Gelişmenin yeni açık kaynak bilgileri ile ilgili kurum ve kuruluşların resmî açıklamaları '
            'doğrultusunda takip edilmesinin uygun olacağı değerlendirilmektedir.'
        )

    endp=doc.add_paragraph()
    endp.paragraph_format.space_before=Pt(8)
    endp.add_run('Arz olunur.')

    if enriched:
        kp=doc.add_paragraph()
        kr=kp.add_run('Kaynak: '); kr.bold=True
        for i,(row,detail) in enumerate(enriched):
            source=_clean_note_text(detail.get('source') or row.get('Kaynak','Açık Kaynak'))
            url=detail.get('canonical') or row.get('Yayıncı_URL') or row.get('URL','')
            if i: kp.add_run('; ')
            kp.add_run(source)
            if url:
                kp.add_run(' ('); _word_hyperlink(kp,url,'Haber linki'); kp.add_run(')')

    bio=BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.getvalue()


# -----------------------------
# V33 — BİLGİ NOTU ADAYLARI + DÜNDEN BERİ NE DEĞİŞTİ?
# V32 çekirdek tarama motoruna dokunmaz.
# -----------------------------
_HISTORY_DB = Path(__file__).resolve().with_name("sanayi_teknoloji_osint_history.db")

def _history_connect():
    conn=sqlite3.connect(str(_HISTORY_DB),timeout=8)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn

def _init_history_db():
    try:
        with _history_connect() as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS scans(
                    scan_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    scanned_at TEXT NOT NULL,
                    period_hours INTEGER,
                    total_news INTEGER,
                    total_events INTEGER
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS event_snapshots(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    scan_id INTEGER NOT NULL,
                    event_id TEXT,
                    title TEXT,
                    source TEXT,
                    url TEXT,
                    category TEXT,
                    summary TEXT,
                    risk_score INTEGER,
                    risk_status TEXT,
                    sentiment TEXT,
                    verification TEXT,
                    source_count INTEGER,
                    event_first_seen TEXT,
                    event_last_seen TEXT,
                    tokens_json TEXT,
                    FOREIGN KEY(scan_id) REFERENCES scans(scan_id)
                )
            """)
            conn.execute("CREATE INDEX IF NOT EXISTS idx_event_snapshots_scan ON event_snapshots(scan_id)")
            conn.execute("""
                CREATE TABLE IF NOT EXISTS shift_marks(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    marked_at TEXT NOT NULL,
                    scan_id INTEGER,
                    label TEXT
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS important_basket(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    added_at TEXT NOT NULL,
                    news_time TEXT,
                    title TEXT NOT NULL,
                    source TEXT,
                    url TEXT,
                    category TEXT,
                    risk_score INTEGER,
                    risk_status TEXT,
                    summary TEXT,
                    UNIQUE(url,title)
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS presentation_basket(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    added_at TEXT NOT NULL,
                    news_time TEXT,
                    title TEXT NOT NULL,
                    source TEXT,
                    url TEXT,
                    category TEXT,
                    summary TEXT,
                    UNIQUE(url,title)
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS osint_report_basket(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    added_at TEXT NOT NULL,
                    news_time TEXT,
                    title TEXT NOT NULL,
                    source TEXT,
                    url TEXT,
                    category TEXT,
                    risk_score INTEGER,
                    risk_status TEXT,
                    summary TEXT,
                    UNIQUE(url,title)
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS app_visits(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    visited_at TEXT NOT NULL
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS note_history(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    created_at TEXT NOT NULL,
                    title TEXT NOT NULL,
                    url TEXT,
                    UNIQUE(url,title)
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS tomorrow_followup(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    added_at TEXT NOT NULL,
                    title TEXT NOT NULL,
                    source TEXT,
                    url TEXT,
                    category TEXT,
                    reason TEXT,
                    UNIQUE(url,title)
                )
            """)
            conn.commit()
        return True
    except Exception:
        return False

def _history_tokens(text):
    try:
        toks=_title_tokens(text)
        return sorted(toks)
    except Exception:
        txt=norm(text)
        return sorted({x for x in re.split(r'\W+',txt) if len(x)>=3})

def _save_scan_history(rows, scanned_at, period_hours):
    """Her taramanın olay özetini yerel SQLite dosyasına kaydeder."""
    if not rows or not _init_history_db():
        return None
    try:
        dfh=pd.DataFrame(rows)
        events=int(dfh['Olay_ID'].nunique()) if 'Olay_ID' in dfh.columns else len(dfh)
        with _history_connect() as conn:
            cur=conn.execute(
                "INSERT INTO scans(scanned_at,period_hours,total_news,total_events) VALUES(?,?,?,?)",
                (scanned_at.isoformat(),int(period_hours),len(dfh),events)
            )
            scan_id=int(cur.lastrowid)

            if 'Olay_ID' in dfh.columns:
                groups=dfh.groupby('Olay_ID',dropna=False)
            else:
                groups=[(f'ROW-{i}',dfh.iloc[[i]]) for i in range(len(dfh))]

            rows_to_insert=[]
            for oid,g in groups:
                g=g.copy()
                if 'Tarih_dt' in g.columns:
                    g['Tarih_dt']=pd.to_datetime(g['Tarih_dt'],utc=True,errors='coerce')
                    g=g.sort_values('Tarih_dt',ascending=False,na_position='last')
                r=g.iloc[0]
                title=str(r.get('Başlık','') or '')
                summary=' '.join(
                    str(x) for x in g.get('İçerik_Özeti',pd.Series(dtype=str)).tolist()
                    if str(x).strip()
                )[:8000]
                domains=set(str(x) for x in g.get('Domain',pd.Series(dtype=str)).tolist() if str(x).strip())
                source_count=max(
                    len(domains),
                    int(r.get('Olay_Kaynak_Sayisi',0) or 0)
                )
                rows_to_insert.append((
                    scan_id,str(oid),title,str(r.get('Kaynak','') or ''),
                    str(r.get('URL','') or ''),str(r.get('Kategori','') or ''),
                    summary,int(g.get('Risk_Skoru',pd.Series([0])).max() or 0),
                    str(r.get('Risk_Durumu','') or ''),str(r.get('Duygu','') or ''),
                    str(r.get('Doğrulama','') or ''),source_count,
                    str(r.get('Olay_İlk_Görülme','') or ''),
                    str(r.get('Olay_Son_Görülme','') or ''),
                    json.dumps(_history_tokens(title),ensure_ascii=False)
                ))

            conn.executemany("""
                INSERT INTO event_snapshots(
                    scan_id,event_id,title,source,url,category,summary,risk_score,risk_status,
                    sentiment,verification,source_count,event_first_seen,event_last_seen,tokens_json
                ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,rows_to_insert)
            conn.commit()
        return scan_id
    except Exception:
        return None

def _previous_scan_id(current_scan_id=None):
    """Öncelik: bugünden önceki en son tarama; yoksa mevcut taramadan önceki en son tarama."""
    if not _init_history_db():
        return None
    try:
        today=datetime.now().astimezone().date().isoformat()
        with _history_connect() as conn:
            if current_scan_id:
                row=conn.execute(
                    "SELECT scan_id FROM scans WHERE scan_id < ? AND substr(scanned_at,1,10) < ? ORDER BY scanned_at DESC LIMIT 1",
                    (int(current_scan_id),today)
                ).fetchone()
                if not row:
                    row=conn.execute(
                        "SELECT scan_id FROM scans WHERE scan_id < ? ORDER BY scanned_at DESC LIMIT 1",
                        (int(current_scan_id),)
                    ).fetchone()
            else:
                row=conn.execute(
                    "SELECT scan_id FROM scans WHERE substr(scanned_at,1,10) < ? ORDER BY scanned_at DESC LIMIT 1",
                    (today,)
                ).fetchone()
            return int(row[0]) if row else None
    except Exception:
        return None

def _load_scan_events(scan_id):
    if not scan_id:
        return pd.DataFrame()
    try:
        with _history_connect() as conn:
            return pd.read_sql_query(
                """SELECT e.*,s.scanned_at,s.period_hours
                   FROM event_snapshots e JOIN scans s ON e.scan_id=s.scan_id
                   WHERE e.scan_id=?""",
                conn,params=(int(scan_id),)
            )
    except Exception:
        return pd.DataFrame()

def _token_jaccard_lists(a,b):
    sa=set(a or []); sb=set(b or [])
    if not sa or not sb:
        return 0.0
    return len(sa&sb)/len(sa|sb)

def _verification_rank(text):
    t=norm(text)
    if 'resmi' in t or 'resmî' in t or 'birincil' in t: return 4
    if 'coklu kaynak' in t or 'çoklu kaynak' in t: return 3
    if 'tek medya' in t: return 2
    if 'sosyal medya' in t: return 1
    return 1

def _risk_rank(status):
    t=norm(status)
    if 'yuksek risk' in t or 'yüksek risk' in t: return 3
    if 'negatif' in t: return 2
    return 1

def _current_event_frame(df):
    if df is None or df.empty:
        return pd.DataFrame()
    items=[]
    group_col='Olay_ID' if 'Olay_ID' in df.columns else None
    groups=df.groupby(group_col,dropna=False) if group_col else [(f'ROW-{i}',df.iloc[[i]]) for i in range(len(df))]
    for oid,g in groups:
        g=g.copy()
        if 'Tarih_dt' in g.columns:
            g['Tarih_dt']=pd.to_datetime(g['Tarih_dt'],utc=True,errors='coerce')
            g=g.sort_values('Tarih_dt',ascending=False,na_position='last')
        r=g.iloc[0]
        summary=' '.join(str(x) for x in g.get('İçerik_Özeti',pd.Series(dtype=str)).tolist() if str(x).strip())[:8000]
        items.append({
            'event_id':str(oid),
            'title':str(r.get('Başlık','') or ''),
            'source':str(r.get('Kaynak','') or ''),
            'url':str(r.get('URL','') or ''),
            'category':str(r.get('Kategori','') or ''),
            'summary':summary,
            'risk_score':int(g.get('Risk_Skoru',pd.Series([0])).max() or 0),
            'risk_status':str(r.get('Risk_Durumu','') or ''),
            'sentiment':str(r.get('Duygu','') or ''),
            'verification':str(r.get('Doğrulama','') or ''),
            'source_count':max(
                int(r.get('Olay_Kaynak_Sayisi',0) or 0),
                len(set(str(x) for x in g.get('Domain',pd.Series(dtype=str)).tolist() if str(x).strip()))
            ),
            'event_first_seen':str(r.get('Olay_İlk_Görülme','') or ''),
            'event_last_seen':str(r.get('Olay_Son_Görülme','') or ''),
            'tokens':_history_tokens(str(r.get('Başlık','') or ''))
        })
    return pd.DataFrame(items)

def _compare_since_previous(df,current_scan_id=None):
    """
    Olay bazında:
    🆕 yeni olay
    🔄 yeni bilgi/güncelleme
    ⚠️ risk arttı
    ✅ teyit güçlendi
    """
    current=_current_event_frame(df)
    prev_id=_previous_scan_id(current_scan_id)
    previous=_load_scan_events(prev_id)
    if current.empty:
        return pd.DataFrame(),None,None
    if previous.empty:
        return pd.DataFrame(),prev_id,None

    prev_records=[]
    for _,p in previous.iterrows():
        try: toks=json.loads(p.get('tokens_json') or '[]')
        except Exception: toks=_history_tokens(p.get('title',''))
        rec=p.to_dict(); rec['tokens']=toks; prev_records.append(rec)

    changes=[]
    for _,c in current.iterrows():
        best=None; best_sim=0.0
        for p in prev_records:
            sim=_token_jaccard_lists(c['tokens'],p['tokens'])
            # Kaynak/URL aynıysa eşleşmeyi kuvvetlendir.
            if c.get('url') and c.get('url')==p.get('url'):
                sim=max(sim,0.95)
            if sim>best_sim:
                best_sim=sim; best=p

        if best is None or best_sim < 0.42:
            changes.append({
                'Değişim':'🆕 YENİ OLAY',
                'Başlık':c['title'],'Kaynak':c['source'],'Kategori':c['category'],
                'Risk':c['risk_score'],'Önceki Risk':'—','Kaynak Sayısı':c['source_count'],
                'Açıklama':'Önceki karşılaştırma taramasında benzer olay tespit edilmedi.',
                'URL':c['url'],'_priority':100+c['risk_score']
            })
            continue

        prev_risk=int(best.get('risk_score') or 0)
        risk_up=(c['risk_score']>=prev_risk+15) or (_risk_rank(c['risk_status'])>_risk_rank(best.get('risk_status','')))
        verify_up=_verification_rank(c['verification'])>_verification_rank(best.get('verification',''))
        sources_up=int(c['source_count'] or 0)>int(best.get('source_count') or 0)

        prev_tokens=set(_history_tokens((best.get('title') or '')+' '+(best.get('summary') or '')))
        cur_tokens=set(_history_tokens(c['title']+' '+c['summary']))
        new_tokens=cur_tokens-prev_tokens
        materially_updated=len(new_tokens)>=6 or sources_up

        if risk_up:
            kind='⚠️ RİSK ARTTI'
            expl=f"Risk {prev_risk}/100 seviyesinden {c['risk_score']}/100 seviyesine yükseldi."
            priority=95+c['risk_score']
        elif verify_up:
            kind='✅ TEYİT GÜÇLENDİ'
            expl=f"Doğrulama seviyesi “{best.get('verification','')}” düzeyinden “{c['verification']}” düzeyine yükseldi."
            priority=90+c['risk_score']
        elif materially_updated:
            kind='🔄 YENİ BİLGİ'
            bits=[]
            if sources_up:
                bits.append(f"kaynak sayısı {int(best.get('source_count') or 0)} → {c['source_count']}")
            if len(new_tokens)>=6:
                sample=', '.join(sorted(list(new_tokens))[:8])
                bits.append(f"yeni içerik unsurları: {sample}")
            expl='; '.join(bits) if bits else 'Olayla ilgili yeni ayrıntılar tespit edildi.'
            priority=80+c['risk_score']
        else:
            continue

        changes.append({
            'Değişim':kind,'Başlık':c['title'],'Kaynak':c['source'],'Kategori':c['category'],
            'Risk':c['risk_score'],'Önceki Risk':prev_risk,'Kaynak Sayısı':c['source_count'],
            'Açıklama':expl,'URL':c['url'],'_priority':priority
        })

    out=pd.DataFrame(changes)
    if not out.empty:
        out=out.sort_values(['_priority','Risk'],ascending=[False,False]).drop(columns=['_priority'])
    prev_time=str(previous.iloc[0].get('scanned_at','')) if not previous.empty else None
    return out,prev_id,prev_time

def _note_candidate_reason(r,change_kind=''):
    reasons=[]
    risk=int(r.get('risk_score',0) or 0)
    if risk>=70: reasons.append('yüksek risk')
    elif risk>=45: reasons.append('dikkat gerektiren risk')
    if str(r.get('sentiment',''))=='Negatif': reasons.append('negatif etki')
    if int(r.get('source_count',0) or 0)>=2: reasons.append('çoklu kaynak')
    vr=norm(r.get('verification',''))
    if 'resmi' in vr or 'resmî' in vr or 'birincil' in vr: reasons.append('birincil/resmî teyit')
    elif 'coklu kaynak' in vr or 'çoklu kaynak' in vr: reasons.append('teyit güçlendi')
    cat=norm(r.get('category',''))
    if any(k in cat for k in ['savunma','yarı iletken','dijital','enerji','sanayi']): reasons.append('stratejik sektör')
    if is_osb_fire(r.get('title',''),r.get('summary','')): reasons.append('OSB yangını/kritik üretim olayı')
    if change_kind:
        reasons.append(change_kind.replace('🆕','').replace('🔄','').replace('⚠️','').replace('✅','').strip().lower())
    return ', '.join(dict.fromkeys(reasons)) or 'güncel ve sektörel önem'

def _information_note_candidates(df,current_scan_id=None,limit=10):
    events=_current_event_frame(df)
    if events.empty:
        return pd.DataFrame()

    changes,_,_=_compare_since_previous(df,current_scan_id)
    change_map={}
    if not changes.empty:
        for _,c in changes.iterrows():
            change_map[c['Başlık']]=c.get('Değişim',c.get('Tür',''))

    rows=[]
    for _,r in events.iterrows():
        score=0
        risk=int(r['risk_score'] or 0)
        score += min(45,int(risk*0.45))
        if r['risk_status']=='Yüksek Risk': score+=18
        elif r['sentiment']=='Negatif': score+=10
        score += min(int(r['source_count'] or 0)*4,16)

        vr=_verification_rank(r['verification'])
        score += {4:12,3:9,2:4,1:0}.get(vr,0)

        title_summary=norm(r['title']+' '+r['summary'])
        if is_osb_fire(r['title'],r['summary']): score+=18
        if any(x in title_summary for x in ['savunma','aselsan','tusaş','tusas','roketsan','baykar','havelsan','füze','iha','siha']): score+=10
        if any(x in title_summary for x in ['yatırım','yeni tesis','kapasite art','ihracat','kritik teknoloji','yarı iletken','çip','nükleer']): score+=9
        if any(x in title_summary for x in ['üretim durdu','fabrika kapandı','yangın','patlama','siber saldırı','ambargo','yaptırım']): score+=12

        change_kind=change_map.get(r['title'],'')
        if change_kind:
            score += 14 if 'YENİ OLAY' in change_kind else 12

        score=min(100,score)
        rows.append({
            'Aday Puanı':score,
            'Başlık':r['title'],
            'Kaynak':r['source'],
            'Kategori':r['category'],
            'Risk':risk,
            'Kaynak Sayısı':r['source_count'],
            'Doğrulama':r['verification'],
            'Değişim':change_kind or '—',
            'Neden Bilgi Notu?':_note_candidate_reason(r,change_kind),
            'URL':r['url']
        })

    out=pd.DataFrame(rows).sort_values(['Aday Puanı','Risk'],ascending=[False,False]).head(limit)
    return out.reset_index(drop=True)


# -----------------------------
# V34 — VARDİYA BAŞLANGIÇ ÖZETİ + ÖNEMLİ GELİŞMELER SEPETİ
# V33 çekirdeğine dokunmaz.
# -----------------------------
def _mark_shift_handover(scan_id=None, label='Devir noktası'):
    if not _init_history_db():
        return False
    try:
        now=datetime.now().astimezone().isoformat()
        with _history_connect() as conn:
            conn.execute(
                "INSERT INTO shift_marks(marked_at,scan_id,label) VALUES(?,?,?)",
                (now,int(scan_id) if scan_id else None,label)
            )
            conn.commit()
        return True
    except Exception:
        return False

def _latest_shift_mark():
    if not _init_history_db():
        return None
    try:
        with _history_connect() as conn:
            row=conn.execute(
                "SELECT marked_at,scan_id,label FROM shift_marks ORDER BY marked_at DESC LIMIT 1"
            ).fetchone()
        return {'marked_at':row[0],'scan_id':row[1],'label':row[2]} if row else None
    except Exception:
        return None

def _shift_baseline(current_scan_id=None):
    """
    Öncelik manuel devir noktasıdır.
    Hiç devir noktası yoksa V33'ün önceki taramasını baseline olarak kullanır.
    """
    mark=_latest_shift_mark()
    if mark:
        try:
            return pd.to_datetime(mark['marked_at'],utc=True),f"Devir noktası: {mark['marked_at']}",mark.get('scan_id')
        except Exception:
            pass

    prev_id=_previous_scan_id(current_scan_id)
    prev=_load_scan_events(prev_id)
    if not prev.empty:
        try:
            ts=pd.to_datetime(str(prev.iloc[0].get('scanned_at','')),utc=True)
            return ts,f"Önceki tarama: {prev.iloc[0].get('scanned_at','')}",prev_id
        except Exception:
            pass
    return None,"Henüz devir noktası yok",None

def _shift_start_summary(df,current_scan_id=None):
    """
    Son devir noktasından bu yana:
    - yeni haber
    - yeni önemli olay
    - yüksek riskli gelişme
    - risk artışı
    - teyit güçlenmesi
    - OSB olayı
    - sabah ilk bakılması gereken 5 gelişme
    """
    if df is None or df.empty:
        return {},pd.DataFrame(),""

    baseline,baseline_label,baseline_scan_id=_shift_baseline(current_scan_id)
    x=df.copy()
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')

    if baseline is not None:
        since=x[(x['Tarih_dt'].isna()) | (x['Tarih_dt']>=baseline)].copy()
    else:
        since=x.copy()

    changes,_,_=_compare_since_previous(df,current_scan_id)
    if not changes.empty:
        new_events=int(changes['Tür'].astype(str).str.contains('YENİ OLAY').sum())
        risk_up=int(changes['Tür'].astype(str).str.contains('RİSK ARTTI').sum())
        verify_up=int(changes['Tür'].astype(str).str.contains('TEYİT').sum())
    else:
        new_events=0; risk_up=0; verify_up=0

    high=int((since.get('Risk_Durumu',pd.Series(dtype=str))=='Yüksek Risk').sum()) if not since.empty else 0
    osb=0
    for _,r in since.iterrows():
        if is_osb_fire(r.get('Başlık',''),r.get('İçerik_Özeti','')):
            osb+=1

    top=_daily_top_events(since,5) if not since.empty else pd.DataFrame()

    stats={
        'new_news':len(since),
        'new_important_events':new_events,
        'high_risk':high,
        'risk_up':risk_up,
        'verify_up':verify_up,
        'osb':osb,
        'baseline_label':baseline_label
    }
    return stats,top,baseline_label

def _add_rows_to_important_basket(rows):
    rows=_v107_enrich_selected_rows(rows)
    if rows is None or len(rows)==0 or not _init_history_db():
        return 0
    added=0
    try:
        with _history_connect() as conn:
            for row in rows:
                title=str(row.get('Başlık','') or '').strip()
                url=str(row.get('URL','') or '').strip()
                if not title:
                    continue
                cur=conn.execute("""
                    INSERT OR IGNORE INTO important_basket(
                        added_at,news_time,title,source,url,category,risk_score,risk_status,summary
                    ) VALUES(?,?,?,?,?,?,?,?,?)
                """,(
                    datetime.now().astimezone().isoformat(),
                    str(row.get('Tarih','') or ''),
                    title,
                    str(row.get('Kaynak','') or ''),
                    url,
                    str(row.get('Kategori','') or ''),
                    int(row.get('Risk_Skoru',0) or 0),
                    str(row.get('Risk_Durumu','') or ''),
                    str(row.get('İçerik_Özeti','') or '')[:8000]
                ))
                if cur.rowcount:
                    added+=1
            conn.commit()
        if added:
            _v73_invalidate_status_cache()
        return added
    except Exception:
        return 0

def _load_important_basket():
    if not _init_history_db():
        return pd.DataFrame()
    try:
        with _history_connect() as conn:
            return pd.read_sql_query(
                "SELECT * FROM important_basket ORDER BY added_at ASC,id ASC",
                conn
            )
    except Exception:
        return pd.DataFrame()


def _add_rows_to_osint_basket(rows):
    rows=_v107_enrich_selected_rows(rows)
    if rows is None or len(rows)==0 or not _init_history_db():
        return 0
    added=0
    try:
        with _history_connect() as conn:
            for row in rows:
                title=str(row.get('Başlık','') or '').strip()
                url=str(row.get('URL','') or '').strip()
                if not title:
                    continue
                cur=conn.execute("""
                    INSERT OR IGNORE INTO osint_report_basket(
                        added_at,news_time,title,source,url,category,risk_score,risk_status,summary
                    ) VALUES(?,?,?,?,?,?,?,?,?)
                """,(
                    datetime.now().astimezone().isoformat(),
                    str(row.get('Tarih','') or ''),
                    title,
                    str(row.get('Kaynak','') or ''),
                    url,
                    str(row.get('Kategori','') or ''),
                    int(row.get('Risk_Skoru',0) or 0),
                    str(row.get('Risk_Durumu','') or ''),
                    str(row.get('İçerik_Özeti','') or '')[:8000]
                ))
                if cur.rowcount:
                    added+=1
            conn.commit()
        if added:
            _v73_invalidate_status_cache()
        return added
    except Exception:
        return 0

def _load_osint_basket():
    if not _init_history_db():
        return pd.DataFrame()
    try:
        with _history_connect() as conn:
            return pd.read_sql_query(
                "SELECT * FROM osint_report_basket ORDER BY added_at ASC,id ASC",
                conn
            )
    except Exception:
        return pd.DataFrame()

def _remove_osint_basket_ids(ids):
    ids=[int(x) for x in ids if str(x).isdigit()]
    if not ids:
        return 0
    try:
        with _history_connect() as conn:
            q="DELETE FROM osint_report_basket WHERE id IN (" + ",".join("?" for _ in ids) + ")"
            cur=conn.execute(q,ids)
            conn.commit()
            if cur.rowcount:
                _v73_invalidate_status_cache()
            return cur.rowcount
    except Exception:
        return 0

def _clear_osint_basket():
    try:
        with _history_connect() as conn:
            cur=conn.execute("DELETE FROM osint_report_basket")
            conn.commit()
            if cur.rowcount:
                _v73_invalidate_status_cache()
            return cur.rowcount
    except Exception:
        return 0

def _remove_basket_ids(ids):
    ids=[int(x) for x in ids if str(x).isdigit()]
    if not ids:
        return 0
    try:
        with _history_connect() as conn:
            q="DELETE FROM important_basket WHERE id IN (" + ",".join("?" for _ in ids) + ")"
            cur=conn.execute(q,ids)
            conn.commit()
            return cur.rowcount
    except Exception:
        return 0

def _clear_important_basket():
    try:
        with _history_connect() as conn:
            cur=conn.execute("DELETE FROM important_basket")
            conn.commit()
            if cur.rowcount:
                _v73_invalidate_status_cache()
            return cur.rowcount
    except Exception:
        return 0

def _v81_sentence_case_title(title):
    t=_clean_note_text(title).strip()
    letters=''.join(c for c in t if c.isalpha())
    if letters and sum(c.isupper() for c in letters)/max(1,len(letters))>.80:
        t=t.lower()
        t=t[:1].upper()+t[1:]
    return t

def _v84_hard_repair_text(text):
    """
    V84: Türkçe olmayan/mojibake karakterleri agresif biçimde temizler.
    Tam onarılamayan bozuk cümleler ÖGN özetine hiç alınmaz.
    """
    t=_clean_note_text(text)

    # Ek yaygın bozulmalar.
    fixes={
        'TÃ¼rkiye':'Türkiye','TÃ¼rk':'Türk','genÃ§':'genç','dÃ¼nya':'dünya',
        'Ã¼lke':'ülke','Ã¼stÃ¼n':'üstün','Ã¶ÄŸrenci':'öğrenci','Ã¶Ärenci':'öğrenci',
        'baÅŸar':'başar','katÄ±lÄ±m':'katılım','mÃ¼cadele':'mücadele',
        'saÄŸladÄ±ÄŸÄ±':'sağladığı','saÄladÄÄ±ÄÄ±':'sağladığı',
        'ettiÄŸi':'ettiği','ettiÄi':'ettiği','TÃ¼rkiyenin':"Türkiye'nin",
        'TÃ¼rkiyeyi':"Türkiye'yi",'Ã§':'ç','ÄŸ':'ğ','Ä±':'ı',
        'Ã¶':'ö','Ã¼':'ü','ÅŸ':'ş','Ã‡':'Ç','Äž':'Ğ','Ä°':'İ','Ã–':'Ö','Ãœ':'Ü','Åž':'Ş'
    }
    for a,b in fixes.items():
        t=t.replace(a,b)

    # Kalan açık mojibake işaretleri varsa cümle güvenilmez kabul edilir.
    return _clean_note_text(t)

def _v84_sentence_is_clean(s):
    bad=('Ã','Ä','Å','Â',' ','\ufffd','','',' ')
    return not any(x in s for x in bad)

def _v84_clean_article_sentences(text):
    """Haber gövdesinden yalnız güvenilir, tam ve kurumsal özetlemeye uygun cümleleri alır."""
    text=_v84_hard_repair_text(text)
    garbage=[
        'çerez','cookie','reklam','devamını oku','tıklayın','anasayfa','son dakika',
        'benzer haber','ilgili haber','foto galeri','video galeri','sıralamayı değiştirmek',
        'kartları yukarı','abone ol','bildirimleri aç','google news','whatsapp kanal',
        'instagram','facebook','twitter','ekonomi gazetesi »','doğru şarj alışkanlıklarını',
        'haberler (','bugün kocaeli gazetesi','açıklaması şöyle','şunları kaydetti',
        'şöyle konuştu','şöyle dedi'
    ]
    out=[]; seen=set()
    for s in _sentence_split_tr(text):
        s=_v84_hard_repair_text(s).strip(" ;:-[]'\"")
        ns=norm(s)
        if not _v84_sentence_is_clean(s):
            continue
        if len(s)<38 or len(s)>480 or any(g in ns for g in garbage):
            continue
        if s.endswith(('…','...')) or re.search(r'\bve k$',s,re.I):
            continue
        # Haber ortasından alınmış doğrudan konuşma/alıntı ile başlama.
        if s.startswith(('"','“',"'",'‘')) or re.match(r'^\d+\s',s):
            continue
        letters=''.join(c for c in s if c.isalpha())
        if letters and len(s)<135 and sum(c.isupper() for c in letters)/max(1,len(letters))>.76:
            continue
        k=title_key(s)
        if not k or k in seen:
            continue
        seen.add(k); out.append(s)
    return out

def _v84_formalize(s):
    """Yalnız cümle sonunu değil, yaygın haber dili kalıntılarını da resmîleştirir."""
    s=_v84_hard_repair_text(s).strip()
    replacements=[
        (r'\bifade etti\b','ifade etmiştir'),(r'\bifade ediyor\b','ifade etmektedir'),
        (r'\bbelirtti\b','belirtmiştir'),(r'\bbelirtiyor\b','belirtmektedir'),
        (r'\baçıkladı\b','açıklamıştır'),(r'\baçıklıyor\b','açıklamaktadır'),
        (r'\bduyurdu\b','duyurmuştur'),(r'\bduyuruyor\b','duyurmaktadır'),
        (r'\bgösterdi\b','göstermiştir'),(r'\bgösteriyor\b','göstermektedir'),
        (r'\bsağladı\b','sağlamıştır'),(r'\bsağlıyor\b','sağlamaktadır'),
        (r'\bhedefliyor\b','hedeflemektedir'),(r'\bplanlıyor\b','planlamaktadır'),
        (r'\bbaşladı\b','başlamıştır'),(r'\bbaşlıyor\b','başlamaktadır'),
        (r'\btamamladı\b','tamamlamıştır'),(r'\btamamladı\b','tamamlamıştır'),
        (r'\bkazandı\b','kazanmıştır'),(r'\bgerçekleşti\b','gerçekleşmiştir'),
        (r'\byükseldi\b','yükselmiştir'),(r'\bgeriledi\b','gerilemiştir'),
        (r'\barttı\b','artmıştır'),(r'\bazaldı\b','azalmıştır'),
        (r'\boldu\b','olmuştur'),(r'\bolacak\b','olacaktır'),
        (r'\byapılacak\b','yapılacaktır'),(r'\bsağlanacak\b','sağlanacaktır'),
        (r'\bbaşlayacak\b','başlayacaktır'),(r'\byer alacak\b','yer alacaktır'),
        (r'\bmücadele edecek\b','mücadele edecektir')
    ]
    for pat,val in replacements:
        s=re.sub(pat,val,s,flags=re.I)
    s=_v66_formalize_sentence_endings(s)
    s=re.sub(r'\bifade ettiği ifade etmiştir\b','ifade etmiştir',s,flags=re.I)
    s=re.sub(r'\bbelirttiği belirtmiştir\b','belirtmiştir',s,flags=re.I)
    return _v84_hard_repair_text(s)

def _v84_score_intro(s):
    ns=norm(s)
    actor=['cumhurbaşkan','bakan','bakanlık','tüik','tübitak','tcmb','tse','türkpatent',
           'ssb','valili','başkan','şirket','üniversite','nasa','ibm','türk telekom',
           'kardemir','togg','gezeravcı','zeytinoğlu']
    action=['açıklad','duyur','başlat','gerçekleştir','tamamla','imzala','yayımla',
            'düzenlen','üret','geliştir','test','ziyaret','göreve','başvuru','yatırım']
    place=['ankara','istanbul','kocaeli','antalya','amasya','astana','pekin','gölcük',
           'türkiye','abd','çin','kazakistan','avustralya','almanya','isveç']
    return 4*sum(x in ns for x in actor)+4*sum(x in ns for x in action)+sum(x in ns for x in place)+min(len(re.findall(r'\d',s)),2)

def _v84_score_detail(s):
    ns=norm(s)
    data=['%','yüzde','milyon','milyar','bin ','adet','mw','gwh','mwh','km','puan',
          'oran','endeks','kapasite','ciro','ihracat','üretim','satış','başvuru','rekor']
    return 4*sum(x in ns for x in data)+min(len(re.findall(r'\d',s)),5)

def _v84_score_result(s):
    ns=norm(s)
    result=['art','azal','gerile','yüksel','ulaş','hedef','plan','beklen','sağla',
            'kazandır','devreye','pilot','kullanıl','rekor','destek','katkı','başarı']
    return 4*sum(x in ns for x in result)+min(len(re.findall(r'\d',s)),3)

def _v80_reference_important_summary(title,summary,full_text=''):
    """
    V84: Önce düzgün bir giriş cümlesi, sonra kritik rakam/detay, sonra sonuç/önem.
    2-3 tam cümle; cümle ortasında kesme yok; Word'de yaklaşık 4 satır hedefi.
    """
    title=_v81_sentence_case_title(_v84_hard_repair_text(title))
    body=_v84_hard_repair_text(full_text or summary)
    good=_v84_clean_article_sentences(body)

    if len(good)<2:
        good=_v84_clean_article_sentences(str(summary)+' '+str(full_text))
    if not good:
        return _v84_formalize(title)

    # Giriş asla haberin ortasından başlamasın: aktör + eylem taşıyan cümleyi seç.
    intro_candidates=good[:12]
    intro=max(intro_candidates,key=lambda s:(_v84_score_intro(s),-good.index(s)))
    if _v84_score_intro(intro)<4:
        # Güçlü giriş bulunamazsa ilk temiz cümleyi kullan.
        intro=good[0]

    chosen=[intro]

    rem=[s for s in good if s not in chosen]
    if rem:
        detail=max(rem,key=lambda s:(_v84_score_detail(s),-good.index(s)))
        if _v84_score_detail(detail)>0:
            chosen.append(detail)

    rem=[s for s in good if s not in chosen]
    if rem:
        result=max(rem,key=lambda s:(_v84_score_result(s),-good.index(s)))
        if _v84_score_result(result)>0:
            chosen.append(result)

    # En az iki cümle olsun.
    if len(chosen)<2:
        for s in good:
            if s not in chosen:
                chosen.append(s)
                break

    chosen=sorted(chosen,key=lambda s:good.index(s))
    formal=[_v84_formalize(s) for s in chosen[:3] if _v84_sentence_is_clean(_v84_formalize(s))]
    text=_clean_note_text(' '.join(formal))

    # Çok uzun cümleler nedeniyle 4 satırı aşmaması için sıkı sınır:
    # 2 veya 3 TAM cümle, yaklaşık 500 karakter.
    sents=_sentence_split_tr(text)
    kept=[]; total=0
    for sent in sents:
        add=len(sent)+(1 if kept else 0)
        if kept and total+add>500:
            break
        kept.append(sent); total+=add
        if len(kept)>=3:
            break

    # Eğer ilk cümle tek başına çok uzunsa, güvenli cümle sınırında sıkıştır.
    if kept and len(' '.join(kept))>520:
        kept=kept[:2]

    result=' '.join(kept).strip()

    # Son güvenlik: bozuk yabancı karakter kalırsa o cümleyi düşür.
    final_sents=[s for s in _sentence_split_tr(result) if _v84_sentence_is_clean(s)]
    return ' '.join(final_sents[:3]).strip()


def _v87_safe_tr(text):
    """Only obvious mojibake repair; never drop the whole item."""
    t=_clean_note_text(text)
    fixes={
        'TÃ¼rkiye':'Türkiye','TÃ¼rk':'Türk','genÃ§':'genç','dÃ¼nya':'dünya',
        'Ã¼lke':'ülke','Ã¼stÃ¼n':'üstün','Ã¶ÄŸrenci':'öğrenci','Ã¶Ärenci':'öğrenci',
        'baÅŸar':'başar','katÄ±lÄ±m':'katılım','mÃ¼cadele':'mücadele',
        'Ä±':'ı','ÄŸ':'ğ','ÅŸ':'ş','Ã§':'ç','Ã¶':'ö','Ã¼':'ü',
        'Ä°':'İ','Äž':'Ğ','Åž':'Ş','Ã‡':'Ç','Ã–':'Ö','Ãœ':'Ü',
        'Â':'','â€™':'’','â€œ':'“','â€':'”','â€“':'–','â€”':'—'
    }
    for a,b in fixes.items():
        t=t.replace(a,b)
    return re.sub(r'\s+',' ',t).strip()


@st.cache_data(ttl=3600,show_spinner=False)
def _v88_cached_article_detail(title,source,url,fallback,news_time):
    """Same article is not fetched again for one hour."""
    try:
        return article_detail({
            'Başlık':title,
            'Kaynak':source,
            'URL':url,
            'Yayıncı_URL':url,
            'İçerik_Özeti':fallback,
            'Tarih':news_time
        })
    except Exception:
        return {
            'title':title,'source':source,'canonical':url,
            'published':news_time,'text':fallback,'images':[]
        }

def _v88_title_core(title,source=''):
    """Remove publisher suffixes and headline clutter."""
    t=_v87_safe_tr(title)
    source=_v87_safe_tr(source)
    # Common Google News/source suffix.
    if source:
        t=re.sub(r'\s*[-–—]\s*'+re.escape(source)+r'\s*$','',t,flags=re.I)
    t=re.sub(r'\s*[-–—]\s*(Haberler|Haber|Son Dakika|Gündem)\s*$','',t,flags=re.I)
    t=re.sub(r'\s+',' ',t).strip(' -–—|')
    return t

def _v88_sentence_bad(s):
    s=_v87_safe_tr(s)
    bad_chars=('Ã','Ä','Å','Â',' ',' ','','')
    if any(x in s for x in bad_chars):
        return True
    n=norm(s)
    noise=[
        'sıralamayı değiştirmek','kartları yukarı','tüvtürk en sık',
        'samsung sevilen modelin','benzer haber','ilgili haber',
        'devamını oku','çerez','cookie','reklam','foto galeri','video galeri',
        'ekonomi gazetesi »','araç sahipleri dikkat'
    ]
    if any(x in n for x in noise):
        return True
    if s.endswith(('…','...')) or re.search(r'\bve k$',s,re.I):
        return True
    return False

def _v88_clean_sentences(text):
    out=[]; seen=set()
    for s in _sentence_chunks(_v87_safe_tr(text)):
        s=_v87_safe_tr(s).strip(" []'\";-:")
        if len(s)<35 or len(s)>430 or _v88_sentence_bad(s):
            continue
        k=title_key(s)
        if not k or k in seen:
            continue
        seen.add(k); out.append(s)
    return out

def _v88_keywords(title):
    stop={'haber','haberi','son','dakika','bugün','yeni','ile','ve','bir','için','olan','oldu',
          'olacak','dedi','açıkladı','duyurdu','türkiye','türk'}
    words=[w for w in re.findall(r'[a-zçğıöşü0-9]+',norm(title)) if len(w)>=4 and w not in stop]
    return set(words[:12])

def _v88_formal(s):
    s=_v87_safe_tr(s)
    pairs=[
        (r'\baçıkladı\b','açıklamıştır'),(r'\bbelirtti\b','belirtmiştir'),
        (r'\bduyurdu\b','duyurmuştur'),(r'\bkaydetti\b','kaydetmiştir'),
        (r'\bifade etti\b','ifade etmiştir'),(r'\bbaşladı\b','başlamıştır'),
        (r'\btamamladı\b','tamamlamıştır'),(r'\bkazandı\b','kazanmıştır'),
        (r'\barttı\b','artmıştır'),(r'\bazaldı\b','azalmıştır'),
        (r'\bgeriledi\b','gerilemiştir'),(r'\byükseldi\b','yükselmiştir'),
        (r'\bulaştı\b','ulaşmıştır'),(r'\bgerçekleşti\b','gerçekleşmiştir'),
        (r'\boldu\b','olmuştur'),(r'\byer alacak\b','yer alacaktır'),
        (r'\bbaşlayacak\b','başlayacaktır'),(r'\bsağlanacak\b','sağlanacaktır'),
        (r'\bverilecek\b','verilecektir'),(r'\bseçilecek\b','seçilecektir'),
        (r'\bkazandırılacak\b','kazandırılacaktır'),(r'\bdevam ediyor\b','devam etmektedir'),
        (r'\bgösteriyor\b','göstermektedir'),(r'\bsağlıyor\b','sağlamaktadır'),
        (r'\bdikkat çekiyor\b','dikkat çekmektedir')
    ]
    for pat,val in pairs:
        s=re.sub(pat,val,s,flags=re.I)
    s=_v66_formalize_sentence_endings(s)
    s=_v87_safe_tr(s).strip()
    if s:
        s=s[0].upper()+s[1:]
    return s

def _v89_normalize_source_title(title,source):
    """Başlıktaki yayıncı/portal eklerini temizler; başlığı çıktı olarak kullanmaz."""
    t=_v88_title_core(title,source)
    t=re.sub(r'\s+[A-Za-zÇĞİÖŞÜçğıöşü0-9_.-]+\.(?:com|com\.tr|net|org|tr)\s*$','',t,flags=re.I)
    return _v87_safe_tr(t).strip(' -–—|')

def _v89_clause_from_sentence(s):
    """
    İkinci bir haber cümlesini ana resmî cümleye eklenebilir bilgi cümleciğine çevirir.
    Tam cümleyi parçalamaz; yalnız son noktayı kaldırır.
    """
    s=_v88_formal(_v87_safe_tr(s)).strip()
    return s.rstrip(' .;:')

def _v89_single_official_sentence(title,source,body,fallback):
    """
    Gerçek STB örneği mantığı:
    Her gelişme için TEK, TAM ve RESMÎ cümle.
    - kim/kurum + ne oldu ana cümlesi,
    - en kritik rakam/yer/tarih aynı cümlede,
    - gerekiyorsa sonuç/hedef ikinci cümlecik olarak noktalı virgülle bağlanır,
    - haber başlığı tek başına çıktı olmaz.
    """
    title=_v89_normalize_source_title(title,source)
    text=_v87_safe_tr(body or fallback)
    sents=_v88_clean_sentences(text)
    if len(sents)<2:
        sents=_v88_clean_sentences(fallback)

    if not sents:
        # Son çare: kaydedilmiş özet varsa onu kullan; sırf başlığı basma.
        fb=_v87_safe_tr(fallback)
        if len(fb)>=60:
            return _v88_formal(fb).rstrip(' .;')+'.'
        return ''

    keywords=_v88_keywords(title)
    actor_terms=['cumhurbaşkan','bakan','bakanlık','başkan','tüik','tübitak','tcmb','tse',
                 'türkpatent','ssb','valili','üniversite','şirket','genel müdür','türk telekom',
                 'kardemir','togg','aselsan','roketsan','gezeravcı','zeytinoğlu','kurum','takım']
    action_terms=['açıkla','duyur','başlat','gerçekleştir','tamamla','imzala','kazan','yatırım',
                  'test','görev','üret','satış','başvuru','düzenlen','ulaş','art','azal','gerile']
    detail_terms=['%','yüzde','milyon','milyar','bin ','adet','mw','gwh','mwh','km','puan',
                  'kapasite','ihracat','üretim','satış','hibe','öğrenci','madalya','rekor','tarih']
    result_terms=['hedef','beklen','sağla','katkı','devreye','plan','başarı','destek','başvuru',
                  'artış','azalış','yüksel','gerile','ulaş']

    def overlap(sent):
        ws=set(re.findall(r'[a-zçğıöşü0-9]+',norm(sent)))
        return len(keywords & ws)

    # Konuyla ilişkisiz "Samsung / TÜVTÜRK / başka haber" parçalarını devreden çıkar.
    related=[x for x in sents if overlap(x)>0]
    pool=related if related else sents[:8]

    def intro_score(x):
        n=norm(x)
        return 7*overlap(x)+4*sum(k in n for k in actor_terms)+4*sum(k in n for k in action_terms)

    intro=max(pool[:8],key=lambda x:(intro_score(x),-sents.index(x)))
    intro_formal=_v89_clause_from_sentence(intro)

    # Başlıkla neredeyse aynıysa, başka giriş ara.
    if title_key(intro_formal)==title_key(title):
        alternatives=[x for x in pool if title_key(x)!=title_key(title)]
        if alternatives:
            intro=max(alternatives,key=lambda x:(intro_score(x),-sents.index(x)))
            intro_formal=_v89_clause_from_sentence(intro)

    rem=[x for x in pool if x!=intro]
    detail=None
    if rem:
        def detail_score(x):
            n=norm(x)
            return 6*overlap(x)+5*sum(k in n for k in detail_terms)+min(len(re.findall(r'\d',x)),6)
        cand=max(rem,key=lambda x:(detail_score(x),-sents.index(x)))
        if detail_score(cand)>0:
            detail=cand

    rem=[x for x in rem if x!=detail]
    result=None
    if rem:
        def result_score(x):
            n=norm(x)
            return 5*overlap(x)+4*sum(k in n for k in result_terms)+2*sum(k in n for k in detail_terms)
        cand=max(rem,key=lambda x:(result_score(x),-sents.index(x)))
        if result_score(cand)>0:
            result=cand

    # Ana cümle doğal biçimde zaten gerekli rakamları içeriyorsa gereksiz tekrar ekleme.
    clauses=[intro_formal]
    intro_digits=set(re.findall(r'\d+(?:[.,]\d+)?',intro_formal))

    for extra in [detail,result]:
        if not extra:
            continue
        ef=_v89_clause_from_sentence(extra)
        if not ef or _v88_sentence_bad(ef):
            continue
        # Aynı olayı/rakamı tekrar eden cümleyi alma.
        nums=set(re.findall(r'\d+(?:[.,]\d+)?',ef))
        if nums and nums.issubset(intro_digits) and title_key(ef)[:80] in title_key(intro_formal):
            continue
        # Başlıkla ilişki şartı: unrelated site snippets cannot enter.
        if keywords and overlap(ef)==0:
            continue
        clauses.append(ef)
        if len(clauses)>=2:  # tek cümlede iki ana bilgi bloğu yeterli
            break

    # Tek resmî cümle: ilk tam cümle + ikinci bilgi bloğu noktalı virgülle.
    if len(clauses)==1:
        out=clauses[0]
    else:
        second=clauses[1]
        # İkinci bloğu küçük harfle doğal bağla; özel isimleri bozma.
        connector='; ayrıca, '
        out=clauses[0]+connector+second

    out=_v87_safe_tr(out).strip(' ;:.')
    # 4 satır hedefi: cümleyi kesmeden 500 karaktere yaklaş.
    if len(out)>500 and len(clauses)>1:
        out=clauses[0].strip(' ;:.')
    if len(out)>520:
        # Çok uzun tek giriş varsa noktalı virgül/virgül sınırından kısalt.
        cut=out[:520]
        candidates=[cut.rfind('; '),cut.rfind(', ')]
        k=max(candidates)
        if k>=300:
            out=cut[:k].rstrip(' ,;')
    return _v87_safe_tr(out)+'.'

# Keep name used by make_important_basket_docx, but route to V89.
def _v88_summary(title,source,body,fallback):
    return _v89_single_official_sentence(title,source,body,fallback)

def _v87_ogn_summary(title, body, fallback):
    """
    Simple, deterministic recovery summarizer:
    - never returns blank when fallback exists,
    - uses stable _akt_formal_summary,
    - keeps 2-3 complete sentences where available,
    - no extra web search and no experimental sentence dropping.
    """
    title=_v87_safe_tr(title)
    body=_v87_safe_tr(body or fallback or title)
    fallback=_v87_safe_tr(fallback)

    try:
        text=_akt_formal_summary(title,body,max_sentences=3,max_chars=700)
    except Exception:
        text=fallback or body or title

    text=_v87_safe_tr(text)
    if not text or title_key(text)==title_key(title):
        text=fallback if len(fallback)>=60 else (body if len(body)>=60 else title)

    # Formalize endings using existing stable routine.
    text=_v66_formalize_sentence_endings(text)

    # Keep max 3 COMPLETE sentences and roughly 4 Word lines.
    sents=_sentence_chunks(text)
    if not sents:
        return text[:520].strip()

    kept=[]; total=0
    for sent in sents:
        sent=_v87_safe_tr(sent).strip()
        if not sent: continue
        add=len(sent)+(1 if kept else 0)
        if kept and total+add>520:
            break
        kept.append(sent)
        total+=add
        if len(kept)>=3:
            break

    out=' '.join(kept).strip()
    return out or (fallback[:520].strip() if fallback else title[:520].strip())


V90_OGN_ENGINE_VERSION='v104_mevcut_bolum_olgunlastirma_20260822_1'

def _v90_clean_title(title,source=''):
    t=_v87_safe_tr(title)
    s=_v87_safe_tr(source)
    if s:
        t=re.sub(r'\s*[-–—]\s*'+re.escape(s)+r'\s*$','',t,flags=re.I)
    t=re.sub(r'\s*[-–—]\s*(Haberler|Haber|Son Dakika|Gündem)\s*$','',t,flags=re.I)
    t=re.sub(r'\s+',' ',t).strip(' -–—|')
    return t

def _v90_clean_sentence(s):
    s=_v87_safe_tr(s).strip(" []'\";-:")
    # Haber portalı / başka başlık / yarım snippet artıkları.
    noise=[
        'sıralamayı değiştirmek','kartları yukarı','devamını oku','benzer haber',
        'ilgili haber','çerez','cookie','reklam','foto galeri','video galeri',
        'google news','whatsapp','instagram','facebook','twitter',
        'araç sahipleri dikkat','samsung sevilen modelin','tüvtürk en sık',
        'ekonomi gazetesi »'
    ]
    ns=norm(s)
    if any(x in ns for x in noise):
        return ''
    if s.endswith(('…','...')) or re.search(r'\bve k$',s,re.I):
        return ''
    if any(x in s for x in ('Ã','Ä','Å',' ',' ','','')):
        return ''
    return s

def _v90_sentences(text):
    out=[]; seen=set()
    for raw in _sentence_chunks(_v87_safe_tr(text)):
        s=_v90_clean_sentence(raw)
        if len(s)<38 or len(s)>520:
            continue
        k=title_key(s)
        if not k or k in seen:
            continue
        seen.add(k)
        out.append(s)
    return out

def _v90_title_words(title):
    stop={
        'haber','haberi','son','dakika','bugün','yeni','ile','ve','bir','için','olan',
        'oldu','olacak','dedi','açıkladı','duyurdu','türkiye','türk','etti','başladı'
    }
    return set(
        w for w in re.findall(r'[a-zçğıöşü0-9]+',norm(title))
        if len(w)>=4 and w not in stop
    )

def _v90_formalize(s):
    s=_v87_safe_tr(s)
    replacements=[
        (r'\baçıkladı\b','açıklamıştır'),(r'\bbelirtti\b','belirtmiştir'),
        (r'\bduyurdu\b','duyurmuştur'),(r'\bkaydetti\b','kaydetmiştir'),
        (r'\bifade etti\b','ifade etmiştir'),(r'\bbaşladı\b','başlamıştır'),
        (r'\btamamladı\b','tamamlamıştır'),(r'\bkazandı\b','kazanmıştır'),
        (r'\barttı\b','artmıştır'),(r'\bazaldı\b','azalmıştır'),
        (r'\bgeriledi\b','gerilemiştir'),(r'\byükseldi\b','yükselmiştir'),
        (r'\bulaştı\b','ulaşmıştır'),(r'\bgerçekleşti\b','gerçekleşmiştir'),
        (r'\boldu\b','olmuştur'),(r'\bkırıldı\b','kırılmıştır'),
        (r'\byer alacak\b','yer alacaktır'),(r'\bbaşlayacak\b','başlayacaktır'),
        (r'\bsağlanacak\b','sağlanacaktır'),(r'\bverilecek\b','verilecektir'),
        (r'\bseçilecek\b','seçilecektir'),(r'\bkazandırılacak\b','kazandırılacaktır'),
        (r'\bdevam ediyor\b','devam etmektedir'),(r'\bgösteriyor\b','göstermektedir'),
        (r'\bsağlıyor\b','sağlamaktadır'),(r'\bdikkat çekiyor\b','dikkat çekmektedir')
    ]
    for pat,val in replacements:
        s=re.sub(pat,val,s,flags=re.I)
    s=_v66_formalize_sentence_endings(s)
    s=_v87_safe_tr(s).strip()
    if s:
        s=s[0].upper()+s[1:]
    return s

def _v90_item_summary(title,source,body,fallback):
    """
    Kurum örneğine yakın TEK, TAM, RESMÎ cümle:
    ana olay + kritik rakam/yer/tarih/kişi + gerekiyorsa tek tamamlayıcı bilgi.
    Başlık doğrudan Word'e yazılmaz.
    """
    title=_v90_clean_title(title,source)
    text=_v87_safe_tr(body or fallback)
    sents=_v90_sentences(text)
    if len(sents)<2:
        sents=_v90_sentences(fallback)

    # Başlık dışında gerçek içerik yoksa fallback'i kullan; yine başlığı tek başına basma.
    if not sents:
        fb=_v87_safe_tr(fallback)
        if len(fb)>=80 and title_key(fb)!=title_key(title):
            return _v90_formalize(fb).rstrip(' .;')+'.'
        return ''

    tw=_v90_title_words(title)

    actor_terms=[
        'cumhurbaşkan','bakan','bakanlık','başkan','tüik','tübitak','tcmb','tse',
        'türkpatent','ssb','valili','üniversite','şirket','genel müdür','türk telekom',
        'kardemir','togg','aselsan','roketsan','gezeravcı','zeytinoğlu','takım'
    ]
    action_terms=[
        'açıkla','duyur','başlat','gerçekleştir','tamamla','imzala','kazan','yatırım',
        'test','görev','üret','satış','başvuru','düzenlen','ulaş','art','azal','gerile'
    ]
    detail_terms=[
        '%','yüzde','milyon','milyar','bin ','adet','mw','gwh','mwh','km','puan',
        'kapasite','ihracat','üretim','satış','hibe','öğrenci','madalya','rekor',
        '2025','2026','2027'
    ]

    def overlap(s):
        words=set(re.findall(r'[a-zçğıöşü0-9]+',norm(s)))
        return len(words & tw)

    # Yalnız haberle ilişkili cümleleri tercih et.
    related=[s for s in sents if overlap(s)>0]
    pool=related if related else sents[:8]

    def intro_score(s):
        n=norm(s)
        return (
            8*overlap(s)
            + 4*sum(x in n for x in actor_terms)
            + 4*sum(x in n for x in action_terms)
            + min(len(re.findall(r'\d',s)),3)
        )

    # İlk cümle haberin ortasından değil, olayı tanımlayan cümle olsun.
    intro=max(pool[:8],key=lambda s:(intro_score(s),-sents.index(s)))
    intro=_v90_formalize(intro).rstrip(' .;')

    # Eğer intro başlıkla neredeyse aynıysa başka gövde cümlesi dene.
    if title_key(intro)==title_key(title):
        alternatives=[s for s in pool if title_key(s)!=title_key(title)]
        if alternatives:
            intro=_v90_formalize(
                max(alternatives,key=lambda s:(intro_score(s),-sents.index(s)))
            ).rstrip(' .;')

    # En kritik ikinci bilgi: rakam/tarih/ölçek; aynı olayla ilişkili olmak zorunda.
    remaining=[s for s in pool if title_key(_v90_formalize(s))!=title_key(intro)]
    detail=''
    if remaining:
        def detail_score(s):
            n=norm(s)
            return (
                7*overlap(s)
                + 5*sum(x in n for x in detail_terms)
                + min(len(re.findall(r'\d',s)),6)
            )
        cand=max(remaining,key=lambda s:(detail_score(s),-sents.index(s)))
        if detail_score(cand)>=5:
            detail=_v90_formalize(cand).rstrip(' .;')

    # Tek resmî cümle oluştur.
    out=intro
    if detail:
        # Aynı rakamları tekrar eden ayrıntıyı ekleme.
        n1=set(re.findall(r'\d+(?:[.,]\d+)?',intro))
        n2=set(re.findall(r'\d+(?:[.,]\d+)?',detail))
        if not (n2 and n2.issubset(n1) and len(detail)<180):
            out += '; ayrıca, ' + detail[0].lower()+detail[1:] if detail else ''

    out=_v87_safe_tr(out).strip(' ;:.')

    # Örnekteki yoğunluk: yaklaşık 4 Word satırı; cümleyi ortadan kesme.
    if len(out)>500 and detail:
        out=intro
    if len(out)>520:
        # Giriş tek başına çok uzunsa en yakın anlamlı virgül/noktalı virgül sınırında kısalt.
        cut=out[:520]
        k=max(cut.rfind('; '),cut.rfind(', '))
        if k>=330:
            out=cut[:k].rstrip(' ,;')

    return out.rstrip(' .;')+'.'

@st.cache_data(ttl=3600,show_spinner=False)
def _v90_fetch_detail(title,source,url,fallback,news_time):
    try:
        return article_detail({
            'Başlık':title,
            'Kaynak':source,
            'URL':url,
            'Yayıncı_URL':url,
            'İçerik_Özeti':fallback,
            'Tarih':news_time
        })
    except Exception:
        return {'title':title,'source':source,'text':fallback,'canonical':url,'images':[]}


def _v92_clean_news_text(text):
    t=_v87_safe_tr(text)
    # Portal artıkları / başka haber başlıkları / navigasyon.
    noise_patterns=[
        r'sıralamayı değiştirmek[^.!?]*[.!?]?',
        r'kartları yukarı[^.!?]*[.!?]?',
        r'devamını oku[^.!?]*[.!?]?',
        r'benzer haber(?:ler)?[^.!?]*[.!?]?',
        r'ilgili haber(?:ler)?[^.!?]*[.!?]?',
        r'google news[^.!?]*[.!?]?',
        r'whatsapp kanal[^.!?]*[.!?]?',
    ]
    for pat in noise_patterns:
        t=re.sub(pat,' ',t,flags=re.I)
    return re.sub(r'\s+',' ',t).strip()

def _v92_formal_sentence(s):
    """Yalnız haber dili yüklemlerini kurumsal geçmiş/şimdiki zamana çevirir."""
    s=_v92_clean_news_text(s).strip()
    pairs=[
        (r'\baçıkladı\b','açıklamıştır'),
        (r'\bbelirtti\b','belirtmiştir'),
        (r'\bduyurdu\b','duyurmuştur'),
        (r'\bkaydetti\b','kaydetmiştir'),
        (r'\bifade etti\b','ifade etmiştir'),
        (r'\bbaşladı\b','başlamıştır'),
        (r'\btamamladı\b','tamamlamıştır'),
        (r'\bkazandı\b','kazanmıştır'),
        (r'\barttı\b','artmıştır'),
        (r'\bazaldı\b','azalmıştır'),
        (r'\bgeriledi\b','gerilemiştir'),
        (r'\byükseldi\b','yükselmiştir'),
        (r'\bulaştı\b','ulaşmıştır'),
        (r'\bgerçekleşti\b','gerçekleşmiştir'),
        (r'\boldu\b','olmuştur'),
        (r'\bkırıldı\b','kırılmıştır'),
        (r'\byayımlandı\b','yayımlanmıştır'),
        (r'\byayınlandı\b','yayımlanmıştır'),
        (r'\bbaşlıyor\b','başlamaktadır'),
        (r'\bdevam ediyor\b','devam etmektedir'),
        (r'\bsağlıyor\b','sağlamaktadır'),
        (r'\bgösteriyor\b','göstermektedir'),
        (r'\bhedefleniyor\b','hedeflenmektedir'),
        (r'\bplanlanıyor\b','planlanmaktadır'),
    ]
    for pat,val in pairs:
        s=re.sub(pat,val,s,flags=re.I)
    return _v66_formalize_sentence_endings(s).strip()

def _v92_good_sentence(s):
    s=_v92_clean_news_text(s)
    if len(s)<35 or len(s)>650:
        return False
    if any(x in s for x in ('Ã','Ä','Å',' ',' ','','')):
        return False
    n=norm(s)
    noise=[
        'çerez','cookie','reklam','foto galeri','video galeri',
        'instagram','facebook','twitter','abone ol','bildirimleri aç',
        'sıralamayı değiştirmek','kartları yukarı','devamını oku',
        'benzer haber','ilgili haber'
    ]
    return not any(x in n for x in noise)

def _v92_summary(title, source, body, fallback):
    """
    STB örneğindeki mantık:
    1) Haberin ana gelişmesini veren giriş cümlesi.
    2) Hemen ardından gelen kritik veri/rakam/açıklama.
    3) Gerekliyse üçüncü ardışık cümle.
    Farklı yerlerden cümle toplayıp yapıştırmaz.
    """
    title=_v90_clean_title(title,source)
    raw=_v92_clean_news_text(body or fallback)

    sents=[]
    seen=set()
    for s in _sentence_chunks(raw):
        s=_v92_clean_news_text(s).strip(" []'\";-:")
        if not _v92_good_sentence(s):
            continue
        k=title_key(s)
        if not k or k in seen:
            continue
        seen.add(k)
        sents.append(s)

    if not sents:
        fb=_v92_clean_news_text(fallback)
        return _v92_formal_sentence(fb) if len(fb)>=50 else ''

    # Başlıktaki ayırt edici sözcükler.
    stop={'haber','haberi','son','dakika','bugün','yeni','ile','ve','bir','için',
          'olan','oldu','olacak','dedi','türkiye','türk'}
    title_words={
        w for w in re.findall(r'[a-zçğıöşü0-9]+',norm(title))
        if len(w)>=4 and w not in stop
    }

    def overlap(s):
        sw=set(re.findall(r'[a-zçğıöşü0-9]+',norm(s)))
        return len(sw & title_words)

    # Giriş: ilk 6 temiz cümle içinde başlıkla en ilişkili olanı bul.
    # Böylece sayfanın ortasından başlamaz; ama başlık tekrarına da mahkûm olmaz.
    first_pool=sents[:6]
    start=max(range(len(first_pool)), key=lambda i:(overlap(first_pool[i]), -i))

    # Eğer ilk cümle zaten makul derecede ilgiliyse onu tercih et.
    if overlap(first_pool[0])>0:
        start=0

    # Ardışık cümleler: bağlam korunur. Maksimum 3 cümle.
    chosen=[]
    total=0
    for s in sents[start:start+3]:
        fs=_v92_formal_sentence(s).strip()
        if not fs:
            continue
        if fs[-1] not in '.!?':
            fs+='.'
        # Yaklaşık 4 satır; cümleyi ortadan kesme.
        add=len(fs)+(1 if chosen else 0)
        if chosen and total+add>560:
            break
        chosen.append(fs)
        total+=add

    # Tek cümle çok kısa kaldıysa bir sonraki ilgili tam cümleyi ekle.
    if len(' '.join(chosen))<220:
        for s in sents[start+len(chosen):]:
            if overlap(s)<=0 and title_words:
                continue
            fs=_v92_formal_sentence(s).strip()
            if not fs:
                continue
            if fs[-1] not in '.!?':
                fs+='.'
            if len(' '.join(chosen+[fs]))<=560:
                chosen.append(fs)
            break

    out=' '.join(chosen).strip()

    # Başlığı tek başına "özet" kabul etme.
    if title_key(out)==title_key(title):
        fb=_v92_clean_news_text(fallback)
        if len(fb)>=80 and title_key(fb)!=title_key(title):
            out=_v92_formal_sentence(fb)

    return _v87_safe_tr(out).strip()


def _v93_sentence_split(text):
    t=_v92_clean_news_text(text)
    # Noktalı virgülü cümle gibi bölme; kurumsal örneklerde bağlı bilgi olabilir.
    return [x.strip() for x in re.split(r'(?<=[.!?])\s+',t) if x.strip()]

def _v93_is_noise(s):
    n=norm(s)
    bad=[
        'sıralamayı değiştirmek','kartları yukarı','benzer haber','ilgili haber',
        'devamını oku','son dakika','reklam','çerez','cookie','instagram',
        'facebook','twitter','whatsapp','youtube','google news',
        'tıklayın','abone ol','yorumlar','etiketler'
    ]
    return len(s)<30 or any(x in n for x in bad) or any(x in s for x in ('Ã','Ä','Å',' '))

def _v93_content_score(s, title_words):
    n=norm(s)
    words=set(re.findall(r'[a-zçğıöşü0-9]+',n))
    score=0
    score += 5*len(words & title_words)
    # Örnek notlarda öne çıkan unsurlar: kurum/kişi, tarih, rakam, oran, yer, karar/eylem.
    if re.search(r'\b\d+(?:[.,]\d+)?\b|%|yüzde|milyon|milyar|bin\b',n): score+=5
    if re.search(r'\b20\d{2}\b|\b\d{1,2}\s+(?:ocak|şubat|mart|nisan|mayıs|haziran|temmuz|ağustos|eylül|ekim|kasım|aralık)\b',n): score+=3
    if any(x in n for x in ['açıklad','duyur','yayımla','başlat','tamamla','imzala','seçilecek',
                             'gerçekleştir','üret','ihrac','yatırım','satın al','devreye al',
                             'denize indir','entegre','belirlen','güncellen','başvuru']): score+=4
    if any(x in n for x in ['bakanlık','tübitak','tüik','aselsan','roketsan','tusaş','havelsan',
                             'cumhurbaşkan','başkanlığı','ajansı','üniversitesi','şirketi','ofisi']): score+=3
    return score

def _v93_build_summary(title, source, body, fallback):
    """
    Tek haber -> tek olay anlatısı.
    Cümleler haberin başından/ortasından rastgele seçilmez:
    ana olay cümlesi bulunur ve yalnız onun çevresindeki aynı bağlam kullanılır.
    """
    title=_v90_clean_title(title,source)
    text=_v92_clean_news_text(body or fallback)
    sentences=[s for s in _v93_sentence_split(text) if not _v93_is_noise(s)]
    if not sentences:
        sentences=[s for s in _v93_sentence_split(fallback) if not _v93_is_noise(s)]
    if not sentences:
        return ''

    stop={'haber','haberi','son','dakika','bugün','yeni','ile','ve','bir','için','olan',
          'oldu','olacak','dedi','türkiye','türk','tarafından','kapsamında'}
    title_words={w for w in re.findall(r'[a-zçğıöşü0-9]+',norm(title)) if len(w)>=4 and w not in stop}

    # Ana olay cümlesi: yalnız ilk 7 cümle içinde aranır.
    # Böylece sayfanın sonundan/başka haberlerden içerik çekilmez.
    pool=sentences[:7]
    scores=[_v93_content_score(s,title_words) for s in pool]
    anchor=max(range(len(pool)),key=lambda i:(scores[i],-i)) if pool else 0

    # Başlangıç cümlesi başlıkla anlamlı örtüşüyorsa girişten başla.
    first_overlap=len(set(re.findall(r'[a-zçğıöşü0-9]+',norm(pool[0]))) & title_words) if pool else 0
    if first_overlap>=1:
        anchor=0

    chosen=[pool[anchor]]
    # Ana olayın hemen devamındaki en fazla iki cümleyi kullan.
    # Uzak cümle avlama kesinlikle yapılmaz.
    for s in sentences[anchor+1:anchor+3]:
        # Çok bariz biçimde yeni bir haber/konu başlıyorsa kes.
        ov=len(set(re.findall(r'[a-zçğıöşü0-9]+',norm(s))) & title_words)
        if ov==0 and len(chosen)>=2 and _v93_content_score(s,title_words)<5:
            break
        chosen.append(s)

    formal=[]
    for s in chosen:
        fs=_v92_formal_sentence(s).strip()
        if not fs:
            continue
        if fs[-1] not in '.!?':
            fs+='.'
        formal.append(fs)

    # 4 satır hedefi: yaklaşık 520 karakter; ASLA kelime/cümle ortasından kesme.
    result=[]
    for s in formal:
        candidate=' '.join(result+[s])
        if result and len(candidate)>520:
            break
        result.append(s)

    out=' '.join(result).strip()

    # Çok kısa ise yalnız bir sonraki ARDIŞIK cümleyi eklemeyi dene.
    if len(out)<170:
        next_i=anchor+len(result)
        if next_i<len(sentences):
            fs=_v92_formal_sentence(sentences[next_i]).strip()
            if fs and fs[-1] not in '.!?': fs+='.'
            if fs and len(out+' '+fs)<=520:
                out=(out+' '+fs).strip()

    # Başlık veya link asla eklenmez.
    out=re.sub(r'https?://\S+','',out)
    out=re.sub(r'\s+',' ',out).strip()
    return _v87_safe_tr(out)


def _v94_formal_summary_text(text):
    """
    Very simple and robust:
    - clean the saved news summary
    - take complete sentences in their original order
    - keep up to ~4 Word lines
    - formalize common journalistic endings
    - never return blank if usable text exists
    """
    t=_v87_safe_tr(text or '')
    t=re.sub(r'https?://\S+',' ',t)
    t=re.sub(r'\s+',' ',t).strip()
    if not t:
        return ''

    # Split only on real sentence endings; preserve original order.
    parts=[x.strip() for x in re.split(r'(?<=[.!?])\s+',t) if x.strip()]
    if not parts:
        parts=[t]

    noise=('devamını oku','benzer haber','ilgili haber','çerez','cookie',
           'reklam','instagram','facebook','twitter','whatsapp','google news')
    clean=[]
    for s in parts:
        ns=norm(s)
        if any(x in ns for x in noise):
            continue
        s=_v92_formal_sentence(s).strip()
        if not s:
            continue
        if s[-1] not in '.!?':
            s+='.'
        clean.append(s)

    if not clean:
        s=_v92_formal_sentence(t).strip()
        return (s.rstrip(' .;')+'.') if s else ''

    # Approx. four lines. Do not cut a sentence.
    chosen=[]
    for s in clean:
        candidate=' '.join(chosen+[s])
        if chosen and len(candidate)>560:
            break
        chosen.append(s)
        if len(chosen)>=3:
            break

    out=' '.join(chosen).strip()
    return out


def _v95_ogn_from_existing_engines(title, body):
    """
    Yeni bir özetleme algoritması yok.
    AKT'de başarılı çalışan:
      _akt_clean_sentences -> _akt_formal_summary
    ve Bilgi Notunda kullanılan kurumsal dil normalizasyonu kullanılır.
    """
    title=_clean_note_text(title)
    body=_clean_note_text(body or title)

    # AKT motoru haberi baştan sona değerlendirir; burada yalnız çıktı uzunluğu kısaltılır.
    text=_akt_formal_summary(
        title,
        body,
        max_sentences=3,
        max_chars=560
    )
    text=_clean_note_text(text)

    # AKT özetindeki ";" akışını ÖGN için tam cümlelere dönüştür.
    clauses=[_clean_note_text(x).strip(' ,;:.') for x in re.split(r'\s*;\s*',text) if _clean_note_text(x)]
    sentences=[]
    for clause in clauses[:3]:
        formal=_v66_formalize_sentence_endings(clause).strip()
        if not formal:
            continue
        if formal[-1] not in '.!?':
            formal+='.'
        sentences.append(formal)

    if not sentences:
        formal=_v66_formalize_sentence_endings(text).strip()
        if formal and formal[-1] not in '.!?':
            formal+='.'
        return formal

    # Yaklaşık 4 satır; tam cümleyi kesme.
    chosen=[]
    for s in sentences:
        candidate=' '.join(chosen+[s])
        if chosen and len(candidate)>560:
            break
        chosen.append(s)

    return ' '.join(chosen).strip()



def _v96_unique_sentences(title, body):
    """Bilgi notu motorundaki temizleme mantığını tek haber için uygular."""
    cleaned=_akt_clean_sentences(
        _clean_note_text(title),
        _clean_note_text(body)
    )
    uniq=[]
    seen=[]
    for sent in cleaned:
        sent=_repair_mojibake_utf8(_clean_note_text(sent)).strip()
        if not sent:
            continue
        toks=set(norm(sent).split())
        if not toks:
            continue
        dup=False
        for old in seen[-30:]:
            union=len(toks|old)
            if union and len(toks&old)/union>=0.78:
                dup=True
                break
        if not dup:
            uniq.append(sent)
            seen.append(toks)
    return uniq

def _v96_has_critical_data(s):
    n=norm(s)
    return bool(
        re.search(r'\b\d+(?:[.,]\d+)?\b|%|yüzde|milyon|milyar|trilyon|bin\b',s,re.I)
        or any(x in n for x in [
            'üretim','ihracat','ithalat','kapasite','yatırım','ciro','satış',
            'başvuru','öğrenci','personel','istihdam','menzil','adet','oran',
            'endeks','bütçe','hibe','destek','maliyet','gelir','zarar'
        ])
    )

def _v96_short_information_note(title, body):
    """
    ÖGN = kısaltılmış bilgi notu.
    En fazla 4 paragraf:
      1) ana gelişme/özet,
      2) kritik veri-rakam-istatistik,
      3) gerekiyorsa tamamlayıcı gelişme/sonuç,
      4) yalnız kaynakta anlamlı bir sonuç/son durum varsa.
    """
    uniq=_v96_unique_sentences(title,body)
    if not uniq:
        fallback=_repair_mojibake_utf8(_clean_note_text(body or title))
        if not fallback:
            return []
        return [_v66_formalize_sentence_endings(fallback)]

    # 1. paragraf: bilgi notundaki gibi ilk 1-2 cümlede olayın özü.
    intro_s=uniq[:2]
    intro=_join_sentences_naturally(intro_s)
    intro=_v66_formalize_sentence_endings(intro)

    # Kritik rakam/veri cümlelerini ASLA sırf kısa özet uğruna atlama.
    critical=[]
    for i,s in enumerate(uniq):
        if i<2:
            continue
        if _v96_has_critical_data(s):
            critical.append((i,_sent_score(s),s))

    # En yüksek bilgi yoğunluklu kritik cümleleri seç, fakat haber sırasını koru.
    chosen_detail_idx=set()
    for i,score,s in sorted(critical,key=lambda x:(x[1],-x[0]),reverse=True)[:6]:
        chosen_detail_idx.add(i)

    # İlk iki cümleden sonra konu akışını tamamlayan yüksek skorlu normal cümleler.
    remaining=[
        (i,_sent_score(s),s) for i,s in enumerate(uniq[2:],start=2)
        if i not in chosen_detail_idx
    ]
    for i,score,s in sorted(remaining,key=lambda x:(x[1],-x[0]),reverse=True)[:3]:
        if score>=3:
            chosen_detail_idx.add(i)

    ordered_details=[uniq[i] for i in sorted(chosen_detail_idx)]

    paragraphs=[intro] if intro else []

    # 2-3. paragraflar: kritik detayları bilgi notu gibi gruplandır.
    if ordered_details:
        if len(ordered_details)<=4:
            p2=_join_sentences_naturally(ordered_details)
            if p2:
                paragraphs.append(_v66_formalize_sentence_endings(p2))
        else:
            split=max(2,min(4,(len(ordered_details)+1)//2))
            p2=_join_sentences_naturally(ordered_details[:split])
            p3=_join_sentences_naturally(ordered_details[split:])
            if p2:
                paragraphs.append(_v66_formalize_sentence_endings(p2))
            if p3:
                paragraphs.append(_v66_formalize_sentence_endings(p3))

    # Son paragraf: generic değerlendirme yazma; kaynakta kalan gerçek son durumdan seç.
    used=set(intro_s+ordered_details)
    tail_candidates=[s for s in uniq[-5:] if s not in used]
    if tail_candidates and len(paragraphs)<4:
        # Sonuç/hedef/son durum taşıyan cümleyi tercih et.
        result_terms=[
            'hedef','beklen','plan','sonuç','bu kapsamda','bu çerçevede','devam',
            'başlayacak','tamamlanacak','uygulanacak','sağlanacak','öngör',
            'artıracak','azaltacak','katkı','etki','takvim'
        ]
        ranked=sorted(
            tail_candidates,
            key=lambda s:(
                sum(x in norm(s) for x in result_terms),
                _sent_score(s)
            ),
            reverse=True
        )
        tail=ranked[0] if ranked else ''
        if tail:
            paragraphs.append(
                _v66_formalize_sentence_endings(_join_sentences_naturally([tail]))
            )

    # Maksimum 4 paragraf; boşları temizle.
    out=[]
    for para in paragraphs[:4]:
        para=_repair_mojibake_utf8(_clean_note_text(para)).strip()
        if para and para not in out:
            out.append(para)
    return out



def _v97_compact_sentence(s,max_chars=260):
    """Çok uzun tek cümleyi, ana özne/eylem ve kritik sayısal parçaları koruyarak kısaltır."""
    s=_repair_mojibake_utf8(_clean_note_text(s)).strip()
    if len(s)<=max_chars:
        return s

    # Virgül/noktalı virgül ile ayrılmış anlamlı parçaları değerlendir.
    parts=[x.strip(' ,;:.') for x in re.split(r'\s*[;,]\s*',s) if x.strip()]
    if not parts:
        return s[:max_chars].rsplit(' ',1)[0].rstrip(' ,;:.')+'.'

    chosen=[parts[0]]
    # Kritik veri/rakam içeren parçaları öncelikle koru.
    critical=[x for x in parts[1:] if _v96_has_critical_data(x)]
    for x in critical:
        cand=', '.join(chosen+[x])
        if len(cand)<=max_chars:
            chosen.append(x)

    # Hâlâ çok kısa ise ikinci parçayı bağlam için ekle.
    if len(chosen)==1 and len(parts)>1:
        cand=', '.join(chosen+[parts[1]])
        if len(cand)<=max_chars:
            chosen.append(parts[1])

    out=', '.join(chosen).strip()
    if out and out[-1] not in '.!?':
        out+='.'
    return out

def _v97_analyst_summary(title,body):
    """
    Gerçek analist mantığıyla kısa ÖGN:
    - tek paragraf,
    - yaklaşık 4 Word satırı,
    - 2-3 tam cümle,
    - ana gelişme + kritik rakam/veri + sonuç/son durum,
    - bilgi notunun kısaltılmış hali.
    """
    uniq=_v96_unique_sentences(title,body)
    if not uniq:
        fallback=_repair_mojibake_utf8(_clean_note_text(body or title))
        fallback=_v66_formalize_sentence_endings(fallback)
        return _v97_compact_sentence(fallback,480)

    # 1) Ana gelişme: haberin ilk anlamlı cümlesi.
    intro=_v66_formalize_sentence_endings(uniq[0]).strip()
    intro=_v97_compact_sentence(intro,260)

    chosen=[intro] if intro else []
    used={0}

    # 2) En kritik veri/rakam/istatistik:
    critical=[]
    for i,s in enumerate(uniq[1:],start=1):
        if _v96_has_critical_data(s):
            # Rakam + kurumsal/sonuç bilgisi daha yüksek puan.
            score=_sent_score(s)
            score+=min(len(re.findall(r'\d+(?:[.,]\d+)?',s)),4)*2
            critical.append((score,i,s))
    if critical:
        _,idx,s=max(critical,key=lambda x:(x[0],-x[1]))
        fs=_v66_formalize_sentence_endings(s).strip()
        fs=_v97_compact_sentence(fs,250)
        if fs and title_key(fs)!=title_key(intro):
            chosen.append(fs); used.add(idx)

    # 3) Sonuç/hedef/sonraki adım taşıyan cümle.
    result_terms=[
        'hedef','beklen','plan','sonuç','başlayacak','tamamlanacak','uygulanacak',
        'sağlanacak','öngör','katkı','etki','devam','rekor','artış','azalış',
        'başvuru','tarih','takvim','yüksel','gerile'
    ]
    result_candidates=[]
    for i,s in enumerate(uniq[1:],start=1):
        if i in used:
            continue
        n=norm(s)
        score=sum(x in n for x in result_terms)*3 + _sent_score(s)
        if score>0:
            result_candidates.append((score,i,s))
    if result_candidates:
        _,idx,s=max(result_candidates,key=lambda x:(x[0],-x[1]))
        fs=_v66_formalize_sentence_endings(s).strip()
        fs=_v97_compact_sentence(fs,220)
        if fs and all(title_key(fs)!=title_key(x) for x in chosen):
            chosen.append(fs)

    # Haber sırasını korumak için seçilen cümleleri tekrar orijinal sıraya koymuyoruz:
    # intro daima ilk, ardından kritik veri, ardından sonuç.
    # Bu, bilgi notunun kısa analist akışıdır.

    # Tek paragraf ve yaklaşık 4 satır: tam cümleyi kesmeden 500 karakter.
    final=[]
    for s in chosen[:3]:
        s=_repair_mojibake_utf8(_clean_note_text(s)).strip()
        if not s:
            continue
        if s[-1] not in '.!?':
            s+='.'
        cand=' '.join(final+[s])
        if final and len(cand)>500:
            break
        final.append(s)

    out=' '.join(final).strip()

    # İlk cümle tek başına çok uzunsa güvenli biçimde sıkıştır.
    if len(out)>500:
        out=_v97_compact_sentence(out,490)

    return _repair_mojibake_utf8(_clean_note_text(out)).strip()



def _v98_strip_site_name(text, source=''):
    """Özetin sonunda kalan yayıncı/site adını temizler."""
    t=_repair_mojibake_utf8(_clean_note_text(text)).strip()
    s=_repair_mojibake_utf8(_clean_note_text(source)).strip()
    if s:
        # Cümle sonunda " - siteadı", " — siteadı", "(siteadı)" vb.
        t=re.sub(r'\s*[-–—|]\s*'+re.escape(s)+r'\s*$', '', t, flags=re.I)
        t=re.sub(r'\s*\(\s*'+re.escape(s)+r'\s*\)\s*$', '', t, flags=re.I)
        t=re.sub(r'\s*'+re.escape(s)+r'\s*$', '', t, flags=re.I)
    # Genel domain/site sonları.
    t=re.sub(r'\s*[-–—|]\s*[\w.-]+\.(?:com|com\.tr|net|org|gov\.tr|edu\.tr)\s*$', '', t, flags=re.I)
    return t.strip(' -–—|')

def _v98_remove_heading_like_text(text, title=''):
    """
    Başlık benzeri ALL CAPS parçalarını ve haber başlığının aynısını özetten çıkarır.
    """
    t=_repair_mojibake_utf8(_clean_note_text(text))
    title_n=title_key(title)

    parts=_sentence_chunks(t)
    out=[]
    for s in parts:
        s=_repair_mojibake_utf8(_clean_note_text(s)).strip()
        if not s:
            continue

        # Haber başlığının aynısı veya çok yakınsa alma.
        if title_n and title_key(s)==title_n:
            continue

        # Kısa ALL CAPS başlıkları alma.
        letters=''.join(ch for ch in s if ch.isalpha())
        if letters and len(s)<160:
            ratio=sum(ch.isupper() for ch in letters)/max(1,len(letters))
            if ratio>0.72:
                continue

        out.append(s)
    return ' '.join(out).strip()

def _v98_safe_tr(text):
    """
    Word'e girmeden önce Türkçe karakterleri güvenli hale getirir.
    Mojibake ve görünmez karakterleri temizler.
    """
    t=_repair_mojibake_utf8(_clean_note_text(text))
    # Kalan yaygın bozukluklar.
    fixes={
        'TÃ¼rkiye':'Türkiye','TÃ¼rk':'Türk','genÃ§':'genç','dÃ¼nya':'dünya',
        'Ã¼lke':'ülke','Ã¼stÃ¼n':'üstün','Ã¶ÄŸrenci':'öğrenci','Ã¶Ärenci':'öğrenci',
        'baÅŸar':'başar','katÄ±lÄ±m':'katılım','mÃ¼cadele':'mücadele',
        'Ä±':'ı','ÄŸ':'ğ','ÅŸ':'ş','Ã§':'ç','Ã¶':'ö','Ã¼':'ü',
        'Ä°':'İ','Äž':'Ğ','Åž':'Ş','Ã‡':'Ç','Ã–':'Ö','Ãœ':'Ü',
        'Â':'','â€™':'’','â€œ':'“','â€':'”','â€“':'–','â€”':'—'
    }
    for a,b in fixes.items():
        t=t.replace(a,b)

    # Word/XML açısından problemli görünmez karakterleri temizle.
    for bad in ('\u00ad','\u200b','\u200c','\u200d','\ufeff'):
        t=t.replace(bad,'')

    # Kalan açık mojibake sembollerini boşlukla değiştir.
    t=re.sub(r'[ÃÄÅÂ ]+',' ',t)
    t=re.sub(r'\s+',' ',t).strip()
    return t

def _v98_exact_four_line_summary(title, source, body):
    """
    V98 ÖGN:
    - her haber tek paragraf,
    - yaklaşık 4 Word satırı,
    - 2-3 tam cümle,
    - resmî dil,
    - kritik rakam/veri korunur,
    - başlık/site adı/bozuk karakter yok.
    """
    # V97 motorunu temel al.
    raw=_v97_analyst_summary(title,body)
    raw=_v98_safe_tr(raw)
    raw=_v98_remove_heading_like_text(raw,title)
    raw=_v98_strip_site_name(raw,source)

    if not raw:
        # Yedek: kayıtlı haber metninden doğrudan kısa resmî özet.
        raw=_v95_ogn_from_existing_engines(title,body)
        raw=_v98_safe_tr(raw)
        raw=_v98_remove_heading_like_text(raw,title)
        raw=_v98_strip_site_name(raw,source)

    # Cümleleri yeniden düzenle, tam cümleler dışında kesme yok.
    sents=[]
    for s in _sentence_chunks(raw):
        s=_v98_safe_tr(s).strip()
        if not s:
            continue
        s=_v66_formalize_sentence_endings(s).strip()
        if s and s[-1] not in '.!?':
            s+='.'
        sents.append(s)

    # Hedef yaklaşık 4 satır = 430-500 karakter bandı.
    chosen=[]
    total=0
    for s in sents[:4]:
        candidate=' '.join(chosen+[s])
        if chosen and len(candidate)>500:
            break
        chosen.append(s)
        total=len(candidate)
        if total>=430:
            break

    # Çok kısa kaldıysa sonraki tam cümleyi ekle.
    if len(' '.join(chosen))<300:
        for s in sents[len(chosen):]:
            candidate=' '.join(chosen+[s])
            if len(candidate)<=500:
                chosen.append(s)
            if len(' '.join(chosen))>=360:
                break

    out=' '.join(chosen).strip()

    # Son güvenlik katmanı.
    out=_v98_safe_tr(out)
    out=_v98_remove_heading_like_text(out,title)
    out=_v98_strip_site_name(out,source)

    # Sonunda nokta olsun, ama site adı olmasın.
    if out and out[-1] not in '.!?':
        out+='.'
    return out



# ========================= V99: ÖGN ANALİST FİLTRESİ =========================
_V99_UI_NOISE = [
    r'sıralamayı değiştirmek için kartları',
    r'kartları yukarı',
    r'kartları aşağı',
    r'view more',
    r'daha fazla göster',
    r'devamını oku',
    r'cookie',
    r'çerez',
    r'reklam',
    r'abone ol',
    r'bildirimleri aç',
    r'ana sayfa',
    r'son dakika',
    r'galeri',
    r'foto galeri',
    r'video galeri',
    r'yorumlar',
    r'paylaş',
]

_V99_NEWS_SPEECH = [
    (r'\bbaşladı\b', 'başlamıştır'),
    (r'\bbaşladıktan\b', 'başladıktan'),
    (r'\baçıkladı\b', 'açıklamıştır'),
    (r'\bbelirtti\b', 'belirtmiştir'),
    (r'\bkaydetti\b', 'kaydetmiştir'),
    (r'\bduyurdu\b', 'duyurmuştur'),
    (r'\bbildirdi\b', 'bildirmiştir'),
    (r'\bifade etti\b', 'ifade etmiştir'),
    (r'\bsöyledi\b', 'belirtmiştir'),
    (r'\bgerçekleşti\b', 'gerçekleşmiştir'),
    (r'\byükseldi\b', 'yükselmiştir'),
    (r'\bgeriledi\b', 'gerilemiştir'),
    (r'\barttı\b', 'artmıştır'),
    (r'\bazaldı\b', 'azalmıştır'),
    (r'\bkırıldı\b', 'kırılmıştır'),
    (r'\btamamlandı\b', 'tamamlanmıştır'),
    (r'\bgirdi\b', 'girmiştir'),
    (r'\bbaşlıyor\b', 'başlamaktadır'),
    (r'\bdevam ediyor\b', 'devam etmektedir'),
    (r'\bhedefliyor\b', 'hedeflemektedir'),
    (r'\bbekleniyor\b', 'beklenmektedir'),
    (r'\bsürüyor\b', 'sürmektedir'),
    (r'\bulaştı\b', 'ulaşmıştır'),
    (r'\bçıktı\b', 'çıkmıştır'),
    (r'\byapıldı\b', 'yapılmıştır'),
    (r'\bseçilecek\b', 'seçilecektir'),
    (r'\byetiştireceğiz\b', 'yetiştirilmesi planlanmaktadır'),
    (r'\bkazandıracağız\b', 'kazandırılması hedeflenmektedir'),
]

def _v99_clean_article_text(text, source=''):
    t=_v98_safe_tr(text)
    t=t.replace('»',' ').replace('›',' ').replace('',"'").replace('','ş').replace('¢','â')
    t=re.sub(r'\s+',' ',t).strip()

    # Web arayüzü / sayfa gürültüsü taşıyan cümleleri tamamen at.
    kept=[]
    for s in _sentence_chunks(t):
        n=norm(s)
        if any(re.search(p,n,re.I) for p in _V99_UI_NOISE):
            continue
        # Kaynak/site navigasyonu gibi çok kısa satırları at.
        if len(s.split()) < 4 and not re.search(r'\d',s):
            continue
        kept.append(s.strip())
    t=' '.join(kept)

    # Kaynak adını cümle başından/sonundan temizle.
    if source:
        ss=re.escape(_v98_safe_tr(source))
        t=re.sub(r'^\s*'+ss+r'\s*[-:»|]*\s*','',t,flags=re.I)
        t=re.sub(r'\s*[-:»|]*\s*'+ss+r'\s*$','',t,flags=re.I)

    # Genel domainleri ve tipik gazete navigasyonunu kaldır.
    t=re.sub(r'\b[\w.-]+\.(?:com\.tr|com|net|org|gov\.tr|edu\.tr)\b',' ',t,flags=re.I)
    t=re.sub(r'^[^.!?]{0,80}\b(?:Gazetesi|Gazete|Haber|Haberleri)\s*[»|:-]+\s*','',t,flags=re.I)
    t=re.sub(r'\s+',' ',t).strip()
    return t

def _v99_is_title_only(title, text):
    a=set(norm(title).split())
    b=set(norm(text).split())
    if not a or not b:
        return False
    return len(a & b)/max(1,len(a | b)) >= 0.72 and len(text) < 220

def _v99_officialize(text):
    t=_v98_safe_tr(text).strip()
    for pat,repl in _V99_NEWS_SPEECH:
        t=re.sub(pat,repl,t,flags=re.I)
    # Haber dilindeki doğrudan alıntı kalıplarını kurumsallaştır.
    t=re.sub(r'\bifadelerini kullandı\b','belirtmiştir',t,flags=re.I)
    t=re.sub(r'\bşunları kaydetti\b','açıklamada bulunmuştur',t,flags=re.I)
    t=re.sub(r'\bşöyle\b\s*:?', '', t, flags=re.I)
    t=_v66_formalize_sentence_endings(t)
    t=re.sub(r'\s+',' ',t).strip()
    return t

def _v99_sentence_value(s):
    n=norm(s)
    score=_sent_score(s)
    nums=len(re.findall(r'\b\d+(?:[.,]\d+)?\b|%',s))
    score += min(nums,5)*3
    # Kurumsal/stratejik bilgi yoğunluğu
    for term in [
        'tüik','tcmb','epdk','bakanlık','cumhurbaşkanı','tübitak','kosgeb','tse','türkpatent',
        'üretim','ihracat','ithalat','yatırım','kapasite','oran','yüzde','milyon','milyar',
        'teslimat','envanter','program','proje','hibe','destek','menzil','adet','öğrenci',
        'araştırmacı','tesis','fabrika','teknoloji','savunma','uzay','yapay zeka'
    ]:
        if term in n:
            score += 2
    return score

def _v99_analyst_ogn(title, source, body, fallback_summary=''):
    """
    İşyeri ÖGN örneğine göre:
    somut özne/kurum + gelişme + kritik rakam/ölçek + sonuç.
    Tek paragraf, 2-3 tam cümle, yaklaşık dört Word satırı.
    """
    title=_v98_safe_tr(title)
    source=_v98_safe_tr(source)
    body=_v99_clean_article_text(body,source)
    fallback_summary=_v99_clean_article_text(fallback_summary,source)

    # Gerçek içerik yoksa başlığı Word'e basmak yerine kayıtlı özeti dene.
    candidate_body=body
    if not candidate_body or _v99_is_title_only(title,candidate_body):
        candidate_body=fallback_summary

    # Hâlâ yalnız başlıksa bunu geçerli ÖGN sayma.
    if not candidate_body or _v99_is_title_only(title,candidate_body):
        return ''

    sentences=_v96_unique_sentences(title,candidate_body)
    clean=[]
    for s in sentences:
        s=_v99_clean_article_text(s,source)
        if not s or _v99_is_title_only(title,s):
            continue
        if any(re.search(p,norm(s),re.I) for p in _V99_UI_NOISE):
            continue
        clean.append(s)

    if not clean:
        return ''

    # İlk cümle: haberin gerçek ana gelişmesi. İlk 4 cümlede en yüksek değerli olanı seç.
    first_pool=list(enumerate(clean[:4]))
    first_idx, first=max(first_pool,key=lambda x:(_v99_sentence_value(x[1]),-x[0]))

    selected=[(first_idx,first)]

    # Kritik rakam/veri: ana cümlede yoksa ayrıca seç.
    critical=[]
    for i,s in enumerate(clean):
        if i==first_idx:
            continue
        if _v96_has_critical_data(s):
            critical.append((_v99_sentence_value(s),i,s))
    if critical:
        _,i,s=max(critical,key=lambda x:(x[0],-x[1]))
        selected.append((i,s))

    # Sonuç/hedef/son durum: yalnız kaynakta varsa.
    result_terms=['hedef','plan','beklen','öngör','teslim','envanter','başvuru','hibe',
                  'artış','azalış','rekor','üretim','faaliyete','tamamlan','başlam']
    results=[]
    used={i for i,_ in selected}
    for i,s in enumerate(clean):
        if i in used:
            continue
        sc=sum(x in norm(s) for x in result_terms)*4 + _v99_sentence_value(s)
        if sc>3:
            results.append((sc,i,s))
    if results:
        _,i,s=max(results,key=lambda x:(x[0],-x[1]))
        selected.append((i,s))

    # Analist akışı: ana gelişme -> veri -> sonuç. En fazla 3 cümle.
    out=[]
    for _,s in selected[:3]:
        s=_v99_officialize(s)
        s=_v98_strip_site_name(s,source)
        s=re.sub(r'^[A-ZÇĞİÖŞÜ0-9\s\-–—:]{8,80}(?=[A-ZÇĞİÖŞÜ][a-zçğıöşü])','',s).strip()
        if not s:
            continue
        if s[-1] not in '.!?':
            s+='.'
        if all(title_key(s)!=title_key(x) for x in out):
            out.append(s)

    # Dört satır hedefi: yaklaşık 560 karakter; cümle ortasında kesme yok.
    chosen=[]
    for s in out:
        cand=' '.join(chosen+[s])
        if chosen and len(cand)>560:
            break
        chosen.append(s)

    text=' '.join(chosen).strip()
    text=_v99_clean_article_text(text,source)
    text=_v99_officialize(text)
    text=_v98_strip_site_name(text,source)
    text=_v98_safe_tr(text)

    if _v99_is_title_only(title,text):
        return ''
    if text and text[-1] not in '.!?':
        text+='.'
    return text


# ========================= V100: KONU BÜTÜNLÜĞÜ =========================

def _v100_title_topic(title, source=''):
    """
    Başlıktan yalnızca konu/özne çekirdeğini çıkarır.
    Başlığı aynen basmaz; ALL CAPS ve site adını temizler.
    """
    t=_v98_safe_tr(title)
    t=_v98_strip_site_name(t,source)
    t=re.sub(r'\s*[-–—|]\s*[\w.-]+\.(?:com|com\.tr|net|org|gov\.tr|edu\.tr)\s*$','',t,flags=re.I)
    t=re.sub(r'!+$','',t).strip()

    # Tamamı büyük harfliyse normal cümle görünümüne getir.
    letters=''.join(ch for ch in t if ch.isalpha())
    if letters and sum(ch.isupper() for ch in letters)/max(1,len(letters)) > 0.72:
        t=t.lower()
        t=t[:1].upper()+t[1:]

    return t.strip(' -–—|:;')

def _v100_subject_hint(title, source=''):
    """
    Başlıktan self-contained paragraf için özne/konu ipucu çıkarır.
    Örn: 'KAAN...' -> 'KAAN projesi'
         'Antalya'da Savunma Sanayine Yatırım Fırsatı' -> 'Antalya'daki savunma sanayii yatırımları'
    """
    t=_v100_title_topic(title,source)
    n=norm(t)

    # Bilinen kalıplar.
    mappings=[
        (r'\bkaan\b', 'KAAN projesi'),
        (r'\bbayraktar kalkan\b', 'Bayraktar Kalkan DİHA'),
        (r'\bt10x\b|\bt10f\b|\btogg\b', 'Togg’un T10X ve T10F modelleri'),
        (r'\belektrikli araç.*şarj\b|\bşarj altyap', 'Türkiye’de elektrikli araç şarj altyapısı'),
        (r'\bkapasite kullanım\b', 'imalat sanayisi kapasite kullanım oranı'),
        (r'\bkardemir\b', 'Kardemir Çelik’in 2026 yılı ilk yarı finansal sonuçları'),
        (r'\bgezeravcı\b', 'Alper Gezeravcı’nın Amasya’daki görevi'),
        (r'\bgoogle.*ai plus\b|\bai plus\b', 'Google’ın üniversite öğrencilerine yönelik AI Plus programı'),
        (r'\bmoğolistan.*mühimmat\b|\bmke.*moğolistan\b', 'MKE’nin Moğolistan’daki mühimmat üretim tesisi'),
        (r'\btekno.*mavi vatan\b|\balkü\b|\bzeronetech\b', 'TEKNOFEST Mavi Vatan kapsamında ALKÜ Zeronetech Takımı'),
        (r'\bsavunma sanay.*yatırım.*antalya\b|\bantalya.*savunma sanay', 'Antalya’daki savunma sanayii yatırım fırsatları'),
        (r'\byapay zeka olimpiyat\b|\bbilim olimpiyat', 'TÜBİTAK Bilim Olimpiyatları kapsamında Türkiye’yi temsil eden öğrenciler'),
        (r'\btürk telekom\b|\bdijitalde hayat kolay\b', 'Türk Telekom’un Dijitalde Hayat Kolay projesi'),
        (r'\b5g.*robotik cerrahi\b|\btcg anadolu.*5g\b', 'TCG ANADOLU’da 5G destekli uzaktan robotik cerrahi uygulaması'),
    ]
    for pat,label in mappings:
        if re.search(pat,n,re.I):
            return label

    # Genel başlık: ilk 8-10 anlamlı kelimeyi konu ipucu olarak kullan.
    words=t.split()
    if not words:
        return ''
    return ' '.join(words[:10]).strip(' ,;:-')

def _v100_is_fragment(s):
    """Özne/bağlam içermeyen kırık veya yarım cümleleri tespit eder."""
    s=_v98_safe_tr(s).strip()
    if not s:
        return True
    n=norm(s)

    # Açık yarım cümleler / UI artıkları.
    if s.endswith((' ve',' ile',' için',' k',';',' :','...','…')):
        return True
    if re.search(r'\b(?:proje|program|şirket|kurum|bu kapsamda|bunun yanında|ayrıca)\b',n) and len(s.split())<7:
        return True
    if re.match(r'^(proje|program|şirket|kurum|bunun yanında|bu kapsamda|ayrıca)\b',n):
        return True
    if re.match(r'^\d+\s+(?:yıl|ay|gün)\b',n):
        return True
    return False

def _v100_contextualize(sentence, subject):
    """
    Paragrafın ilk cümlesi kendi başına anlaşılmıyorsa başlıktan gelen konu çekirdeğiyle
    bağlamı tamamlar. Başlığı olduğu gibi eklemez.
    """
    s=_v99_officialize(sentence).strip()
    if not s:
        return ''
    n=norm(s)

    weak_starts=(
        'proje ','program ','şirket ','kurum ','bunun yanında ',
        'bu kapsamda ','ayrıca ','ilk olarak ','daha sonra '
    )
    if any(n.startswith(x) for x in weak_starts) or _v100_is_fragment(s):
        if subject:
            # "Proje, Mart..." -> "KAAN projesinde üretim faaliyetleri Mart..."
            if n.startswith('proje '):
                s=re.sub(r'^\s*Proje\s*,?\s*', subject+' kapsamında ', s, flags=re.I)
            elif n.startswith('program '):
                s=re.sub(r'^\s*Program\s*,?\s*', subject+' kapsamında ', s, flags=re.I)
            elif n.startswith('şirket '):
                s=re.sub(r'^\s*Şirket\s*,?\s*', subject+' kapsamında ilgili şirket ', s, flags=re.I)
            else:
                s=subject.rstrip(' .')+' kapsamında '+s[:1].lower()+s[1:]
    return s

def _v100_pick_summary_sentences(title, source, body, fallback):
    """
    Self-contained analist özeti:
    1) ilk cümle mutlaka konuyu/özneyi açıklar,
    2) ikinci cümle kritik veri/rakam,
    3) üçüncü cümle sonuç/hedef/takvim.
    """
    subject=_v100_subject_hint(title,source)
    body=_v99_clean_article_text(body,source)
    fallback=_v99_clean_article_text(fallback,source)

    candidate=body
    if not candidate or _v99_is_title_only(title,candidate):
        candidate=fallback

    sents=_v96_unique_sentences(title,candidate)
    clean=[]
    for s in sents:
        s=_v99_clean_article_text(s,source)
        if not s or _v100_is_fragment(s):
            continue
        if _v99_is_title_only(title,s):
            continue
        clean.append(s)

    if not clean:
        return ''

    # 1) Ana olay: ilk 5 cümle içinde subject/title overlap + kurumsal eylem.
    tw=set(re.findall(r'[a-zçğıöşü0-9]+',norm(subject or title)))
    def overlap(s):
        sw=set(re.findall(r'[a-zçğıöşü0-9]+',norm(s)))
        return len(sw & tw)

    action_terms=['açıkla','duyur','başlat','gerçekleştir','tamamla','imzala','üret',
                  'satış','yatırım','test','görev','teslim','envanter','faaliyete','seç']
    def intro_score(s,idx):
        n=norm(s)
        return 7*overlap(s)+4*sum(x in n for x in action_terms)+_sent_score(s)-idx

    first_pool=list(enumerate(clean[:5]))
    idx0, intro=max(first_pool,key=lambda z:intro_score(z[1],z[0]))
    intro=_v100_contextualize(intro,subject)

    # Eğer ilk cümlede subject hiç yoksa, doğal konu cümlesiyle bağla.
    if subject and overlap(intro)==0:
        # Başlıktan birebir kopya değil, konu bağlamı.
        if intro:
            intro=subject.rstrip(' .')+' hakkında, '+intro[:1].lower()+intro[1:]

    selected=[intro] if intro else []
    used={idx0}

    # 2) Kritik veri/rakam.
    critical=[]
    for i,s in enumerate(clean):
        if i in used:
            continue
        if _v96_has_critical_data(s):
            critical.append((_v99_sentence_value(s)+3*overlap(s),i,s))
    if critical:
        _,i,s=max(critical,key=lambda x:(x[0],-x[1]))
        fs=_v99_officialize(s)
        if fs and not _v100_is_fragment(fs):
            selected.append(fs); used.add(i)

    # 3) Sonuç/hedef/takvim.
    result_terms=['hedef','beklen','plan','başvuru','teslim','envanter','rekor',
                  'katkı','artış','azalış','faaliyete','tamamlan','başlam','uygulan']
    results=[]
    for i,s in enumerate(clean):
        if i in used:
            continue
        n=norm(s)
        sc=4*sum(x in n for x in result_terms)+_sent_score(s)+2*overlap(s)
        if sc>3:
            results.append((sc,i,s))
    if results:
        _,i,s=max(results,key=lambda x:(x[0],-x[1]))
        fs=_v99_officialize(s)
        if fs and not _v100_is_fragment(fs):
            selected.append(fs)

    # Tekrarlı sayısal cümleleri azalt.
    final=[]
    seen_nums=[]
    for s in selected[:3]:
        s=_v98_safe_tr(_v99_officialize(s)).strip()
        if not s:
            continue
        nums=set(re.findall(r'\d+(?:[.,]\d+)?',s))
        duplicate=False
        for prev,prev_nums in zip(final,seen_nums):
            if nums and nums==prev_nums and len(nums)>=1 and title_key(s)==title_key(prev):
                duplicate=True
                break
        if duplicate:
            continue
        if s[-1] not in '.!?':
            s+='.'
        final.append(s)
        seen_nums.append(nums)

    # 4 satır hedefi ~560 karakter; tam cümle kesilmez.
    out=[]
    for s in final:
        cand=' '.join(out+[s])
        if out and len(cand)>560:
            break
        out.append(s)

    text=' '.join(out).strip()
    text=_v99_clean_article_text(text,source)
    text=_v99_officialize(text)
    text=_v98_strip_site_name(text,source)
    text=_v98_safe_tr(text)

    # Son güvenlik: paragraf yine başlık seviyesindeyse geçersiz say.
    if _v99_is_title_only(title,text):
        return ''
    return text


# ========================= V101: ÖGN KONU BÜTÜNLÜĞÜ 2.0 =========================

def _v101_clean_unicode(text):
    """Türkçe metindeki mojibake/control karakterlerini Word öncesi agresif biçimde temizler."""
    import unicodedata
    t=_repair_mojibake_utf8(_clean_note_text(text))

    # Kullanıcı çıktısında görülen tek-byte/control bozulmaları.
    fixes={
        '\x9c':'Ü','\x9e':'Ş','\x9f':'ş','\x96':'Ö','\x91':"'",'\x92':"'",'\x93':'“','\x94':'”',
        'T BİTAK':'TÜBİTAK','T BİTAK':'TÜBİTAK',
        'ö şrenci':'öğrenci','ö şrenc':'öğrenc','yarı ş':'yarış','etti şi':'ettiği',
        'ba şarı':'başarı','gümü ş':'gümüş',' zekinci':' Özekinci',
        'Yapay Zek â':'Yapay Zekâ','veyaşağı':'veya aşağı'
    }
    for a,b in fixes.items():
        t=t.replace(a,b)

    # Harflerin arasına sızmış boşluklu ş/ğ/ü bozulmaları.
    t=re.sub(r'\bö\s+şrenc', 'öğrenc', t, flags=re.I)
    t=re.sub(r'\byarı\s+ş', 'yarış', t, flags=re.I)
    t=re.sub(r'\bba\s+şar', 'başar', t, flags=re.I)
    t=re.sub(r'\bgümü\s+ş', 'gümüş', t, flags=re.I)
    t=re.sub(r'\betti\s+şi\b', 'ettiği', t, flags=re.I)

    # Kontrol/görünmez karakterleri kaldır.
    t=''.join(ch for ch in t if unicodedata.category(ch) not in {'Cc','Cf'} or ch in '\t\n\r')
    t=unicodedata.normalize('NFC',t)
    t=re.sub(r'\s+',' ',t).strip()
    return t

def _v101_title_terms(title):
    stop={'ve','ile','için','bir','bu','da','de','ile','olan','olarak','son','yeni',
          'türkiye','türk','haber','haberi','başladı','oldu','edildi','açıkladı'}
    return [w for w in re.findall(r'[a-zçğıöşü0-9]+',norm(title)) if len(w)>2 and w not in stop]

def _v101_sentence_overlap(title,s):
    tw=set(_v101_title_terms(title))
    sw=set(re.findall(r'[a-zçğıöşü0-9]+',norm(s)))
    return len(tw & sw)

def _v101_bad_sentence(s):
    s=_v101_clean_unicode(s).strip()
    n=norm(s)
    if not s or len(s)<35:
        return True
    if _v100_is_fragment(s):
        return True
    if any(re.search(p,n,re.I) for p in _V99_UI_NOISE):
        return True
    # Başlık/navigation artıkları.
    if re.search(r'\b(?:yarışıyor|tükendi|fırsatı|büyüyor)\s*!?\s*$',n) and len(s)<120:
        return True
    if re.search(r'\b(?:gazetesi|gazete|haberleri?)\s*[»|:-]',n):
        return True
    return False

def _v101_formal_sentence(s):
    """Haber dilini resmî bilgi notu diline yaklaştırır; doğrudan alıntı kırıntılarını atar."""
    s=_v101_clean_unicode(s)
    # Açık/kapanmamış tırnakları temizle.
    s=s.replace('"','').replace('“','').replace('”','')
    s=re.sub(r'\b(?:dedi|diyor|diye konuştu)\b','belirtmiştir',s,flags=re.I)
    s=_v99_officialize(s)
    # Kalan yaygın haber dili.
    repl=[
        (r'\bedecek\b','edecektir'),(r'\bolacak\b','olacaktır'),
        (r'\byarışacak\b','yarışacaktır'),(r'\bsunuluyor\b','sunulmaktadır'),
        (r'\büretiyor\b','üretmektedir'),(r'\bilerliyor\b','ilerlemektedir'),
        (r'\bdevreye alındı\b','devreye alınmıştır'),
        (r'\bonay alındı\b','onay alınmıştır'),
        (r'\bgerçekleştirildi\b','gerçekleştirilmiştir'),
        (r'\bdahil etti\b','dahil etmiştir'),
        (r'\btest etti\b','test etmiştir'),
        (r'\badım attı\b','adım atmıştır'),
    ]
    for a,b in repl:
        s=re.sub(a,b,s,flags=re.I)
    return re.sub(r'\s+',' ',s).strip()

def _v101_build_intro(title,source,sents):
    """
    İlk cümle mutlaka 'kim/ne + ne oldu' bilgisini taşır.
    Başlığın kendisini yazmaz; haber gövdesindeki en iyi bağlam cümlesini seçer.
    """
    subject=_v100_subject_hint(title,source)
    scored=[]
    for i,s in enumerate(sents[:8]):
        n=norm(s)
        action=sum(x in n for x in [
            'açıkla','duyur','başlat','başvur','gerçekleştir','üret','teslim','envanter',
            'faaliyete','satış','seç','program','proje','oran','veri','rekor','görev'
        ])
        sc=8*_v101_sentence_overlap(title,s)+4*action+_v99_sentence_value(s)-i
        scored.append((sc,i,s))
    if not scored:
        return '',-1
    _,idx,s=max(scored,key=lambda x:x[0])
    s=_v101_formal_sentence(s)

    # "Proje...", "Tüketim...", "Bunun yanında..." gibi referansı belirsiz başlangıçları konuya bağla.
    n=norm(s)
    ambiguous=re.match(r'^(proje|program|tüketim|şirket|bu kapsamda|bunun yanında|ayrıca|yoğun ilgi)',n)
    if subject and (ambiguous or _v101_sentence_overlap(title,s)==0):
        if n.startswith('proje'):
            s=re.sub(r'^\s*Proje\s*,?\s*',subject+' kapsamında ',s,flags=re.I)
        elif n.startswith('program'):
            s=re.sub(r'^\s*Program\s*,?\s*',subject+' kapsamında ',s,flags=re.I)
        elif n.startswith('tüketim'):
            s=subject.rstrip(' .')+' kapsamında '+s[:1].lower()+s[1:]
        else:
            s=subject.rstrip(' .')+' kapsamında '+s[:1].lower()+s[1:]

    return s,idx

def _v101_semantic_key(s):
    """Yakın tekrarları, özellikle aynı rakamı tekrarlayan cümleleri azaltır."""
    n=norm(s)
    nums=tuple(re.findall(r'\d+(?:[.,]\d+)?',n))
    words=[w for w in re.findall(r'[a-zçğıöşü]+',n) if len(w)>4]
    return set(words),set(nums)

def _v101_is_near_duplicate(s,chosen):
    sw,sn=_v101_semantic_key(s)
    for prev in chosen:
        pw,pn=_v101_semantic_key(prev)
        word_sim=len(sw&pw)/max(1,len(sw|pw))
        num_sim=(bool(sn) and bool(pn) and len(sn&pn)/max(1,len(sn|pn))>=0.75)
        if word_sim>=0.48 or (num_sim and word_sim>=0.25):
            return True
    return False

def _v101_analyst_paragraph(title,source,body,fallback=''):
    """
    Her haber için bağımsız okunabilen 4 satırlık mini bilgi notu:
    GİRİŞ: kim/ne, hangi gelişme
    GELİŞME: kritik veri/rakam/yer/tarih
    SONUÇ: hedef, sonuç, mevcut durum veya sonraki aşama
    """
    title=_v101_clean_unicode(title)
    source=_v101_clean_unicode(source)
    body=_v99_clean_article_text(_v101_clean_unicode(body),source)
    fallback=_v99_clean_article_text(_v101_clean_unicode(fallback),source)

    candidate=body
    if not candidate or _v99_is_title_only(title,candidate):
        candidate=fallback

    raw=_v96_unique_sentences(title,candidate)
    sents=[]
    for s in raw:
        s=_v99_clean_article_text(_v101_clean_unicode(s),source)
        if _v101_bad_sentence(s) or _v99_is_title_only(title,s):
            continue
        sents.append(s)

    # İçerik gerçekten yoksa başlığı "özet" diye basma.
    if not sents:
        return ''

    intro,intro_idx=_v101_build_intro(title,source,sents)
    if not intro:
        return ''

    chosen=[intro]
    used={intro_idx}

    # Gelişme: rakam/istatistik/ölçek/yer/tarih taşıyan en güçlü cümle.
    detail_candidates=[]
    for i,s in enumerate(sents):
        if i in used: continue
        n=norm(s)
        data=bool(re.search(r'\d|%|yüzde|milyon|milyar|bin|adet|oran|tarih',s,re.I))
        sc=_v99_sentence_value(s)+5*data+2*_v101_sentence_overlap(title,s)
        if sc>3:
            detail_candidates.append((sc,i,s))
    for _,i,s in sorted(detail_candidates,reverse=True):
        fs=_v101_formal_sentence(s)
        if fs and not _v101_is_near_duplicate(fs,chosen):
            chosen.append(fs); used.add(i); break

    # Sonuç: hedef/son durum/takvim/etki.
    result_candidates=[]
    for i,s in enumerate(sents):
        if i in used: continue
        n=norm(s)
        hits=sum(x in n for x in [
            'hedef','plan','beklen','teslim','başvuru','envanter','faaliyete',
            'tamamlan','başlam','katkı','sonuç','rekor','artış','azalış','dönem',
            'seviye','ulaş','oluştur','sağla'
        ])
        sc=5*hits+_v99_sentence_value(s)+_v101_sentence_overlap(title,s)
        if hits:
            result_candidates.append((sc,i,s))
    for _,i,s in sorted(result_candidates,reverse=True):
        fs=_v101_formal_sentence(s)
        if fs and not _v101_is_near_duplicate(fs,chosen):
            chosen.append(fs); break

    # 2-3 tam cümle; 4 Word satırı hedefi. Cümle ortasında kesme yapılmaz.
    final=[]
    for s in chosen[:3]:
        s=_v101_formal_sentence(s).strip()
        if not s: continue
        if s[-1] not in '.!?': s+='.'
        cand=' '.join(final+[s])
        if final and len(cand)>620:
            break
        final.append(s)

    text=' '.join(final)
    text=_v101_clean_unicode(_v98_strip_site_name(text,source))
    # Cümle sonunda kalan başlık/site artıkları.
    text=re.sub(r'\s+[A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜ0-9\s\'’-]{8,}\s*!?\s*(?=\.|$)','',text)
    text=re.sub(r'\s+',' ',text).strip()
    return text

def make_important_basket_docx_v101(basket_df):
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))
    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_v101_clean_unicode(r.get('title',''))
        source=_v101_clean_unicode(r.get('source',''))
        summary=_v101_clean_unicode(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_v101_clean_unicode(r.get('news_time',''))

        detail={}
        try:
            detail=article_detail({
                'Başlık':title,'Kaynak':source,'URL':url,'Yayıncı_URL':url,
                'İçerik_Özeti':summary,'Tarih':news_time
            }) or {}
        except Exception:
            pass

        text=_v101_analyst_paragraph(title,source,detail.get('text') or '',summary)
        if not text and summary:
            text=_v101_analyst_paragraph(title,source,summary,summary)
        return text

    outputs=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try: outputs[idx]=fut.result()
                except Exception: outputs[idx]=''

    for idx,r in enumerate(rows):
        text=_v101_clean_unicode(outputs[idx] if idx<len(outputs) else '')
        if not text:
            # İçeriksiz başlığı sahte bir 4 satırlık özet haline getirmiyoruz.
            # Kullanıcıya Word içinde başlık kalıntısı göstermek yerine bu kayıt atlanır.
            continue
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.0
        p.paragraph_format.space_after=Pt(7)
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')
    bio=BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

# ======================= /V101: ÖGN KONU BÜTÜNLÜĞÜ 2.0 =========================
def make_important_basket_docx_v100(basket_df):
    """
    V100: Her haber kendi içinde anlamlı bir bütün oluşturur.
    Başlıktan konu bağlamı alınır ama başlık Word'e ayrı yazılmaz.
    'Proje...' gibi bağlamsız başlangıçlar düzeltilir.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_v98_safe_tr(r.get('title',''))
        source=_v98_safe_tr(r.get('source',''))
        summary=_v98_safe_tr(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_v98_safe_tr(r.get('news_time',''))

        detail={}
        try:
            detail=article_detail({
                'Başlık':title,'Kaynak':source,'URL':url,'Yayıncı_URL':url,
                'İçerik_Özeti':summary,'Tarih':news_time
            }) or {}
        except Exception:
            pass

        body=detail.get('text') or ''
        text=_v100_pick_summary_sentences(title,source,body,summary)

        # Tam metin sorunluysa yalnız kayıtlı özetle tekrar dene.
        if not text and summary:
            text=_v100_pick_summary_sentences(title,source,summary,summary)

        # Son çare: başlığı "başlık" olarak değil, konu cümlesine dönüştür.
        if not text:
            subject=_v100_subject_hint(title,source)
            if subject:
                text=f'{subject} kapsamında gelişmenin ayrıntılarına ilişkin haber içeriği sınırlı olduğundan, mevcut açık kaynak metninden doğrulanabilir ek bilgi çıkarılamamıştır.'
                text=_v99_officialize(text)
        return text

    outputs=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try:
                    outputs[idx]=fut.result()
                except Exception:
                    outputs[idx]=''

    for idx,r in enumerate(rows):
        text=_v98_safe_tr(outputs[idx] if idx<len(outputs) else '')
        if not text:
            continue

        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.0
        p.paragraph_format.space_after=Pt(7)
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')
    bio=BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

# ======================= /V100: KONU BÜTÜNLÜĞÜ =========================
def make_important_basket_docx_v99(basket_df):
    """
    V99 ÖGN Word:
    - sepetteki her haber ayrı işlenir;
    - web sayfası/UI gürültüsü atılır;
    - başlık/site adı basılmaz;
    - içerik bulunamayan haber başlık olarak Word'e sokulmaz;
    - resmî, analitik, 4 satıra yakın özet üretilir;
    - Türkçe mojibake temizliği son katmanda tekrar uygulanır.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_v98_safe_tr(r.get('title',''))
        source=_v98_safe_tr(r.get('source',''))
        summary=_v98_safe_tr(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_v98_safe_tr(r.get('news_time',''))

        detail={}
        try:
            detail=article_detail({
                'Başlık':title,'Kaynak':source,'URL':url,'Yayıncı_URL':url,
                'İçerik_Özeti':summary,'Tarih':news_time
            }) or {}
        except Exception:
            pass

        body=detail.get('text') or ''
        text=_v99_analyst_ogn(title,source,body,summary)

        # İçerik çekilememişse bir kez daha mevcut summary ile dene.
        if not text and summary:
            text=_v99_analyst_ogn(title,source,summary,summary)

        return text

    outputs=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try:
                    outputs[idx]=fut.result()
                except Exception:
                    outputs[idx]=''

    for idx,r in enumerate(rows):
        text=_v98_safe_tr(outputs[idx] if idx<len(outputs) else '')
        if not text:
            # Başlığı çıktı olarak kullanma. İçeriği olmayan haberi uyarı metniyle bozmak yerine atla.
            continue

        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.0
        p.paragraph_format.space_after=Pt(7)
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
# ======================= /V99: ÖGN ANALİST FİLTRESİ =========================
def make_important_basket_docx_v98(basket_df):
    """
    ÖGN Word:
    1) Sepetteki HER haber Word'e aktarılır.
    2) Her haber tek paragraf / yaklaşık 4 satır.
    3) Resmî dil.
    4) Türkçe karakter temizliği.
    5) Başlık yok.
    6) Site/yayıncı adı yok.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_v98_safe_tr(r.get('title',''))
        source=_v98_safe_tr(r.get('source',''))
        summary=_v98_safe_tr(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_v98_safe_tr(r.get('news_time',''))

        try:
            detail=article_detail({
                'Başlık':title,
                'Kaynak':source,
                'URL':url,
                'Yayıncı_URL':url,
                'İçerik_Özeti':summary,
                'Tarih':news_time
            })
        except Exception:
            detail={}

        body=_v98_safe_tr(detail.get('text') or summary or title)
        out=_v98_exact_four_line_summary(title,source,body)

        # Her haber mutlaka çıksın.
        if not out:
            out=_v98_exact_four_line_summary(title,source,summary or title)

        # Son çare: başlığı doğrudan yazmak yerine kısa resmî cümleye çevir.
        if not out:
            clean_title=_v98_remove_heading_like_text(title,title)
            if not clean_title:
                clean_title=title.capitalize() if title else 'Gelişme'
            out=_v66_formalize_sentence_endings(clean_title).strip()
            if out and out[-1] not in '.!?':
                out+='.'

        return out

    outputs=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try:
                    outputs[idx]=fut.result()
                except Exception:
                    r=rows[idx]
                    outputs[idx]=_v98_exact_four_line_summary(
                        r.get('title',''),
                        r.get('source',''),
                        r.get('summary','') or r.get('title','')
                    )

    # Her sepet haberi sırayla Word'e yazılır.
    for idx,r in enumerate(rows):
        text=outputs[idx] if idx<len(outputs) else ''
        text=_v98_safe_tr(text)

        # Hiçbir haber sessizce atlanmasın.
        if not text:
            title=_v98_safe_tr(r.get('title',''))
            text=_v66_formalize_sentence_endings(title.capitalize()).strip()
            if text and text[-1] not in '.!?':
                text+='.'

        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.0
        p.paragraph_format.space_after=Pt(7)

        # Başlık/link/site ayrıca yazılmaz.
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v97(basket_df):
    """
    Önemli Gelişmeler Word:
    Her haber TEK paragraf, yaklaşık 4 satır.
    Ana gelişme + kritik veri/rakam + sonuç.
    Başlık ve haber linki yazılmaz.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_clean_note_text(r.get('title',''))
        source=_clean_note_text(r.get('source',''))
        summary=_clean_note_text(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_clean_note_text(r.get('news_time',''))

        try:
            detail=article_detail({
                'Başlık':title,
                'Kaynak':source,
                'URL':url,
                'Yayıncı_URL':url,
                'İçerik_Özeti':summary,
                'Tarih':news_time
            })
        except Exception:
            detail={}

        body=_clean_note_text(detail.get('text') or summary or title)
        out=_v97_analyst_summary(title,body)

        if not out:
            out=_v97_analyst_summary(title,summary or title)

        return out

    outputs=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try:
                    outputs[idx]=fut.result()
                except Exception:
                    r=rows[idx]
                    outputs[idx]=_v97_analyst_summary(
                        r.get('title',''),
                        r.get('summary','') or r.get('title','')
                    )

    for text in outputs:
        text=_repair_mojibake_utf8(_clean_note_text(text)).strip()
        if not text:
            continue
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.0
        p.paragraph_format.space_after=Pt(7)
        # Her haber tek bütünleşik paragraf.
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v96(basket_df):
    """
    Önemli Gelişmeler Word = her haber için ayrı ayrı kısaltılmış bilgi notu.
    Başlık/link yazılmaz. Her haber en fazla 4 paragraftır.
    Kritik veri, rakam ve istatistikler yüksek öncelikle korunur.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_clean_note_text(r.get('title',''))
        source=_clean_note_text(r.get('source',''))
        summary=_clean_note_text(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_clean_note_text(r.get('news_time',''))

        # Bilgi Notu ile aynı içerik alma sistemi.
        try:
            detail=article_detail({
                'Başlık':title,
                'Kaynak':source,
                'URL':url,
                'Yayıncı_URL':url,
                'İçerik_Özeti':summary,
                'Tarih':news_time
            })
        except Exception:
            detail={}

        body=_clean_note_text(detail.get('text') or summary or title)
        paras=_v96_short_information_note(title,body)

        # Tam metin sorunluysa kayıtlı özet üzerinde aynı motoru tekrar çalıştır.
        if not paras:
            paras=_v96_short_information_note(title,summary or title)

        return paras

    outputs=[[] for _ in rows]
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try:
                    outputs[idx]=fut.result()
                except Exception:
                    r=rows[idx]
                    outputs[idx]=_v96_short_information_note(
                        r.get('title',''),
                        r.get('summary','') or r.get('title','')
                    )

    for item_index,paras in enumerate(outputs):
        if not paras:
            continue

        for para_index,text in enumerate(paras[:4]):
            text=_repair_mojibake_utf8(_clean_note_text(text)).strip()
            if not text:
                continue
            p=doc.add_paragraph()
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent=Cm(1.25)
            p.paragraph_format.line_spacing=1.0
            p.paragraph_format.space_after=Pt(4 if para_index<len(paras)-1 else 8)
            # (STB) yalnız haberin son paragrafında.
            suffix=' (STB).' if para_index==len(paras[:4])-1 else ''
            p.add_run(text.rstrip(' .;')+suffix)

    doc.add_paragraph('Arz olunur.')

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v95(basket_df):
    """
    Önemli Gelişmeler Word:
    - Haber işleme: Bilgi Notu gibi article_detail()
    - Özetleme: AKT gibi _akt_formal_summary()
    - Dil: mevcut V66 resmî dil normalizasyonu
    - Çıktı: başlıksız, linksiz, yaklaşık 4 satır.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def process_row(r):
        title=_clean_note_text(r.get('title',''))
        source=_clean_note_text(r.get('source',''))
        summary=_clean_note_text(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_clean_note_text(r.get('news_time',''))

        # Bilgi Notu ile aynı yaklaşım: mümkünse gerçek haber metnini al.
        try:
            detail=article_detail({
                'Başlık':title,
                'Kaynak':source,
                'URL':url,
                'Yayıncı_URL':url,
                'İçerik_Özeti':summary,
                'Tarih':news_time
            })
        except Exception:
            detail={}

        body=_clean_note_text(detail.get('text') or summary or title)

        # AKT'nin çalışan özet motorunu doğrudan kullan.
        result=_v95_ogn_from_existing_engines(title,body)

        # Tam metin tarafı sonuç vermezse AKT motorunu kayıtlı özet üzerinde çalıştır.
        if not result or len(result)<50:
            result=_v95_ogn_from_existing_engines(title,summary or title)

        return result

    # Bilgi notu/AKT içerik yaklaşımı korunurken Word beklemesini azaltmak için paralel oku.
    outputs=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(6,len(rows))) as ex:
            jobs={ex.submit(process_row,r):i for i,r in enumerate(rows)}
            for fut in concurrent.futures.as_completed(jobs):
                idx=jobs[fut]
                try:
                    outputs[idx]=fut.result()
                except Exception:
                    r=rows[idx]
                    outputs[idx]=_v95_ogn_from_existing_engines(
                        r.get('title',''),
                        r.get('summary','') or r.get('title','')
                    )

    for text in outputs:
        text=_clean_note_text(text)
        if not text:
            continue
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.line_spacing=1.0
        # Haber başlığı, kaynak adı veya URL ayrıca yazılmaz.
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v94(basket_df):
    """
    Reliable ÖGN Word generator.
    No headline, no URL, no bullet.
    Every basket row produces one paragraph.
    Primary source = summary already stored when the news was scanned.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    stl=doc.styles['Normal']
    stl.font.name='Times New Roman'
    stl.font.size=Pt(12)
    stl._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p0=doc.add_paragraph()
    p0.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p0.add_run(now.strftime('%d/%m/%Y'))

    p1=doc.add_paragraph()
    p1.add_run('Konu: ').bold=True
    p1.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    for r in rows:
        title=_v87_safe_tr(r.get('title',''))
        source=_v87_safe_tr(r.get('source',''))
        summary=_v87_safe_tr(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_v87_safe_tr(r.get('news_time',''))

        # Use the text already saved in the basket first: fast and stable.
        text=summary

        # Only if the stored summary is genuinely empty/too short, try the article.
        if len(text.strip())<60 and url:
            try:
                d=article_detail({
                    'Başlık':title,'Kaynak':source,'URL':url,'Yayıncı_URL':url,
                    'İçerik_Özeti':summary,'Tarih':news_time
                })
                fetched=_v87_safe_tr((d or {}).get('text',''))
                if len(fetched)>len(text):
                    text=fetched
            except Exception:
                pass

        out=_v94_formal_summary_text(text)

        # Absolute fallback: do not create an empty Word document.
        # If scan stored no summary and article could not be read, use a cleaned
        # sentence from the title rather than silently omitting the news.
        if not out:
            out=_v92_formal_sentence(title).strip()
            if out and out[-1] not in '.!?':
                out+='.'

        if not out:
            continue

        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.line_spacing=1.0
        p.add_run(out.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')
    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v93(basket_df):
    """STB referansına göre sıfırdan yazılmış Önemli Gelişmeler Word motoru."""
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    style=doc.styles['Normal']
    style.font.name='Times New Roman'
    style.font.size=Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y'))

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    rows=[] if basket_df is None else basket_df.to_dict('records')

    def one(r):
        title=_v87_safe_tr(r.get('title',''))
        source=_v87_safe_tr(r.get('source',''))
        fallback=_v87_safe_tr(r.get('summary',''))
        url=str(r.get('url','') or '')
        news_time=_v87_safe_tr(r.get('news_time',''))
        try:
            d=_v90_fetch_detail(title,source,url,fallback,news_time)
            body=_v87_safe_tr((d or {}).get('text','') or fallback)
        except Exception:
            body=fallback
        return _v93_build_summary(title,source,body,fallback)

    summaries=['']*len(rows)
    if rows:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(8,len(rows))) as ex:
            jobs={ex.submit(one,r):i for i,r in enumerate(rows)}
            for f in concurrent.futures.as_completed(jobs):
                i=jobs[f]
                try: summaries[i]=f.result()
                except Exception: summaries[i]=''

    for text in summaries:
        if not text:
            continue
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.line_spacing=1.0
        # Başlık yok, URL yok, madde imi yok.
        p.add_run(text.rstrip(' .;')+' (STB).')

    doc.add_paragraph('Arz olunur.')
    bio=BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v92(basket_df):
    """
    V92: ÖGN için sade ve izlenebilir akış.
    Haberleri paralel alır; her haber için tek gövde + ardışık 2-3 cümle.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y')).bold=True

    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    records=[] if basket_df is None else basket_df.to_dict('records')

    def process(item):
        title=_v87_safe_tr(item.get('title',''))
        source=_v87_safe_tr(item.get('source',''))
        fallback=_v87_safe_tr(item.get('summary',''))
        url=str(item.get('url','') or '')
        news_time=_v87_safe_tr(item.get('news_time',''))

        detail=_v90_fetch_detail(title,source,url,fallback,news_time)
        body=_v87_safe_tr((detail or {}).get('text','') or fallback)

        txt=_v92_summary(title,source,body,fallback)
        if not txt:
            txt=_v92_summary(title,source,fallback,fallback)
        return txt

    summaries=['']*len(records)
    if records:
        workers=min(8,max(1,len(records)))
        with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as ex:
            fmap={ex.submit(process,r):i for i,r in enumerate(records)}
            for fut in concurrent.futures.as_completed(fmap):
                idx=fmap[fut]
                try:
                    summaries[idx]=fut.result()
                except Exception:
                    summaries[idx]=''

    for txt in summaries:
        if not txt:
            continue
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.line_spacing=1.0
        p.add_run(_v87_safe_tr(txt).rstrip(' .;')+' (STB).')

    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8)
    p.add_run('Arz olunur.')

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_important_basket_docx_v90(basket_df):
    """
    V90 ÖGN motoru. Eski Word baytlarını/fonksiyonlarını kullanmaz.
    Her haber için gerçek metni paralel alır, sırayı korur.
    """
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'; normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y')).bold=True
    p=doc.add_paragraph()
    p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    records=[] if basket_df is None else basket_df.to_dict('records')
    if not records:
        doc.add_paragraph('Kayıtlı önemli gelişme bulunmamaktadır.')
    else:
        def process(item):
            title=_v87_safe_tr(item.get('title',''))
            source=_v87_safe_tr(item.get('source',''))
            fallback=_v87_safe_tr(item.get('summary',''))
            url=str(item.get('url','') or '')
            news_time=_v87_safe_tr(item.get('news_time',''))

            detail=_v90_fetch_detail(title,source,url,fallback,news_time)
            body=_v87_safe_tr((detail or {}).get('text','') or fallback)
            txt=_v90_item_summary(title,source,body,fallback)

            # Asla eski başlık-çıktı davranışına dönme.
            if not txt:
                # fallback gövdesinden tek resmî cümle oluşturmayı tekrar dene.
                txt=_v90_item_summary(title,source,fallback,fallback)
            if not txt:
                # Son çare: başlığı değil, açıklayıcı bir kurum cümlesi oluştur.
                clean_title=_v90_clean_title(title,source)
                txt=f'{clean_title} konusuna ilişkin gelişme açık kaynaklarda yer almıştır.'
                txt=_v90_formalize(txt)
            return txt

        summaries=['']*len(records)
        workers=min(8,max(1,len(records)))
        with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as ex:
            fmap={ex.submit(process,r):i for i,r in enumerate(records)}
            for fut in concurrent.futures.as_completed(fmap):
                idx=fmap[fut]
                try:
                    summaries[idx]=fut.result()
                except Exception:
                    summaries[idx]=''

        for rr,txt in zip(records,summaries):
            if not txt:
                continue
            p=doc.add_paragraph()
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after=Pt(6)
            p.paragraph_format.line_spacing=1.0
            p.add_run(_v87_safe_tr(txt).rstrip(' .;')+' (STB).')

    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8)
    p.add_run('Arz olunur.')

    bio=BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()
def make_important_basket_docx(basket_df):
    """
    V88:
    - article fetches run in parallel instead of one-by-one,
    - results cached for 1 hour,
    - output order remains basket order,
    - no item is silently dropped.
    """
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    normal=doc.styles['Normal']
    normal.font.name='Times New Roman'; normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    now=datetime.now().astimezone()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(now.strftime('%d/%m/%Y')).bold=True
    p=doc.add_paragraph(); p.add_run('Konu: ').bold=True
    p.add_run('STB Temsilciliği Önemli Gelişmeler Notu')

    if basket_df is None or basket_df.empty:
        doc.add_paragraph('Kayıtlı önemli gelişme bulunmamaktadır.')
    else:
        records=basket_df.to_dict('records')

        def fetch_one(item):
            title=_v87_safe_tr(item.get('title',''))
            source=_v87_safe_tr(item.get('source',''))
            fallback=_v87_safe_tr(item.get('summary',''))
            url=str(item.get('url','') or '')
            news_time=_v87_safe_tr(item.get('news_time',''))

            # If saved summary is already substantial, don't delay Word just to fetch again.
            # Full article is requested mainly for short/snippet-like summaries.
            detail={}
            if len(fallback)<380:
                detail=_v88_cached_article_detail(title,source,url,fallback,news_time)
            body=_v87_safe_tr((detail or {}).get('text','') or fallback)
            return _v88_summary(title,source,body,fallback)

        summaries=['']*len(records)
        max_workers=min(6,max(1,len(records)))
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
            futmap={ex.submit(fetch_one,r):idx for idx,r in enumerate(records)}
            for fut in concurrent.futures.as_completed(futmap):
                idx=futmap[fut]
                try:
                    summaries[idx]=fut.result()
                except Exception:
                    rr=records[idx]
                    summaries[idx]=_v88_summary(
                        rr.get('title',''),rr.get('source',''),
                        rr.get('summary',''),rr.get('summary','')
                    )

        for rr,txt in zip(records,summaries):
            if not txt:
                txt=_v88_formal(_v87_safe_tr(rr.get('summary','') or rr.get('title','')))
            p=doc.add_paragraph()
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after=Pt(5)
            p.paragraph_format.line_spacing=1.0
            p.add_run(_v87_safe_tr(txt).rstrip(' .;')+' (STB).')

    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    p.add_run('Arz olunur.')
    bio=BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

# -----------------------------
# GÜNLÜK DURUM ÖZETİ — V32 EK MODÜL
# V31 çekirdek tarama / risk / alarm / bilgi notu fonksiyonlarına dokunmaz.
# -----------------------------
def _daily_summary_stats(df):
    x=df.copy()
    if x.empty:
        return {}

    neg=int((x['Duygu']=='Negatif').sum()) if 'Duygu' in x else 0
    high=int((x['Risk_Durumu']=='Yüksek Risk').sum()) if 'Risk_Durumu' in x else 0

    osb=0
    if 'Başlık' in x:
        for _,r in x.iterrows():
            if is_osb_fire(r.get('Başlık',''),r.get('İçerik_Özeti','')):
                osb+=1

    def count_terms(terms):
        c=0
        for _,r in x.iterrows():
            text=norm(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')} {r.get('Kategori','')}")
            if any(t in text for t in terms):
                c+=1
        return c

    investment=count_terms(['yatırım','yatirim','fabrika aç','tesis aç','kapasite art','yeni tesis','teşvik','tesvik'])
    defence=count_terms(['savunma','aselsan','tusaş','tusas','roketsan','baykar','havelsan','saha expo','iha','siha','füze','fuze'])
    cyber=count_terms(['siber','veri sızınt','veri sizint','fidye yazılım','fidye yazilim','hack','siber saldır','siber saldir'])

    return {
        'total':len(x),
        'negative':neg,
        'high_risk':high,
        'osb_fire':osb,
        'investment':investment,
        'defence':defence,
        'cyber':cyber
    }


def _daily_top_events(df, n=5):
    """
    V62: Sabah ilk bakılacak gelişmeleri seçer.
    Negatiflik tek başına belirleyici değildir. Stratejik sanayi-teknoloji ilgisi,
    ekonomik/kurumsal etki, resmî teyit, çoklu kaynak, yenilik ve risk birlikte puanlanır.
    """
    if df.empty:
        return df.copy()

    x=df.copy()
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')
    strategic_terms=[
        'yatırım','üretim','ihracat','ithalat','kapasite','fabrika','tesis','osb',
        'savunma','tusaş','aselsan','roketsan','havelsan','baykar','kaan',
        'yapay zeka','yapay zekâ','çip','yarı iletken','siber','teknoloji',
        'arge','ar-ge','tübitak','kosgeb','patent','togg','otomotiv','enerji',
        'kritik mineral','uzay','uydu','teknofest','sanayi üretimi'
    ]
    high_value_terms=[
        'milyar','milyon','rekor','anlaşma','sözleşme','yatırım','teşvik',
        'ihracat','üretim','kapasite','lansman','ilk kez','yeni tesis',
        'stratejik','program','eylem planı','resmi gazete','resmî gazete'
    ]
    low_relevance_terms=[
        'trafik kazası','magazin','spor','dualarla anıldı','hayatını kaybeden muhabir'
    ]

    def importance(r):
        text=norm(f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')} {r.get('Kategori','')}")
        score=0

        # Sanayi-teknoloji alanına doğrudan ilgi en güçlü ölçüt.
        score += min(sum(1 for k in strategic_terms if k in text)*8,40)
        score += min(sum(1 for k in high_value_terms if k in text)*5,20)

        cat=norm(r.get('Kategori',''))
        if any(k in cat for k in ['savunma','sanayi','üretim','dijital','yapay zeka','yapay zekâ',
                                  'otomotiv','uzay','enerji','teknoloji']):
            score+=18

        # Risk önemlidir ama negatiflik listeyi ele geçirmez.
        risk=int(r.get('Risk_Skoru',0) or 0)
        score+=min(risk//4,20)
        if r.get('Risk_Durumu')=='Yüksek Risk':
            score+=12
        if r.get('Duygu')=='Negatif':
            score+=5

        if critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti','')):
            score+=25

        try:
            score+=min(int(r.get('Olay_Kaynak_Sayisi',0) or 0)*5,20)
        except Exception:
            pass

        verification=norm(r.get('Doğrulama',''))
        if 'resmi' in verification or 'resmî' in verification or 'birincil' in verification:
            score+=22
        elif 'çoklu kaynak' in verification or 'coklu kaynak' in verification:
            score+=14

        if any(k in text for k in low_relevance_terms):
            score-=30

        return score

    x['_Önem']=x.apply(importance,axis=1)

    if 'Olay_ID' in x.columns:
        x=x.sort_values(['_Önem','Tarih_dt'],ascending=[False,False],na_position='last')
        x=x.drop_duplicates(subset=['Olay_ID'],keep='first')
    else:
        x=x.sort_values(['_Önem','Tarih_dt'],ascending=[False,False],na_position='last')

    return x.head(n).drop(columns=['_Önem'],errors='ignore')


def _daily_summary_text(df):
    stats=_daily_summary_stats(df)
    top=_daily_top_events(df,5)
    if not stats:
        return '',top,stats

    intro=(
        f"Sanayi ve teknoloji alanında gerçekleştirilen güncel açık kaynak taramasında toplam {stats['total']} haber tespit edilmiştir. "
        f"Bunların {stats['negative']} adedi negatif içerik, {stats['high_risk']} adedi yüksek riskli gelişme olarak sınıflandırılmıştır. "
        f"Tarama kapsamında {stats['osb_fire']} organize sanayi bölgesi yangını, {stats['investment']} yatırım/kapasite gelişmesi, "
        f"{stats['defence']} savunma sanayii bağlantılı içerik ve {stats['cyber']} siber güvenlik bağlantılı içerik belirlenmiştir."
    )

    paras=[intro]
    if not top.empty:
        paras.append(
            "Günün genel görünümünde öne çıkan gelişmeler; güncellik, risk düzeyi, kaynak teyidi ve sanayi-teknoloji alanına muhtemel etkileri "
            "birlikte dikkate alınarak aşağıda özetlenmiştir."
        )
        for i,(_,r) in enumerate(top.iterrows(),1):
            title=_clean_note_text(r.get('Başlık',''))
            source=_clean_note_text(r.get('Kaynak','Açık Kaynak'))
            when=_clean_note_text(r.get('Tarih',''))
            content=_clean_note_text(r.get('İçerik_Özeti',''))

            # Başlığı tekrar etmek yerine içerikten anlamlı cümleleri seç.
            sents=_detail_sentences(content,title)
            useful=[]
            seen=set()
            for s in sents:
                s=_clean_note_text(s)
                key=norm(s)
                if not s or len(s)<35 or key in seen:
                    continue
                seen.add(key)
                useful.append(s)
                if len(useful)>=4:
                    break

            detail=_join_sentences_naturally(useful) if useful else content[:700].strip()
            risk=int(r.get('Risk_Skoru',0) or 0)
            status=_clean_note_text(r.get('Risk_Durumu',''))
            category=_clean_note_text(r.get('Kategori',''))

            p=f"{i}. {when} tarihinde {source} kaynaklı gelişmede, {detail}" if detail else f"{i}. {when} tarihinde {source} kaynaklı “{title}” başlıklı gelişme öne çıkmıştır."
            if p and p[-1] not in '.!?':
                p+='.'
            if category:
                p+=f" Gelişme sistemde {category} başlığı altında izlenmektedir."
            if risk:
                p+=f" Risk puanı {risk}/100"
                if status:
                    p+=f" ve risk durumu {status}"
                p+=" olarak değerlendirilmiştir."
            paras.append(p)

    # Günlük tabloya dair kısa analitik kapanış.
    emphasis=[]
    if stats['high_risk']:
        emphasis.append(f"{stats['high_risk']} yüksek riskli gelişmenin")
    if stats['negative']:
        emphasis.append(f"{stats['negative']} negatif içeriğin")
    if stats['investment']:
        emphasis.append(f"{stats['investment']} yatırım/kapasite gelişmesinin")
    if stats['defence']:
        emphasis.append(f"{stats['defence']} savunma sanayii gelişmesinin")
    if stats['cyber']:
        emphasis.append(f"{stats['cyber']} siber güvenlik gelişmesinin")

    if emphasis:
        focus=', '.join(emphasis[:-1]) + ((' ve '+emphasis[-1]) if len(emphasis)>1 else emphasis[0])
        conclusion=(
            f"Günlük görünümde özellikle {focus} takip edilmesi gereken başlıklar arasında bulunduğu değerlendirilmektedir. "
            "Yeni resmî açıklamalar, üretim ve tedarik zincirine olası etkiler ile farklı açık kaynaklardan gelecek teyitlerin izlenmesi önem taşımaktadır."
        )
    else:
        conclusion=(
            "Günlük görünümde belirgin bir yüksek risk yoğunlaşması görülmemekle birlikte, yeni resmî açıklamalar ile üretim, yatırım, "
            "tedarik zinciri ve teknoloji alanındaki gelişmelerin izlenmesinin sürdürülmesi önem taşımaktadır."
        )
    paras.append(conclusion)
    return '\n\n'.join(paras),top,stats


def make_daily_summary_docx(df):
    text,top,stats=_daily_summary_text(df)

    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    doc.styles['Normal'].font.name='Times New Roman'
    doc.styles['Normal'].font.size=Pt(11)

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('GÜNLÜK SANAYİ VE TEKNOLOJİ DURUM ÖZETİ')
    r.bold=True; r.font.size=Pt(14)

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().astimezone().strftime('%d.%m.%Y %H:%M'))

    for block in text.split('\n\n'):
        bp=doc.add_paragraph()
        bp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.first_line_indent=Cm(1.25)
        bp.paragraph_format.line_spacing=1.15
        bp.paragraph_format.space_after=Pt(8)
        bp.add_run(block)

    if not top.empty:
        hp=doc.add_paragraph()
        rr=hp.add_run('ÖNE ÇIKAN GELİŞMELERİN KAYNAKLARI')
        rr.bold=True
        for i,(_,row) in enumerate(top.iterrows(),1):
            p=doc.add_paragraph()
            p.add_run(f"{i}. {_clean_note_text(row.get('Kaynak','Açık Kaynak'))} — {_clean_note_text(row.get('Başlık',''))}")
            if row.get('URL'):
                p.add_run(' — ')
                _word_hyperlink(p,row['URL'],'Haber linki')

    bio=BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.getvalue()

# -----------------------------
# DOCX — AKT / Açık Kaynak Taraması formatı
# Tarama motoru korunur. Yalnızca seçilen haberlerin rapora aktarılması değiştirilmiştir.
# -----------------------------
@st.cache_data(ttl=1800, show_spinner=False)
def article_detail(row):
    """
    Seçilen kayıt için gerçek yayıncı URL'sini ve gerçek haber sayfasını bulur.
    Google News'in kodlanmış RSS bağlantıları doğrudan yayıncı adresi değilse
    sırasıyla decoder, HTTP redirect, GDELT ve DuckDuckGo üzerinden çözülür.
    """
    if isinstance(row, str):
        row = {"URL": row}

    original_url = str(row.get("URL") or "").strip()
    fallback_title = str(row.get("Başlık") or "").strip()
    fallback_snippet = str(row.get("İçerik_Özeti") or "").strip()
    publisher_url = str(row.get("Yayıncı_URL") or "").strip()
    publisher_name = str(row.get("Yayıncı") or row.get("Kaynak") or "").strip()

    out = {
        "title": fallback_title,
        "canonical": original_url,
        "published": str(row.get("Tarih") or ""),
        "text": fallback_snippet,
        "images": [],
        "source": publisher_name,
    }

    def is_google(u):
        try:
            h = urlparse(u).netloc.lower()
            return h == "news.google.com" or h.endswith(".google.com")
        except Exception:
            return False

    def valid_article_url(u):
        if not u or not u.startswith("http"):
            return False
        h = urlparse(u).netloc.lower()
        return h not in {"news.google.com", "www.google.com", "google.com"} and "google.com" not in h

    def fetch_page(u):
        try:
            rr = requests.get(
                u,
                headers={
                    **HEADERS,
                    "Accept-Language": "tr-TR,tr;q=0.9,en;q=0.7",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                },
                timeout=12,
                allow_redirects=True,
            )
            if rr.status_code >= 400 or not rr.text:
                return None, None
            return rr, BeautifulSoup(rr.text, "html.parser")
        except Exception:
            return None, None

    def decode_with_package(u):
        try:
            from googlenewsdecoder import gnewsdecoder
            result = gnewsdecoder(u, interval=0.2)
            if isinstance(result, dict) and result.get("status"):
                decoded = result.get("decoded_url")
                if valid_article_url(decoded):
                    return decoded
        except Exception:
            pass
        return ""

    def decode_with_http(u):
        rr, soup = fetch_page(u)
        if rr and valid_article_url(rr.url):
            return rr.url

        if soup:
            for attrs in (
                {"property": "og:url"},
                {"name": "twitter:url"},
            ):
                tag = soup.find("meta", attrs=attrs)
                if tag and valid_article_url(tag.get("content", "")):
                    return requests.compat.urljoin(rr.url, tag["content"])

            tag = soup.find("link", rel=lambda x: x and "canonical" in str(x).lower())
            if tag and valid_article_url(requests.compat.urljoin(rr.url, tag.get("href", ""))):
                return requests.compat.urljoin(rr.url, tag.get("href"))

        return ""

    def decode_with_search(title):
        if not title:
            return ""

        # Önce GDELT: sonuçlar doğrudan yayıncı URL'si verir.
        try:
            q = '"' + title.replace('"', " ")[:240] + '"'
            r = requests.get(
                "https://api.gdeltproject.org/api/v2/doc/doc",
                params={
                    "query": q,
                    "mode": "artlist",
                    "maxrecords": 20,
                    "format": "json",
                    "sort": "HybridRel",
                    "timespan": "30d",
                },
                headers=HEADERS,
                timeout=8,
            )
            if r.ok:
                arts = r.json().get("articles", []) or []
                target = norm(title)
                for art in arts:
                    u = art.get("url") or ""
                    t = norm(art.get("title") or "")
                    if valid_article_url(u):
                        # Exact/near exact başlık eşleşmesi öncelikli.
                        if target and (target in t or t in target):
                            return u
                for art in arts:
                    u = art.get("url") or ""
                    if valid_article_url(u):
                        return u
        except Exception:
            pass

        # Son fallback: DuckDuckGo doğrudan yayıncı URL'si döndürebilir.
        try:
            from ddgs import DDGS
        except Exception:
            try:
                from duckduckgo_search import DDGS
            except Exception:
                DDGS = None

        if DDGS:
            try:
                with DDGS() as d:
                    results = list(d.text(f'"{title}"', region="tr-tr", timelimit="m", max_results=8))
                target = norm(title)
                for item in results:
                    u = item.get("href") or item.get("url") or ""
                    t = norm(item.get("title") or "")
                    if valid_article_url(u) and target and (target in t or t in target):
                        return u
                for item in results:
                    u = item.get("href") or item.get("url") or ""
                    if valid_article_url(u):
                        return u
            except Exception:
                pass

        return ""

    # 1) Google News bağlantısını çöz.
    real_url = ""
    if is_google(original_url):
        real_url = decode_with_package(original_url)
        if not real_url:
            real_url = decode_with_http(original_url)
        if not real_url:
            real_url = decode_with_search(fallback_title)
    elif valid_article_url(original_url):
        real_url = original_url
    else:
        real_url = decode_with_search(fallback_title)

    # 2) Gerçek sayfayı indir.
    rr, soup = fetch_page(real_url) if real_url else (None, None)

    if rr and soup:
        out["canonical"] = real_url or rr.url

        # Canonical
        can = soup.find("link", rel=lambda x: x and "canonical" in str(x).lower())
        if can and can.get("href"):
            out["canonical"] = requests.compat.urljoin(rr.url, can["href"])
        else:
            ogurl = soup.find("meta", attrs={"property": "og:url"})
            if ogurl and ogurl.get("content"):
                out["canonical"] = requests.compat.urljoin(rr.url, ogurl["content"])

        # Başlık
        for attrs in (
            {"property": "og:title"},
            {"name": "twitter:title"},
        ):
            t = soup.find("meta", attrs=attrs)
            if t and t.get("content"):
                out["title"] = t["content"].strip()
                break
        if not out["title"] and soup.title:
            out["title"] = soup.title.get_text(" ", strip=True)

        # Yayıncı
        for attrs in (
            {"property": "og:site_name"},
            {"name": "application-name"},
        ):
            t = soup.find("meta", attrs=attrs)
            if t and t.get("content"):
                out["source"] = t["content"].strip()
                break

        # Tarih
        for attrs in (
            {"property": "article:published_time"},
            {"itemprop": "datePublished"},
            {"name": "date"},
            {"name": "pubdate"},
        ):
            t = soup.find("meta", attrs=attrs)
            if t and t.get("content"):
                out["published"] = t["content"].strip()
                break

        bodies = []
        images = []

        def walk_json(obj):
            if isinstance(obj, dict):
                typ = str(obj.get("@type", "")).lower()
                if "article" in typ or "news" in typ:
                    if obj.get("headline"):
                        out["title"] = str(obj["headline"])
                    if obj.get("datePublished"):
                        out["published"] = str(obj["datePublished"])
                    if obj.get("articleBody"):
                        bodies.append(str(obj["articleBody"]))
                    pub = obj.get("publisher")
                    if isinstance(pub, dict) and pub.get("name"):
                        out["source"] = str(pub["name"])
                    im = obj.get("image") or obj.get("thumbnailUrl")
                    if isinstance(im, str):
                        images.append(im)
                    elif isinstance(im, list):
                        for x in im:
                            if isinstance(x, str):
                                images.append(x)
                            elif isinstance(x, dict) and x.get("url"):
                                images.append(str(x["url"]))
                    elif isinstance(im, dict) and im.get("url"):
                        images.append(str(im["url"]))
                for v in obj.values():
                    walk_json(v)
            elif isinstance(obj, list):
                for x in obj:
                    walk_json(x)

        for tag in soup.find_all("script", attrs={"type": re.compile(r"application/ld\+json", re.I)}):
            try:
                raw = tag.string or tag.get_text()
                if raw:
                    walk_json(json.loads(raw))
            except Exception:
                pass

        for attrs in (
            {"property": "og:image"},
            {"property": "og:image:url"},
            {"name": "twitter:image"},
            {"name": "twitter:image:src"},
        ):
            t = soup.find("meta", attrs=attrs)
            if t and t.get("content"):
                images.append(requests.compat.urljoin(rr.url, t["content"]))

        selectors = [
            '[itemprop="articleBody"]',
            "article",
            '[class*="article-body"]',
            '[class*="article-content"]',
            '[class*="news-content"]',
            '[class*="news-detail"]',
            '[class*="story-body"]',
            '[class*="post-content"]',
            '[class*="entry-content"]',
            '[class*="content-body"]',
            "main",
        ]
        for selector in selectors:
            for node in soup.select(selector)[:4]:
                parts = []
                for p in node.find_all(["p", "h2", "h3", "li"]):
                    txt = p.get_text(" ", strip=True)
                    if len(txt) >= 40:
                        parts.append(txt)
                if parts:
                    candidate = " ".join(parts)
                    if len(candidate) >= 250:
                        bodies.append(candidate)

        if not bodies:
            for p in soup.find_all("p"):
                txt = p.get_text(" ", strip=True)
                if len(txt) >= 45:
                    bodies.append(txt)

        for img in soup.find_all("img"):
            for attr in ("src", "data-src", "data-lazy-src", "data-original", "data-image"):
                value = img.get(attr)
                if value:
                    images.append(requests.compat.urljoin(rr.url, value))

        # Temizle
        seen=set()
        out["images"]=[]
        for u in images:
            if not isinstance(u,str): continue
            u=u.strip()
            if not u or u in seen: continue
            if any(x in u.lower() for x in ("favicon","sprite","avatar","logo")): continue
            seen.add(u); out["images"].append(u)
            if len(out["images"]) >= 20: break

        texts=[]
        seen_t=set()
        for body in bodies:
            body=re.sub(r"\s+"," ",html.unescape(body)).strip()
            if len(body)<120: continue
            key=norm(body[:700])
            if key in seen_t: continue
            seen_t.add(key); texts.append(body)
        # V92: Birden fazla article/main/content bloğunu BİRLEŞTİRME.
        # Aynı sayfadaki önerilen haberler ve tekrar blokları ÖGN'ye karışmasın.
        # Tek, en kapsamlı gövdeyi kullan.
        texts.sort(key=len, reverse=True)
        if texts:
            out["text"]=texts[0][:18000]

    # 3) Sayfa erişilemediyse bile RSS kaydını çöp etmiyoruz.
    # Generic Google Haberler adını asla gerçek yayıncı olarak rapora yazma.
    generic = {"google haberler","google news","google","google news rss","rss"}
    if norm(out["source"]) in generic:
        if publisher_name and norm(publisher_name) not in generic:
            out["source"] = publisher_name
        elif publisher_url:
            out["source"] = urlparse(publisher_url).netloc.replace("www.", "")
        else:
            out["source"] = "Açık Kaynak"

    # Başlık generic ise snippet/ekran başlığı kullan.
    if norm(out["title"]) in generic or not out["title"]:
        out["title"] = fallback_title or fallback_snippet

    if not out["text"] or len(out["text"]) < 250:
        out["text"] = fallback_snippet or out["title"]

    # Eğer gerçek URL çözüldüyse onu kullan; çözülmediyse Google News linkini rapora koyma.
    if not valid_article_url(out["canonical"]):
        out["canonical"] = publisher_url or original_url

    return out

def _download_report_image(url):
    if not url:
        return None
    try:
        rr = requests.get(
            url,
            headers={
                **HEADERS,
                "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.7",
            },
            timeout=12,
        )
        if rr.status_code != 200 or len(rr.content) < 1200:
            return None

        im = Image.open(BytesIO(rr.content))
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        im.thumbnail((1600, 1200), Image.LANCZOS)

        bio = BytesIO()
        im.save(bio, "JPEG", quality=88)
        bio.seek(0)
        return bio
    except Exception:
        return None


def _word_hyperlink(paragraph, url, label):
    if not url:
        paragraph.add_run(label)
        return

    try:
        rid = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rid)

        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "Hyperlink")
        rpr.append(rstyle)
        run.append(rpr)

        text = OxmlElement("w:t")
        text.text = label
        run.append(text)
        hyperlink.append(run)

        paragraph._p.append(hyperlink)
    except Exception:
        paragraph.add_run(url)


# -----------------------------
# V43 — TAM HABER METNİNE GÖRE NEGATİF/RİSK ANALİZİ
# -----------------------------
def _deep_negative_reclassify(rows, max_workers=14):
    """
    Her haberi mümkünse gerçek haber sayfasındaki tam metinle yeniden sınıflandırır.
    Sayfaya erişilemezse mevcut başlık + kısa içerik fallback olur.

    Yalnızca negatif/risk alanları güncellenir; kategori, olay kümeleri ve diğer
    çalışan modüller korunur.
    """
    if not rows:
        return rows, {'tam_metin':0,'kisa_icerik':0,'hata':0}

    results=[None]*len(rows)
    stats={'tam_metin':0,'kisa_icerik':0,'hata':0}

    def one(idx,row):
        try:
            detail=article_detail(row)
            full_text=re.sub(r'\s+',' ',str(detail.get('text') or '')).strip()
            snippet=re.sub(r'\s+',' ',str(row.get('İçerik_Özeti') or '')).strip()

            # article_detail erişemezse fallback olarak snippet döndürebilir.
            is_full=bool(full_text) and len(full_text)>=max(450,len(snippet)+180)
            analysis_text=full_text if is_full else (snippet or full_text or row.get('Başlık',''))

            sentiment,score,status,neg,risk,_cat,reasons=classify(
                row.get('Başlık',''),
                analysis_text,
                row.get('Domain','')
            )

            return idx,{
                'Duygu':sentiment,
                'Skor':score,
                'Risk_Skoru':score,
                'Risk_Durumu':status,
                'Risk_Gerekçesi':'; '.join(reasons),
                'Negatif_Sinyaller':neg,
                'Risk_Sinyalleri':risk,
                'Negatif_Analiz_Kapsamı':'Tam haber metni' if is_full else 'Başlık + kısa içerik',
                '_is_full':is_full
            }
        except Exception:
            sentiment,score,status,neg,risk,_cat,reasons=classify(
                row.get('Başlık',''),
                row.get('İçerik_Özeti',''),
                row.get('Domain','')
            )
            return idx,{
                'Duygu':sentiment,
                'Skor':score,
                'Risk_Skoru':score,
                'Risk_Durumu':status,
                'Risk_Gerekçesi':'; '.join(reasons),
                'Negatif_Sinyaller':neg,
                'Risk_Sinyalleri':risk,
                'Negatif_Analiz_Kapsamı':'Başlık + kısa içerik',
                '_is_full':False,
                '_error':True
            }

    with concurrent.futures.ThreadPoolExecutor(max_workers=min(max_workers,len(rows))) as ex:
        futures=[ex.submit(one,i,r.copy()) for i,r in enumerate(rows)]
        for fut in concurrent.futures.as_completed(futures):
            try:
                idx,data=fut.result()
                results[idx]=data
            except Exception:
                pass

    out=[]
    for i,row in enumerate(rows):
        r=row.copy()
        data=results[i]
        if data:
            if data.pop('_is_full',False):
                stats['tam_metin']+=1
            else:
                stats['kisa_icerik']+=1
            if data.pop('_error',False):
                stats['hata']+=1
            r.update(data)
        else:
            stats['kisa_icerik']+=1
            r['Negatif_Analiz_Kapsamı']='Başlık + kısa içerik'
        out.append(r)

    return out,stats


def _real_source(row, detail, real_url):
    generic = {"google haberler", "google news", "google", "google news rss", "rss"}

    for value in (
        detail.get("source"),
        row.get("Yayıncı"),
        row.get("Kaynak"),
    ):
        value = str(value or "").strip()
        if value and norm(value) not in generic:
            return value

    for value in (row.get("Yayıncı_URL"), real_url):
        value = str(value or "").strip()
        if valid_host := (urlparse(value).netloc.lower().replace("www.", "") if value else ""):
            if "google.com" not in valid_host:
                return valid_host

    return "Açık Kaynak"

def _akt_clean_sentences(title, body):
    text=BeautifulSoup(str(body or ''),'html.parser').get_text(' ',strip=True)
    text=re.sub(r'\s+',' ',text).strip()
    if not text:
        return []

    raw=re.split(r'(?<=[.!?])\s+',text)
    title_n=norm(title)
    boiler=[
        'çerez','cookie','abonelik','abone ol','reklam','tüm hakları saklıdır',
        'gizlilik politikası','kullanım koşulları','google news','bildirimleri aç',
        'uygulamamızı indirin','facebook','instagram','whatsapp','twitter',
        'son dakika haberleri için','haberlerimizi takip','ilgili haberler',
        'öne çıkan haberler','etiketler','yorumlar'
    ]

    kept=[]
    token_sets=[]
    for s in raw:
        s=re.sub(r'\s+',' ',s).strip()
        sn=norm(s)
        if len(s)<32 or sn==title_n:
            continue
        if any(b in sn for b in boiler):
            continue

        toks={x for x in re.findall(r'\w+',sn) if len(x)>2}
        if not toks:
            continue

        duplicate=False
        for old in token_sets[-20:]:
            inter=len(toks & old); union=len(toks | old)
            if union and inter/union>=0.78:
                duplicate=True
                break
        if duplicate:
            continue

        kept.append(s)
        token_sets.append(toks)

    return kept

def _akt_sentence_score(s):
    n=norm(s)
    score=0
    if re.search(r'\b\d+(?:[.,]\d+)?\b',s): score+=4
    if '%' in s or 'yüzde' in n: score+=3
    if any(x in n for x in ['açıkladı','belirtti','bildirdi','kaydetti','duyurdu','ifade etti','vurguladı']): score+=2
    if any(x in n for x in ['arttı','azaldı','geriledi','yükseldi','düştü','ulaştı','çıktı','indi','daraldı','büyüdü']): score+=3
    if any(x in n for x in ['üretim','ihracat','ithalat','istihdam','kapasite','yatırım','hasar','etkilendi','müşteri','tesis','fabrika']): score+=2
    if any(x in n for x in ['nedeni','sonucu','buna göre','bu kapsamda','öte yandan','ayrıca','son olarak']): score+=1
    return score

def _akt_formal_summary(title, body, max_sentences=10, max_chars=2800):
    """
    Haber başından sonuna okunur:
    - tekrar/menü temizlenir,
    - başlangıçtan ilk önemli bilgiler,
    - ortadaki en güçlü veri/açıklamalar,
    - sondaki sonuç/son durum birlikte seçilir,
    - orijinal haber sırası korunur.
    """
    sentences=_akt_clean_sentences(title,body)
    if not sentences:
        fallback=re.sub(r'\s+',' ',str(body or title or '')).strip()
        return fallback[:max_chars].rstrip(' .;')

    n=len(sentences)
    chosen=set(range(min(2,n)))  # başlangıç

    # son durum / sonuç
    for i in range(max(0,n-2),n):
        chosen.add(i)

    # gövdedeki en vurucu sayısal/kurumsal bilgiler
    ranked=sorted(
        [(i,_akt_sentence_score(s)) for i,s in enumerate(sentences)],
        key=lambda z:(z[1],-z[0]),
        reverse=True
    )
    for i,_ in ranked:
        if len(chosen)>=max_sentences:
            break
        chosen.add(i)

    ordered=[sentences[i] for i in sorted(chosen)]

    clauses=[]
    total=0
    for s in ordered:
        s=s.strip().rstrip(' .;:')
        if not s:
            continue
        if total+len(s)>max_chars and clauses:
            break
        clauses.append(s)
        total+=len(s)+2

    if not clauses:
        clauses=[sentences[0].strip().rstrip(' .;:')]

    # Örnekteki resmî AKT anlatımına yakın tek akış.
    text='; '.join(clauses)
    if text:
        first=text[0]
        if first.isalpha() and not text[:5].isupper():
            text=first.lower()+text[1:]
    return text

def _expanded_report_text(title, body):
    # Geriye dönük uyumluluk: AKT artık ham tam metni değil, resmî ve tekrarsız özeti kullanır.
    return _akt_formal_summary(title,body)


# -----------------------------
# V66 — KURUMSAL RESMÎ DİL NORMALİZASYONU
# -----------------------------
def _v66_formalize_sentence_endings(text):
    """
    V67: Önemli Gelişmeler ve Bilgi Notunda cümle sonlarındaki haber dili
    (-yor/-dı) yerine kurumsal resmî dil (-maktadır/-miştir) kullanılır.
    """
    t=re.sub(r'\s+',' ',str(text or '')).strip()
    if not t:
        return t

    exact=[
        ('açıklıyor','açıklamaktadır'),('belirtiyor','belirtmektedir'),
        ('bildiriyor','bildirmektedir'),('duyuruyor','duyurmaktadır'),
        ('söylüyor','söylemektedir'),('ifade ediyor','ifade etmektedir'),
        ('vurguluyor','vurgulamaktadır'),('gösteriyor','göstermektedir'),
        ('işaret ediyor','işaret etmektedir'),('ortaya koyuyor','ortaya koymaktadır'),
        ('öne çıkarıyor','öne çıkarmaktadır'),('öne çıkıyor','öne çıkmaktadır'),
        ('yer alıyor','yer almaktadır'),('devam ediyor','devam etmektedir'),
        ('sürüyor','sürmektedir'),('yürütülüyor','yürütülmektedir'),
        ('sürdürülüyor','sürdürülmektedir'),('yapılıyor','yapılmaktadır'),
        ('gerçekleştiriliyor','gerçekleştirilmektedir'),('kullanılıyor','kullanılmaktadır'),
        ('sayılıyor','sayılmaktadır'),('belirtiliyor','belirtilmektedir'),
        ('açıklanıyor','açıklanmaktadır'),('bildiriliyor','bildirilmektedir'),
        ('duyuruluyor','duyurulmaktadır'),('değerlendiriliyor','değerlendirilmektedir'),
        ('bekleniyor','beklenmektedir'),('planlanıyor','planlanmaktadır'),
        ('hedefleniyor','hedeflenmektedir'),('öngörülüyor','öngörülmektedir'),
        ('çalışılıyor','çalışılmaktadır'),('gerçekleşiyor','gerçekleşmektedir'),
        ('sağlıyor','sağlamaktadır'),('oluşturuyor','oluşturmaktadır'),
        ('taşıyor','taşımaktadır'),('sunuyor','sunmaktadır'),('koruyor','korumaktadır'),
        ('dolduruyor','doldurmaktadır'),('geçiyor','geçmektedir'),
        ('vuruyor','vurmaktadır'),('tamamlıyor','tamamlamaktadır'),
        ('artıyor','artmaktadır'),('azalıyor','azalmaktadır'),
    ]
    past=[
        ('yapıldı','yapılmıştır'),('gerçekleştirildi','gerçekleştirilmiştir'),
        ('açıklandı','açıklanmıştır'),('duyuruldu','duyurulmuştur'),
        ('yayımlandı','yayımlanmıştır'),('yayınlandı','yayımlanmıştır'),
        ('başladı','başlamıştır'),('tamamlandı','tamamlanmıştır'),
        ('sona erdi','sona ermiştir'),('arttı','artmıştır'),('azaldı','azalmıştır'),
        ('düştü','düşmüştür'),('yükseldi','yükselmiştir'),('geriledi','gerilemiştir'),
        ('ulaştı','ulaşmıştır'),('çıktı','çıkmıştır'),('geldi','gelmiştir'),
        ('verildi','verilmiştir'),('belirlendi','belirlenmiştir'),
        ('kaydedildi','kaydedilmiştir'),('tespit edildi','tespit edilmiştir'),
        ('bildirildi','bildirilmiştir'),('belirtildi','belirtilmiştir'),
        ('ifade edildi','ifade edilmiştir'),('vurgulandı','vurgulanmıştır'),
        ('kararlaştırıldı','kararlaştırılmıştır'),('onaylandı','onaylanmıştır'),
        ('imzalandı','imzalanmıştır'),('kuruldu','kurulmuştur'),
        ('devreye alındı','devreye alınmıştır'),('duyurdu','duyurmuştur'),
        ('açıkladı','açıklamıştır'),('belirtti','belirtmiştir'),
        ('bildirdi','bildirmiştir'),('gösterdi','göstermiştir'),
        ('sağladı','sağlamıştır'),('geçti','geçmiştir'),('vurdu','vurmuştur'),
    ]
    pairs=exact+past
    parts=re.split(r'(?<=[.!?])\s+',t)
    out=[]
    for s in parts:
        s=s.strip()
        if not s: continue
        punct=s[-1] if s[-1] in '.!?' else '.'
        core=s[:-1].rstrip() if s[-1] in '.!?' else s
        low=core.lower()
        for old,newv in sorted(pairs,key=lambda x:len(x[0]),reverse=True):
            if low.endswith(old):
                core=core[:-len(old)]+newv
                break
        out.append(core.rstrip(' .;:')+punct)
    return ' '.join(out)


def _v66_limit_important_paragraph(text,max_chars=520,max_sentences=3):
    """
    Önemli gelişmeler notunda her gelişmeyi Word üzerinde yaklaşık dört satırı
    aşmayacak yoğunlukta tutar. Öncelik ilk bilgi taşıyan cümlelere verilir.
    """
    clean=_v66_formalize_sentence_endings(text)
    sents=_sentence_chunks(clean)
    chosen=[]
    total=0
    for s in sents:
        s=s.strip()
        if not s: continue
        if total+len(s)>max_chars and chosen:
            break
        chosen.append(s)
        total+=len(s)+1
        if len(chosen)>=max_sentences:
            break
    result=' '.join(chosen).strip()
    if len(result)>max_chars:
        cut=result[:max_chars].rsplit(' ',1)[0].rstrip(' ,;:')
        # Kurumsal kapanış; kesilmiş yarım yüklem bırakma.
        if cut and cut[-1] not in '.!?':
            cut+='.'
        result=cut
    return result

def _akt_topic_labels(rows):
    joined=norm(' '.join(
        f"{r.get('Başlık','')} {r.get('İçerik_Özeti','')} {r.get('Kategori','')}"
        for r in rows
    ))
    mapping=[
        ('istihdam','istihdam'),
        ('sanayi üret','sanayi üretimi'),
        ('otomotiv','otomotiv üretimi'),
        ('yapay zeka','yapay zeka'),
        ('yapay zekâ','yapay zeka'),
        ('veri sızınt','veri sızıntısı'),
        ('siber saldır','siber güvenlik'),
        ('ihracat','ihracat'),
        ('yatırım','yatırım'),
        ('kapasite kullanım','kapasite kullanım oranı'),
        ('savunma','savunma sanayii'),
        ('enerji','enerji'),
        ('yangın','sanayi tesisi yangını'),
        ('patlama','sanayi tesisi patlaması'),
        ('ar-ge','Ar-Ge'),
        ('arge','Ar-Ge')
    ]
    out=[]
    for key,label in mapping:
        if key in joined and label not in out:
            out.append(label)
        if len(out)>=6:
            break
    return out

def _akt_findings_intro(rows):
    topics=_akt_topic_labels(rows)
    if topics:
        if len(topics)==1:
            topic_text=f'“{topics[0]}”'
        else:
            topic_text=', '.join(f'“{x}”' for x in topics[:-1]) + f' ve “{topics[-1]}”'
        return (
            "Sanayi ve Teknoloji alanlarında yapılan açık kaynak taraması neticesinde bazı haber "
            f"bültenlerinde {topic_text} konu başlıklarıyla ilgili içerikler hazırlandığı tespit edilmiştir. "
            "İçeriklerin hangi internet sitesinde yer aldığı, başlığı, bağlantı adresi, içeriğin detaylı özeti "
            "ve görseli aşağıda yer almaktadır."
        )
    return (
        "Sanayi ve Teknoloji alanlarında yapılan açık kaynak taraması neticesinde seçilen haber içerikleri "
        "tespit edilmiştir. İçeriklerin hangi internet sitesinde yer aldığı, başlığı, bağlantı adresi, "
        "içeriğin detaylı özeti ve görseli aşağıda yer almaktadır."
    )


def _v67_akt_reported_content(text):
    """
    AKT'de haber içeriğini dolaylı anlatı biçimine çevirir:
    açıklıyor -> açıkladığı, duyurdu -> duyurduğu, belirtiyor -> belirttiği vb.
    Son kapanış tek kez 'hususları ifade edilmektedir.' olur.
    """
    t=re.sub(r'\s+',' ',str(text or '')).strip().rstrip(' .;:')
    if not t: return t

    conv=[
        ('ifade ediyor','ifade ettiği'),('ifade etti','ifade ettiği'),
        ('açıklıyor','açıkladığı'),('açıkladı','açıkladığı'),
        ('belirtiyor','belirttiği'),('belirtti','belirttiği'),
        ('bildiriyor','bildirdiği'),('bildirdi','bildirdiği'),
        ('duyuruyor','duyurduğu'),('duyurdu','duyurduğu'),
        ('vurguluyor','vurguladığı'),('vurguladı','vurguladığı'),
        ('gösteriyor','gösterdiği'),('gösterdi','gösterdiği'),
        ('işaret ediyor','işaret ettiği'),('işaret etti','işaret ettiği'),
        ('ortaya koyuyor','ortaya koyduğu'),('ortaya koydu','ortaya koyduğu'),
        ('sağlıyor','sağladığı'),('sağladı','sağladığı'),
        ('dolduruyor','doldurduğu'),('doldurdu','doldurduğu'),
        ('yer alıyor','yer aldığı'),('yer aldı','yer aldığı'),
        ('devam ediyor','devam ettiği'),('devam etti','devam ettiği'),
        ('sürüyor','sürdüğü'),('sürdü','sürdüğü'),
        ('tamamladı','tamamladığı'),('tamamlıyor','tamamladığı'),
        ('vuruyor','vurduğu'),('vurdu','vurduğu'),
        ('geçiyor','geçtiği'),('geçti','geçtiği'),
        ('yapıldı','yapıldığı'),('gerçekleştirildi','gerçekleştirildiği'),
        ('açıklandı','açıklandığı'),('duyuruldu','duyurulduğu'),
        ('yayımlandı','yayımlandığı'),('başladı','başladığı'),
        ('tamamlandı','tamamlandığı'),('ulaştı','ulaştığı'),
        ('arttı','arttığı'),('azaldı','azaldığı'),
        ('oldu','olduğu'),('oluyor','olduğu'),
        ('sahiptir','sahip olduğu'),('dayanmaktadır','dayandığı'),
        ('değişebilir','değişebileceği'),
    ]

    clauses=[x.strip(' ,;:.') for x in re.split(r'\s*;\s*',t) if x.strip()]
    out=[]
    for c in clauses:
        low=c.lower()
        changed=False
        for old,newv in sorted(conv,key=lambda x:len(x[0]),reverse=True):
            # Haber özetindeki yüklem çoğunlukla cümlecik sonundadır.
            if low.endswith(old):
                c=c[:-len(old)]+newv
                changed=True
                break
        # Nokta ile birleşmiş kısa cümlelerde de son yüklemi dönüştür.
        if not changed:
            for old,newv in sorted(conv,key=lambda x:len(x[0]),reverse=True):
                c=re.sub(r'\b'+re.escape(old)+r'(?=\s*$)',newv,c,flags=re.I)
        out.append(c.rstrip(' .;:'))
    return '; '.join(out)

def make_docx(rows):
    """
    Kullanıcının ilettiği STB AKT örneğine yakın resmî format:
    Başlık -> görev alanı -> tarih -> bulgular -> numaralı haber/özet/link -> görsel -> Arz olunur.
    """
    doc=Document()
    section=doc.sections[0]
    section.top_margin=Cm(2.0)
    section.bottom_margin=Cm(2.0)
    section.left_margin=Cm(2.5)
    section.right_margin=Cm(2.5)

    normal=doc.styles["Normal"]
    normal.font.name="Times New Roman"
    normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Times New Roman")

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(10)
    r=p.add_run("AÇIK KAYNAK TARAMA ÇALIŞMASI")
    r.bold=True
    r.font.name="Times New Roman"
    r.font.size=Pt(14)

    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(0)
    p.add_run("Tarama Yapılan Görev Alanı: ").bold=True
    p.add_run("Sanayi ve Teknoloji")

    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(8)
    p.add_run("Tarih: ").bold=True
    p.add_run(datetime.now().astimezone().strftime("%d.%m.%Y"))

    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(3)
    p.add_run("Bulgular: ").bold=True
    p.add_run(_akt_findings_intro(rows))

    for i,row in enumerate(rows,1):
        detail=article_detail(row)

        real_url=detail.get("canonical") or row.get("Yayıncı_URL") or row.get("URL","")
        title=(detail.get("title") or row.get("Başlık") or "").strip()
        source=_real_source(row,detail,real_url)
        body=detail.get("text") or row.get("İçerik_Özeti") or title
        summary=_v67_akt_reported_content(_akt_formal_summary(title,body))

        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent=Cm(0.75)
        p.paragraph_format.space_before=Pt(4)
        p.paragraph_format.space_after=Pt(6)

        nr=p.add_run(f"{i}. ")
        nr.bold=True

        sr=p.add_run(f'“{source}”')
        sr.bold=True
        p.add_run(' isimli internet sitesinde, ')
        tr=p.add_run(f'“{title}”')
        tr.bold=True
        p.add_run(' başlığıyla bir haber yayımlanmıştır. (')
        _word_hyperlink(p,real_url,real_url if real_url else "Haber Linki")
        p.add_run(') Söz konusu haber içeriğinde, ')
        p.add_run(summary)
        p.add_run(' hususları ifade edilmiştir.')

        image_stream=None
        image_url=""
        for candidate in detail.get("images",[]):
            image_stream=_download_report_image(candidate)
            if image_stream:
                image_url=candidate
                break

        if image_stream or detail.get("images"):
            cap=doc.add_paragraph()
            cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before=Pt(4)
            cap.paragraph_format.space_after=Pt(4)
            cr=cap.add_run(f'Görsel {i}: “{source}” Sitesinde Yer Alan Görsel')
            cr.bold=True
            cr.font.name="Times New Roman"
            cr.font.size=Pt(11)

        if image_stream:
            ip=doc.add_paragraph()
            ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_after=Pt(10)
            ip.add_run().add_picture(image_stream,width=Cm(14.5))
        elif detail.get("images"):
            # Örnekte görsel esas; indirilemediyse raporu gereksiz teknik metinle doldurma.
            lp=doc.add_paragraph()
            lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_after=Pt(8)
            _word_hyperlink(lp,detail["images"][0],"Görseli Aç")

    endp=doc.add_paragraph()
    endp.paragraph_format.space_before=Pt(8)
    endp.add_run("Arz olunur.")

    bio=BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# -----------------------------
# V63 — İŞ AKIŞI HAFIZASI / İKİNCİ GÖZ / YARINA TAKİP
# -----------------------------
def _v63_mark_notes(rows):
    if rows is None or len(rows)==0 or not _init_history_db():
        return
    try:
        with _history_connect() as conn:
            for row in rows:
                title=str(row.get('Başlık','') or '').strip()
                if not title: continue
                conn.execute(
                    "INSERT OR IGNORE INTO note_history(created_at,title,url) VALUES(?,?,?)",
                    (datetime.now().astimezone().isoformat(),title,str(row.get('URL','') or '').strip())
                )
            conn.commit()
        _v73_invalidate_status_cache()
    except Exception:
        pass

def _v73_invalidate_status_cache():
    st.session_state.pop('_v73_status_sets_cache',None)

def _v63_status_sets():
    """
    V102 — Tek sorguda dört işlem durumunu okur:
    Önemli Gelişmeler, AKT, Sunum ve hazırlanmış Bilgi Notu.
    """
    cached=st.session_state.get('_v73_status_sets_cache')
    # V102 geçiş güvenliği: V101 açık oturumlarında cache 3 elemanlıdır.
    # Yeni sürüm 4 durum kümesi kullandığı için eski cache'i otomatik geçersiz kıl.
    if cached is not None:
        try:
            if isinstance(cached,(tuple,list)) and len(cached)==4:
                return cached
        except Exception:
            pass
        st.session_state.pop('_v73_status_sets_cache',None)

    imp=set(); akt=set(); notes=set(); pres=set()
    if not _init_history_db():
        return imp,akt,notes,pres
    try:
        with _history_connect() as conn:
            for table,target in [
                ('important_basket',imp),
                ('osint_report_basket',akt),
                ('note_history',notes),
                ('presentation_basket',pres)
            ]:
                rows=conn.execute(f"SELECT title,url FROM {table}").fetchall()
                for title,url in rows:
                    target.add(str(url).strip() if str(url or '').strip() else title_key(str(title or '')))
    except Exception:
        pass
    result=(imp,akt,notes,pres)
    st.session_state['_v73_status_sets_cache']=result
    return result

def _v63_add_status_badges(df):
    """Her haber tablosuna tek bir Durum sütunu ekler."""
    if df is None or df.empty:
        return df
    out=df.copy()
    imp,akt,notes,pres=_v63_status_sets()

    def badge(r):
        # Ana tarama tabloları Türkçe; sepet tabloları İngilizce kolon adları kullanır.
        url=str(r.get('URL',r.get('url','')) or '').strip()
        title=str(r.get('Başlık',r.get('title','')) or '')
        k=url or title_key(title)
        b=[]
        if k in pres:  b.append('🖥️ Sunum Sepetinde')
        if k in imp:   b.append('📌 Önemli Gelişmelerde')
        if k in notes: b.append('📝 Bilgi Notu Yapıldı')
        if k in akt:   b.append('📁 AKT Sepetinde')
        return ' • '.join(b) if b else '—'

    out['Durum']=out.apply(badge,axis=1)
    return out

def _v63_missed_candidates(df,limit=12):
    """Yüksek değerli fakat iki sepette de olmayan olayları ikinci göz olarak gösterir."""
    if df is None or df.empty: return pd.DataFrame()
    value=_v52_event_value_table(df,max(30,limit*2))
    if value.empty: return value
    imp,akt,notes,pres=_v63_status_sets()
    rows=[]
    for _,v in value.iterrows():
        url=str(v.get('URL','') or '').strip()
        key=url or title_key(str(v.get('Gelişme','')))
        if key in imp or key in akt: continue
        # İkinci göz eşiği: güçlü değer skoru veya belirgin risk.
        if int(v.get('Değer_Skoru',0) or 0)<55 and int(v.get('Risk',0) or 0)<60:
            continue
        rows.append(v.to_dict())
        if len(rows)>=limit: break
    return pd.DataFrame(rows)

def _v63_tomorrow_candidates(df,limit=15):
    """Sonuçlanmamış, stratejik/riskli ve takip değeri olan olayları yarın için önerir."""
    if df is None or df.empty: return pd.DataFrame()
    life=_v58_event_lifecycle_table(df,40)
    if life.empty: return pd.DataFrame()
    out=life[life['Aşama']!='✅ Sonuçlandı'].copy()
    out=out[(pd.to_numeric(out['Risk_Skoru'],errors='coerce').fillna(0)>=35) |
            (pd.to_numeric(out['Kaynak_Sayısı'],errors='coerce').fillna(0)>=2)]
    if out.empty: return out
    out['Takip_Gerekçesi']=out.apply(
        lambda r:(
            'Olay gelişiyor; yeni açıklama/sonuç bekleniyor.'
            if 'Gelişiyor' in str(r.get('Aşama','')) else
            'Teyit edildi; uygulama/sonuç etkisi izlenmeli.'
            if 'Teyit' in str(r.get('Aşama','')) else
            'İlk sinyal; ikinci kaynak veya resmî teyit izlenmeli.'
        ),axis=1
    )
    return out.head(limit)

def _v63_add_tomorrow(rows):
    if rows is None or len(rows)==0 or not _init_history_db(): return 0
    added=0
    try:
        with _history_connect() as conn:
            for row in rows:
                title=str(row.get('Başlık','') or '').strip()
                if not title: continue
                cur=conn.execute("""
                    INSERT OR IGNORE INTO tomorrow_followup(
                        added_at,title,source,url,category,reason
                    ) VALUES(?,?,?,?,?,?)
                """,(
                    datetime.now().astimezone().isoformat(),title,
                    str(row.get('Kaynak','') or ''),str(row.get('URL','') or ''),
                    str(row.get('Kategori','') or ''),str(row.get('Takip_Gerekçesi','') or '')
                ))
                added+=int(bool(cur.rowcount))
            conn.commit()
    except Exception: pass
    return added

def _v63_load_tomorrow():
    if not _init_history_db(): return pd.DataFrame()
    try:
        with _history_connect() as conn:
            return pd.read_sql_query("SELECT * FROM tomorrow_followup ORDER BY added_at DESC",conn)
    except Exception:
        return pd.DataFrame()



# -----------------------------
# V68 — ANALİST KOMUTA MERKEZİ / SONRAKİ EN İYİ İŞLEM
# -----------------------------
def _v68_analyst_command_center(df,limit=8):
    """
    V69 Analist Komuta Merkezi:
    - 09:00–17:30 Bilgi Notu: veri/istatistik, resmî açıklama, ürün/teknoloji tanıtımı vb.
    - 09:00–17:30 AKT: negatif, eleştirel, yapısal eleştiri, propaganda/dezenformasyon niteliği taşıyan olumsuz içerikler.
    - 17:30 sonrası: yalnız kritik/acil gelişmeler.
    - Sunum: resmî veri/istatistik, resmî açıklama veya resmî teyitli bilgi.
    """
    cols=['Öncelik','Önerilen_İşlem','Tarih','Başlık','Neden','Durum','Değer_Skoru','Risk_Skoru','URL']
    if df is None or df.empty:
        return pd.DataFrame(columns=cols), 'Veri Yok', ''

    try:
        from zoneinfo import ZoneInfo
        now_tr=datetime.now(ZoneInfo('Europe/Istanbul'))
    except Exception:
        now_tr=datetime.now().astimezone()
    hour=now_tr.hour + now_tr.minute/60

    if 9 <= hour < 14:
        phase='09:00–14:00 | Bilgi notu • AKT hazırlığı • sunum/veri kontrolü'
        phase_hint='Bilgi notunda resmî/veri odaklı içerikler; AKT’de negatif-eleştirel içerikler; sunumda ise resmî ve teyitli bilgiler önceliklendirilmektedir.'
    elif 14 <= hour < 17.5:
        phase='14:00–17:30 | Bilgi notu • sunum • önemli gelişmeleri zenginleştirme'
        phase_hint='Yeni resmî veri/açıklamalar bilgi notu ve sunum için; negatif-eleştirel içerikler AKT takibi için değerlendirilmektedir.'
    else:
        phase='17:30 sonrası | Kritik takip modu'
        phase_hint='Rutin bilgi notu ve sunum önerileri durdurulmakta; yalnızca kritik/acil gelişmeler öne çıkarılmaktadır.'

    value=_v52_event_value_table(df,max(40,limit*5))
    if value.empty:
        return pd.DataFrame(columns=cols),phase,phase_hint

    imp,akt,notes,pres=_v63_status_sets()
    actions=[]

    data_terms=[
        'istatistik','veri','oran','endeks','sanayi üretimi','kapasite kullanım',
        'ihracat','ithalat','ciro','istihdam','işsizlik','büyüme','yatırım teşvik',
        'arge','ar-ge','patent','başvuru','milyar','milyon','yüzde','%'
    ]
    product_terms=[
        'ürün tanıt','tanıtıldı','tanıttı','yeni ürün','yeni teknoloji','prototip',
        'seri üretim','ilk teslimat','envantere','platform','sistem geliştir',
        'füze','uydu','çip','yarı iletken','yapay zeka','yapay zekâ'
    ]
    propaganda_terms=[
        'propaganda','dezenformasyon','manipülasyon','iddia','suçlama','eleştiri',
        'eleştirel','tepki','kriz','başarısız','skandal','zarar','kayıp','çöküş',
        'iflas','işten çıkar','üretim durdu','üretimi durdur','gecikme','yaptırım',
        'ambargo','boykot','bağımlılık','risk','tehdit'
    ]

    for _,v in value.iterrows():
        row=_v53_find_event_row(df,v)
        if row is None:
            continue

        title=str(row.get('Başlık','') or v.get('Gelişme',''))
        url=str(row.get('URL','') or v.get('URL','')).strip()
        key=url or title_key(title)
        score=int(v.get('Değer_Skoru',0) or 0)
        risk=int(row.get('Risk_Skoru',v.get('Risk',0)) or 0)
        text=norm(f"{title} {row.get('İçerik_Özeti','')} {row.get('Kategori','')} {row.get('Doğrulama','')}")
        critical=bool(critical_industrial_incident(title,row.get('İçerik_Özeti','')))
        official=_is_official_radar_row(row)
        verification=norm(row.get('Doğrulama',''))
        officially_verified=official or any(x in verification for x in ['resmi','resmî','birincil','teyit'])
        negative=(str(row.get('Duygu',''))=='Negatif' or
                  str(row.get('Risk_Durumu',''))=='Yüksek Risk' or
                  any(x in text for x in propaganda_terms))
        data_stat=any(x in text for x in data_terms)
        product_intro=any(x in text for x in product_terms)
        multi=int(v.get('Kaynak_Sayısı',0) or 0)>=2

        badges=[]
        if key in imp: badges.append('📌 Önemli Gelişmelerde')
        if key in akt: badges.append('📁 AKT’de')
        if key in notes: badges.append('📝 Bilgi Notu Hazırlandı')
        status=' • '.join(badges) if badges else 'Henüz işleme alınmadı'

        proposals=[]

        # 17:30 sonrası: yalnız kritik gelişme.
        if hour >= 17.5 or hour < 9:
            if critical or risk>=75 or score>=88:
                why=[]
                if critical: why.append('kritik sanayi olayı')
                if risk>=75: why.append('çok yüksek risk')
                if score>=88: why.append('çok yüksek analitik değer')
                if officially_verified: why.append('resmî/teyitli bilgi')
                proposals.append((120+risk,'🚨 KRİTİK GELİŞME — ACİL DEĞERLENDİR',why))
        else:
            # Bilgi Notu: veri/istatistik, resmî açıklama, ürün/teknoloji tanıtımı.
            if key not in notes and (data_stat or official or product_intro):
                why=[]
                if data_stat: why.append('veri/istatistiki bilgi')
                if official: why.append('resmî açıklama/birincil kaynak')
                if product_intro: why.append('ürün/teknoloji tanıtımı veya somut teknolojik gelişme')
                if multi: why.append(f"{int(v.get('Kaynak_Sayısı',0) or 0)} farklı kaynak")
                proposals.append((105+score,'📝 Bilgi Notu Değerlendir',why))

            # AKT: negatif, eleştirel, propaganda/dezenformasyon/olumsuz içerik.
            if key not in akt and negative:
                why=[]
                if str(row.get('Duygu',''))=='Negatif': why.append('negatif/olumsuz içerik')
                if any(x in text for x in ['eleştiri','eleştirel','tepki','suçlama']): why.append('eleştirel dil/yapısal eleştiri')
                if any(x in text for x in ['propaganda','dezenformasyon','manipülasyon','iddia']): why.append('propaganda/manipülasyon iddiası veya niteliği')
                if risk>=55: why.append('dikkat gerektiren risk/etki')
                proposals.append((100+score+risk//5,'📁 AKT Sepetine Almayı Değerlendir',why))

            # Sunum: yalnız resmî veri/istatistik veya resmî/teyitli bilgi.
            if (data_stat and officially_verified) or official or (officially_verified and score>=55):
                why=[]
                if data_stat: why.append('resmî/teyitli veri veya istatistik')
                if official: why.append('resmî açıklama')
                elif officially_verified: why.append('resmî teyitli bilgi')
                proposals.append((85+score,'🖥️ Sunuma Eklemeyi Değerlendir',why))

            # Mevcut önemli gelişme yeni kaynaklarla zenginleşmişse ayrıca hatırlat.
            if key in imp and multi:
                proposals.append((78+score,'🔄 Önemli Gelişmeyi Zenginleştir',
                                  ['önemli gelişme sepetinde','yeni/çoklu kaynak desteği mevcut']))

        for priority,action,reason in proposals:
            actions.append({
                'Öncelik':priority,
                'Önerilen_İşlem':action,
                'Tarih':row.get('Tarih',''),
                'Başlık':title,
                'Neden':' • '.join(dict.fromkeys(reason)) if reason else 'analist değerlendirmesi önerilmektedir',
                'Durum':status,
                'Değer_Skoru':score,
                'Risk_Skoru':risk,
                'URL':url
            })

    if not actions:
        return pd.DataFrame(columns=cols),phase,phase_hint

    out=pd.DataFrame(actions)
    # Aynı haber aynı işlem için yalnız bir kez gösterilsin.
    out=out.sort_values(['Öncelik','Değer_Skoru','Risk_Skoru'],ascending=[False,False,False])
    out=out.drop_duplicates(subset=['Önerilen_İşlem','URL','Başlık'],keep='first').head(limit).reset_index(drop=True)
    out['Öncelik']=range(1,len(out)+1)
    return out[cols],phase,phase_hint


def _v73_row_keys(df):
    """apply(axis=1) yerine hızlı, vektörize haber anahtarı üretir."""
    if df is None or df.empty:
        return pd.Series(dtype=str)
    urls=df['URL'].fillna('').astype(str).str.strip() if 'URL' in df.columns else pd.Series('',index=df.index)
    titles=df['Başlık'].fillna('').astype(str) if 'Başlık' in df.columns else pd.Series('',index=df.index)
    # title_key yalnız URL'siz satırlarda çalışır.
    fallback=titles.map(title_key)
    return urls.where(urls.ne(''),fallback)



# ============================================================
# V107 — OLAY BAZLI KAYNAK ZENGİNLEŞTİRME
# Kullanıcı yerel/kısa bir haberi seçse bile aynı taramadaki aynı olayın
# ana akım/resmî/daha ayrıntılı sürümleri birleştirilerek sepete aktarılır.
# Ek web isteği YOKTUR; yalnız mevcut tarama havuzu kullanılır.
# ============================================================

def _v107_source_quality(row):
    """Bir olay kümesinde hangi haber sürümünün ana taşıyıcı olacağını puanlar."""
    d=str(row.get('Domain','') or '')
    source=str(row.get('Kaynak','') or '')
    text=_clean_note_text(row.get('İçerik_Özeti',''))
    score=0
    if d in TR_OFFICIAL:
        score+=80
    elif d in TR_MAIN:
        score+=55
    elif d in TR_TECH:
        score+=45
    elif d in SOCIAL:
        score+=5
    else:
        score+=20
    # Ayrıntılı gövdeyi ödüllendir; aşırı uzun portal artıklarına sınırsız puan verme.
    score+=min(len(text)//180,30)
    score+=min(len(re.findall(r'\b\d+(?:[.,]\d+)?\b',text))*2,16)
    if _verification_rank(row.get('Doğrulama',''))>=3:
        score+=12
    if source and norm(source) not in {'google','google news','google haberler','açık kaynak'}:
        score+=4
    return score

def _v107_same_event(selected, candidate):
    """V104 olay mantığını kullanarak seçili haberle gerçekten aynı olayı sınar."""
    su=str(selected.get('URL','') or '').strip()
    cu=str(candidate.get('URL','') or '').strip()
    if su and cu and su==cu:
        return True
    soid=str(selected.get('Olay_ID','') or '').strip()
    coid=str(candidate.get('Olay_ID','') or '').strip()
    if soid and coid and soid==coid:
        return True
    sim=_v104_event_similarity(
        selected.get('Başlık',''),selected.get('İçerik_Özeti',''),
        candidate.get('Başlık',''),candidate.get('İçerik_Özeti','')
    )
    return sim>=0.50

def _v107_unique_sentences(rows, max_chars=8000):
    """
    En iyi kaynaklardan gelen tamamlayıcı cümleleri birleştirir.
    Aynı cümleyi/çok benzer bilgiyi tekrar eklemez; kritik rakamları korur.
    """
    out=[]; seen=set()
    for row in rows:
        raw=_v84_hard_repair_text(row.get('İçerik_Özeti',''))
        sentences=_v84_clean_article_sentences(raw)
        if not sentences and raw:
            sentences=[raw]
        for s in sentences:
            s=_clean_note_text(s).strip()
            if len(s)<25:
                continue
            k=title_key(s)
            toks=set(_history_tokens(s))
            duplicate=False
            for oldk,oldtoks in seen:
                if k==oldk:
                    duplicate=True; break
                oldtoks=set(oldtoks)
                if toks and oldtoks:
                    jac=len(toks&oldtoks)/max(1,len(toks|oldtoks))
                    if jac>=0.72:
                        duplicate=True; break
            if duplicate:
                continue
            out.append(s)
            # V108 düzeltmesi: set nesnesi hashlenemez; frozenset olarak saklanır.
            seen.add((k,frozenset(toks)))
            if len(' '.join(out))>=max_chars:
                break
        if len(' '.join(out))>=max_chars:
            break
    return ' '.join(out)[:max_chars]

def _v107_enrich_selected_rows(rows):
    """
    Seçilen her haberi, st.session_state['rows'] içindeki aynı olayın diğer
    sürümleriyle zenginleştirir. Bir olaydan sepete yalnız bir kayıt gider.
    """
    if rows is None or len(rows)==0:
        return []
    selected=[dict(r) for r in rows]
    pool=st.session_state.get('rows') or []
    if not pool:
        return selected

    enriched=[]
    used_event_sigs=set()

    for sel in selected:
        matches=[]
        for cand in pool:
            try:
                if _v107_same_event(sel,cand):
                    matches.append(dict(cand))
            except Exception:
                continue
        if not matches:
            matches=[sel]

        # Yanlış geniş kümeyi engelle: en fazla en güçlü 8 kaynak.
        matches=sorted(matches,key=_v107_source_quality,reverse=True)[:8]
        best=matches[0].copy()

        # Seçilen kaydın işlem/risk bağlamını kaybetme.
        best['Risk_Skoru']=max([int(x.get('Risk_Skoru',0) or 0) for x in matches] or [int(sel.get('Risk_Skoru',0) or 0)])
        if sel.get('Risk_Durumu'):
            best['Risk_Durumu']=sel.get('Risk_Durumu')
        if sel.get('Kategori'):
            best['Kategori']=sel.get('Kategori')

        merged=_v107_unique_sentences(matches)
        if merged:
            best['İçerik_Özeti']=merged

        domains=[]
        sources=[]
        urls=[]
        for x in matches:
            d=str(x.get('Domain','') or '').strip()
            s=_clean_note_text(x.get('Kaynak','')).strip()
            u=str(x.get('URL','') or '').strip()
            if d and d not in domains: domains.append(d)
            if s and s not in sources: sources.append(s)
            if u and u not in urls: urls.append(u)

        best['Olay_Kaynak_Sayisi']=max(len(domains),len(sources),int(best.get('Olay_Kaynak_Sayisi',0) or 0))
        best['Zenginleştirme_Kaynakları']=' | '.join(sources[:8])
        best['Zenginleştirme_URLleri']=' | '.join(urls[:8])
        best['Zenginleştirildi']='Evet' if len(matches)>1 else 'Hayır'

        # Aynı olay kullanıcı tarafından iki farklı satırdan seçildiyse sepete iki kez gitmesin.
        sig=' '.join(sorted(_v104_event_tokens(best.get('Başlık',''),best.get('İçerik_Özeti',''))))
        if sig and sig in used_event_sigs:
            continue
        if sig: used_event_sigs.add(sig)
        enriched.append(best)

    return enriched

def _v74_bulk_add_basket(rows,table_name):
    """
    V74: Kronoloji hızlı işlemleri için tek SQLite executemany çağrısı.
    Satır satır execute yerine toplu INSERT OR IGNORE kullanır.
    """
    if rows is None or len(rows)==0 or not _init_history_db():
        return 0
    if table_name not in ('important_basket','osint_report_basket'):
        return 0
    payload=[]
    now_iso=datetime.now().astimezone().isoformat()
    for row in rows:
        title=str(row.get('Başlık','') or '').strip()
        if not title:
            continue
        payload.append((
            now_iso,
            str(row.get('Tarih','') or ''),
            title,
            str(row.get('Kaynak','') or ''),
            str(row.get('URL','') or '').strip(),
            str(row.get('Kategori','') or ''),
            int(row.get('Risk_Skoru',0) or 0),
            str(row.get('Risk_Durumu','') or ''),
            str(row.get('İçerik_Özeti','') or '')[:8000]
        ))
    if not payload:
        return 0
    try:
        with _history_connect() as conn:
            before=conn.total_changes
            conn.executemany(f"""
                INSERT OR IGNORE INTO {table_name}(
                    added_at,news_time,title,source,url,category,risk_score,risk_status,summary
                ) VALUES(?,?,?,?,?,?,?,?,?)
            """,payload)
            conn.commit()
            added=conn.total_changes-before
        if added:
            _v73_invalidate_status_cache()
        return int(added)
    except Exception:
        return 0

def _v74_fast_add_important(rows):
    return _v74_bulk_add_basket(_v107_enrich_selected_rows(rows),'important_basket')

def _v74_fast_add_osint(rows):
    return _v74_bulk_add_basket(_v107_enrich_selected_rows(rows),'osint_report_basket')

def _v80_add_presentation(rows):
    """Her bölümden seçilen haberleri sunum sepetine toplu ekler."""
    if rows is None or len(rows)==0 or not _init_history_db():
        return 0
    payload=[]
    now_iso=datetime.now().astimezone().isoformat()
    for row in rows:
        title=_clean_note_text(row.get('Başlık',''))
        if not title:
            continue
        payload.append((
            now_iso,
            str(row.get('Tarih','') or ''),
            title,
            _clean_note_text(row.get('Kaynak','')),
            str(row.get('URL','') or '').strip(),
            _clean_note_text(row.get('Kategori','')),
            _clean_note_text(row.get('İçerik_Özeti',''))[:5000]
        ))
    try:
        with _history_connect() as conn:
            before=conn.total_changes
            conn.executemany("""
                INSERT OR IGNORE INTO presentation_basket(
                    added_at,news_time,title,source,url,category,summary
                ) VALUES(?,?,?,?,?,?,?)
            """,payload)
            conn.commit()
            added=int(conn.total_changes-before)
        if added:
            _v73_invalidate_status_cache()
        return added
    except Exception:
        return 0

def _v80_load_presentation():
    if not _init_history_db():
        return pd.DataFrame()
    try:
        with _history_connect() as conn:
            return pd.read_sql_query(
                "SELECT * FROM presentation_basket ORDER BY id DESC",conn
            )
    except Exception:
        return pd.DataFrame()

def _v80_clear_presentation():
    try:
        with _history_connect() as conn:
            cur=conn.execute("DELETE FROM presentation_basket")
            conn.commit()
            removed=cur.rowcount
        if removed:
            _v73_invalidate_status_cache()
        return removed
    except Exception:
        return 0

def _v81_remove_presentation_ids(ids):
    ids=[int(x) for x in ids if str(x).isdigit()]
    if not ids: return 0
    try:
        with _history_connect() as conn:
            marks=','.join('?' for _ in ids)
            cur=conn.execute(f"DELETE FROM presentation_basket WHERE id IN ({marks})",ids)
            conn.commit()
            removed=cur.rowcount
        if removed:
            _v73_invalidate_status_cache()
        return removed
    except Exception:
        return 0

def _v81_basket_to_rows(bdf):
    rows=[]
    if bdf is None or bdf.empty: return rows
    for _,r in bdf.iterrows():
        rows.append({'Tarih':_clean_note_text(r.get('news_time','')),'Kaynak':_clean_note_text(r.get('source','')),
        'Başlık':_clean_note_text(r.get('title','')),'İçerik_Özeti':_clean_note_text(r.get('summary','')),
        'URL':str(r.get('url','') or ''),'Kategori':_clean_note_text(r.get('category','')),
        'Risk_Skoru':r.get('risk_score',0),'Risk_Durumu':_clean_note_text(r.get('risk_status',''))})
    return rows


def _v73_main_selected(selected_keys):
    """
    Ana tarama DataFrame'ini yalnız kullanıcı gerçekten bir işlem butonuna bastığında oluşturur/eşleştirir.
    Checkbox işaretlemek artık yüzlerce satır üzerinde gereksiz tekrar filtrelemesi başlatmaz.
    """
    if not selected_keys:
        return pd.DataFrame()
    main_rows=st.session_state.get('rows') or []
    if not main_rows:
        return pd.DataFrame()
    main_df=pd.DataFrame(main_rows)
    keys=_v73_row_keys(main_df)
    return main_df[keys.isin(selected_keys)].copy()

def _section_select_table(section_key, data, columns, height=420):
    """
    V75 ULTRA HIZ:
    Tüm bölüm tablolarında checkbox değişikliği form içinde kalır.
    Streamlit yalnız kullanıcı işlem düğmesine bastığında rerun yapar.
    """
    if data is None or data.empty:
        return pd.DataFrame()

    tbl=_v63_add_status_badges(data.copy())
    if 'Durum' not in columns:
        columns=list(columns)
        insert_at=columns.index('Başlık')+1 if 'Başlık' in columns else 0
        columns.insert(insert_at,'Durum')

    tbl['_row_key']=_v73_row_keys(tbl).values
    selected_map=st.session_state.section_selections.get(section_key,{})
    tbl.insert(0,'Seç',[bool(selected_map.get(k,False)) for k in tbl['_row_key']]) if 'Seç' not in tbl.columns else None
    if 'Seç' in tbl.columns:
        tbl['Seç']=[bool(selected_map.get(k,bool(v))) for k,v in zip(tbl['_row_key'],tbl['Seç'].tolist())]

    show_cols=['Seç']+[c for c in columns if c in tbl.columns and c!='Seç']

    with st.form(key=f'v75_fast_section_form_{section_key}',clear_on_submit=False):
        edited=st.data_editor(
            tbl[show_cols+['_row_key']],
            column_config={
                'Seç':st.column_config.CheckboxColumn('Seç'),
                'URL':st.column_config.LinkColumn('Haber Linki'),
                'Medya_URL':st.column_config.LinkColumn('Medya'),
                'Resmî_URL':st.column_config.LinkColumn('Resmî Açıklama'),
                'Eşleşme':st.column_config.NumberColumn('Eşleşme',format='%d%%'),
                'Değer_Skoru':st.column_config.ProgressColumn('Değer Skoru',min_value=0,max_value=100,format='%d/100'),
                'Risk':st.column_config.NumberColumn('Risk',format='%d/100'),
                'Risk_Skoru':st.column_config.NumberColumn('Risk',format='%d/100'),
                'Durum':st.column_config.TextColumn('Durum',width='large'),
                '_row_key':None
            },
            disabled=[c for c in show_cols if c!='Seç']+['_row_key'],
            hide_index=True,use_container_width=True,height=height,
            key=f'v75_section_editor_{section_key}'
        )
        a1,a2,a3,a4=st.columns(4)
        with a1: do_imp=st.form_submit_button('📌 Önemli Gelişmelere Ekle',use_container_width=True)
        with a2: do_akt=st.form_submit_button('🗂️ AKT Sepetine Ekle',use_container_width=True)
        with a3: do_pres=st.form_submit_button('🖥️ Sunum Sepetine Ekle',use_container_width=True)
        with a4: do_note=st.form_submit_button('📝 Bilgi Notu Oluştur',use_container_width=True)

    selected_keys=set(edited.loc[edited['Seç'].astype(bool),'_row_key'].astype(str))
    st.session_state.section_selections[section_key]={k:(k in selected_keys) for k in edited['_row_key'].astype(str)}
    selected=data[_v73_row_keys(data).isin(selected_keys)].copy()

    if do_imp or do_akt or do_note or do_pres:
        if not selected_keys:
            st.warning('Önce en az bir haberi işaretleyin.')
        else:
            # Önce görünür bölüm verisini kullan: ana dataframe eşleştirmesine çoğu işlemde gerek yok.
            action_rows=selected.copy()
            if do_imp:
                n=_v74_fast_add_important(action_rows.to_dict('records'))
                st.success(f'✅ {n} yeni haber Önemli Gelişmeler Sepeti’ne eklenmiştir.')
            elif do_akt:
                n=_v74_fast_add_osint(action_rows.to_dict('records'))
                st.success(f'✅ {n} yeni haber AKT Sepeti’ne eklenmiştir.')
            elif do_pres:
                n=_v80_add_presentation(action_rows.to_dict('records'))
                st.success(f'✅ {n} yeni haber Sunum Sepeti’ne eklenmiştir.')
            elif do_note:
                # Bilgi notunda tam içerik gerekiyorsa yalnız burada ana tabloya dön.
                full=_v73_main_selected(selected_keys)
                if full.empty: full=action_rows
                with st.spinner(f'{len(full)} seçili haber için bilgi notu hazırlanmaktadır...'):
                    try:
                        st.session_state[f'section_note_bytes_{section_key}']=make_analyst_docx(
                            full,title='SANAYİ & TEKNOLOJİ BİLGİ NOTU'
                        )
                        _v63_mark_notes(full.to_dict('records'))
                        _v73_invalidate_status_cache()
                    except Exception as e:
                        st.session_state[f'section_note_bytes_{section_key}']=None
                        st.error(f'Bilgi notu hazırlanamadı: {e}')

    section_note_bytes=st.session_state.get(f'section_note_bytes_{section_key}')
    if section_note_bytes:
        st.download_button(
            '⬇️ Hazırlanan Bilgi Notunu İndir',
            data=section_note_bytes,
            file_name=f'Sanayi_Teknoloji_Bilgi_Notu_{section_key}_{date.today()}.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            use_container_width=True,
            key=f'v75_note_download_{section_key}'
        )
    return selected

def _collect_section_selected_from_main_df(df):
    if df is None or df.empty:
        return pd.DataFrame()
    keys=set()
    for selmap in st.session_state.section_selections.values():
        for k,v in selmap.items():
            if v:
                keys.add(str(k))
    if not keys:
        return pd.DataFrame()
    mask=_v73_row_keys(df).isin(keys)
    return df[mask].copy()


# -----------------------------
# V60 — OTOMATİK GERİ DÖNÜŞ / ANOMALİ / GÜN SONU
# -----------------------------
def _v60_register_visit_once():
    """
    Yeni browser oturumunda bir kez çalışır.
    Önceki giriş zamanını alır, mevcut girişi kaydeder.
    Streamlit rerun'larında baseline değişmez.
    """
    if st.session_state.get('_v60_visit_initialized',False):
        return st.session_state.get('_v60_previous_visit')

    previous=None
    now=datetime.now().astimezone()
    if _init_history_db():
        try:
            with _history_connect() as conn:
                row=conn.execute(
                    "SELECT visited_at FROM app_visits ORDER BY visited_at DESC LIMIT 1"
                ).fetchone()
                if row:
                    previous=pd.to_datetime(row[0],utc=True,errors='coerce')
                conn.execute(
                    "INSERT INTO app_visits(visited_at) VALUES(?)",
                    (now.isoformat(),)
                )
                conn.commit()
        except Exception:
            previous=None

    st.session_state['_v60_visit_initialized']=True
    st.session_state['_v60_previous_visit']=previous
    st.session_state['_v60_this_visit']=now
    return previous

def _v60_auto_catchup(previous_visit,user_query=''):
    """
    Kullanıcı yeniden giriş yaptığında manuel buton gerektirmeden,
    son girişten bu yana gelişmeleri hafif bir sorgu setiyle kontrol eder.
    Tam tarama değildir; yalnızca dönüş brifingi içindir.
    """
    if previous_visit is None or pd.isna(previous_visit):
        return [],None

    now_utc=datetime.now(timezone.utc)
    prev_utc=previous_visit.to_pydatetime() if hasattr(previous_visit,'to_pydatetime') else previous_visit
    if prev_utc.tzinfo is None:
        prev_utc=prev_utc.replace(tzinfo=timezone.utc)
    else:
        prev_utc=prev_utc.astimezone(timezone.utc)

    delta_h=max(0.25,(now_utc-prev_utc).total_seconds()/3600)
    # Google/RSS tarafında geniş pencere kullanılır; kesin filtre aşağıda previous_visit ile yapılır.
    when=period_window(max(3,delta_h))

    queries=[
        f'Türkiye (sanayi OR teknoloji OR üretim OR fabrika OR tesis OR yatırım OR OSB) when:{when}',
        f'Türkiye (savunma OR ASELSAN OR TUSAŞ OR ROKETSAN OR HAVELSAN OR Baykar OR otomotiv OR TOGG) when:{when}',
        f'Türkiye ("yapay zeka" OR "yarı iletken" OR çip OR siber OR Ar-Ge OR TÜBİTAK OR KOSGEB) when:{when}',
    ]
    queries += build_negative_queries(when)
    queries += build_official_radar_queries(when)

    raw=[]
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(7,len(queries))) as ex:
        futs=[ex.submit(rss,q) for q in queries]
        for f in concurrent.futures.as_completed(futs):
            try:
                raw.extend(f.result() or [])
            except Exception:
                pass

    rows,_=normalize_rows(raw,prev_utc,'turkish',user_query)
    rows=dedupe(rows)
    if rows:
        rows=enrich_rows(rows)
    return rows,delta_h

def _v60_now_to_know_table(rows,n=5):
    if not rows:
        return pd.DataFrame()
    df=pd.DataFrame(rows)
    if df.empty:
        return df
    value=_v52_event_value_table(df,max(n,10))
    if value.empty:
        return value
    return value.head(n).copy()

def _v60_anomaly_radar(df,current_hours,lookback_days=14):
    """
    Mevcut taramadaki kategori olay hızını, geçmiş günlerin son taramalarındaki
    saatlik olay hızıyla karşılaştırır. Ek web isteği yoktur.
    """
    cols=['Kategori','Şimdi','Beklenen','Normalin_Katı','Durum']
    if df is None or df.empty or not _init_history_db():
        return pd.DataFrame(columns=cols)

    try:
        cutoff=(datetime.now().astimezone()-timedelta(days=lookback_days)).isoformat()
        with _history_connect() as conn:
            hist=pd.read_sql_query("""
                SELECT s.scan_id,s.scanned_at,s.period_hours,e.category
                FROM scans s
                JOIN event_snapshots e ON e.scan_id=s.scan_id
                WHERE s.scanned_at>=?
                ORDER BY s.scanned_at DESC
            """,conn,params=(cutoff,))
    except Exception:
        return pd.DataFrame(columns=cols)

    if hist.empty:
        return pd.DataFrame(columns=cols)

    hist['day']=hist['scanned_at'].astype(str).str.slice(0,10)
    # Aynı gün çok tarama varsa yalnız o günün en son taraması baseline olur.
    last_scan_per_day=(
        hist[['day','scan_id','scanned_at']]
        .drop_duplicates()
        .sort_values('scanned_at')
        .groupby('day',as_index=False)
        .tail(1)[['day','scan_id']]
    )
    hist=hist.merge(last_scan_per_day,on=['day','scan_id'],how='inner')
    if hist['day'].nunique()<2:
        return pd.DataFrame(columns=cols)

    scan_hours=hist[['scan_id','period_hours']].drop_duplicates().set_index('scan_id')['period_hours'].to_dict()
    hc=hist.groupby(['scan_id','category']).size().reset_index(name='events')
    hc['rate']=hc.apply(
        lambda r:r['events']/max(1,float(scan_hours.get(r['scan_id'],24) or 24)),axis=1
    )
    baseline=hc.groupby('category')['rate'].agg(['mean','std','count']).reset_index()

    cur=df.groupby('Kategori')['Olay_ID'].nunique() if 'Olay_ID' in df.columns else df.groupby('Kategori').size()
    rows=[]
    for cat,current in cur.items():
        b=baseline[baseline['category']==cat]
        if b.empty:
            continue
        mean_rate=float(b.iloc[0]['mean'] or 0)
        expected=max(0.1,mean_rate*max(1,float(current_hours)))
        ratio=float(current)/expected if expected else 0
        # Hem göreli hem mutlak fark arıyoruz; küçük bazlarda sahte alarmı azaltır.
        if current>=3 and ratio>=1.8 and (current-expected)>=2:
            level='🔴 Çok Olağandışı' if ratio>=3 else '🟠 Olağandışı'
            rows.append({
                'Kategori':cat,
                'Şimdi':int(current),
                'Beklenen':round(expected,1),
                'Normalin_Katı':round(ratio,1),
                'Durum':level
            })

    if not rows:
        return pd.DataFrame(columns=cols)
    return pd.DataFrame(rows).sort_values(['Normalin_Katı','Şimdi'],ascending=[False,False])

def _v60_day_end_performance(df=None):
    """Bugünün operasyonel üretimini yerel geçmiş/sepet kayıtlarından özetler."""
    today=datetime.now().astimezone().date().isoformat()
    result={
        'Taramalar':0,'Benzersiz Olay':0,'Negatif':0,'Yüksek Risk':0,
        'Önemli Sepete Eklenen':0,'AKT Sepete Eklenen':0,'Kritik Sanayi':0
    }
    if _init_history_db():
        try:
            with _history_connect() as conn:
                result['Taramalar']=int(conn.execute(
                    "SELECT COUNT(*) FROM scans WHERE substr(scanned_at,1,10)=?",(today,)
                ).fetchone()[0] or 0)

                q="""SELECT COUNT(DISTINCT e.title)
                     FROM event_snapshots e JOIN scans s ON e.scan_id=s.scan_id
                     WHERE substr(s.scanned_at,1,10)=?"""
                result['Benzersiz Olay']=int(conn.execute(q,(today,)).fetchone()[0] or 0)

                qn="""SELECT COUNT(DISTINCT e.title)
                      FROM event_snapshots e JOIN scans s ON e.scan_id=s.scan_id
                      WHERE substr(s.scanned_at,1,10)=? AND e.sentiment='Negatif'"""
                result['Negatif']=int(conn.execute(qn,(today,)).fetchone()[0] or 0)

                qh="""SELECT COUNT(DISTINCT e.title)
                      FROM event_snapshots e JOIN scans s ON e.scan_id=s.scan_id
                      WHERE substr(s.scanned_at,1,10)=? AND e.risk_status='Yüksek Risk'"""
                result['Yüksek Risk']=int(conn.execute(qh,(today,)).fetchone()[0] or 0)

                result['Önemli Sepete Eklenen']=int(conn.execute(
                    "SELECT COUNT(*) FROM important_basket WHERE substr(added_at,1,10)=?",(today,)
                ).fetchone()[0] or 0)
                result['AKT Sepete Eklenen']=int(conn.execute(
                    "SELECT COUNT(*) FROM osint_report_basket WHERE substr(added_at,1,10)=?",(today,)
                ).fetchone()[0] or 0)
        except Exception:
            pass

    if df is not None and not df.empty:
        try:
            result['Kritik Sanayi']=int(df.apply(
                lambda r:bool(critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti',''))),
                axis=1
            ).sum())
        except Exception:
            pass
    return result


# ============================================================
# V110 — V109 PANEL HATA DÜZELTMESİ
# NOT: Bu override'lar bütün eski fonksiyon tanımlarından SONRA,
# UI başlamadan hemen önce tanımlanır. Böylece eski V33/V51 fonksiyonları
# yeni davranışı ezemez.
# ============================================================

def _official_radar_rows(df):
    """V110 — Resmî Kaynak Radarı her zaman 'Kurum Türü' sütununu üretir."""
    if df is None or df.empty:
        return pd.DataFrame()
    x=df[df.apply(_is_official_radar_row,axis=1)].copy()
    if x.empty:
        # UI KeyError vermesin.
        x['Kurum Türü']=pd.Series(dtype=str)
        return x
    x['Kurum Türü']=x.apply(_v109_official_source_type,axis=1)
    x['Tarih_dt']=pd.to_datetime(x.get('Tarih_dt'),utc=True,errors='coerce')
    x=x.sort_values('Tarih_dt',ascending=False,na_position='last')
    return x.drop_duplicates(subset=['URL','Başlık'])

def _v110_merge_change_rows(out):
    """
    Aynı yeni olay farklı başlıklarla geldiyse 'Dünden beri' tablosunda tek satıra indirir.
    Özellikle aynı yer + aynı olay kelimelerini taşıyan haberlerde daha toleranslıdır.
    """
    if out is None or out.empty:
        return out
    rows=[]
    for _,r in out.iterrows():
        rd=r.to_dict()
        merged=False
        for i,old in enumerate(rows):
            # Yalnız aynı değişim sınıfında birleştir.
            if str(old.get('Değişim',''))!=str(rd.get('Değişim','')):
                continue
            sim=_v104_event_similarity(
                old.get('Başlık',''),old.get('Ne Değişti?',''),
                rd.get('Başlık',''),rd.get('Ne Değişti?','')
            )
            # Aynı olay türü/yer için başlık token örtüşmesi.
            a=set(_title_tokens(old.get('Başlık','')))
            b=set(_title_tokens(rd.get('Başlık','')))
            overlap=len(a&b)/max(1,min(len(a),len(b))) if a and b else 0
            if sim>=0.42 or overlap>=0.55:
                # Daha yüksek riskli / daha çok kaynaklı satırı temsilci tut.
                old_score=(int(old.get('Risk',0) or 0),int(old.get('Kaynak Sayısı',0) or 0))
                new_score=(int(rd.get('Risk',0) or 0),int(rd.get('Kaynak Sayısı',0) or 0))
                if new_score>old_score:
                    rd['Kaynak Sayısı']=max(int(rd.get('Kaynak Sayısı',1) or 1),int(old.get('Kaynak Sayısı',1) or 1))
                    rows[i]=rd
                else:
                    rows[i]['Kaynak Sayısı']=max(int(old.get('Kaynak Sayısı',1) or 1),int(rd.get('Kaynak Sayısı',1) or 1))
                merged=True
                break
        if not merged:
            rows.append(rd)
    return pd.DataFrame(rows)

def _compare_since_previous(df,current_scan_id=None):
    """
    V110 — gerçek değişiklikleri gösterir ve kullanıcıya doğrudan ne değiştiğini söyler.
    'Yeni Olay' yalnız sınıflandırma türüdür; asıl bilgi 'Ne Değişti?' sütunundadır.
    """
    current=_v104_event_representatives(df)
    prev_id=_previous_scan_id(current_scan_id)
    previous=_load_scan_events(prev_id)
    if current is None or current.empty:
        return pd.DataFrame(),None,None
    if previous.empty:
        return pd.DataFrame(),prev_id,None

    prev_records=[p.to_dict() for _,p in previous.iterrows()]
    changes=[]

    for _,r in current.iterrows():
        c={
            'title':str(r.get('Başlık','') or ''),
            'source':str(r.get('Kaynak','') or ''),
            'url':str(r.get('URL','') or ''),
            'category':str(r.get('Kategori','') or ''),
            'summary':str(r.get('İçerik_Özeti','') or ''),
            'risk_score':int(r.get('Risk_Skoru',0) or 0),
            'risk_status':str(r.get('Risk_Durumu','') or ''),
            'verification':str(r.get('Doğrulama','') or ''),
            'source_count':int(r.get('Olay_Kaynak_Sayisi',1) or 1)
        }

        best=None; best_sim=0.0
        for pr in prev_records:
            sim=_v104_event_similarity(
                c['title'],c['summary'],
                pr.get('title',''),pr.get('summary','')
            )
            if c['url'] and c['url']==str(pr.get('url','') or ''):
                sim=max(sim,0.98)
            if sim>best_sim:
                best_sim=sim; best=pr

        if best is None or best_sim<0.50:
            kind='🆕 YENİ OLAY'
            priority=100+c['risk_score']
            prev_risk='—'
            # Genel metin yerine olayın kendisini doğrudan söyle.
            concise=_clean_note_text(c['summary'])
            first=_v109_sentences(concise)
            if first:
                detail=_v66_formalize_sentence_endings(first[0]).strip()
                if detail and detail[-1] not in '.!?': detail+='.'
                diff='Önceki taramada bulunmayan yeni gelişme tespit edilmiştir: '+detail
            else:
                diff='Önceki taramada bulunmayan yeni gelişme tespit edilmiştir: '+_clean_note_text(c['title']).rstrip('.')+'.'
        else:
            risk_up,verify_up,material,_,_=_v104_material_change(best,c)
            if risk_up:
                kind='⚠️ RİSK ARTTI'; priority=95+c['risk_score']
            elif verify_up:
                kind='✅ TEYİT GÜÇLENDİ'; priority=90+c['risk_score']
            elif material:
                kind='🔄 YENİ BİLGİ'; priority=80+c['risk_score']
            else:
                continue
            prev_risk=int(best.get('risk_score') or 0)
            diff=_v109_direct_difference(best,c,kind)

        changes.append({
            'Ne Değişti?':diff,
            'Tür':kind,
            'Başlık':c['title'],
            'Kaynak':c['source'],
            'Kategori':c['category'],
            'Risk':c['risk_score'],
            'Önceki Risk':prev_risk,
            'Kaynak Sayısı':c['source_count'],
            'URL':c['url'],
            '_priority':priority
        })

    out=pd.DataFrame(changes)
    if not out.empty:
        # Birleştirme fonksiyonunun mevcut isimle çalışması için geçici Değişim alanı.
        out['Değişim']=out['Tür']
        out=_v110_merge_change_rows(out)
        if not out.empty:
            if 'Tür' not in out.columns and 'Değişim' in out.columns:
                out['Tür']=out['Değişim']
            out=out.sort_values(['_priority','Risk'],ascending=[False,False],na_position='last')
            out=out.drop(columns=['_priority','Değişim'],errors='ignore')

    prev_time=str(previous.iloc[0].get('scanned_at','')) if not previous.empty else None
    return out,prev_id,prev_time

# ============================================================
# /V110
# ============================================================

# V110 — V109 düzeltmesi: override sırası düzeltildi; Resmî Kaynak Radarı KeyError giderildi;
# Dünden Beri Ne Değişti tablosunda 'Ne Değişti?' ana sütun haline getirildi ve benzer yeni olaylar birleştirildi.


# ============================================================
# V111 — DÜNDEN BERİ / PANEL DEVAMLILIĞI DÜZELTMESİ
# ============================================================

def _compare_since_previous(df,current_scan_id=None):
    """
    V111:
    - 'Ne Değişti?' başlığı tekrar etmez.
    - Yeni olayda yalnızca olayın önceki taramada bulunmadığını söyler;
      olayın kendisi zaten Başlık sütununda görülür.
    - Eski modüller için 'Değişim' alanı korunur; yeni UI için 'Tür' de bulunur.
      Böylece Bilgi Notu Adayları ve panelin devamı KeyError ile durmaz.
    """
    current=_v104_event_representatives(df)
    prev_id=_previous_scan_id(current_scan_id)
    previous=_load_scan_events(prev_id)

    empty_cols=['Ne Değişti?','Tür','Değişim','Başlık','Kaynak','Kategori',
                'Risk','Önceki Risk','Kaynak Sayısı','URL']

    if current is None or current.empty:
        return pd.DataFrame(columns=empty_cols),None,None
    if previous.empty:
        return pd.DataFrame(columns=empty_cols),prev_id,None

    prev_records=[p.to_dict() for _,p in previous.iterrows()]
    changes=[]

    for _,r in current.iterrows():
        c={
            'title':str(r.get('Başlık','') or ''),
            'source':str(r.get('Kaynak','') or ''),
            'url':str(r.get('URL','') or ''),
            'category':str(r.get('Kategori','') or ''),
            'summary':str(r.get('İçerik_Özeti','') or ''),
            'risk_score':int(r.get('Risk_Skoru',0) or 0),
            'risk_status':str(r.get('Risk_Durumu','') or ''),
            'verification':str(r.get('Doğrulama','') or ''),
            'source_count':int(r.get('Olay_Kaynak_Sayisi',1) or 1)
        }

        best=None
        best_sim=0.0
        for pr in prev_records:
            sim=_v104_event_similarity(
                c['title'],c['summary'],
                pr.get('title',''),pr.get('summary','')
            )
            if c['url'] and c['url']==str(pr.get('url','') or ''):
                sim=max(sim,0.98)
            if sim>best_sim:
                best_sim=sim
                best=pr

        if best is None or best_sim<0.50:
            kind='🆕 YENİ OLAY'
            priority=100+c['risk_score']
            prev_risk='—'
            # Başlık zaten ayrı sütunda. Burada onu yeniden yazma.
            if c['source']:
                diff=(
                    f"Bu olaya ilişkin kayıt önceki taramada bulunmamaktadır; "
                    f"gelişme mevcut taramada ilk kez {c['source']} kaynağında tespit edilmiştir."
                )
            else:
                diff=(
                    "Bu olaya ilişkin kayıt önceki taramada bulunmamaktadır; "
                    "gelişme mevcut taramada ilk kez tespit edilmiştir."
                )
        else:
            risk_up,verify_up,material,_,_=_v104_material_change(best,c)

            if risk_up:
                kind='⚠️ RİSK ARTTI'
                priority=95+c['risk_score']
            elif verify_up:
                kind='✅ TEYİT GÜÇLENDİ'
                priority=90+c['risk_score']
            elif material:
                kind='🔄 YENİ BİLGİ'
                priority=80+c['risk_score']
            else:
                continue

            prev_risk=int(best.get('risk_score') or 0)
            diff=_v109_direct_difference(best,c,kind)

        changes.append({
            'Ne Değişti?':diff,
            'Tür':kind,
            # Geriye dönük uyumluluk: mevcut aday/özet fonksiyonları bunu kullanıyor.
            'Değişim':kind,
            'Başlık':c['title'],
            'Kaynak':c['source'],
            'Kategori':c['category'],
            'Risk':c['risk_score'],
            'Önceki Risk':prev_risk,
            'Kaynak Sayısı':c['source_count'],
            'URL':c['url'],
            '_priority':priority
        })

    out=pd.DataFrame(changes)
    if out.empty:
        return pd.DataFrame(columns=empty_cols),prev_id,str(previous.iloc[0].get('scanned_at',''))

    # Aynı yeni olayın farklı kaynaklarını tek satırda birleştir.
    out=_v110_merge_change_rows(out)

    # Birleştirme sonrası iki isim de kesinlikle bulunsun.
    if 'Tür' not in out.columns and 'Değişim' in out.columns:
        out['Tür']=out['Değişim']
    if 'Değişim' not in out.columns and 'Tür' in out.columns:
        out['Değişim']=out['Tür']

    out=out.sort_values(['_priority','Risk'],ascending=[False,False],na_position='last')
    out=out.drop(columns=['_priority'],errors='ignore')

    # UI başlığı tekrar etmesin; fakat eski fonksiyonlar Değişim'i kullanabilsin.
    ordered=[c for c in empty_cols if c in out.columns]
    extras=[c for c in out.columns if c not in ordered]
    out=out[ordered+extras]

    prev_time=str(previous.iloc[0].get('scanned_at','')) if not previous.empty else None
    return out,prev_id,prev_time

# ============================================================
# /V111
# ============================================================

# V111 — Bilgi Notu Adayları KeyError düzeltildi; panelin devamı artık yüklenir.
# 'Ne Değişti?' yeni olaylarda başlığı tekrar etmez; Değişim/Tür çift alanı ile geriye dönük uyumluluk sağlanmıştır.


# ============================================================
# V112 — DURUM TARİHİ + HIZ OPTİMİZASYONU
# 1) Durum rozetlerinde işlemin yapıldığı tarih/saat gösterilir.
# 2) Bilgi notu üretiminde mevcut olay zenginleştirmesi + detay cache kullanılır.
# 3) Sepet silme işlemleri tek hızlı SQLite transaction ile yapılır.
# 4) Aynı taramadaki pahalı "Dünden beri" karşılaştırması rerun'larda cache'lenir.
# ============================================================

def _v112_status_key_variants(title='',url='',summary=''):
    keys=set()
    url=str(url or '').strip()
    title=str(title or '')
    summary=str(summary or '')
    if url:
        keys.add('U:'+url)
    tk=title_key(title)
    if tk:
        keys.add('T:'+tk)
    try:
        sig=' '.join(sorted(_v104_event_tokens(title,summary)))
        if sig:
            keys.add('E:'+sig)
    except Exception:
        pass
    return keys

def _v112_parse_status_time(value):
    try:
        dt=pd.to_datetime(value,utc=True,errors='coerce')
        if pd.isna(dt):
            return None
        return dt
    except Exception:
        return None

def _v112_format_status_time(value):
    dt=_v112_parse_status_time(value)
    if dt is None:
        return ''
    try:
        local_tz=datetime.now().astimezone().tzinfo
        dt=dt.tz_convert(local_tz)
        return dt.strftime('%d.%m.%Y %H:%M')
    except Exception:
        try:
            return dt.strftime('%d.%m.%Y %H:%M')
        except Exception:
            return ''

def _v112_status_history():
    """
    Her işlem için olay/URL/başlık anahtarını EN SON işlem tarihine eşler.
    Tek sorgu grubu + session cache: tüm paneller aynı veriyi tekrar tekrar okumaz.
    """
    cached=st.session_state.get('_v112_status_history_cache')
    if cached is not None:
        return cached

    result={'imp':{},'akt':{},'notes':{},'pres':{}}
    if not _init_history_db():
        return result

    specs=[
        ('important_basket','added_at','imp'),
        ('osint_report_basket','added_at','akt'),
        ('note_history','created_at','notes'),
        ('presentation_basket','added_at','pres'),
    ]
    try:
        with _history_connect() as conn:
            for table,time_col,key in specs:
                rows=conn.execute(
                    f"SELECT {time_col},title,url FROM {table} ORDER BY {time_col} DESC"
                ).fetchall()
                for ts,title,url in rows:
                    for k in _v112_status_key_variants(title,url,''):
                        # Sorgu DESC olduğu için ilk değer en yeni tarihtir.
                        if k not in result[key]:
                            result[key][k]=str(ts or '')
    except Exception:
        pass

    st.session_state['_v112_status_history_cache']=result
    return result

def _v63_status_sets():
    """Eski kodlarla uyumluluk için durum anahtarlarını set olarak döndürür."""
    h=_v112_status_history()
    return set(h['imp']),set(h['akt']),set(h['notes']),set(h['pres'])

def _v104_event_status_sets():
    """V104 olay bazlı durum altyapısının V112 tarihli cache ile uyumlu hali."""
    h=_v112_status_history()
    return {
        'imp':set(h['imp']),
        'akt':set(h['akt']),
        'notes':set(h['notes']),
        'pres':set(h['pres'])
    }

def _v73_invalidate_status_cache():
    st.session_state.pop('_v73_status_sets_cache',None)
    st.session_state.pop('_v104_status_cache',None)
    st.session_state.pop('_v112_status_history_cache',None)

def _v63_add_status_badges(df):
    """
    Durum örneği:
    📌 ÖGN — 22.08.2026 13:42 • 📝 Bilgi Notu — 22.08.2026 14:05
    """
    if df is None or df.empty:
        return df
    out=df.copy()
    hist=_v112_status_history()

    def badge(r):
        title=str(r.get('Başlık',r.get('title','')) or '')
        url=str(r.get('URL',r.get('url','')) or '').strip()
        summary=str(r.get('İçerik_Özeti',r.get('summary','')) or '')
        keys=_v112_status_key_variants(title,url,summary)

        def latest(bucket):
            vals=[hist[bucket].get(k) for k in keys if hist[bucket].get(k)]
            if not vals:
                return ''
            parsed=[(_v112_parse_status_time(v),v) for v in vals]
            parsed=[x for x in parsed if x[0] is not None]
            if parsed:
                return max(parsed,key=lambda x:x[0])[1]
            return vals[0]

        b=[]
        ts=latest('pres')
        if ts: b.append(f"🖥️ Sunum — {_v112_format_status_time(ts)}")
        ts=latest('imp')
        if ts: b.append(f"📌 ÖGN — {_v112_format_status_time(ts)}")
        ts=latest('notes')
        if ts: b.append(f"📝 Bilgi Notu — {_v112_format_status_time(ts)}")
        ts=latest('akt')
        if ts: b.append(f"📁 AKT — {_v112_format_status_time(ts)}")
        return ' • '.join(b) if b else '—'

    out['Durum']=out.apply(badge,axis=1)
    return out

def _v112_fast_delete(table, ids=None):
    """Silme için tek transaction; satır satır işlem ve gereksiz sorgu yoktur."""
    allowed={'important_basket','osint_report_basket','presentation_basket'}
    if table not in allowed or not _init_history_db():
        return 0
    try:
        with _history_connect() as conn:
            # Web uygulamasında silme gecikmesini azaltmak için küçük transaction.
            conn.execute("PRAGMA synchronous=NORMAL")
            if ids is None:
                cur=conn.execute(f"DELETE FROM {table}")
            else:
                ids=[int(x) for x in ids if str(x).isdigit()]
                if not ids:
                    return 0
                marks=','.join('?' for _ in ids)
                cur=conn.execute(f"DELETE FROM {table} WHERE id IN ({marks})",ids)
            conn.commit()
            removed=max(int(cur.rowcount or 0),0)
        if removed:
            _v73_invalidate_status_cache()
        return removed
    except Exception:
        return 0

def _remove_basket_ids(ids):
    return _v112_fast_delete('important_basket',ids)

def _clear_important_basket():
    return _v112_fast_delete('important_basket',None)

def _remove_osint_basket_ids(ids):
    return _v112_fast_delete('osint_report_basket',ids)

def _clear_osint_basket():
    return _v112_fast_delete('osint_report_basket',None)

def _v81_remove_presentation_ids(ids):
    return _v112_fast_delete('presentation_basket',ids)

def _v80_clear_presentation():
    return _v112_fast_delete('presentation_basket',None)

def _v112_detail_cache_key(row):
    url=str(row.get('URL',row.get('url','')) or '').strip()
    if url:
        return 'U:'+url
    return 'T:'+title_key(str(row.get('Başlık',row.get('title','')) or ''))

def _v112_cached_article_detail(row):
    """
    Aynı haber için tekrar Word üretildiğinde sayfayı yeniden indirmez.
    Cache yalnız mevcut kullanıcı oturumunda tutulur.
    """
    cache=st.session_state.setdefault('_v112_article_detail_cache',{})
    key=_v112_detail_cache_key(row)
    if key in cache:
        return cache[key]
    try:
        detail=article_detail(row) or {}
    except Exception:
        detail={}
    # Cache'in sınırsız büyümesini önle.
    if len(cache)>160:
        try:
            for old in list(cache.keys())[:40]:
                cache.pop(old,None)
        except Exception:
            cache={}
            st.session_state['_v112_article_detail_cache']=cache
    cache[key]=detail
    return detail

def make_analyst_docx(df, title='BİLGİ NOTU'):
    """
    V112 hızlı bilgi notu:
    - Önce V107 olay zenginleştirmesini kullanır.
    - Yeterince zengin mevcut özet varsa yeniden web isteği yapmaz.
    - Gerekli haber detaylarını paralel ve oturum cache'li alır.
    - V66 resmî dil / belge yapısı korunur.
    """
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    styles=doc.styles
    styles['Normal'].font.name='Times New Roman'
    styles['Normal'].font.size=Pt(12)
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(_clean_note_text(title)); r.bold=True; r.font.size=Pt(14)
    p=doc.add_paragraph(); p.add_run('Tarih: ').bold=True
    p.add_run(datetime.now().astimezone().strftime('%d.%m.%Y'))

    x=df.copy() if df is not None else pd.DataFrame()
    if x.empty:
        rows=[]
    else:
        if 'Tarih_dt' in x.columns:
            x['Tarih_dt']=pd.to_datetime(x['Tarih_dt'],utc=True,errors='coerce')
            x=x.sort_values('Tarih_dt',ascending=True,na_position='last')
        rows=x.to_dict('records')

    # Aynı olayın mevcut taramadaki daha iyi kaynaklarını önce birleştir.
    try:
        enriched_rows=_v107_enrich_selected_rows(rows)
        if enriched_rows:
            rows=enriched_rows
    except Exception:
        pass

    enriched=[None]*len(rows)

    def get_one(i,row):
        summary=_clean_note_text(row.get('İçerik_Özeti',''))
        # V107 zenginleştirmesi yeterli içerik sağladıysa web fetch'i atla.
        rich_enough=(
            len(summary)>=650
            and len(_sentence_chunks(summary))>=3
        )
        if rich_enough:
            return i,row,{}
        return i,row,_v112_cached_article_detail(row)

    if rows:
        workers=min(6,len(rows))
        if workers<=1:
            for i,row in enumerate(rows):
                _,rr,dd=get_one(i,row)
                enriched[i]=(rr,dd)
        else:
            with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as ex:
                futs=[ex.submit(get_one,i,row) for i,row in enumerate(rows)]
                for fut in concurrent.futures.as_completed(futs):
                    try:
                        i,rr,dd=fut.result()
                        enriched[i]=(rr,dd)
                    except Exception:
                        pass

    enriched=[x for x in enriched if x is not None]

    all_sent=[]
    for row,detail in enriched:
        title_text=_clean_note_text(detail.get('title') or row.get('Başlık',''))
        body=_clean_note_text(detail.get('text') or row.get('İçerik_Özeti') or title_text)
        all_sent.extend(_akt_clean_sentences(title_text,body))

    uniq=[]; seen=[]
    for sent in all_sent:
        sent=_clean_note_text(sent)
        key=norm(sent)
        toks=set(key.split())
        if not key: continue
        dup=False
        for old in seen[-35:]:
            union=len(toks|old)
            if union and len(toks&old)/union>=0.78:
                dup=True; break
        if not dup:
            uniq.append(sent.strip()); seen.append(toks)

    def add_body(text):
        bp=doc.add_paragraph()
        bp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.first_line_indent=Cm(1.25)
        bp.paragraph_format.line_spacing=1.15
        bp.paragraph_format.space_after=Pt(8)
        safe_text=_repair_mojibake_utf8(_clean_note_text(text))
        bp.add_run(_v66_formalize_sentence_endings(safe_text))

    if uniq:
        intro=_join_sentences_naturally(uniq[:2])
        add_body(intro)

        detail_s=uniq[2:] or uniq
        detail_s=detail_s[:18]
        if len(detail_s)<=9:
            add_body(_join_sentences_naturally(detail_s))
        else:
            add_body(_join_sentences_naturally(detail_s[:9]))
            add_body(_join_sentences_naturally(detail_s[9:18]))

        tail=_join_sentences_naturally(uniq[-3:])
        if tail:
            conclusion=(
                f"Mevcut bilgiler çerçevesinde, {tail[0].lower()+tail[1:]} "
                "Gelişmenin sanayi ve teknoloji alanındaki muhtemel etkilerinin, ilgili kurum ve kuruluşların "
                "yeni açıklamaları ile resmî veriler doğrultusunda takip edilmesinin uygun olacağı değerlendirilmektedir."
            )
        else:
            conclusion=(
                "Mevcut bilgiler çerçevesinde gelişmenin sanayi ve teknoloji alanındaki etkilerinin, ilgili kurum "
                "ve kuruluşların yeni açıklamaları ile resmî veriler doğrultusunda takip edilmesinin uygun olacağı değerlendirilmektedir."
            )
        add_body(conclusion)
    else:
        add_body('Seçilen habere ilişkin ayrıntılı içerik temin edilememiştir.')
        add_body(
            'Gelişmenin yeni açık kaynak bilgileri ile ilgili kurum ve kuruluşların resmî açıklamaları '
            'doğrultusunda takip edilmesinin uygun olacağı değerlendirilmektedir.'
        )

    endp=doc.add_paragraph()
    endp.paragraph_format.space_before=Pt(8)
    endp.add_run('Arz olunur.')

    if enriched:
        kp=doc.add_paragraph()
        kr=kp.add_run('Kaynak: '); kr.bold=True
        for i,(row,detail) in enumerate(enriched):
            source=_clean_note_text(detail.get('source') or row.get('Kaynak','Açık Kaynak'))
            url=detail.get('canonical') or row.get('Yayıncı_URL') or row.get('URL','')
            if i: kp.add_run('; ')
            kp.add_run(source)
            if url:
                kp.add_run(' ('); _word_hyperlink(kp,url,'Haber linki'); kp.add_run(')')

    bio=BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.getvalue()

# Aynı taramadaki pahalı karşılaştırmayı sepet silme gibi rerun'larda tekrar hesaplama.
_v112_compare_impl=_compare_since_previous

def _v112_scan_cache_key(df,current_scan_id=None):
    try:
        n=len(df)
    except Exception:
        n=0
    sid=current_scan_id or st.session_state.get('current_scan_id') or 'none'
    return f'{sid}:{n}'

def _compare_since_previous(df,current_scan_id=None):
    key=_v112_scan_cache_key(df,current_scan_id)
    cache=st.session_state.setdefault('_v112_compare_cache',{})
    if key in cache:
        out,prev_id,prev_time=cache[key]
        return out.copy(),prev_id,prev_time
    result=_v112_compare_impl(df,current_scan_id)
    out,prev_id,prev_time=result
    cache.clear()
    cache[key]=(out.copy(),prev_id,prev_time)
    return out,prev_id,prev_time

# ============================================================
# /V112
# ============================================================

# V112 — Durum alanında işlem tarih/saatleri gösterilir; bilgi notu ve sepet silme akışları hızlandırılmıştır.

# -----------------------------
# UI
# -----------------------------
st.title('🛡️ T.C. Sanayi ve Teknoloji Bakanlığı Açık Kaynak Tarama Merkezi')
st.caption('Hızlı ilk bakış · olay kümeleri · risk/negatif ayrımı · Türk medya önceliği · Yunan/Türk savunma · kaynak güvenilirliği · trend · alarm · seçilen haberlerden DOCX')
with st.sidebar:
    st.header('⚙️ Tarama Ayarları')
    default=('sanayi OR teknoloji OR üretim OR imalat OR fabrika OR OSB OR makine OR otomasyon OR robotik OR Ar-Ge OR patent OR yapay zeka OR yazılım OR siber güvenlik OR çip OR yarı iletken OR elektronik OR telekom OR kuantum OR biyoteknoloji OR nanoteknoloji OR savunma sanayii OR ASELSAN OR TUSAŞ OR ROKETSAN OR HAVELSAN OR Baykar OR İHA OR SİHA OR KAAN OR havacılık OR uzay OR uydu OR otomotiv OR TOGG OR batarya OR enerji OR hidrojen OR kimya OR petrokimya OR demir çelik OR madencilik OR tekstil OR gıda teknolojisi OR tarım teknolojisi OR lojistik OR tedarik zinciri OR TÜBİTAK OR KOSGEB OR teknopark OR yatırım teşvik OR yerlileştirme')
    query=st.text_area('Geniş sanayi / teknoloji sorgusu:',default,height=190)
    watch=st.text_area('⭐ Takip listesi (virgül / satır sonu):','ASELSAN, TUSAŞ, ROKETSAN, HAVELSAN, Baykar, TOGG, TÜBİTAK',height=90)
    neg=st.checkbox('⚠️ Negatif haberleri ayrıca tespit et',True)
    greek=st.checkbox('🇬🇷 Yunan medyası — yalnızca Türk savunma sanayii',True)
    social=st.checkbox('📱 Türk açık sosyal / indeks kaynakları',True)
    global_on=st.checkbox('🌍 Global basın (opsiyonel)',False)
    instant_alerts=st.checkbox('🔔 Tarama sırasında negatif/yüksek risk bildirimi göster',True,
                               help='Tarama devam ederken yeni negatif veya yüksek riskli içerik yakalanırsa ekranda anlık bildirim gösterir.')
    period=st.selectbox('🕒 Haber dönemi',['⚡ Son 3 saat','📅 Son 24 saat','📆 Son 48 saat','📆 Son 1 hafta','🗓️ Son 1 ay'],index=1)
    hours={'⚡ Son 3 saat':3,'📅 Son 24 saat':24,'📆 Son 48 saat':48,'📆 Son 1 hafta':168,'🗓️ Son 1 ay':720}[period]
    run=st.button('🔍 TARAMAYI BAŞLAT / YENİLE',type='primary',use_container_width=True)

if 'rows' not in st.session_state: st.session_state.rows=None
if 'scan_time' not in st.session_state: st.session_state.scan_time=None
if 'stats' not in st.session_state: st.session_state.stats={}
if 'last_scan_alerts' not in st.session_state: st.session_state.last_scan_alerts=[]
if 'docx_bytes' not in st.session_state: st.session_state.docx_bytes=None
if 'note_bytes' not in st.session_state: st.session_state.note_bytes=None

if 'current_scan_id' not in st.session_state: st.session_state.current_scan_id=None
if 'history_status' not in st.session_state: st.session_state.history_status=_init_history_db()
if 'basket_docx_bytes' not in st.session_state: st.session_state.basket_docx_bytes=None
if 'section_selections' not in st.session_state: st.session_state.section_selections={}

# V60: Yeni browser oturumunda önceki giriş zamanı otomatik belirlenir.
_v60_previous_visit=_v60_register_visit_once()
if '_v60_catchup_done' not in st.session_state:
    st.session_state['_v60_catchup_done']=False
if '_v60_catchup_rows' not in st.session_state:
    st.session_state['_v60_catchup_rows']=[]
if '_v60_catchup_hours' not in st.session_state:
    st.session_state['_v60_catchup_hours']=None

if not st.session_state['_v60_catchup_done']:
    st.session_state['_v60_catchup_done']=True
    if _v60_previous_visit is not None and not pd.isna(_v60_previous_visit):
        with st.spinner('⏱️ Son girişinizden bu yana gelişmeler otomatik kontrol ediliyor...'):
            _catch_rows,_catch_hours=_v60_auto_catchup(_v60_previous_visit,query)
            st.session_state['_v60_catchup_rows']=_catch_rows
            st.session_state['_v60_catchup_hours']=_catch_hours


if run:
    cutoff=(datetime.now(timezone.utc)-timedelta(hours=hours)).astimezone(timezone.utc)
    when=period_window(hours)
    batches=[('🇹🇷 Türk medya / sanayi-teknoloji',build_turkish_queries(when,query),'turkish')]
    # V41 bağımsız katmanları: yalnızca 4 ek sorgu; mevcut paralel havuzda çalışır.
    batches.append(('🏛️ Resmî kaynak radarı',build_official_radar_queries(when),'official'))
    batches.append(('📊 Önemli istatistik radarı',build_statistics_queries(when),'statistics'))
    if neg: batches.append(('⚠️ Negatif haber taraması',build_negative_queries(when),'negative'))
    if greek: batches.append(('🇬🇷 Yunan medyası / Türk savunma',build_greek_queries(when),'greek'))
    if social: batches.append(('📱 Açık sosyal / indeks',build_social_queries(when),'social'))
    if global_on: batches.append(('🌍 Global basın',[
        f'(Turkey OR Türkiye) (industry OR manufacturing OR technology OR semiconductor OR defense OR aerospace OR automotive) timespan:{when}',
        f'(Turkey OR Turkish) (Baykar OR ASELSAN OR TUSAŞ OR ROKETSAN OR HAVELSAN OR KAAN OR drone OR missile) timespan:{when}'
    ],'global'))
    all_rows=[]; stat={'Ham sonuç':0,'Zaman dışı':0,'Konu dışı':0,'Yunan dışı':0,'Kaynak dışı':0,'Sonuç':0,'Olay':0}
    live_alarm_box=st.empty()
    status_box=st.status('🔎 Tarama başlıyor...',expanded=True)

    alerted_keys=set()
    live_alerts=[]
    toast_count=0
    MAX_TOASTS_PER_SCAN=2

    def _alert_key(row):
        return row.get('URL') or title_key(row.get('Başlık',''))

    def _register_alert(row):
        key=_alert_key(row)
        if not key or key in alerted_keys:
            return False
        alerted_keys.add(key)
        risk_score=int(row.get('Risk_Skoru',row.get('Skor',0)) or 0)
        critical_label=critical_industrial_incident(row.get('Başlık',''),row.get('İçerik_Özeti',''))
        is_high=row.get('Risk_Durumu')=='Yüksek Risk' or risk_score>=70 or bool(critical_label)
        live_alerts.insert(0,{
            'Tarih':str(row.get('Tarih','')),
            'Seviye':critical_label if critical_label else ('YÜKSEK RİSK' if is_high else 'NEGATİF'),
            'Kaynak':str(row.get('Kaynak','Açık Kaynak')),
            'Başlık':str(row.get('Başlık','')),
            'Risk':risk_score,
            'URL':row.get('URL','')
        })
        del live_alerts[25:]
        return True

    def _merge_batch(raw,mode):
        nonlocal_dummy=None
        norm_rows,reasons=normalize_rows(raw,cutoff,mode,query)
        stat['Zaman dışı']+=reasons['zaman']
        stat['Konu dışı']+=reasons['konu']
        stat['Yunan dışı']+=reasons['yunan']
        stat['Kaynak dışı']+=reasons['kaynak']
        return norm_rows

    # 1) Türk ana taraması önce: kullanıcı ilk sonuçları en kısa sürede görsün.
    primary_label,primary_queries,primary_mode=batches[0]
    status_box.write(f'{primary_label} — {len(primary_queries)} sorgu / 10 eşzamanlı')
    primary_raw=[]
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(10,len(primary_queries))) as ex:
        futures=[ex.submit(rss,q) for q in primary_queries]
        for f in concurrent.futures.as_completed(futures):
            try:
                primary_raw.extend(f.result() or [])
            except Exception:
                pass
    stat['Ham sonuç']+=len(primary_raw)
    all_rows=dedupe(_merge_batch(primary_raw,primary_mode))
    stat['Sonuç']=len(all_rows)

    # V44: Hızlı İlk Bakış kaldırıldı.
    # Tarama sonuçları doğrudan aşağıdaki ana Görünüm ekranında (Kronolojik/Negatif/Yüksek Risk vb.) açılır.

    # 2) Negatif + Yunan + sosyal + global sorgularını TEK HAVUZDA paralel çalıştır.
    supplemental=batches[1:]
    jobs=[]
    for label,queries,mode in supplemental:
        for q in queries:
            jobs.append((label,q,mode))

    supplemental_raw_by_mode={}
    if jobs:
        status_box.write(f'⚡ Tamamlayıcı kaynaklar — {len(jobs)} sorgu / 12 eşzamanlı')
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(12,len(jobs))) as ex:
            future_map={ex.submit(rss,q):(label,mode) for label,q,mode in jobs}
            for fut in concurrent.futures.as_completed(future_map):
                label,mode=future_map[fut]
                try:
                    chunk=fut.result() or []
                except Exception:
                    chunk=[]
                stat['Ham sonuç']+=len(chunk)
                supplemental_raw_by_mode.setdefault(mode,[]).extend(chunk)

                # ÖZEL KRİTİK SANAYİ OLAYI ALARMI:
                # OSB/OSB dışı fabrika-tesis yangın ve patlamalarında sorgu döner dönmez bildir.
                if instant_alerts and chunk:
                    quick_rows,_quick_reasons=normalize_rows(chunk,cutoff,mode,query)
                    for qr in quick_rows:
                        critical_label=critical_industrial_incident(qr.get('Başlık',''),qr.get('İçerik_Özeti',''))
                        if critical_label:
                            qkey=_alert_key(qr)
                            if qkey not in alerted_keys:
                                _register_alert(qr)
                                icon='💥' if 'PATLAMA' in critical_label else '🔥'
                                st.toast(
                                    f'{critical_label}: {str(qr.get("Başlık",""))[:105]}',
                                    icon=icon
                                )
                                live_alarm_box.error(
                                    f'{icon} **KRİTİK SANAYİ OLAYI ALARMI — {critical_label}** — '
                                    f'{qr.get("Tarih","")} · {qr.get("Kaynak","Açık Kaynak")} · '
                                    f'{str(qr.get("Başlık",""))[:140]}'
                                )

    # Mode bazlı normalize + birleştirme.
    for mode,raw in supplemental_raw_by_mode.items():
        incoming=_merge_batch(raw,mode)
        old_keys={_alert_key(x) for x in all_rows}
        all_rows=dedupe(all_rows+incoming)
        # Genel negatif/yüksek risk alarmı bu aşamada verilmez.
        # Önce aşağıda gerçek haber sayfasının tam metni okunarak nihai sınıflandırma yapılır.
        # Kritik sanayi yangın/patlama anlık alarmı yukarıdaki özel blokta aynen devam eder.
        stat['Sonuç']=len(all_rows)

    # 3) Analitik katman — V44 performans düzenlemesi.
    # V43'teki her haber sayfasını tek tek indiren tam-metin negatif analizi kaldırıldı.
    # V42'deki hızlı ve bağlam duyarlı Başlık + RSS İçerik/Özet sınıflandırması kullanılır.
    if all_rows:
        status_box.write('🧩 Hızlı olay analizi hazırlanıyor...')
        all_rows=enrich_rows(all_rows)
        stat['Olay']=len({r.get('Olay_ID') for r in all_rows})
    else:
        stat['Olay']=0

    # Nihai alarm listesi mevcut hızlı sınıflandırmadan oluşturulur.
    # Kritik sanayi yangın/patlama alarmı aynen korunur.
    live_alerts=[]
    alerted_keys=set()
    final_toast_count=0
    for ar in all_rows:
        critical_label=critical_industrial_incident(ar.get('Başlık',''),ar.get('İçerik_Özeti',''))
        is_negative=(ar.get('Duygu')=='Negatif')
        is_high=(ar.get('Risk_Durumu')=='Yüksek Risk' or int(ar.get('Risk_Skoru',0) or 0)>=70)
        if critical_label or is_negative or is_high:
            if _register_alert(ar):
                if instant_alerts and not critical_label and final_toast_count < MAX_TOASTS_PER_SCAN:
                    st.toast(
                        f'{"🚨 YÜKSEK RİSK" if is_high else "⚠️ NEGATİF"}: {str(ar.get("Başlık",""))[:100]}',
                        icon='🚨' if is_high else '⚠️'
                    )
                    final_toast_count+=1

    if live_alerts:
        live_alarm_box.warning(
            f'🔔 {len(live_alerts)} negatif/riskli içerik yakalandı. Son: {live_alerts[0]["Başlık"][:100]}'
        )

    status_box.update(
        label=f'✅ Tarama tamamlandı — {len(all_rows)} haber / {stat["Olay"]} olay',
        state='complete'
    )
    # V101 — Tarama sonucu daha session_state'e yazılmadan gerçek yayın tarih-saatine göre sıralanır.
    # Böylece özellikle Son 24 Saat taramasında Kronolojik ekran ilk açılışta en yeni -> en eski gelir.
    def _v101_row_dt(_r):
        _d=_to_utc_datetime(_r.get('Tarih_dt'))
        if _d is None:
            _d=_to_utc_datetime(_r.get('Tarih'))
        return _d or datetime.min.replace(tzinfo=timezone.utc)

    all_rows=sorted(all_rows,key=_v101_row_dt,reverse=True)
    st.session_state.rows=all_rows
    st.session_state.scan_time=datetime.now().astimezone()
    st.session_state.stats=stat
    st.session_state.last_scan_alerts=live_alerts

    # V33 geçmiş karşılaştırma katmanı: tarama bittikten SONRA olay özetini kaydeder.
    # Tarama motoruna veya sıralamaya müdahale etmez.
    st.session_state.current_scan_id=_save_scan_history(
        all_rows,
        st.session_state.scan_time,
        hours
    )


# V60 — ŞU AN BİLMEN GEREKENLER: manuel çalışmaz, yeni oturumda otomatik hazırlanır.
st.subheader('⚡ Şu An Bilmen Gerekenler')
_prev=st.session_state.get('_v60_previous_visit')
_catch_rows=st.session_state.get('_v60_catchup_rows') or []
_catch_hours=st.session_state.get('_v60_catchup_hours')

if _prev is None or pd.isna(_prev):
    st.info('İlk giriş kaydı oluşturuldu. Bir sonraki girişinizde bu alan son girişinizden sonraki gelişmeleri otomatik gösterecek.')
else:
    try:
        _prev_local=pd.to_datetime(_prev,utc=True).tz_convert(datetime.now().astimezone().tzinfo)
        st.caption(f'Son giriş: {_prev_local.strftime("%d.%m.%Y %H:%M")} — bu tarihten sonraki gelişmeler otomatik kontrol edildi.')
    except Exception:
        pass

    _now5=_v60_now_to_know_table(_catch_rows,5)
    if _now5.empty:
        st.success('Son girişinizden bu yana öncelikli yeni bir gelişme tespit edilmedi.')
    else:
        st.warning(f'Son girişinizden bu yana dikkat gerektiren {_now5.shape[0]} gelişme öne çıkıyor.')

        # V61: Bu bölümden doğrudan seçim/sepet/bilgi notu işlemleri yapılabilir.
        _catch_df=pd.DataFrame(_catch_rows)
        _know_rows=[]
        for _,_v in _now5.iterrows():
            _url=str(_v.get('URL','') or '')
            _title=norm(_v.get('Gelişme',''))
            _match=pd.DataFrame()
            if not _catch_df.empty and _url and 'URL' in _catch_df.columns:
                _match=_catch_df[_catch_df['URL'].astype(str)==_url]
            if _match.empty and not _catch_df.empty and 'Başlık' in _catch_df.columns:
                _match=_catch_df[_catch_df['Başlık'].astype(str).map(norm)==_title]
            if not _match.empty:
                _r=_match.iloc[0].to_dict()
            else:
                _r={
                    'Tarih':_v.get('Tarih',''),'Başlık':_v.get('Gelişme',''),
                    'URL':_v.get('URL',''),'Risk_Skoru':_v.get('Risk',0),
                    'İçerik_Özeti':'','Kaynak':'','Kategori':''
                }
            _r['Değer_Skoru']=int(_v.get('Değer_Skoru',0) or 0)
            _r['Neden_Değerli']=_v.get('Neden_Değerli','')
            _r['Kaynak_Sayısı']=int(_v.get('Kaynak_Sayısı',0) or 0)
            _know_rows.append(_r)

        _know_select=pd.DataFrame(_know_rows)
        if 'Seç' not in _know_select.columns:
            _know_select.insert(0,'Seç',False)

        _edited_know=st.data_editor(
            _know_select[['Seç','Tarih','Başlık','İçerik_Özeti','Değer_Skoru',
                          'Neden_Değerli','Kaynak_Sayısı','Risk_Skoru','URL']],
            column_config={
                'Seç':st.column_config.CheckboxColumn('Seç'),
                'Değer_Skoru':st.column_config.ProgressColumn('Değer Skoru',min_value=0,max_value=100,format='%d/100'),
                'Risk_Skoru':st.column_config.NumberColumn('Risk',format='%d/100'),
                'URL':st.column_config.LinkColumn('Haber Linki'),
                'İçerik_Özeti':st.column_config.TextColumn('Kısa İçerik',width='large')
            },
            disabled=['Tarih','Başlık','İçerik_Özeti','Değer_Skoru','Neden_Değerli',
                      'Kaynak_Sayısı','Risk_Skoru','URL'],
            hide_index=True,use_container_width=True,
            height=min(480,100+62*len(_know_select)),
            key='v61_now_to_know_editor'
        )

        _selected_idx=_edited_know.index[_edited_know['Seç'].astype(bool)].tolist()
        _selected_know=_know_select.loc[_selected_idx].copy() if _selected_idx else pd.DataFrame()

        k1,k2,k3,k4=st.columns(4)
        with k1:
            if st.button('📌 Önemli Gelişmelere Ekle',key='v82_know_imp',use_container_width=True):
                if _selected_know.empty:
                    st.warning('Önce en az bir gelişmeyi seçin.')
                else:
                    _n=_add_rows_to_important_basket(_selected_know.to_dict('records'))
                    st.success(f'{_n} haber önemli gelişmeler sepetine eklendi.')
        with k2:
            if st.button('🗂️ AKT Sepetine Ekle',key='v82_know_akt',use_container_width=True):
                if _selected_know.empty:
                    st.warning('Önce en az bir gelişmeyi seçin.')
                else:
                    _n=_add_rows_to_osint_basket(_selected_know.to_dict('records'))
                    st.success(f'{_n} haber açık kaynak tarama sepetine eklendi.')
        with k3:
            if st.button('🖥️ Sunum Sepetine Ekle',key='v82_know_pres',use_container_width=True):
                if _selected_know.empty:
                    st.warning('Önce en az bir gelişmeyi seçin.')
                else:
                    _n=_v80_add_presentation(_selected_know.to_dict('records'))
                    st.success(f'{_n} haber sunum sepetine eklendi.')
        with k4:
            if st.button('📝 Detaylı Bilgi Notu Oluştur',key='v82_know_note',use_container_width=True):
                if _selected_know.empty:
                    st.warning('Önce en az bir gelişmeyi seçin.')
                else:
                    with st.spinner(f'{len(_selected_know)} seçili gelişmenin ayrıntılı bilgi notu hazırlanıyor...'):
                        try:
                            st.session_state['v61_know_note_bytes']=make_analyst_docx(
                                _selected_know,
                                title='SANAYİ & TEKNOLOJİ BİLGİ NOTU'
                            )
                            _v63_mark_notes(_selected_know.to_dict('records'))
                        except Exception as _e:
                            st.session_state['v61_know_note_bytes']=None
                            st.error(f'Bilgi notu hazırlanamadı: {_e}')

        if st.session_state.get('v61_know_note_bytes'):
            st.download_button(
                '⬇️ Hazırlanan Bilgi Notunu İndir',
                data=st.session_state['v61_know_note_bytes'],
                file_name=f'Sanayi_Teknoloji_Bilgi_Notu_Su_An_Bilmen_Gerekenler_{date.today()}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True,
                key='v61_know_note_download'
            )

st.markdown('---')

# ============================================================
# V68 — ANALİST KOMUTA MERKEZİ
# ============================================================
st.caption('⚡ V75 ultra hızlı mod: checkbox işlemleri form içinde tutulmakta; tik atmak tek başına uygulamayı yeniden çalıştırmamaktadır.')
st.subheader('🎛️ Analist Komuta Merkezi')
st.caption(
    'Bu alan çalışma saatine ve içeriğin niteliğine göre işlem önermektedir: Bilgi Notu için veri/istatistik, '
    'resmî açıklama ve ürün/teknoloji gelişmeleri; AKT için negatif/eleştirel/olumsuz veya propaganda niteliğindeki '
    'içerikler; sunum için resmî veri, resmî açıklama ve teyitli bilgiler esas alınmaktadır.'
)

_cmd_rows=st.session_state.get('rows')
if _cmd_rows:
    _cmd_df=pd.DataFrame(_cmd_rows)
    if not _cmd_df.empty and 'Tarih_dt' in _cmd_df.columns:
        _cmd_df['Tarih_dt']=pd.to_datetime(_cmd_df['Tarih_dt'],utc=True,errors='coerce')

    _cmd,_phase,_phase_hint=_v68_analyst_command_center(_cmd_df,8)

    cphase1,cphase2=st.columns([1,2])
    with cphase1:
        st.info(f'**Çalışma Fazı**\n\n{_phase}')
    with cphase2:
        st.info(f'**Sistem Önceliği**\n\n{_phase_hint}')

    if _cmd.empty:
        st.success('Şu anda ayrıca işlem önerilecek yüksek öncelikli bir gelişme bulunmamaktadır.')
    else:
        st.markdown(f'**Şimdi yapılması önerilen {len(_cmd)} işlem**')
        _cmd_for_select=_cmd.copy()
        _section_select_table(
            'v68_command_center',
            _cmd_for_select,
            ['Öncelik','Önerilen_İşlem','Tarih','Başlık','Neden','Durum',
             'Değer_Skoru','Risk_Skoru','URL'],
            height=min(620,105+58*len(_cmd_for_select))
        )
        st.caption(
            'Öneriler karar yerine geçmemektedir. Haberi seçerek doğrudan Önemli Gelişmeler, AKT veya '
            'Bilgi Notu işlemlerini aynı bölümden uygulayabilirsiniz.'
        )
else:
    st.info('İlk ana tarama tamamlandığında Analist Komuta Merkezi otomatik olarak işlem önerileri oluşturacaktır.')

st.markdown('---')

rows=st.session_state.rows
if rows is None:
    st.info('👋 Hazır. Tarama başlamaz. Zaman aralığını seçip **TARAMAYI BAŞLAT / YENİLE** düğmesine basın.')
else:
    # Tarama sırasında satırlar zaten enrich_rows() ile zenginleştiriliyor.
    # Checkbox / sekme / buton gibi UI etkileşimlerinde pahalı analizi tekrar çalıştırmıyoruz.
    df=pd.DataFrame(rows)
    if not df.empty:
        df['Tarih_dt']=pd.to_datetime(df['Tarih_dt'],utc=True,errors='coerce')
        df=df.sort_values('Tarih_dt',ascending=False,na_position='last').reset_index(drop=True)
    st.caption(f'Son tarama: {st.session_state.scan_time.strftime("%d.%m.%Y %H:%M:%S") if st.session_state.scan_time else "-"}')
    with st.expander('🧪 Tarama teşhisi',False): st.json(st.session_state.stats)
    if df.empty:
        st.warning('Sonuç bulunamadı. Tarama teşhisini açarak hangi aşamada sonuçların azaldığını görebilirsiniz.')
    else:
        total=len(df); negc=int((df.Duygu=='Negatif').sum()); riskc=int((df.Risk_Durumu=='Yüksek Risk').sum()); trc=int(df.Kaynak_Grubu.astype(str).str.startswith('🇹🇷').sum()); grc=int(df.Kaynak_Grubu.astype(str).str.startswith('🇬🇷').sum()); events=df['Olay_ID'].nunique()
        a,b,c,d,e,f=st.columns(6); a.metric('Toplam',total); b.metric('Olay',events); c.metric('Negatif',negc); d.metric('Yüksek Risk',riskc); e.metric('🇹🇷 Türk',trc); f.metric('🇬🇷 Yunan',grc)


        # ---------------------------------------------------------
        # V34 — VARDİYA BAŞLANGIÇ ÖZETİ
        # ---------------------------------------------------------
        # V46 — ANA GÖRÜNÜM EN ÜSTTE
        # Tarama tamamlandığında ilk bölüm doğrudan Kronolojik / Negatif / Yüksek Risk vb. ana haber görünümüdür.
        # Performans: Streamlit tabs içindeki TÜM içerikleri arka planda çalıştırır.
        # Bu nedenle tek seferde yalnızca seçilen görünümü üretiriz.
        st.subheader('🌅 Vardiya Başlangıç Özeti')
        st.caption('Sabah ilk analitik bakış: aynı olay tekilleştirilir; resmî veri/açıklama, stratejik sanayi-teknoloji gelişmesi, kritik negatif, savunma/uzay/teknoloji programı ve yüksek teyitli yeni gelişmeler önceliklendirilir.')
        shift_stats,shift_top,shift_baseline_label=_shift_start_summary(
            df,
            st.session_state.get('current_scan_id')
        )
        if shift_stats:
            st.caption(shift_stats.get('baseline_label',''))
            s1,s2,s3,s4,s5,s6=st.columns(6)
            s1.metric('Son devirden beri yeni haber',shift_stats['new_news'])
            s2.metric('Yeni önemli olay',shift_stats['new_important_events'])
            s3.metric('Yüksek riskli gelişme',shift_stats['high_risk'])
            s4.metric('Risk artışı',shift_stats['risk_up'])
            s5.metric('Teyit güçlenmesi',shift_stats['verify_up'])
            s6.metric('OSB olayı',shift_stats['osb'])

            st.markdown('**Sabah ilk bakılması gereken 5–8 gelişme**')
            if shift_top.empty:
                st.info('Öne çıkan gelişme bulunamadı.')
            else:
                _section_select_table(
                    'shift_top',
                    shift_top,
                    ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Risk_Durumu','Doğrulama','URL'],
                    height=min(330,70+45*len(shift_top))
                )

        c_shift1,c_shift2=st.columns([2,1])
        with c_shift1:
            st.caption('Devir noktası, bir sonraki vardiya başlangıç özetinin başlangıç zamanını belirler.')
        with c_shift2:
            if st.button('📍 ŞİMDİYİ DEVİR NOKTASI OLARAK KAYDET',use_container_width=True):
                if _mark_shift_handover(st.session_state.get('current_scan_id'),'Manuel devir noktası'):
                    st.success('Devir noktası kaydedildi.')
                else:
                    st.error('Devir noktası kaydedilemedi.')

        # ---------------------------------------------------------
        # V33 — DÜNDEN BERİ NE DEĞİŞTİ?
        # ---------------------------------------------------------
        st.subheader('🆕 Dünden Beri Ne Değişti?')
        changes,previous_scan_id,previous_scan_time=_compare_since_previous(
            df,
            st.session_state.get('current_scan_id')
        )
        if previous_scan_id is None:
            st.info(
                'Henüz karşılaştırılabilecek eski tarama bulunmuyor. '
                'Bu tarama yerel geçmişe kaydedildi; sonraki taramalarda yeni olaylar ve değişiklikler otomatik gösterilecek.'
            )
        else:
            if previous_scan_time:
                st.caption(f'Karşılaştırılan önceki tarama: {previous_scan_time}')
            if changes.empty:
                st.success('Önceki taramaya göre anlamlı yeni olay, risk artışı, teyit artışı veya içerik güncellemesi tespit edilmedi.')
            else:
                _change_type_col='Tür' if 'Tür' in changes.columns else 'Değişim'
                new_n=int(changes[_change_type_col].astype(str).str.contains('YENİ OLAY').sum())
                upd_n=int(changes[_change_type_col].astype(str).str.contains('YENİ BİLGİ').sum())
                risk_n=int(changes[_change_type_col].astype(str).str.contains('RİSK ARTTI').sum())
                ver_n=int(changes[_change_type_col].astype(str).str.contains('TEYİT').sum())
                q1,q2,q3,q4=st.columns(4)
                q1.metric('Yeni Olay',new_n)
                q2.metric('Yeni Bilgi',upd_n)
                q3.metric('Risk Artışı',risk_n)
                q4.metric('Teyit Güçlendi',ver_n)
                changes_view=changes.head(25).copy()
                _section_select_table(
                    'changes',
                    changes_view,
                    ['Ne Değişti?','Tür','Başlık','Kaynak','Kategori','Risk','Önceki Risk','Kaynak Sayısı','URL'],
                    height=min(560,70+35*min(len(changes_view),25))
                )

        # Son taramada anlık yakalanan bildirimlerin kalıcı özeti
        recent_alerts=st.session_state.get('last_scan_alerts',[])
        if recent_alerts:
            with st.expander(f'🔔 Son taramada yakalanan yeni negatif/riskli içerikler ({len(recent_alerts)})',False):
                alert_df=pd.DataFrame(recent_alerts)
                alert_view=alert_df.copy()
                if 'Risk' in alert_view.columns and 'Risk_Skoru' not in alert_view.columns:
                    alert_view['Risk_Skoru']=alert_view['Risk']
                _section_select_table(
                    'recent_alerts',
                    alert_view,
                    ['Tarih','Seviye','Kaynak','Başlık','Risk_Skoru','URL'],
                    height=min(420,42+35*len(alert_view))
                )

        # Alarm bandı
        alarms=df[(df.Risk_Skoru>=70) | (df.Duygu=='Negatif')].sort_values(['Risk_Skoru','Tarih_dt'],ascending=[False,False])
        if not alarms.empty:
            st.subheader('🚨 Yeni / Öncelikli Alarmlar')
            alarm_view=alarms.head(10).copy()
            _section_select_table(
                'priority_alarms',
                alarm_view,
                ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Risk_Gerekçesi','Doğrulama','URL'],
                height=min(470,70+38*len(alarm_view))
            )

        # Kritik sanayi olayları için SABİT bölüm.
        # Her zaman görünür; olay varsa içerik dolar, yoksa boş durum gösterilir.
        st.subheader('🚨 Kritik Sanayi Olayları — OSB / OSB Dışı Yangın ve Patlama')
        st.caption('OSB ve OSB dışındaki fabrika, tesis ve sanayi alanlarında tespit edilen yangın/patlama olayları burada sürekli izlenir.')

        critical_mask=df.apply(
            lambda r:bool(critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti',''))),
            axis=1
        )
        critical_events=df[critical_mask].copy().sort_values('Tarih_dt',ascending=False)

        if not critical_events.empty:
            critical_events['Kritik_Olay']=critical_events.apply(
                lambda r:critical_industrial_incident(r.get('Başlık',''),r.get('İçerik_Özeti','')) or '',
                axis=1
            )
            st.error(f'🚨 **KRİTİK SANAYİ OLAYI ALARMI — {len(critical_events)} içerik tespit edildi**')
            _section_select_table(
                'critical_industrial_events',
                critical_events,
                ['Tarih','Kritik_Olay','Kaynak','Başlık','Risk_Skoru','URL'],
                height=min(340,70+36*len(critical_events))
            )
        else:
            st.info('Bu tarama döneminde OSB / OSB dışı sanayi tesisi yangını veya patlaması tespit edilmedi.')

        st.markdown('---')
        st.subheader('🏆 Günün En Değerli 10 Gelişmesi')
        st.caption(
            'Aynı olaya ait haberlar tek gelişmede birleştirilir. Değer Skoru; önem/risk, farklı kaynak sayısı, '
            'resmî teyit, güncellik, stratejik sanayi-teknoloji önemi, negatif/eleştirel etki ve haber yoğunluğunu birlikte değerlendirir. '
            'Kaynak gerçek okunma/tıklanma verisi sağlıyorsa ileride ayrıca eklenebilir; mevcut sistem erişilemeyen okunma sayılarını tahmin etmez.'
        )
        value10=_v52_event_value_table(df,10)
        if value10.empty:
            st.info('Bu taramada sıralanabilecek gelişme bulunamadı.')
        else:
            _section_select_table(
                'daily_top10_value',
                value10,
                ['Sıra','Değer_Skoru','Tarih','Gelişme','Neden_Değerli',
                 'Kaynak_Sayısı','Haber_Sayısı','Resmî_Teyit','Risk','URL'],
                height=min(680,105+55*len(value10))
            )

            if st.button('📊 BUGÜNÜN DURUM ÖZETİNİ OLUŞTUR',use_container_width=True,key='v54_top10_summary_btn'):
                with st.spinner('Yalnızca en değerli 10 gelişmenin haber içerikleri okunuyor ve özetleniyor...'):
                    summary_text=_v54_deep_top10_summary(df,value10,45)
                    st.session_state.daily_summary_text=summary_text
                    st.session_state.daily_summary_bytes=make_v54_top10_summary_docx(df,value10,summary_text)

            if st.session_state.get('daily_summary_text'):
                st.text_area(
                    'Bugünün Durum Özeti — En Değerli 10 Gelişme',
                    st.session_state.daily_summary_text,
                    height=520,
                    key='v54_daily_summary_preview'
                )
                st.caption(
                    'Özet yalnızca yukarıdaki 10 gelişmenin haber içeriğini anlatır; değer skoru, kaynak sayısı ve '
                    'sıralama gerekçeleri metne eklenmez. Toplam çıktı 45 satırı geçmez.'
                )
                if st.session_state.get('daily_summary_bytes'):
                    st.download_button(
                        '⬇️ BUGÜNÜN DURUM ÖZETİNİ WORD OLARAK İNDİR',
                        data=st.session_state.daily_summary_bytes,
                        file_name=f'bugunun_durum_ozeti_top10_{datetime.now().strftime("%Y%m%d_%H%M")}.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True,
                        key='v54_top10_summary_download'
                    )

        # V33 — BİLGİ NOTU ADAYLARI
        # ---------------------------------------------------------
        st.subheader('🎯 Bilgi Notu Adayları')
        st.caption('Mevcut taramadaki olayları risk, teyit, kaynak sayısı, stratejik önem ve yenilik açısından puanlar.')
        candidate_count=st.slider('Gösterilecek aday sayısı',5,15,10,1,key='candidate_count')
        candidates=_information_note_candidates(
            df,
            st.session_state.get('current_scan_id'),
            candidate_count
        )
        if candidates.empty:
            st.info('Bu taramada bilgi notu adayı oluşturulamadı.')
        else:
            _section_select_table(
                'candidates',
                candidates,
                ['Aday Puanı','Başlık','Kaynak','Kategori','Risk','Kaynak Sayısı','Doğrulama','Değişim','Neden Bilgi Notu?','URL'],
                height=min(470,65+36*len(candidates))
            )

        # ---------------------------------------------------------

        st.markdown('---')
        st.subheader('👀 Kaçırıyor Olabilir Miyim? — İkinci Göz')
        st.caption(
            'Mevcut taramada yüksek değer/risk taşıdığı hâlde henüz Önemli Gelişmeler veya '
            'Açık Kaynak Tarama sepetine alınmamış olayları otomatik gösterir.'
        )
        _missed=_v63_missed_candidates(df,12)
        if _missed.empty:
            st.success('Şu anda sepetler dışında kalan belirgin yüksek değerli bir gelişme görünmüyor.')
        else:
            st.warning(f'Henüz hiçbir sepete alınmamış {_missed.shape[0]} dikkat çekici gelişme var.')
            _section_select_table(
                'v63_missed',
                _missed.rename(columns={'Gelişme':'Başlık'}),
                ['Tarih','Başlık','Değer_Skoru','Neden_Değerli','Kaynak_Sayısı','Risk','URL'],
                height=min(620,100+48*len(_missed))
            )

        view=st.radio(
            'Görünüm',
            ['📰 Kronolojik','⚠️ Negatif','🚨 Yüksek Risk','🇹🇷 Türk','🇬🇷 Yunan','🧩 Olaylar','📈 Trend / Analiz','⭐ Takip Listesi'],
            horizontal=True,
            key='main_view'
        )

        cols=['Seç','Tarih','Kaynak_Grubu','Kaynak','Kategori','Başlık','İçerik_Özeti','Duygu','Risk_Skoru','Risk_Durumu','Kaynak_Güvenilirliği','Doğrulama','URL']

        if view=='📰 Kronolojik':
            st.caption(
                '☑️ Hızlı işlem modu: kutucuklara tıklarken sayfa yeniden çalıştırılmaz. '
                'Seçiminizi yaptıktan sonra aşağıdaki işlem düğmelerinden birine basmanız yeterlidir.'
            )
            group_events=st.toggle(
                '🧩 Aynı olayı tek satırda göster',
                value=True,
                key='v109_chron_group_events',
                help='Aynı gelişmenin farklı kaynaklardaki haberlerini tek olay satırında birleştirir.'
            )
            chronology_base=_v109_chronology_events(df) if group_events else df.copy()

            page_size=40
            total_pages=max(1,(len(chronology_base)+page_size-1)//page_size)
            page_no=st.number_input(
                'Sayfa',min_value=1,max_value=total_pages,value=1,step=1,
                key='news_page'
            )
            start_i=(int(page_no)-1)*page_size
            end_i=min(start_i+page_size,len(chronology_base))
            page_df=chronology_base.iloc[start_i:end_i].copy()

            page_df['İçerik_Özeti']=page_df['İçerik_Özeti'].astype(str).str.slice(0,220)
            page_df=_v63_add_status_badges(page_df)
            chron_cols=[
                'Seç','Tarih','Kaynak_Grubu','Kaynak','Kaynak Sayısı','Haber Sayısı',
                'Kategori','Başlık','Durum','İçerik_Özeti','Duygu','Risk_Skoru',
                'Risk_Durumu','Kaynak_Güvenilirliği','Doğrulama','URL'
            ]
            chron_cols=[c for c in chron_cols if c in page_df.columns]

            st.caption(f'{start_i+1}-{end_i} / {len(chronology_base)} kayıt' + (' (olay bazlı)' if group_events else ' haber'))

            # FORM: checkbox tıklamaları rerun yapmaz. Yalnız işlem butonuna basınca tek rerun olur.
            with st.form(
                key=f'v74_chronology_fast_form_{int(page_no)}',
                clear_on_submit=False
            ):
                edited=st.data_editor(
                    page_df[chron_cols],
                    column_config={
                        'Seç':st.column_config.CheckboxColumn('Seç'),
                        'URL':st.column_config.LinkColumn('Haber Linki'),
                        'İçerik_Özeti':st.column_config.TextColumn('Kısa İçerik',width='large'),
                        'Risk_Skoru':st.column_config.NumberColumn('Risk',format='%d/100'),
                        'Kaynak Sayısı':st.column_config.NumberColumn('Kaynak',format='%d'),
                        'Haber Sayısı':st.column_config.NumberColumn('Haber',format='%d'),
                        'Durum':st.column_config.TextColumn('Durum',width='large')
                    },
                    disabled=[x for x in chron_cols if x!='Seç'],
                    hide_index=True,
                    use_container_width=True,
                    height=535,
                    key=f'v74_chron_editor_{int(page_no)}'
                )

                st.markdown('### ⚡ Seçilen Haberlerle Hızlı İşlem')
                c1,c2,c3,c4=st.columns(4)
                with c1: do_imp=st.form_submit_button('📌 Önemli Gelişmelere Ekle',use_container_width=True)
                with c2: do_akt=st.form_submit_button('🗂️ AKT Sepetine Ekle',use_container_width=True)
                with c3: do_pres=st.form_submit_button('🖥️ Sunum Sepetine Ekle',use_container_width=True)
                with c4: do_note=st.form_submit_button('📝 Detaylı Bilgi Notu Oluştur',use_container_width=True)

            if do_imp or do_akt or do_note or do_pres:
                selected_mask=edited['Seç'].astype(bool).to_numpy()
                selected_page=page_df.loc[selected_mask].copy()

                if selected_page.empty:
                    st.warning('Önce en az bir haberi işaretleyin.')
                elif do_imp:
                    n=_v74_fast_add_important(selected_page.to_dict('records'))
                    st.success(f'✅ {n} yeni haber Önemli Gelişmeler Sepeti’ne eklenmiştir.')
                elif do_akt:
                    n=_v74_fast_add_osint(selected_page.to_dict('records'))
                    st.success(f'✅ {n} yeni haber Açık Kaynak Tarama Sepeti’ne eklenmiştir.')
                elif do_pres:
                    n=_v80_add_presentation(selected_page.to_dict('records'))
                    st.success(f'✅ {n} yeni haber Sunum Sepeti’ne eklenmiştir.')
                elif do_note:
                    with st.spinner(
                        f'{len(selected_page)} seçili haber için ayrıntılı bilgi notu hazırlanmaktadır...'
                    ):
                        try:
                            # Tam içerik için kısa page_df yerine ana df'deki aynı URL'leri kullan.
                            selected_urls=set(selected_page['URL'].fillna('').astype(str))
                            full_selected=df[df['URL'].fillna('').astype(str).isin(selected_urls)].copy()
                            if full_selected.empty:
                                full_selected=selected_page.copy()
                            st.session_state['v74_chron_note_bytes']=make_analyst_docx(
                                full_selected,
                                title='SANAYİ & TEKNOLOJİ BİLGİ NOTU'
                            )
                            _v63_mark_notes(full_selected.to_dict('records'))
                            _v73_invalidate_status_cache()
                            st.success('✅ Bilgi notu hazırlanmıştır.')
                        except Exception as e:
                            st.session_state['v74_chron_note_bytes']=None
                            st.error(f'Bilgi notu hazırlanamadı: {e}')

            if st.session_state.get('v74_chron_note_bytes'):
                st.download_button(
                    '⬇️ KRONOLOJİDEN HAZIRLANAN BİLGİ NOTUNU İNDİR',
                    data=st.session_state['v74_chron_note_bytes'],
                    file_name=f'Sanayi_Teknoloji_Bilgi_Notu_{date.today()}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True,
                    key='v74_chron_note_download'
                )


            if group_events and not page_df.empty and '_Olay_ID' in page_df.columns:
                with st.expander('🔎 Olayın tüm kaynaklarını aç',False):
                    _event_options={
                        f"{str(r.get('Başlık',''))[:120]} — {int(r.get('Kaynak Sayısı',1) or 1)} kaynak":str(r.get('_Olay_ID',''))
                        for _,r in page_df.iterrows()
                    }
                    if _event_options:
                        _event_label=st.selectbox(
                            'Kaynaklarını görmek istediğiniz olay',
                            list(_event_options.keys()),
                            key=f'v109_event_source_select_{int(page_no)}'
                        )
                        _event_sources=_v109_event_sources(df,_event_options.get(_event_label))
                        if _event_sources.empty:
                            st.info('Bu olay için ayrıntılı kaynak kaydı bulunamadı.')
                        else:
                            st.dataframe(
                                _event_sources[['Tarih','Kaynak','Başlık','İçerik_Özeti','URL']].head(20),
                                column_config={
                                    'URL':st.column_config.LinkColumn('Haber Linki'),
                                    'İçerik_Özeti':st.column_config.TextColumn('Kısa İçerik',width='large')
                                },
                                hide_index=True,use_container_width=True,
                                height=min(520,80+42*len(_event_sources.head(20)))
                            )

        elif view=='⚠️ Negatif':
            _section_select_table(
                'negative_view',
                df[df.Duygu=='Negatif'],
                ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Risk_Gerekçesi','Doğrulama','URL'],
                height=600
            )

        elif view=='🚨 Yüksek Risk':
            _section_select_table(
                'highrisk_view',
                df[df.Risk_Durumu=='Yüksek Risk'],
                ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Risk_Gerekçesi','Doğrulama','URL'],
                height=600
            )

        elif view=='🇹🇷 Türk':
            _section_select_table(
                'turkish_view',
                df[df.Kaynak_Grubu.astype(str).str.startswith('🇹🇷')],
                ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Duygu','URL'],
                height=600
            )

        elif view=='🇬🇷 Yunan':
            _section_select_table(
                'greek_view',
                df[df.Kaynak_Grubu.astype(str).str.startswith('🇬🇷')],
                ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Duygu','URL'],
                height=600
            )

        elif view=='🧩 Olaylar':
            ev=build_event_summary(df)
            st.dataframe(ev,hide_index=True,use_container_width=True,height=480)
            chosen=st.selectbox('Olay zaman çizelgesini göster:',ev['Olay_ID'].tolist() if not ev.empty else [])
            if chosen:
                g=df[df.Olay_ID==chosen].sort_values('Tarih_dt',ascending=True)
                _section_select_table(
                    f'event_{chosen}',
                    g,
                    ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Doğrulama','URL'],
                    height=min(500,80+40*len(g))
                )

        elif view=='📈 Trend / Analiz':
            st.subheader('📊 Konu yoğunluğu')
            tr=trend_table(df)
            if not tr.empty:
                st.bar_chart(tr.set_index('Kategori')['Haber'])
            st.subheader('📈 Gündem yoğunluğu')
            tmp=df[df['Tarih_dt'].notna()].copy()
            tmp['Saat']=tmp['Tarih_dt'].dt.strftime('%Y-%m-%d %H:00')
            if not tmp.empty:
                st.line_chart(tmp.groupby('Saat').size())
            st.subheader('🧭 Yoğun konular')
            for _,r in tr.head(10).iterrows():
                st.write(f"**{r['Kategori']}** — {int(r['Haber'])} haber")

        elif view=='⭐ Takip Listesi':
            hits=watchlist_hits(df,watch)
            st.write(f'Listede eşleşen: **{len(hits)}** haber')
            if not hits.empty:
                _section_select_table(
                    'watchlist_view',
                    hits,
                    ['Tarih','Kaynak','Kategori','Başlık','Risk_Skoru','Duygu','URL'],
                    height=550
                )




        # V34 — ÖNEMLİ GELİŞMELER SEPETİ
        # ---------------------------------------------------------
        st.subheader('📌 24 Saatlik Önemli Gelişmeler Sepeti')
        st.caption('Gün boyunca önemli gördüğünüz haberleri burada biriktirin; vardiya sonunda Word olarak alın.')

        # V74: Kronoloji artık kendi hızlı işlem düğmelerine sahiptir.
        # Burada yalnız diğer bölümlerdeki seçili kayıtlar toplanır.
        selected_from_sections=_collect_section_selected_from_main_df(df)

        if st.button('➕ BÖLÜMLERDE İŞARETLEDİKLERİMİ ÖNEMLİ GELİŞMELER SEPETİNE EKLE',use_container_width=True):
            if selected_from_sections.empty:
                st.warning('Önce herhangi bir bölümde haberlerin yanındaki kutucuklardan seçim yapın.')
            else:
                added=_add_rows_to_important_basket(selected_from_sections.to_dict('records'))
                st.success(f'{added} yeni gelişme sepete eklendi.')

        basket=_load_important_basket()
        if basket.empty:
            st.info('Önemli gelişmeler sepeti şu anda boş.')
        else:
            # -----------------------------------------------------
            # V78 — ÖGN SEPETİ: SİLME ve BİLGİ NOTU TAMAMEN AYRI
            # -----------------------------------------------------
            basket_view=basket[['id','news_time','source','category','title','risk_score','risk_status','url']].copy()
            basket_view=_v63_add_status_badges(basket_view)

            # A) Sadece silme işlemi için checkbox.
            delete_view=basket_view.copy()
            delete_view.insert(0,'Sil',False)
            with st.form('v78_important_basket_delete_form',clear_on_submit=False):
                edited_delete=st.data_editor(
                    delete_view,
                    column_config={
                        'Sil':st.column_config.CheckboxColumn('Sil'),
                        'url':st.column_config.LinkColumn('Haber Linki'),
                        'risk_score':st.column_config.NumberColumn('Risk',format='%d/100'),
                        'Durum':st.column_config.TextColumn('Durum',width='large')
                    },
                    disabled=[c for c in delete_view.columns if c!='Sil'],
                    hide_index=True,use_container_width=True,
                    height=min(430,80+36*len(delete_view)),
                    key='v78_important_basket_delete_editor'
                )
                remove_btn=st.form_submit_button(
                    '🗑️ İŞARETLENENLERİ SEPETTEN ÇIKAR',
                    use_container_width=True
                )

            if remove_btn:
                ids=edited_delete.loc[edited_delete['Sil']==True,'id'].astype(int).tolist()
                removed=_remove_basket_ids(ids)
                st.success(f'{removed} kayıt sepetten çıkarıldı.')

            # B) Bilgi notunda TEK HABER seçilir. Sepetin tamamı hiçbir şekilde
            # make_analyst_docx'e gönderilmez.
            st.markdown('### 📝 Sepetten Seçilen Tek Haberden Detaylı Bilgi Notu')
            option_rows=[]
            for _,r in basket.iterrows():
                clean_title=_clean_note_text(r.get('title',''))
                option_rows.append((
                    int(r.get('id')),
                    f"{clean_title} — {_clean_note_text(r.get('source',''))}"
                ))

            label_to_id={label:rid for rid,label in option_rows}
            selected_label=st.selectbox(
                'Bilgi notu oluşturulacak haber',
                options=list(label_to_id.keys()),
                key='v78_ogn_note_single_select'
            ) if option_rows else None

            if st.button(
                '📝 SEÇİLEN TEK HABERDEN DETAYLI BİLGİ NOTU OLUŞTUR',
                use_container_width=True,
                key='v78_ogn_note_single_button'
            ):
                if not selected_label:
                    st.warning('Bilgi notu için bir haber seçin.')
                else:
                    selected_id=int(label_to_id[selected_label])
                    # Kesin tek satır: ID eşleşmesi + head(1).
                    selected_basket=basket[basket['id'].astype(int)==selected_id].head(1).copy()

                    if selected_basket.empty:
                        st.error('Seçilen haber sepette bulunamadı.')
                    else:
                        r=selected_basket.iloc[0]
                        important_note_rows=pd.DataFrame([{
                            'Tarih':_clean_note_text(r.get('news_time','')),
                            'Kaynak':_clean_note_text(r.get('source','')),
                            'Başlık':_clean_note_text(r.get('title','')),
                            'İçerik_Özeti':_clean_note_text(r.get('summary','')),
                            'URL':str(r.get('url','') or ''),
                            'Kategori':_clean_note_text(r.get('category','')),
                            'Risk_Skoru':r.get('risk_score',0),
                            'Risk_Durumu':_clean_note_text(r.get('risk_status',''))
                        }])

                        # Güvenlik kontrolü: make_analyst_docx'e asla 1'den fazla satır gitmesin.
                        important_note_rows=important_note_rows.head(1)

                        with st.spinner('Seçilen tek haberin tam metni okunuyor ve detaylı bilgi notu hazırlanıyor...'):
                            try:
                                st.session_state['v78_ogn_note_bytes']=make_analyst_docx(
                                    important_note_rows,
                                    title='SANAYİ & TEKNOLOJİ BİLGİ NOTU'
                                )
                                st.session_state['v78_ogn_note_title']=important_note_rows.iloc[0]['Başlık']
                                _v63_mark_notes(important_note_rows.to_dict('records'))
                                _v73_invalidate_status_cache()
                                st.success('✅ Bilgi notu yalnızca seçilen tek haberden hazırlanmıştır.')
                            except Exception as e:
                                st.session_state['v78_ogn_note_bytes']=None
                                st.error(f'Bilgi notu hazırlanamadı: {e}')

            if st.session_state.get('v78_ogn_note_bytes'):
                st.info(
                    'Bilgi notuna alınan tek haber: '
                    + _clean_note_text(st.session_state.get('v78_ogn_note_title',''))
                )
                st.download_button(
                    '⬇️ SEÇİLEN TEK HABERİN DETAYLI BİLGİ NOTUNU İNDİR',
                    data=st.session_state['v78_ogn_note_bytes'],
                    file_name=f'OGN_Secilen_Haber_Bilgi_Notu_{date.today()}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True,
                    key='v78_ogn_note_download'
                )

            if selected_label and st.button('🖥️ SEÇİLEN ÖNEMLİ GELİŞMEYİ SUNUM SEPETİNE EKLE',use_container_width=True,key='v81_ogn_to_pres'):
                _one=basket[basket['id'].astype(int)==int(label_to_id[selected_label])].head(1)
                st.success(f"✅ {_v80_add_presentation(_v81_basket_to_rows(_one))} haber Sunum Sepeti’ne eklenmiştir.")

            # V90: önceki sürümlerden kalan Word bytes kesinlikle kullanılmaz.
            if st.session_state.get('_ogn_engine_version') != V90_OGN_ENGINE_VERSION:
                st.session_state['_ogn_engine_version']=V90_OGN_ENGINE_VERSION
                st.session_state.pop('v90_ogn_docx_bytes',None)
                st.session_state.pop('basket_docx_bytes',None)

            b1,b2=st.columns(2)
            with b1:
                if st.button('📄 ÖNEMLİ GELİŞMELER WORD OLUŞTUR',use_container_width=True,key='v90_make_ogn_word'):
                    # Her basışta eski çıktı silinir ve V90 motoruyla baştan hazırlanır.
                    st.session_state.pop('v90_ogn_docx_bytes',None)
                    with st.spinner('Önemli gelişmeler gerçek haber içeriklerinden resmî biçimde özetleniyor...'):
                        st.session_state['v90_ogn_docx_bytes']=make_important_basket_docx_v101(basket)
                if st.session_state.get('v90_ogn_docx_bytes'):
                    st.download_button(
                        '⬇️ 24 SAATLİK ÖNEMLİ GELİŞMELER / WORD — V102',
                        st.session_state['v90_ogn_docx_bytes'],
                        file_name=f'24_Saatlik_Onemli_Gelismeler_V102_{date.today()}.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True,
                        key='v90_download_ogn_word'
                    )
            with b2:
                if st.button('🧹 SEPETİ TAMAMEN TEMİZLE',use_container_width=True):
                    removed=_clear_important_basket()
                    st.success(f'{removed} kayıt silindi.')


        st.markdown('---')
        st.subheader('🗂️ Açık Kaynak Tarama Çalışması Sepeti')
        st.caption('14:00 açık kaynak tarama raporuna girecek haberleri gün boyunca ayrı bir sepette biriktirin.')

        osint_selected_now=df[df.get('Seç',False)==True] if 'Seç' in df.columns else pd.DataFrame()
        osint_selected_sections=_collect_section_selected_from_main_df(df)

        o1,o2=st.columns(2)
        with o1:
            if st.button('➕ KRONOLOJİDE SEÇİLİ HABERLERİ AKT SEPETİNE EKLE',use_container_width=True):
                if osint_selected_now.empty:
                    st.warning('Önce kronolojik görünümden haber seçin ve seçimleri kaydedin.')
                else:
                    added=_add_rows_to_osint_basket(osint_selected_now.to_dict('records'))
                    st.success(f'{added} haber AKT sepetine eklendi.')
        with o2:
            if st.button('➕ BÖLÜMLERDE İŞARETLEDİKLERİMİ AKT SEPETİNE EKLE',use_container_width=True):
                if osint_selected_sections.empty:
                    st.warning('Önce herhangi bir bölümde seçim yapın.')
                else:
                    added=_add_rows_to_osint_basket(osint_selected_sections.to_dict('records'))
                    st.success(f'{added} haber AKT sepetine eklendi.')

        osint_basket=_load_osint_basket()
        if osint_basket.empty:
            st.info('Açık kaynak tarama çalışması sepeti boş.')
        else:
            # -----------------------------------------------------
            # V79 — AKT SEPETİ: ÖGN İLE AYNI TEK HABER BİLGİ NOTU MANTIĞI
            # -----------------------------------------------------
            osint_view=osint_basket[['id','news_time','source','category','title','risk_score','risk_status','url']].copy()
            osint_view=_v63_add_status_badges(osint_view)

            # A) Silme işlemi ayrı checkbox formunda kalır.
            delete_osint_view=osint_view.copy()
            delete_osint_view.insert(0,'Sil',False)
            with st.form('v79_osint_basket_delete_form',clear_on_submit=False):
                edited_osint=st.data_editor(
                    delete_osint_view,
                    column_config={
                        'Sil':st.column_config.CheckboxColumn('Sil'),
                        'url':st.column_config.LinkColumn('Haber Linki'),
                        'risk_score':st.column_config.NumberColumn('Risk',format='%d/100'),
                        'Durum':st.column_config.TextColumn('Durum',width='large')
                    },
                    disabled=[c for c in delete_osint_view.columns if c!='Sil'],
                    hide_index=True,use_container_width=True,
                    height=min(430,80+36*len(delete_osint_view)),
                    key='v79_osint_basket_delete_editor'
                )
                remove_osint=st.form_submit_button(
                    '🗑️ İŞARETLENENLERİ AKT SEPETİNDEN ÇIKAR',
                    use_container_width=True
                )

            if remove_osint:
                ids=edited_osint.loc[edited_osint['Sil']==True,'id'].astype(int).tolist()
                removed=_remove_osint_basket_ids(ids)
                st.success(f'{removed} kayıt AKT sepetinden çıkarıldı.')

            # AKT raporu sepetin tamamından hazırlanabilir; bu davranış korunur.
            osint_rows=[]
            for _,r in osint_basket.iterrows():
                osint_rows.append({
                    'Tarih':_clean_note_text(r.get('news_time','')),
                    'Kaynak':_clean_note_text(r.get('source','')),
                    'Başlık':_clean_note_text(r.get('title','')),
                    'İçerik_Özeti':_clean_note_text(r.get('summary','')),
                    'URL':str(r.get('url','') or ''),
                    'Kategori':_clean_note_text(r.get('category','')),
                    'Risk_Skoru':r.get('risk_score',0),
                    'Risk_Durumu':_clean_note_text(r.get('risk_status','')),
                    'Yayıncı':_clean_note_text(r.get('source','')),
                    'Yayıncı_URL':''
                })

            _akt_pres_opts={f"{_clean_note_text(r.get('title',''))} — {_clean_note_text(r.get('source',''))}":int(r.get('id')) for _,r in osint_basket.iterrows()}
            _akt_pres_label=st.selectbox('Sunuma eklenecek AKT haberi',list(_akt_pres_opts.keys()),key='v81_akt_pres_select') if _akt_pres_opts else None
            if st.button('🖥️ SEÇİLEN AKT HABERİNİ SUNUM SEPETİNE EKLE',use_container_width=True,key='v81_akt_to_pres'):
                if _akt_pres_label:
                    _one=osint_basket[osint_basket['id'].astype(int)==_akt_pres_opts[_akt_pres_label]].head(1)
                    st.success(f"✅ {_v80_add_presentation(_v81_basket_to_rows(_one))} haber Sunum Sepeti’ne eklenmiştir.")

            ob1,ob2=st.columns(2)
            with ob1:
                if st.button('📝 AKT SEPETİNDEN WORD HAZIRLA',use_container_width=True,key='v79_akt_report'):
                    with st.spinner('AKT sepetindeki haberler rapora hazırlanıyor...'):
                        st.session_state.docx_bytes=make_docx(osint_rows)
                if st.session_state.get('docx_bytes'):
                    st.download_button(
                        '⬇️ AKT SEPETİNDEN AÇIK KAYNAK RAPORU / WORD',
                        st.session_state.docx_bytes,
                        file_name=f'Sanayi_Teknoloji_Acik_Kaynak_Sepet_{date.today()}.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True,
                        key='v79_akt_report_download'
                    )
            with ob2:
                if st.button('🧹 AKT SEPETİNİ TAMAMEN TEMİZLE',use_container_width=True,key='v79_clear_akt'):
                    removed=_clear_osint_basket()
                    st.success(f'{removed} kayıt silindi.')

            # B) Bilgi notu için yalnız TEK HABER seçilir.
            st.markdown('### 📝 AKT Sepetinden Seçilen Tek Haberden Detaylı Bilgi Notu')

            akt_option_rows=[]
            for _,r in osint_basket.iterrows():
                clean_title=_clean_note_text(r.get('title',''))
                akt_option_rows.append((
                    int(r.get('id')),
                    f"{clean_title} — {_clean_note_text(r.get('source',''))}"
                ))

            akt_label_to_id={label:rid for rid,label in akt_option_rows}
            selected_akt_label=st.selectbox(
                'Bilgi notu oluşturulacak AKT haberi',
                options=list(akt_label_to_id.keys()),
                key='v79_akt_note_single_select'
            ) if akt_option_rows else None

            if st.button(
                '📝 SEÇİLEN TEK AKT HABERİNDEN DETAYLI BİLGİ NOTU OLUŞTUR',
                use_container_width=True,
                key='v79_akt_note_single_button'
            ):
                if not selected_akt_label:
                    st.warning('Bilgi notu için bir AKT haberi seçin.')
                else:
                    selected_akt_id=int(akt_label_to_id[selected_akt_label])
                    # Kesin tek satır: ID eşleşmesi ve head(1).
                    selected_akt=osint_basket[
                        osint_basket['id'].astype(int)==selected_akt_id
                    ].head(1).copy()

                    if selected_akt.empty:
                        st.error('Seçilen AKT haberi sepette bulunamadı.')
                    else:
                        r=selected_akt.iloc[0]
                        akt_note_df=pd.DataFrame([{
                            'Tarih':_clean_note_text(r.get('news_time','')),
                            'Kaynak':_clean_note_text(r.get('source','')),
                            'Başlık':_clean_note_text(r.get('title','')),
                            'İçerik_Özeti':_clean_note_text(r.get('summary','')),
                            'URL':str(r.get('url','') or ''),
                            'Kategori':_clean_note_text(r.get('category','')),
                            'Risk_Skoru':r.get('risk_score',0),
                            'Risk_Durumu':_clean_note_text(r.get('risk_status',''))
                        }]).head(1)

                        with st.spinner('Seçilen tek AKT haberinin tam metni okunuyor ve detaylı bilgi notu hazırlanıyor...'):
                            try:
                                st.session_state['v79_akt_note_bytes']=make_analyst_docx(
                                    akt_note_df,
                                    title='SANAYİ & TEKNOLOJİ BİLGİ NOTU'
                                )
                                st.session_state['v79_akt_note_title']=akt_note_df.iloc[0]['Başlık']
                                _v63_mark_notes(akt_note_df.to_dict('records'))
                                _v73_invalidate_status_cache()
                                st.success('✅ Bilgi notu yalnızca seçilen tek AKT haberinden hazırlanmıştır.')
                            except Exception as e:
                                st.session_state['v79_akt_note_bytes']=None
                                st.error(f'Bilgi notu hazırlanamadı: {e}')

            if st.session_state.get('v79_akt_note_bytes'):
                st.info(
                    'Bilgi notuna alınan tek AKT haberi: '
                    + _clean_note_text(st.session_state.get('v79_akt_note_title',''))
                )
                st.download_button(
                    '⬇️ SEÇİLEN TEK AKT HABERİNİN DETAYLI BİLGİ NOTUNU İNDİR',
                    data=st.session_state['v79_akt_note_bytes'],
                    file_name=f'AKT_Secilen_Haber_Bilgi_Notu_{date.today()}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True,
                    key='v79_akt_note_download'
                )


        st.markdown('---')
        st.subheader('🖥️ Sunum Sepeti')
        st.caption('Önemli Gelişmeler ve AKT sepetlerinin hemen altında yer almaktadır.')
        _pb=_v80_load_presentation()
        if _pb.empty:
            st.info('Sunum sepeti boş.')
        else:
            _pv=_pb[['id','news_time','source','title','url']].copy()
            _pv=_v63_add_status_badges(_pv)
            _pv.insert(0,'Seç',False)
            with st.form('v81_presentation_basket_form',clear_on_submit=False):
                _ped=st.data_editor(_pv,column_config={
                    'Seç':st.column_config.CheckboxColumn('Seç'),
                    'url':st.column_config.LinkColumn('Haber Linki'),
                    'Durum':st.column_config.TextColumn('Durum',width='large')
                },
                    disabled=[c for c in _pv.columns if c!='Seç'],hide_index=True,use_container_width=True,height=min(420,80+36*len(_pv)))
                p1,p2,p3,p4=st.columns(4)
                with p1: _toimp=st.form_submit_button('📌 Önemli Gelişmelere Ekle',use_container_width=True)
                with p2: _toakt=st.form_submit_button('🗂️ AKT Sepetine Ekle',use_container_width=True)
                with p3: _pnote=st.form_submit_button('📝 Bilgi Notu Oluştur',use_container_width=True)
                with p4: _prem=st.form_submit_button('🗑️ Sepetten Çıkar',use_container_width=True)
            _ids=_ped.loc[_ped['Seç']==True,'id'].astype(int).tolist()
            _sel=_pb[_pb['id'].astype(int).isin(_ids)]
            _rows=_v81_basket_to_rows(_sel)
            if _toimp:
                if _rows: st.success(f"✅ {_v74_fast_add_important(_rows)} haber Önemli Gelişmeler Sepeti’ne eklenmiştir.")
                else: st.warning('Önce haber seçin.')
            if _toakt:
                if _rows: st.success(f"✅ {_v74_fast_add_osint(_rows)} haber AKT Sepeti’ne eklenmiştir.")
                else: st.warning('Önce haber seçin.')
            if _prem:
                st.success(f"✅ {_v81_remove_presentation_ids(_ids)} haber çıkarılmıştır.")
            if _pnote:
                if len(_rows)!=1: st.warning('Detaylı bilgi notu için yalnızca bir haber seçin.')
                else:
                    with st.spinner('Seçilen sunum haberinden detaylı bilgi notu hazırlanıyor...'):
                        st.session_state['v81_pres_note_bytes']=make_analyst_docx(pd.DataFrame(_rows).head(1),title='SANAYİ & TEKNOLOJİ BİLGİ NOTU')
            if st.session_state.get('v81_pres_note_bytes'):
                st.download_button('⬇️ SUNUM SEPETİ BİLGİ NOTUNU İNDİR',st.session_state['v81_pres_note_bytes'],
                    file_name=f'Sunum_Sepeti_Bilgi_Notu_{date.today()}.docx',mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True,key='v81_pres_note_download')

        st.markdown('---')
        st.subheader('🏛️ Resmî Kaynak Radarı')
        st.caption('Sanayi ve Teknoloji Bakanlığı, TÜBİTAK, KOSGEB, TÜRKPATENT, TSE, SSB, TÜİK ve diğer birincil kamu kaynaklarından gelen içerikleri ayrı gösterir.')
        official_radar=_official_radar_rows(df)
        if official_radar.empty:
            st.info('Bu taramada resmî/birincil kaynaklardan eşleşen yeni içerik bulunamadı.')
        else:
            if 'Kurum Türü' not in official_radar.columns:
                official_radar=official_radar.copy()
                official_radar['Kurum Türü']=official_radar.apply(_v109_official_source_type,axis=1)
            _types=['Tümü']+[
                x for x in ['Bakanlık','TÜİK','TÜBİTAK','KOSGEB','TÜRKPATENT','TSE','SSB','Resmî Gazete','Diğer Resmî']
                if x in set(official_radar['Kurum Türü'].astype(str))
            ]
            _official_type=st.radio('Kaynak türü',_types,horizontal=True,key='v109_official_source_type')
            _official_show=official_radar if _official_type=='Tümü' else official_radar[official_radar['Kurum Türü']==_official_type]
            _section_select_table(
                'official_radar_'+re.sub(r'[^a-zA-Z0-9]+','_',norm(_official_type)),
                _official_show.head(30),
                ['Tarih','Kurum Türü','Kaynak','Kategori','Başlık','İçerik_Özeti','Risk_Skoru','Doğrulama','URL'],
                height=min(600,90+38*min(len(_official_show),30))
            )

        st.markdown('---')
        st.subheader('🧭 Olay Yaşam Döngüsü')
        st.caption(
            'Aynı olayın mevcut taramadaki gelişim aşamasını otomatik gösterir: '
            'İlk Sinyal → Gelişiyor → Teyit Edildi → Sonuçlandı. Bu alan sabittir.'
        )
        lifecycle=_v58_event_lifecycle_table(df,25)
        if lifecycle.empty:
            st.info('Bu taramada yaşam döngüsü oluşturulabilecek olay bulunamadı.')
        else:
            _section_select_table(
                'v58_event_lifecycle',
                lifecycle,
                ['Tarih','Aşama','Başlık','Kategori','Kaynak_Sayısı','Haber_Sayısı',
                 'Doğrulama','Risk_Skoru','Aşama_Gerekçesi','URL'],
                height=min(700,100+40*len(lifecycle))
            )


        st.markdown('---')
        st.subheader('📋 Gün Sonu Performans Özeti')
        st.caption('Bugün sistemde oluşan tarama ve çalışma çıktılarının operasyonel özeti.')
        _perf=_v60_day_end_performance(df)
        p1,p2,p3,p4,p5,p6,p7=st.columns(7)
        p1.metric('Tarama',_perf['Taramalar'])
        p2.metric('Benzersiz Olay',_perf['Benzersiz Olay'])
        p3.metric('Negatif',_perf['Negatif'])
        p4.metric('Yüksek Risk',_perf['Yüksek Risk'])
        p5.metric('Önemli Sepet',_perf['Önemli Sepete Eklenen'])
        p6.metric('AKT Sepet',_perf['AKT Sepete Eklenen'])
        p7.metric('Kritik Sanayi',_perf['Kritik Sanayi'])

        st.write(
            f"Bugün {_perf['Taramalar']} tarama gerçekleştirilmiş; geçmiş kayıtlarında "
            f"{_perf['Benzersiz Olay']} benzersiz olay, {_perf['Negatif']} negatif ve "
            f"{_perf['Yüksek Risk']} yüksek riskli gelişme kaydedilmiştir. "
            f"{_perf['Önemli Sepete Eklenen']} içerik önemli gelişmeler sepetine, "
            f"{_perf['AKT Sepete Eklenen']} içerik açık kaynak tarama sepetine eklenmiştir."
        )

        st.markdown('---'); st.subheader('📝 Seçilen haberlerden çıktı üret')
        # Form gönderildiyse session_state güncellenmiştir; aksi halde mevcut kayıtlı seçimleri kullan.
        current_rows=st.session_state.rows or []
        selected_df=pd.DataFrame(current_rows)
        if not selected_df.empty and 'Tarih_dt' in selected_df.columns:
            selected_df['Tarih_dt']=pd.to_datetime(selected_df['Tarih_dt'],utc=True,errors='coerce')
        selected=selected_df[selected_df.get('Seç',False)==True] if not selected_df.empty and 'Seç' in selected_df.columns else pd.DataFrame()
        st.write(f'{len(selected)} haber seçildi. Tarama sırasında görsel/tam metin indirilmez; yalnızca seçtiğiniz içerikler için derin zenginleştirme yapılır.')
        c1,c2=st.columns(2)
        with c1:
            if st.button('📝 AÇIK KAYNAK RAPORU / WORD',type='primary',use_container_width=True):
                if selected.empty: st.warning('Önce haber seçin.')
                else:
                    with st.spinner(f'{len(selected)} haber zenginleştiriliyor...'): st.session_state.docx_bytes=make_docx(selected.to_dict('records'))
            if st.session_state.docx_bytes: st.download_button('⬇️ Açık Kaynak Raporu DOCX',st.session_state.docx_bytes,file_name=f'Sanayi_Teknoloji_Acik_Kaynak_{date.today()}.docx',mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',use_container_width=True)
        with c2:
            if st.button('📌 AYRINTILI BİLGİ NOTU / WORD',use_container_width=True):
                if selected.empty: st.warning('Önce haber seçin.')
                else:
                    with st.spinner(f'{len(selected)} haberin tam haber metni okunuyor ve ayrıntılı bilgi notu hazırlanıyor...'):
                        st.session_state.note_bytes=make_analyst_docx(selected,title='SANAYİ & TEKNOLOJİ BİLGİ NOTU')
            if st.session_state.note_bytes: st.download_button('⬇️ Bilgi Notu DOCX',st.session_state.note_bytes,file_name=f'Sanayi_Teknoloji_Bilgi_Notu_{date.today()}.docx',mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',use_container_width=True)

st.caption('İlk açılışta otomatik tarama yoktur. Her yenileme yeni ağ taraması yapar. Haberler en yeni → en eski sıralanır; olay kümeleri, risk gerekçesi, kaynak güvenilirliği, doğrulama, trend ve takip listesi tarama sonucunda yer alır. DOCX aşamasında seçilen haberlerin gerçek yayıncı sayfası, görseli, linki ve geniş içeriği alınır.')
